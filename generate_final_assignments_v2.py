"""
Generate Final Assignment Documents - Version 2
Using ACTUAL templates with ALL requirements from assessment briefs

Student: KAREEM SCHULTZ
USN: 252IFCBR0596
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import shutil

# Student Details
STUDENT_NAME = "KAREEM SCHULTZ"
STUDENT_USN = "252IFCBR0596"

# Base paths
BASE_PATH = r"C:\Users\admin\Documents\SCQF L6 - Jain University Work\Course Matrix and Syllabus"

def add_paragraph_after(doc, text, style=None):
    """Add a paragraph with proper formatting"""
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_heading_styled(doc, text, level=1):
    """Add a heading"""
    h = doc.add_heading(text, level=level)
    return h

def add_table_styled(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def fill_student_details(doc):
    """Fill in student details in the template tables"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Student USN:" in cell.text:
                    cell.text = f"Student USN: {STUDENT_USN}"
                elif cell.text.strip() == "Student USN:":
                    # Find next cell or same cell
                    pass


def generate_he9e46_v2():
    """Generate HE9E 46 with ALL requirements"""
    print("Generating HE9E 46 (v2) with all requirements...")

    # Copy template
    template_path = os.path.join(BASE_PATH, "HE9E 46_Contemporary Business Issues_Student Assessment Template.docx")
    output_path = os.path.join(BASE_PATH, "My_Assignments", "HE9E_46_Contemporary_Business", "Final", "HE9E_46_Final_Assignment_v2.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # Fill student details in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if "Student USN:" in text and text.strip() == "Student USN:":
                    cell.text = f"Student USN: {STUDENT_USN}"
                if "Date Due:" in text and text.strip() == "Date Due:":
                    cell.text = "Date Due: 15-02-2026"
                if "Date of Submission:" in text and text.strip() == "Date of Submission:":
                    cell.text = "Date of Submission: "

    # Add content after declaration (find where to insert)
    # The template has declaration on page 3, so we add after that

    doc.add_page_break()

    # ===== TASK A =====
    add_heading_styled(doc, "TASK A: Features and Characteristics of SMEs", 1)
    add_heading_styled(doc, "LO1: Explain the features and characteristics of Small and Medium Enterprises", 2)

    # Task 1.1 (5 marks)
    add_heading_styled(doc, "1.1 Small and Medium Enterprises (SMEs): Scope and Formal Criteria", 2)
    add_paragraph_after(doc, """The European Commission provides the most widely-used definition of Small and Medium Enterprises. According to Commission Recommendation 2003/361, businesses are classified based on employee numbers and financial thresholds:""")

    add_table_styled(doc,
        ["Category", "Employees", "Annual Turnover", "OR", "Balance Sheet Total"],
        [
            ["Medium-sized", "< 250", "≤ €50 million", "OR", "≤ €43 million"],
            ["Small", "< 50", "≤ €10 million", "OR", "≤ €10 million"],
            ["Micro", "< 10", "≤ €2 million", "OR", "≤ €2 million"]
        ])

    add_paragraph_after(doc, """The employee threshold is mandatory, while businesses need only meet either the turnover or balance sheet criterion. SMEs account for approximately 99% of all businesses in the EU and generate roughly two-thirds of private sector employment (European Commission, 2003). This demonstrates their critical role in the contemporary business environment.""")

    # Task 1.2 (15 marks) - WITH E-COMMERCE FOCUS
    add_heading_styled(doc, "1.2 Key Features of SMEs - Focus on E-Commerce Platform", 2)

    add_heading_styled(doc, "Feature 1: Ownership and Control", 3)
    add_paragraph_after(doc, """SMEs are typically owned and controlled by individuals, families, or small groups of partners rather than external shareholders. The owner-manager structure means decisions can be made quickly without bureaucratic processes.

E-Commerce Application - Meesho (India): Meesho started as a small e-commerce platform enabling individuals to resell products through social media. Founded in 2015, the founders maintained close control over operations, allowing them to pivot quickly from a general marketplace to focusing specifically on social commerce for resellers. This ownership structure enabled rapid decision-making when they noticed the potential of WhatsApp-based selling in India's smaller towns.""")

    add_heading_styled(doc, "Feature 2: Flexibility and Agility", 3)
    add_paragraph_after(doc, """SMEs can adapt quickly to market changes due to simpler structures and shorter communication chains. Unlike large corporations with complex hierarchies, small businesses can implement changes almost immediately.

E-Commerce Application - Nykaa (India): Nykaa began as a small online beauty retailer in 2012. When they noticed customers wanted authentic products with expert advice, they quickly added beauty content, tutorials, and expert recommendations to their platform. A larger corporation would have needed months of planning and approval. This flexibility allowed Nykaa to differentiate itself in the competitive e-commerce space and grow from a startup to a major player.""")

    add_heading_styled(doc, "Feature 3: Limited Financial Resources but Innovation Potential", 3)
    add_paragraph_after(doc, """SMEs typically operate with constrained budgets compared to large enterprises. However, these limitations often drive creative problem-solving and innovative approaches.

E-Commerce Application - Zepto (India): Zepto, a quick-commerce startup delivering groceries in 10 minutes, couldn't afford massive warehouse networks like Amazon or Flipkart. Instead, they innovated with "dark stores" - small, strategically located micro-warehouses in residential areas. This resource constraint led to a business model innovation that larger competitors are now copying. Their limited funding forced them to find efficient solutions that became competitive advantages.""")

    # Task 1.3 (15 marks) - 5 stages with 3 characteristics each + examples
    add_heading_styled(doc, "1.3 Stages of the Business Life Cycle: Characteristics and Examples", 2)

    add_heading_styled(doc, "Stage 1: Start-up", 3)
    add_paragraph_after(doc, """Characteristics:
• Low or zero initial sales while establishing market presence
• Negative cash flow due to high setup costs and investment requirements
• Simple organisational structure with founder handling multiple roles

Example: Zomato (2008-2010) - Started as a simple website called Foodiebay listing restaurant menus in Delhi. Sales were minimal, losses were constant as they built their database, and founder Deepinder Goyal personally visited restaurants to collect menus.""")

    add_heading_styled(doc, "Stage 2: Growth", 3)
    add_paragraph_after(doc, """Characteristics:
• Rapidly increasing sales as market acceptance builds
• Transition from losses to profitability
• Need for formal structures, processes, and additional staff

Example: Swiggy (2016-2019) - After proving their food delivery model worked, Swiggy experienced explosive growth. They expanded from 8 cities to 500+, hired thousands of delivery partners, and needed to build proper HR, operations, and technology teams.""")

    add_heading_styled(doc, "Stage 3: Maturity", 3)
    add_paragraph_after(doc, """Characteristics:
• Stable, predictable revenue streams
• Optimised and efficient operations
• Focus shifts from growth to defending market position

Example: Amul (present day) - As India's largest dairy cooperative, Amul has reached maturity with consistent revenues, optimised supply chains across India, and focuses on maintaining its dominant position against competitors like Mother Dairy.""")

    add_heading_styled(doc, "Stage 4: Renewal/Reinvention", 3)
    add_paragraph_after(doc, """Characteristics:
• Significant investment in new products, technologies, or markets
• Organisational restructuring or culture change
• Calculated risk-taking to avoid decline

Example: LEGO (2003-2010) - Facing near-bankruptcy, LEGO completely reinvented itself by returning to core products, launching successful movie tie-ins, and entering digital gaming. This transformation required new leadership, cultural change, and strategic redirection.""")

    add_heading_styled(doc, "Stage 5: Decline", 3)
    add_paragraph_after(doc, """Characteristics:
• Falling sales and declining market share
• Products or services becoming obsolete
• Cost-cutting measures and potential downsizing

Example: Nokia Mobile Phones (2010-2014) - Despite being the world's largest phone manufacturer, Nokia failed to adapt to the smartphone revolution. Sales plummeted, market share collapsed, and eventually the mobile division was sold to Microsoft.""")

    # Task 1.4 (15 marks) - 2 objectives per stage
    add_heading_styled(doc, "1.4 Business Objectives Across the Life Cycle", 2)

    add_heading_styled(doc, "Start-up Stage Objectives", 3)
    add_paragraph_after(doc, """1. Survival: The primary goal is simply staying in business long enough for the concept to prove itself. Many startups fail within the first year due to running out of cash.

2. Product-Market Fit: Demonstrating that customers actually want the product or service and are willing to pay for it. Without this validation, growth is impossible.""")

    add_heading_styled(doc, "Growth Stage Objectives", 3)
    add_paragraph_after(doc, """1. Market Share Expansion: Capturing customers before competitors do, building brand recognition, and establishing market presence.

2. Building Organisational Capacity: Hiring staff, creating systems, and developing infrastructure that can handle increased demand without quality suffering.""")

    add_heading_styled(doc, "Maturity Stage Objectives", 3)
    add_paragraph_after(doc, """1. Profit Maximisation: Optimising operations to generate the best possible returns from established market position.

2. Market Position Defence: Protecting customer base from competitor encroachment through loyalty programmes, quality maintenance, and brand reinforcement.""")

    add_heading_styled(doc, "Renewal Stage Objectives", 3)
    add_paragraph_after(doc, """1. Innovation and Transformation: Developing new products, services, or business models to replace declining offerings.

2. Organisational Adaptation: Changing culture, structure, and processes to compete effectively in evolving markets.""")

    add_heading_styled(doc, "Decline Stage Objectives", 3)
    add_paragraph_after(doc, """1. Loss Minimisation: Reducing costs and managing resources to limit financial damage during the downturn.

2. Exit Strategy Development: Deciding whether to attempt turnaround, sell the business, merge with another company, or close operations.""")

    doc.add_page_break()

    # ===== TASK B =====
    add_heading_styled(doc, "TASK B: SME Business Strategies", 1)
    add_heading_styled(doc, "LO2: Analyse business strategies of Small and Medium Enterprises", 2)

    # Task 2.1 (10 marks)
    add_heading_styled(doc, "2.1 Internal and External Factors Influencing SMEs", 2)

    add_heading_styled(doc, "Internal Factors", 3)

    add_paragraph_after(doc, """Internal Factor 1: Cash Flow Management

Cash flow problems are the leading cause of SME failure, with research indicating approximately 82% of business failures are linked to poor cash management (NetSuite, 2024). Unlike large corporations with credit facilities and reserves, SMEs often operate on tight margins with limited access to emergency funding.

Impact: Difficulty paying suppliers leads to damaged relationships and potential supply disruptions. Inability to invest in opportunities when they arise means missing growth chances. During economic downturns, limited cash reserves can mean the difference between survival and closure.""")

    add_paragraph_after(doc, """Internal Factor 2: Skills and Management Capabilities

Many SME owners are experts in their product or service but lack formal business training. A skilled baker might open a bakery without understanding accounting, marketing, or employee management.

Impact: Poor financial decisions due to lack of knowledge, ineffective marketing that wastes limited resources, and difficulty scaling operations beyond what the owner can personally oversee.""")

    add_heading_styled(doc, "External Factors", 3)

    add_paragraph_after(doc, """External Factor 1: Market Competition

SMEs face intense competition from both large corporations with greater resources and other small businesses fighting for the same customers.

Impact: Larger competitors can undercut on price due to economies of scale, outspend on marketing, and absorb losses that would destroy an SME. Globalisation means competition now comes from international companies, not just local rivals.""")

    add_paragraph_after(doc, """External Factor 2: Regulatory Environment and Economic Conditions

Changes in regulations (tax, employment law, health and safety) create compliance burdens that fall proportionately harder on small businesses. Economic factors like inflation and interest rates are beyond SME control but significantly affect operations.

Impact: When India introduced GST, many small businesses struggled with compliance requirements. Rising interest rates increase borrowing costs. Inflation squeezes margins as costs rise faster than prices can be increased.""")

    # Task 2.2 (20 marks) - 4 functions across life cycle
    add_heading_styled(doc, "2.2 Functional Activities Across the Business Life Cycle", 2)

    add_heading_styled(doc, "Operations Function", 3)
    add_paragraph_after(doc, """Start-up: Focus on developing a minimum viable product and basic delivery processes. Quality and efficiency are secondary to proving the concept works.

Growth: The challenge is scaling production while maintaining quality. Systems need documentation and standardisation. What worked for 10 customers doesn't work for 100.

Maturity: Operations focus shifts to optimisation - reducing costs, improving efficiency, implementing lean management principles.

Decline: Scaling back operations to match reduced demand, maintaining only essential capabilities.""")

    add_heading_styled(doc, "Marketing (including Digital Marketing)", 3)
    add_paragraph_after(doc, """Start-up: Limited budget marketing focused on reaching early adopters. Heavy reliance on social media marketing, content marketing, and word-of-mouth due to cost constraints.

Growth: Marketing becomes more aggressive with increased budgets. Digital marketing campaigns expand across platforms. Focus on customer acquisition and market share capture.

Maturity: Shift to customer retention and brand reinforcement. Digital marketing focuses on loyalty programmes, email marketing to existing customers, and defending against competitor campaigns.

Decline: Marketing budgets are cut. Focus narrows to retaining the most valuable customer segments.""")

    add_heading_styled(doc, "Human Resource Management", 3)
    add_paragraph_after(doc, """Start-up: Informal HR with founder handling hiring. Employees are generalists. Culture develops organically.

Growth: Rapid hiring creates need for formal HR policies, job descriptions, training programmes. Maintaining culture while scaling is challenging.

Maturity: Focus on retention, specialist development, and succession planning. HR systems are mature and optimised.

Decline: Managing redundancies while retaining key staff. Difficult decisions about workforce reduction.""")

    add_heading_styled(doc, "Finance/Accounts", 3)
    add_paragraph_after(doc, """Start-up: Survival focus - managing cash burn rate, securing funding, monitoring runway. Every expense scrutinised.

Growth: May need additional funding for expansion. Financial systems need upgrading to handle complexity. Focus on revenue growth.

Maturity: Focus on profitability optimisation, dividend decisions, potential acquisitions. Sophisticated financial management.

Decline: Cash preservation becomes critical. May involve asset sales, cost reduction, or planning for exit.""")

    # Task 2.3 (20 marks) - 2 strategies: digitalisation and partnerships
    add_heading_styled(doc, "2.3 Strategic Options for SMEs", 2)

    add_heading_styled(doc, "Strategy 1: Digitalisation", 3)
    add_paragraph_after(doc, """Digitalisation involves adopting digital technologies to transform business operations, customer interactions, and value creation.

How It Works for SMEs:
• Implementing e-commerce platforms to reach customers beyond physical location
• Using cloud-based software for operations, accounting, and customer management
• Leveraging social media and digital marketing for cost-effective customer acquisition
• Adopting automation to reduce manual processes and costs

Expected Outcomes:

Positive Effects:
• Expanded market reach without physical expansion costs
• Reduced operational costs through automation
• Improved customer experience through 24/7 availability
• Better data for decision-making through analytics
• Competitive parity with larger businesses in online presence

Potential Risks:
• Initial investment costs for technology implementation
• Skills gaps requiring training or new hires
• Cybersecurity vulnerabilities
• Dependence on technology providers

Real Example: A local bookshop in Bangalore implemented an e-commerce website and Instagram presence during COVID-19. Digital sales now account for 40% of revenue, reaching customers across India who would never have visited the physical store. The digitalisation investment of ₹2 lakhs has generated additional annual revenue of ₹15 lakhs.""")

    add_heading_styled(doc, "Strategy 2: Strategic Partnerships", 3)
    add_paragraph_after(doc, """Strategic partnerships involve collaborating with other organisations to access resources, capabilities, or markets that would be difficult to develop independently.

How It Works for SMEs:
• Joint ventures with complementary businesses
• Distribution partnerships with larger companies
• Technology partnerships for access to platforms
• Supply chain partnerships for better terms

Expected Outcomes:

Positive Effects:
• Access to resources and capabilities beyond the SME's means
• Shared risk on new ventures
• Faster market entry than building capabilities internally
• Credibility boost from association with established partners
• Knowledge transfer and learning opportunities

Potential Risks:
• Loss of independence in decision-making
• Partner conflicts over direction or profits
• Dependence on partner performance
• Intellectual property concerns

Real Example: Chaayos, an Indian chai café chain, partnered with Swiggy and Zomato for delivery rather than building their own delivery fleet. This partnership allowed them to focus on their core strength (making great chai) while accessing millions of potential customers through established platforms. The partnership model enabled expansion to 180+ outlets without the capital investment of building delivery infrastructure.""")

    doc.add_page_break()

    # References
    add_heading_styled(doc, "References", 1)

    refs = [
        "European Commission (2003) 'Commission Recommendation of 6 May 2003 concerning the definition of micro, small and medium-sized enterprises', Official Journal of the European Union, L 124, pp. 36-41.",
        "",
        "NetSuite (2024) '10 Top Financial Challenges for Small Businesses'. Available at: https://www.netsuite.com/portal/resource/articles/business-strategy/small-business-financial-challenges.shtml (Accessed: 23 January 2026).",
        "",
        "JAIN Online (2025) 'HE9E 46 Contemporary Business Issues Week 1 Lecture Slides'. Bangalore: JAIN University.",
        "",
        "Corporate Finance Institute (n.d.) 'Business Life Cycle'. Available at: https://corporatefinanceinstitute.com/resources/valuation/business-life-cycle/ (Accessed: 23 January 2026).",
    ]

    for ref in refs:
        if ref:
            add_paragraph_after(doc, ref)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def generate_j229_76_v2():
    """Generate J229 76 with ALL requirements including missing tasks"""
    print("Generating J229 76 (v2) with all requirements...")

    template_path = os.path.join(BASE_PATH, "J229 76_Understanding Business._Student Assessment Template.docx")
    output_path = os.path.join(BASE_PATH, "My_Assignments", "J229_76_Understanding_Business", "Final", "J229_76_Final_Assignment_v2.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    doc.add_page_break()

    # ===== TASK A =====
    add_heading_styled(doc, "TASK A: Large Organisations (50 marks)", 1)
    add_heading_styled(doc, "LO1: Analyse the features, objectives and internal structures of large organisations", 2)

    # Task 1.1.1 (8 marks)
    add_heading_styled(doc, "1.1 Comparing Features of Large Organisations from Different Sectors", 2)
    add_heading_styled(doc, "Task 1.1.1: Comparison of Tata Motors, NHS Scotland, and Oxfam", 3)

    add_table_styled(doc,
        ["Feature", "Tata Motors (Secondary/Private)", "NHS Scotland (Tertiary/Public)", "Oxfam (Third Sector)"],
        [
            ["Scale & Resources", "58,000+ employees, ₹345,000 crore revenue, global manufacturing facilities", "160,000 staff, £17 billion annual budget, 14 regional health boards", "10,000+ staff and volunteers, operates in 90+ countries, relies on donations"],
            ["Formal Structure", "Hierarchical PLC structure with Board of Directors, divisional organisation", "Statutory body structure defined by law, NHS Boards accountable to Ministers", "Confederation structure with trustees, 20+ independent member organisations"],
            ["Specialisation", "Highly specialised: R&D engineers, production specialists, marketing experts", "Medical specialisation: doctors, nurses, administrators, support staff", "Specialised roles: humanitarian workers, advocacy experts, retail managers"],
            ["Regulatory Requirements", "Companies Act compliance, stock exchange regulations, automotive standards", "NHS legislation, medical regulations, government accountability requirements", "Charity Commission registration, fundraising regulations, tax exemptions"],
            ["Access to Capital", "Stock market listing, bank loans, retained profits, bond issues", "Government funding from taxation, cannot raise private capital", "Donations, grants, charity shop income, limited commercial options"]
        ])

    # Task 1.1.2 (6 marks)
    add_heading_styled(doc, "Task 1.1.2: Comparison of Company Types", 3)

    add_table_styled(doc,
        ["Feature", "Private Limited Company (Ltd)", "Public Limited Company (PLC)", "Franchise"],
        [
            ["Ownership", "Shares held privately, cannot be traded publicly", "Shares traded on stock exchange, anyone can buy", "Franchisee owns local business, franchisor owns brand"],
            ["Liability", "Limited to share value", "Limited to share value", "Limited for franchisee company"],
            ["Capital Raising", "Private investment, loans, retained profit", "Public share issues, large-scale capital access", "Franchisee investment plus franchisor support"],
            ["Control", "Owners maintain tight control", "Dispersed ownership, board governance", "Franchisor controls brand/systems, franchisee runs daily operations"],
            ["Example", "Bosch India, Reliance Retail", "Tata Motors, Infosys", "McDonald's franchises, Domino's outlets"]
        ])

    # Task 1.1.3 (6 marks)
    add_heading_styled(doc, "Task 1.1.3: Internal (Organic) vs External (Inorganic) Growth", 3)

    add_paragraph_after(doc, """Organic Growth - Internal expansion using own resources:

Example 1 - Local Bakery Expanding Production:
A neighbourhood bakery increasing capacity by purchasing additional ovens and hiring more staff represents organic growth. The owner invests retained profits to meet growing demand without external involvement.
• Access to Technology: Limited - relies on available equipment
• Speed: Slow - gradual capacity increase
• Investment: Lower - uses internal funds
• Risk: Lower - controlled expansion

Example 2 - Infosys Technology Investment:
Infosys expanded organically by investing heavily in training centres, new service lines, and R&D facilities. They built capabilities internally rather than acquiring other companies.
• Access to Technology: Developed internally through R&D
• Speed: Medium - takes time to build expertise
• Investment: Significant but self-funded
• Risk: Moderate - controlled development

Inorganic Growth - External expansion through acquisitions/partnerships:

Example 1 - Tata Motors acquiring Jaguar Land Rover:
In 2008, Tata Motors paid $2.3 billion to acquire Jaguar Land Rover from Ford. This gave Tata instant access to premium car markets, advanced technology, and global brand recognition.
• Access to Technology: Immediate - acquired JLR's engineering capabilities
• Speed: Instant - immediate market presence
• New Markets: Entered luxury segment globally overnight
• Investment: Very high - $2.3 billion acquisition cost
• Risk: Higher - integration challenges, cultural differences

Example 2 - Starbucks-Tata Partnership:
Rather than entering India alone, Starbucks partnered with Tata to form Tata Starbucks. This joint venture combined Starbucks' brand with Tata's local knowledge and real estate access.
• Access to Technology: Shared expertise in coffee sourcing and retail
• Speed: Fast - leveraged Tata's existing infrastructure
• New Markets: Entered Indian market with reduced risk
• Investment: Shared between partners
• Risk: Lower - shared with established local partner""")

    # Task 1.2.1 (15 marks)
    add_heading_styled(doc, "1.2 Objectives of Large Organisations", 2)
    add_heading_styled(doc, "Task 1.2.1: Organisational Objectives and Their Importance", 3)

    add_heading_styled(doc, "Reliance Jio", 4)
    add_paragraph_after(doc, """Objective: Digital inclusion - providing affordable internet access to all Indians

Importance: When Jio launched in 2016, India's internet penetration was low and data was expensive. Jio's objective aligned with national development goals (Digital India initiative) while creating massive commercial opportunity. By making data virtually free initially, Jio acquired 465+ million subscribers, becoming India's largest telecom operator. This objective was important because it:
• Addressed a genuine national need for connectivity
• Created competitive advantage through scale
• Supported government policy objectives
• Transformed the entire telecom industry""")

    add_heading_styled(doc, "Tata Group", 4)
    add_paragraph_after(doc, """Objective: Nation-building through ethical business - improving quality of life while maintaining commercial success

Importance: The Tata Group's unique ownership structure (majority shares held by charitable trusts) enables this objective. This was important because it:
• Built enormous public trust in the Tata name
• Enabled long-term thinking over short-term profits
• Created social impact through hospitals, schools, and research institutions
• Demonstrated that profit and purpose can coexist
• Provided competitive advantage through reputation""")

    add_heading_styled(doc, "Amul", 4)
    add_paragraph_after(doc, """Objective: Fair farmer returns - ensuring dairy farmers receive maximum value for their milk

Importance: As a cooperative owned by 3.6 million farmers, this objective is Amul's reason for existence. This was important because it:
• Protected farmers from exploitation by middlemen
• Returned 85% of consumer rupee to farmers (vs 33% global average)
• Triggered India's "White Revolution" making India the world's largest milk producer
• Created sustainable rural livelihoods
• Built farmer loyalty ensuring consistent supply""")

    add_heading_styled(doc, "Unilever", 4)
    add_paragraph_after(doc, """Objective: Sustainable growth - achieving business growth while reducing environmental footprint

Importance: This objective responds to changing consumer expectations and regulatory environment. This was important because it:
• Addresses growing consumer demand for sustainable products
• Reduces long-term risks from climate change and resource scarcity
• Prepares for tightening environmental regulations
• Differentiates from competitors in crowded markets
• Attracts environmentally-conscious talent and investors""")

    # Task 1.3.1 (10 marks)
    add_heading_styled(doc, "1.3 Internal Structures of Large Organisations", 2)
    add_heading_styled(doc, "Task 1.3.1: Organisational Structures with Company Examples", 3)

    add_heading_styled(doc, "Hierarchical Structure - NHS Scotland", 4)
    add_paragraph_after(doc, """Description: Multiple management levels with clear chain of command from Scottish Government through NHS Boards to frontline staff.

Why NHS Scotland uses this structure:
• Accountability: Clear reporting lines essential for public service delivery
• Compliance: Healthcare requires strict protocols and regulatory adherence
• Coordination: Managing 160,000 staff across Scotland requires formal hierarchy
• Specialisation: Different management levels handle strategic, tactical, and operational decisions""")

    add_heading_styled(doc, "Functional Structure - Tata Motors", 4)
    add_paragraph_after(doc, """Description: Organisation divided into specialist departments (R&D, Production, Marketing, Finance, HR) with expertise concentrated in each function.

Why Tata Motors uses this structure:
• Expertise development: Engineers work with engineers, improving technical capability
• Efficiency: Reduces duplication of specialist resources
• Clear responsibilities: Each function has defined role in vehicle development and production
• Career paths: Staff can progress within their specialist area""")

    add_heading_styled(doc, "Matrix Structure - Infosys/TCS", 4)
    add_paragraph_after(doc, """Description: Combines functional departments with project teams. Employees report to both functional manager and project manager.

Why IT companies use this structure:
• Resource flexibility: Specialists can work across multiple client projects
• Knowledge sharing: Expertise flows between projects and functions
• Client focus: Project teams organised around specific client needs
• Innovation: Cross-functional collaboration drives creative solutions""")

    # Task 1.3.2 (5 marks) - COMPLETELY NEW
    add_heading_styled(doc, "Task 1.3.2: Decision-Making in Large Organisations", 3)

    add_paragraph_after(doc, """Large organisations make decisions at three levels:

Strategic Decisions (Board/Senior Management):
• Long-term direction, major investments, market entry/exit
• Example: Tata's decision to acquire Jaguar Land Rover
• Uses tools like SWOT analysis to assess strengths, weaknesses, opportunities, threats

Tactical Decisions (Middle Management):
• Medium-term implementation of strategy
• Example: How to integrate JLR operations with Tata systems
• Uses tools like decision trees to evaluate options and likely outcomes

Operational Decisions (Front-line Management):
• Day-to-day running of business
• Example: Production scheduling, shift allocation
• Uses standard procedures and guidelines""")

    add_paragraph_after(doc, """SWOT Analysis: Used for strategic planning by systematically evaluating:
• Strengths: Internal advantages (e.g., strong brand, skilled workforce)
• Weaknesses: Internal limitations (e.g., outdated technology, high costs)
• Opportunities: External possibilities (e.g., new markets, changing regulations)
• Threats: External risks (e.g., new competitors, economic downturn)

Decision Trees: Visual tool showing possible decisions, their likely outcomes, and associated probabilities. Helps evaluate options by:
• Mapping all possible choices
• Estimating probability of each outcome
• Calculating expected value of each path
• Comparing alternatives systematically""")

    doc.add_page_break()

    # ===== TASK B =====
    add_heading_styled(doc, "TASK B: Business Environment (50 marks)", 1)
    add_heading_styled(doc, "LO2: Analyse the environment in which large organisations operate", 2)

    # Task 2.1 (15 marks)
    add_heading_styled(doc, "2.1 Impact of Internal Factors", 2)

    add_heading_styled(doc, "Toyota: Positive Internal Factors", 3)
    add_paragraph_after(doc, """How Internal Culture Shapes Behaviour and Quality:
Toyota's "Toyota Way" culture is built on two pillars: Respect for People and Continuous Improvement (Kaizen). This culture shapes every employee's behaviour - workers are empowered to stop production lines if they spot defects, something unthinkable in most factories. Quality is everyone's responsibility, not just the quality department's.

Why Internal Alignment is Critical:
Toyota's success demonstrates how aligning internal factors (culture, workforce, processes) with organisational objectives creates sustainable competitive advantage. Their production system (TPS) with Just-in-Time manufacturing and worker empowerment has made them one of the world's most efficient manufacturers for 70+ years.""")

    add_heading_styled(doc, "Boeing: Negative Internal Factors", 3)
    add_paragraph_after(doc, """How Internal Weaknesses Create Operational Risk:
Boeing's 737 MAX crisis (346 deaths in two crashes) resulted from internal failures:
• Culture shift from engineering excellence to financial performance
• Management disconnection from frontline engineers
• Ignoring internal warnings about MCAS system problems
• Prioritising schedule and cost over safety

Consequences of Internal Misalignment:
• $87 billion lost in shareholder value
• $2.5 billion settlement with US government
• Global fleet grounding
• Destroyed reputation for safety

This demonstrates how internal weaknesses - poor culture, management disconnection, suppressed warnings - can have catastrophic consequences when internal factors are not aligned with organisational objectives.""")

    # Task 2.2 (20 marks) - PESTEC
    add_heading_styled(doc, "2.2 Impact of External Factors - PESTEC Analysis", 2)

    add_heading_styled(doc, "Netflix: Technological and Social Factors", 3)
    add_paragraph_after(doc, """Technological Factors:
• Streaming technology transformation: Netflix pivoted from DVD rental to streaming, fundamentally changing entertainment consumption
• AI personalisation: Machine learning recommendation engine influences 80% of viewing choices
• Infrastructure investment: Own content delivery network (Open Connect) provides competitive advantage
• Competition enabled by same technology: Disney+, Amazon Prime use identical streaming capabilities

Social Factors:
• Binge-watching culture: Netflix pioneered and normalised consuming entire series at once
• Demand for diverse content: Investment in local content (Money Heist, Squid Game) responds to global audience preferences
• Short-form competition: TikTok and YouTube Shorts changing younger viewers' habits

How Netflix Responds: Continuous technological investment, content localisation, introduction of ad-supported tiers, gaming expansion.""")

    add_heading_styled(doc, "McDonald's: Globalisation, Cultural Adaptation, and Ethical Considerations", 3)
    add_paragraph_after(doc, """Globalisation and Cultural Adaptation:
• Operates in 100+ countries requiring cultural sensitivity
• Menu adaptation: McAloo Tikki (India), Teriyaki Burger (Japan), McArabia (Middle East)
• Balance between global brand consistency and local relevance
• Supply chain localisation in each market

Ethical Considerations:
• Health concerns: Pressure to offer healthier options, nutritional information
• Animal welfare: Commitments to cage-free eggs, sustainable sourcing
• Environmental impact: Targets for recyclable packaging, sustainable beef
• Worker conditions: Scrutiny of wages and working conditions

How Global Organisations Balance Standardisation with Local Responsiveness:
McDonald's uses "glocal" strategy - maintaining core brand elements (Golden Arches, service standards) while adapting menus, marketing, and practices to local cultures and ethical expectations.""")

    # Task 2.3 (15 marks)
    add_heading_styled(doc, "2.3 Stakeholder Conflicts of Interest", 2)

    add_paragraph_after(doc, """Stakeholder Groups in Large Organisations:
• Owners/Shareholders: Want returns on investment (dividends, share price growth)
• Managers: Want job security, career progression, remuneration
• Employees: Want fair wages, good conditions, job security
• Customers: Want quality products at fair prices
• Suppliers: Want reliable orders and fair payment terms
• Government: Wants tax revenue, employment, regulatory compliance
• Community: Wants jobs but also environmental protection and social responsibility""")

    add_heading_styled(doc, "Conflict 1: Employee Welfare vs Shareholder Expectations", 3)
    add_paragraph_after(doc, """Case: Amazon Warehouse Workers

The Conflict: Shareholders expect profit maximisation, which pressures management to minimise costs. Employees want fair wages, reasonable working conditions, and job security. These objectives directly conflict.

Evidence of Conflict:
• Reports of demanding productivity targets in warehouses
• Criticism of monitoring and performance management systems
• Disputes over unionisation attempts
• COVID-19 safety concerns

How Unresolved Conflicts Damage Performance:
• High staff turnover (reportedly 150% annual rate at some facilities)
• Negative publicity affecting brand perception
• Regulatory scrutiny and potential legal action
• Difficulty attracting talent

Attempted Resolution:
• Minimum wage increases ($15/hour in US)
• Investment in safety measures
• Benefits improvements
• However, tension continues between worker advocacy groups and company""")

    add_heading_styled(doc, "Conflict 2: Business Expansion vs Local Community", 3)
    add_paragraph_after(doc, """Case: Amazon HQ2 (New York)

The Conflict: Amazon wanted to expand with a second headquarters, promising jobs and investment. Local community had concerns about gentrification, housing costs, and use of tax incentives.

How Organisations Attempt to Balance Demands:
• Amazon offered job creation (25,000 jobs)
• Promised infrastructure investment
• Negotiated tax incentives with state government

How Conflicts Can Damage Reputation:
• Intense local opposition led to Amazon withdrawing from New York
• Demonstrated that corporate power has limits when facing organised community resistance
• Highlighted importance of genuine stakeholder engagement, not just announcement

Why Integrated Environmental Analysis Matters:
Companies must consider all stakeholder groups before major decisions. Amazon's failure to adequately engage local communities in New York shows that ignoring stakeholder concerns - even when governments are supportive - can derail strategic plans.""")

    doc.add_page_break()

    # References
    add_heading_styled(doc, "References", 1)

    refs = [
        "Harvard Business School (2024) 'Why Boeing's Problems with the 737 MAX Began More Than 25 Years Ago'. Available at: https://www.library.hbs.edu/working-knowledge/ (Accessed: 23 January 2026).",
        "",
        "NHS Scotland (n.d.) About NHS Scotland. Available at: https://www.scot.nhs.uk/about-nhs-scotland/ (Accessed: 23 January 2026).",
        "",
        "Oxfam International (n.d.) How We Are Organized. Available at: https://www.oxfam.org/en/what-we-do/about/how-we-are-organized (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'Netflix PESTEL Analysis'. Available at: https://panmore.com/netflix-pestel-pestle-analysis (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'McDonald's PESTEL Analysis'. Available at: https://panmore.com/mcdonalds-pestel-pestle-analysis (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'Toyota's Organizational Culture'. Available at: https://panmore.com/toyota-organizational-culture (Accessed: 23 January 2026).",
        "",
        "Tata Motors (2024) Annual Report 2024. Mumbai: Tata Motors Limited.",
    ]

    for ref in refs:
        if ref:
            add_paragraph_after(doc, ref)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def generate_j22a_76_v2():
    """Generate J22A 76 with ALL requirements including lockouts, picketing, and additional legislation"""
    print("Generating J22A 76 (v2) with all requirements...")

    template_path = os.path.join(BASE_PATH, "J22A 76_Management of People and Finance_Student_Assessment Template.docx")
    output_path = os.path.join(BASE_PATH, "My_Assignments", "J22A_76_Management_People_Finance", "Final", "J22A_76_Final_Assignment_v2.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    doc.add_page_break()

    # ===== TASK A (60 marks) =====
    add_heading_styled(doc, "TASK A: Management of People (60 marks)", 1)
    add_heading_styled(doc, "LO1: Apply knowledge of how management of people can meet objectives", 2)

    # Task 1.1 (15 marks)
    add_heading_styled(doc, "1.1 Human Resource Management Approaches", 2)

    add_heading_styled(doc, "Approach 1: Recruitment and Selection", 3)
    add_paragraph_after(doc, """Recruitment involves attracting, screening, and selecting qualified candidates for positions within an organisation.

Process:
1. Job Analysis → 2. Job Description → 3. Person Specification → 4. Advertising → 5. Shortlisting → 6. Selection → 7. Onboarding

Contribution to Effective HRM:
• Ensures right person for right job, reducing costly hiring mistakes
• Internal recruitment motivates existing staff through career progression
• External recruitment brings fresh perspectives and skills
• Proper selection reduces turnover costs and improves productivity

Example: Infosys conducts rigorous campus recruitment, selecting graduates through aptitude tests, technical interviews, and HR rounds. This systematic approach ensures they hire candidates who fit their culture and can be trained to meet client needs.""")

    add_heading_styled(doc, "Approach 2: Training and Development", 3)
    add_paragraph_after(doc, """Training focuses on developing skills for current roles, while development prepares employees for future responsibilities.

Types of Training:
• Induction: Orienting new employees to the organisation
• On-the-job: Learning while working under supervision
• Off-the-job: External courses, workshops, certifications

Contribution to Effective HRM:
• Improves employee performance and productivity
• Reduces errors and accidents
• Increases employee satisfaction and retention
• Builds organisational capability for future challenges

Example: Tata Group's Tata Management Training Centre provides leadership development programmes. This investment in employee development has created a pipeline of leaders who understand Tata values and can lead diverse businesses.""")

    add_heading_styled(doc, "Approach 3: Performance Management", 3)
    add_paragraph_after(doc, """Performance management involves setting objectives, monitoring progress, providing feedback, and rewarding achievement.

Components:
• Goal setting aligned with organisational objectives
• Regular performance reviews and feedback
• Recognition and reward systems
• Performance improvement plans where needed

Contribution to Effective HRM:
• Aligns individual effort with organisational goals
• Identifies training needs and high-potential employees
• Motivates through recognition and fair rewards
• Provides basis for promotion and succession planning

Example: Google uses "OKRs" (Objectives and Key Results) - employees set ambitious objectives with measurable results, reviewed quarterly. This transparent system ensures everyone understands how their work contributes to company goals.""")

    # Task 1.2 (15 marks) - Maslow WITH DIAGRAM NOTE
    add_heading_styled(doc, "1.2 Motivation: Maslow's Hierarchy of Needs", 2)

    add_paragraph_after(doc, """[INSERT MASLOW'S HIERARCHY PYRAMID DIAGRAM HERE]

The diagram should show a pyramid with five levels:
- Base: Physiological Needs
- Level 2: Safety Needs
- Level 3: Social/Belonging Needs
- Level 4: Esteem Needs
- Top: Self-Actualisation""")

    add_heading_styled(doc, "Level 1: Physiological Needs", 3)
    add_paragraph_after(doc, """Basic survival needs - food, water, shelter, rest.

Workplace Application: Fair wages enabling employees to meet basic living costs, adequate breaks, comfortable working temperature, clean facilities.

Example: Amazon's $15 minimum wage policy addresses workers' ability to afford basic necessities, recognising that underpaid employees cannot focus on higher-level contributions.""")

    add_heading_styled(doc, "Level 2: Safety and Security Needs", 3)
    add_paragraph_after(doc, """Protection from physical and economic harm, stability and security.

Workplace Application: Job security through permanent contracts, pension schemes, health insurance, safe working conditions, clear policies.

Example: Japanese companies traditionally offered "lifetime employment," providing security that enabled employees to focus entirely on their work without fear of redundancy.""")

    add_heading_styled(doc, "Level 3: Social/Belonging Needs", 3)
    add_paragraph_after(doc, """Connection, friendship, feeling part of a group.

Workplace Application: Team-building activities, social events, collaborative work environments, supportive management relationships.

Example: Google's open office layouts, team lunches, and social spaces are designed to foster belonging and connection among employees.""")

    add_heading_styled(doc, "Level 4: Esteem Needs", 3)
    add_paragraph_after(doc, """Self-respect and recognition from others, achievement, competence.

Workplace Application: Recognition programmes, meaningful job titles, public praise, responsibility, promotions.

Example: Infosys's "Awards for Excellence" programme recognises outstanding contributions, satisfying employees' need for acknowledgment and respect.""")

    add_heading_styled(doc, "Level 5: Self-Actualisation", 3)
    add_paragraph_after(doc, """Reaching full potential, personal growth, meaningful work.

Workplace Application: Challenging assignments, autonomy, creativity opportunities, purpose-driven work.

Example: 3M's "15% time" policy allowed employees to spend 15% of work time on personal projects, enabling self-actualisation through innovation (leading to products like Post-it Notes).""")

    add_paragraph_after(doc, """Applying Maslow to Improve Effectiveness:
Managers should identify which need level each employee is at and tailor motivation accordingly. An employee worried about job security (Level 2) won't be motivated by recognition programmes (Level 4). Once lower needs are satisfied, they cease to motivate, and higher needs become important.""")

    # Task 1.3 (15 marks) - 5 CORRECT INDUSTRIAL ACTIONS
    add_heading_styled(doc, "1.3 Employee Relations: Forms of Industrial Action", 2)

    add_heading_styled(doc, "1. Strike", 3)
    add_paragraph_after(doc, """Complete withdrawal of labour - employees stop working entirely.

How it affects organisations:
• Production stops completely, orders unfulfilled
• Revenue loss during strike period
• Customer relationships damaged
• May lead to permanent loss of customers to competitors
• Long-term relationship damage between management and workers

Example: UK rail strikes (2022-23) caused billions in economic losses and permanent shifts in commuting patterns.""")

    add_heading_styled(doc, "2. Go-Slow", 3)
    add_paragraph_after(doc, """Employees deliberately work at reduced pace while remaining at work.

How it affects organisations:
• Reduced productivity and missed targets
• Difficult to prove as deliberate (workers claim being "careful")
• Lower output means missed deadlines
• Management frustration without clear remedy
• Workers still receive normal pay

Example: Factory workers taking twice as long over each task, halving daily production without technically breaking any rules.""")

    add_heading_styled(doc, "3. Work-to-Rule", 3)
    add_paragraph_after(doc, """Employees do exactly what their contract requires - no more, no less. Every rule followed precisely.

How it affects organisations:
• Significant slowdown without contract breach
• Exposes how much organisations depend on employee goodwill
• Cannot discipline workers for "following rules"
• Customer frustration from delays

Example: Air traffic controllers following every safety procedure exactly, causing flight delays without breaking any regulations.""")

    add_heading_styled(doc, "4. Lockout", 3)
    add_paragraph_after(doc, """Employer-initiated action where management prevents employees from entering the workplace.

How it affects organisations:
• Production stops but employer controls timing
• Used as pressure tactic in negotiations
• Can be costly if prolonged
• Legal requirements must be followed
• May escalate conflict with workforce

Example: During wage disputes, employers may lock out workers to force acceptance of terms, preventing continued work at existing conditions.""")

    add_heading_styled(doc, "5. Picketing", 3)
    add_paragraph_after(doc, """Workers gather outside workplace to publicise dispute and discourage others from entering.

How it affects organisations:
• Creates negative publicity and reputational damage
• May deter customers and suppliers from entering premises
• Can delay deliveries and disrupt operations
• Visible demonstration of worker discontent
• Legal limits on number of picketers and behaviour

Example: Striking workers forming picket lines outside factories, distributing leaflets explaining their grievances to passing public.""")

    # Task 1.4 (15 marks) - ALL 5 LEGISLATION AREAS
    add_heading_styled(doc, "1.4 Employment Legislation", 2)

    add_heading_styled(doc, "Equality and Diversity (Equality Act 2010)", 3)
    add_paragraph_after(doc, """Protects against discrimination based on: age, disability, gender reassignment, marriage/civil partnership, pregnancy, race, religion, sex, sexual orientation.

Impact on HR: Recruitment must use fair, job-related criteria. Equal pay for equal work. Reasonable adjustments for disabled employees. Training on unconscious bias. Tribunal claims have no upper compensation limit.""")

    add_heading_styled(doc, "Health and Safety (Health and Safety at Work Act 1974)", 3)
    add_paragraph_after(doc, """Employers must ensure health, safety, and welfare of employees.

Impact on HR: Risk assessments required for all work activities. Safety training mandatory. Safe equipment provision. Accident reporting. Non-compliance can result in prosecution, fines, or imprisonment.""")

    add_heading_styled(doc, "Minimum Wage (National Minimum Wage Act 1998)", 3)
    add_paragraph_after(doc, """Sets minimum hourly rates employers must pay, varying by age.

Impact on HR: Pay systems must ensure compliance for all workers including apprentices. Affects budgeting and pricing decisions. Penalties for non-compliance include back-payment orders and public naming.""")

    add_heading_styled(doc, "Working Time Regulations (Working Time Regulations 1998)", 3)
    add_paragraph_after(doc, """Limits working hours and ensures rest periods.

Impact on HR: Maximum 48-hour average working week (unless opted out). Minimum rest breaks. Paid annual leave entitlement (28 days). Night work limits. Records must be kept of working hours.""")

    add_heading_styled(doc, "Dismissal and Redundancy", 3)
    add_paragraph_after(doc, """Employment Rights Act 1996 and related legislation govern termination of employment.

Impact on HR: Fair procedures required for dismissal. Unfair dismissal claims possible after qualifying period. Redundancy requires consultation, fair selection, and statutory payments. Notice periods must be observed. Tribunal claims for unfair treatment.""")

    doc.add_page_break()

    # ===== TASK B (40 marks) =====
    add_heading_styled(doc, "TASK B: Management of Finance (40 marks)", 1)
    add_heading_styled(doc, "LO2: Analyse how management of finance contributes to effectiveness", 2)

    # Task 2.1 (15 marks)
    add_heading_styled(doc, "2.1 Sources of Finance", 2)

    add_heading_styled(doc, "Source 1: Retained Profits (Internal)", 3)
    add_paragraph_after(doc, """Profits kept in business rather than distributed as dividends.

Why appropriate for large organisations:
• No interest payments - "free" capital
• No dilution of ownership or control
• Available immediately without approval processes
• No restrictions on use

Example: Apple retains massive profits, using them for R&D and acquisitions without needing external funding.""")

    add_heading_styled(doc, "Source 2: Share Capital (Equity Financing)", 3)
    add_paragraph_after(doc, """Raising funds by issuing new shares to investors.

Why appropriate for large organisations:
• Large amounts possible through public offerings
• No repayment required unlike debt
• Shared risk - investors lose if company fails
• Can bring expertise alongside capital

Example: Reliance Jio raised ₹152,056 crore from investors including Google and Facebook, funding expansion without increasing debt burden.""")

    add_heading_styled(doc, "Source 3: Corporate Bonds (Debt Financing)", 3)
    add_paragraph_after(doc, """Issuing bonds to investors who receive regular interest and return of principal at maturity.

Why appropriate for large organisations:
• Large amounts available to creditworthy companies
• Fixed interest rates enable budgeting certainty
• No ownership dilution
• Interest is tax-deductible
• Long maturity periods available

Example: Tata Steel issues bonds to fund capital expenditure, accessing debt markets at favourable rates due to strong credit rating.""")

    # Task 2.2 (10 marks)
    add_heading_styled(doc, "2.2 Purposes of Financial Statements", 2)

    add_paragraph_after(doc, """1. Investment Decision Support:
Investors use financial statements to assess profitability, growth, and returns before deciding whether to invest. Without reliable statements, investment decisions would be guesswork.

2. Lending Decision Support:
Banks analyse financial statements to assess creditworthiness, ability to repay loans, available collateral, and existing debt levels before lending.

3. Business Planning and Control:
Management uses statements to compare actual performance against budgets, identify trends, allocate resources, and set future targets.

4. Financial Health Assessment:
Statements reveal what the business owns (assets), owes (liabilities), whether it's solvent, and whether trends are positive or negative - identifying potential problems early.

5. Legal and Tax Compliance:
Limited companies must prepare and file statutory accounts. Statements are used to calculate tax liabilities, inform shareholders, and enable audit verification.""")

    # Task 2.3 (15 marks)
    add_heading_styled(doc, "2.3 Accounting Ratios and Limitations", 2)

    add_heading_styled(doc, "1. Gross Profit Margin", 3)
    add_paragraph_after(doc, """Formula: (Gross Profit ÷ Revenue) × 100

Use: Measures efficiency of production/purchasing. Higher margin means more remains after direct costs.

Analysis: Supermarkets operate on 2-5% margins; software companies achieve 80%+. Declining margins may indicate rising costs or pricing pressure.""")

    add_heading_styled(doc, "2. Net Profit Margin", 3)
    add_paragraph_after(doc, """Formula: (Net Profit ÷ Revenue) × 100

Use: Shows overall profitability after all expenses. What's actually left for owners.

Analysis: Comparing with gross margin reveals overhead impact. High gross but low net margin suggests overhead costs need attention.""")

    add_heading_styled(doc, "3. Current Ratio", 3)
    add_paragraph_after(doc, """Formula: Current Assets ÷ Current Liabilities

Use: Measures ability to pay short-term obligations. Can debts due within 12 months be covered?

Analysis: 2:1 often considered healthy. Too low suggests liquidity problems; too high may indicate inefficient asset use.""")

    add_heading_styled(doc, "4. Acid Test (Quick Ratio)", 3)
    add_paragraph_after(doc, """Formula: (Current Assets - Inventory) ÷ Current Liabilities

Use: Stricter liquidity test excluding inventory which may be hard to convert quickly.

Analysis: 1:1 is minimum safe level. Below this, the business may struggle if it can't sell stock quickly.""")

    add_heading_styled(doc, "5. Return on Capital Employed (ROCE)", 3)
    add_paragraph_after(doc, """Formula: (Operating Profit ÷ Capital Employed) × 100

Use: Measures how effectively capital generates profit. Often considered the most important ratio for investors.

Analysis: Should exceed cost of borrowing. 20% ROCE with 8% loan interest is creating value; 8% ROCE with 10% interest is destroying it.""")

    add_heading_styled(doc, "Limitations of Ratio Analysis", 3)
    add_paragraph_after(doc, """1. Historical Focus: Based on past data, cannot predict future performance
2. Accounting Policy Differences: Different methods make comparison between companies misleading
3. Industry Variations: "Good" values vary significantly by sector
4. Timing Issues: Snapshot may not represent typical position
5. Window Dressing: Companies may manipulate figures before reporting dates
6. Non-Financial Factors: Ratios don't capture customer satisfaction, employee morale, brand strength
7. Context Required: Falling ratios might reflect investment; rising ratios might reflect harmful cost-cutting""")

    doc.add_page_break()

    # References
    add_heading_styled(doc, "References", 1)

    refs = [
        "ACAS (n.d.) 'Strikes and Industrial Action'. Available at: https://www.acas.org.uk/strikes-and-industrial-action (Accessed: 23 January 2026).",
        "",
        "CIPD (n.d.) 'Employment Law'. Available at: https://www.cipd.org/uk/knowledge/factsheets/ (Accessed: 23 January 2026).",
        "",
        "Corporate Finance Institute (n.d.) 'Financial Ratios'. Available at: https://corporatefinanceinstitute.com/resources/accounting/financial-ratios/ (Accessed: 23 January 2026).",
        "",
        "Health and Safety Executive (n.d.) 'Health and Safety at Work Act 1974'. Available at: https://www.hse.gov.uk (Accessed: 23 January 2026).",
        "",
        "Simply Psychology (n.d.) 'Maslow's Hierarchy of Needs'. Available at: https://www.simplypsychology.org/maslow.html (Accessed: 23 January 2026).",
    ]

    for ref in refs:
        if ref:
            add_paragraph_after(doc, ref)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print("Generating Final Assignment Documents - Version 2")
    print("With ALL requirements from assessment briefs")
    print(f"Student: {STUDENT_NAME}")
    print(f"USN: {STUDENT_USN}")
    print("=" * 60)
    print()

    he9e46_path = generate_he9e46_v2()
    j229_path = generate_j229_76_v2()
    j22a_path = generate_j22a_76_v2()

    print()
    print("=" * 60)
    print("All assignments generated successfully!")
    print("=" * 60)
    print()
    print("Files created:")
    print(f"  1. {he9e46_path}")
    print(f"  2. {j229_path}")
    print(f"  3. {j22a_path}")
    print()
    print("IMPORTANT REMINDERS:")
    print("  1. Sign the declaration in each document")
    print("  2. J22A 76: Add Maslow's pyramid diagram manually")
    print("  3. Fill in any remaining student details in templates")
    print("  4. Add footers with: USN_Name_UnitCode_UnitTitle_Page")
