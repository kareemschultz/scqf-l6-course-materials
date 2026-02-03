#!/usr/bin/env python3
"""
Build SCQF Level 6 Assignment DOCX files for Kareem Schultz.
Uses python-docx to create properly formatted Word documents matching the official template.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Student Info ──
STUDENT_NAME = "KAREEM SCHULTZ"
STUDENT_USN = "252IFCBR0596"
COLLEGE = "JAIN College"
PROGRAMME = "Foundation Diploma in Business / CS & IT"
PROGRAMME_FULL = "Foundation Diploma in Business / CS & IT\n(SCQF Level 6 – Equivalent to A Level)"
ACADEMIC_YEAR = "Academic Year 2026"

# ── Unit Definitions ──
UNITS = {
    "J229_76": {
        "code": "J229 76",
        "title": "Understanding Business",
        "assignment_title": "Assignment 1\nLO1 - Task A: Report (Individual)\nLO2 - Task B: Report (Individual)",
        "results": "Total Marks: 100\nAssignment 1 Task A: 50%\nAssignment 1 Task B: 50%",
        "date_issue": "15-01-2026",
        "date_due": "15-02-2026",
        "heading": "Assignment 1: Report on Understanding Business",
        "draft_path": "/home/karetech/clawd/temp/SCQF-L6-Course-Materials/My_Assignments/J229_76_Understanding_Business/Drafts/J229_76_Draft_v2_Humanized.md",
        "footer_code": "J229 76",
        "footer_title": "Understanding Business",
        "filename": "252IFCBR0596_KareemSchultz_J229_76_Understanding_Business.docx",
    },
    "J22A_76": {
        "code": "J22A 76",
        "title": "Management of People and Finance",
        "assignment_title": "Assignment 1\nLO1 - Task A: Report (Individual)\nLO2 - Task B: Report (Individual)",
        "results": "Total Marks: 100\nAssignment 1 Task A: 60%\nAssignment 1 Task B: 40%",
        "date_issue": "15-01-2026",
        "date_due": "15-02-2026",
        "heading": "Assignment 1: Report on Management of People and Finance",
        "draft_path": "/home/karetech/clawd/temp/SCQF-L6-Course-Materials/My_Assignments/J22A_76_Management_People_Finance/Drafts/J22A_76_Draft_v2_Humanized.md",
        "footer_code": "J22A 76",
        "footer_title": "Management of People and Finance",
        "filename": "252IFCBR0596_KareemSchultz_J22A_76_Management_People_Finance.docx",
    },
    "HE9E_46": {
        "code": "HE9E 46",
        "title": "Contemporary Business Issues",
        "assignment_title": "Assignment 1\nLO1 - Task A: Report (Individual)\nLO2 - Task B: Report (Individual)",
        "results": "Total Marks: 100\nAssignment 1 Task A: 50%\nAssignment 1 Task B: 50%",
        "date_issue": "15-01-2026",
        "date_due": "15-02-2026",
        "heading": "Assignment 1: Report on Contemporary Business Issues",
        "draft_path": "/home/karetech/clawd/temp/SCQF-L6-Course-Materials/My_Assignments/HE9E_46_Contemporary_Business/Drafts/HE9E_46_Draft_v2_Humanized.md",
        "footer_code": "HE9E 46",
        "footer_title": "Contemporary Business Issues",
        "filename": "252IFCBR0596_KareemSchultz_HE9E_46_Contemporary_Business_Issues.docx",
    },
}

OUTPUT_DIR = "/home/karetech/clawd/temp/SCQF_Final"

# ── Constants ──
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)   # Template uses 12pt TNR
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)
LINE_SPACING = 1.5


# ═══════════════════════════════════════════════════════════════
# Helper Functions
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" w:space="0" w:color="{val.get("color", "000000")}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, style=None, alignment=None, font_name=FONT_NAME,
                  font_size=FONT_SIZE, bold=False, italic=False, space_after=Pt(6),
                  space_before=Pt(0), line_spacing=LINE_SPACING, color=None,
                  first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    pf = p.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    return p


def add_run_to_para(para, text, font_name=FONT_NAME, font_size=FONT_SIZE,
                    bold=False, italic=False, color=None):
    """Add a run to an existing paragraph."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def set_table_font(table, font_name=FONT_NAME, font_size=Pt(11)):
    """Set font for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = font_size


def add_table_with_borders(doc, data, col_widths=None, header_row=True,
                           font_size=Pt(11), alignment=WD_TABLE_ALIGNMENT.CENTER):
    """Add a table with borders and optional header shading."""
    rows = len(data)
    cols = len(data[0]) if data else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = alignment

    # Apply borders to all cells
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            # Clear default empty paragraph
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = font_size
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)

            if header_row and ri == 0:
                run.bold = True
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for ci, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = width

    return table


def setup_footer(section, footer_text):
    """Setup footer with student info and page number."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(footer_text + "_")
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._r.append(fldChar1)
    
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2.font.name = FONT_NAME
    run2.font.size = Pt(9)
    run2._r.append(instrText)
    
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._r.append(fldChar2)


# ═══════════════════════════════════════════════════════════════
# Cover Page Builder
# ═══════════════════════════════════════════════════════════════

def build_cover_page(doc, unit):
    """Build the cover page matching the template exactly."""
    section = doc.sections[0]
    section.page_width = Inches(7.5)  # Match template (A4-ish)
    section.page_height = Inches(10.6)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Blank line
    add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))

    # "Assignment Brief" header
    add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    add_paragraph(doc, "Assignment Brief", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=FONT_SIZE, bold=True, space_after=Pt(6))

    # Unit title
    add_paragraph(doc, unit["title"], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=FONT_SIZE, bold=True, space_after=Pt(6))

    # Academic Year
    add_paragraph(doc, ACADEMIC_YEAR, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=FONT_SIZE, bold=True, space_after=Pt(12))

    # Blank line
    add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))

    # ── TABLE 1: Unit Information ──
    table1 = doc.add_table(rows=5, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply borders
    tbl = table1._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Row 0: Unit Information (header, merged)
    cell = table1.cell(0, 0)
    cell.merge(table1.cell(0, 1))
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run("Unit Information:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    set_cell_shading(cell, "D9E2F3")

    # Row 1: Qualification
    c0 = table1.cell(1, 0)
    c0.paragraphs[0].clear()
    run = c0.paragraphs[0].add_run("Qualification:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    c1 = table1.cell(1, 1)
    c1.paragraphs[0].clear()
    run = c1.paragraphs[0].add_run(PROGRAMME_FULL)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True

    # Row 2: Unit Code & Title
    c0 = table1.cell(2, 0)
    c0.paragraphs[0].clear()
    run = c0.paragraphs[0].add_run("Unit Code & Title:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    c1 = table1.cell(2, 1)
    c1.paragraphs[0].clear()
    run = c1.paragraphs[0].add_run(f"{unit['code']}\n{unit['title']}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True

    # Row 3: Assignment Title
    c0 = table1.cell(3, 0)
    c0.paragraphs[0].clear()
    run = c0.paragraphs[0].add_run("Assignment Title:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    c1 = table1.cell(3, 1)
    c1.paragraphs[0].clear()
    run = c1.paragraphs[0].add_run(unit["assignment_title"])
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    # Row 4: Results Reporting
    c0 = table1.cell(4, 0)
    c0.paragraphs[0].clear()
    run = c0.paragraphs[0].add_run("Results Reporting:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    c1 = table1.cell(4, 1)
    c1.paragraphs[0].clear()
    run = c1.paragraphs[0].add_run(unit["results"])
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    add_paragraph(doc, "", space_after=Pt(0))

    # ── TABLE 2: Dates ──
    table2 = doc.add_table(rows=1, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2 = table2._tbl
    tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders2 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr2.append(borders2)

    c0 = table2.cell(0, 0)
    c0.paragraphs[0].clear()
    p = c0.paragraphs[0]
    run = p.add_run("Date of Issue: ")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    run2 = p.add_run(unit["date_issue"])
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE

    c1 = table2.cell(0, 1)
    c1.paragraphs[0].clear()
    p = c1.paragraphs[0]
    run = p.add_run("Due Date: ")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    run2 = p.add_run(unit["date_due"])
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE

    add_paragraph(doc, "", space_after=Pt(0))

    # ── TABLE 3: Student Info ──
    table3 = doc.add_table(rows=6, cols=1)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3 = table3._tbl
    tblPr3 = tbl3.tblPr if tbl3.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders3 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr3.append(borders3)

    rows_data = [
        ("To be Filled by the Student", True, "D9E2F3"),
        (f"Student USN: {STUDENT_USN}", False, None),
        (f"College Name & Site: {COLLEGE}", False, None),
        ("Tutor:", False, None),
        (f"Date Due: {unit['date_due']}", False, None),
        ("Date of Submission:", False, None),
    ]
    for ri, (text, is_bold, shade) in enumerate(rows_data):
        cell = table3.cell(ri, 0)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.bold = is_bold
        if shade:
            set_cell_shading(cell, shade)

    add_paragraph(doc, "", space_after=Pt(0))

    # ── TABLE 4: Markers ──
    table4 = doc.add_table(rows=3, cols=2)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4 = table4._tbl
    tblPr4 = tbl4.tblPr if tbl4.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders4 = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr4.append(borders4)

    markers_data = [
        ("First Marker:", "Second Marker:"),
        ("", ""),
        ("Agreed Mark:", "Refer: Yes / No"),
    ]
    for ri, (c0_text, c1_text) in enumerate(markers_data):
        for ci, text in enumerate([c0_text, c1_text]):
            cell = table4.cell(ri, ci)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            if ri == 0:
                run.bold = True


def build_guidelines(doc):
    """Build the General Guidelines and Declaration sections."""
    # Submission notice
    add_paragraph(doc, "", space_after=Pt(0))
    p = add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
    run = p.runs[0]
    run.text = "*All work must be submitted on or before the due date. If an extension of time to submit work is required, a Mitigating Circumstance Form must be submitted along with supporting evidence."
    run.italic = True
    run.font.size = FONT_SIZE
    run.font.name = FONT_NAME

    # Extension question
    add_paragraph(doc, "Has an extension been approved?                Yes              No", 
                  space_after=Pt(3))
    add_paragraph(doc, "If yes, please provide the new submission date ….…/.…. /……., and affix appropriate evidence.",
                  space_after=Pt(12))

    # ── General Guidelines ──
    add_paragraph(doc, "General Guidelines for Assessment", bold=True, font_size=Pt(13),
                  space_after=Pt(6), space_before=Pt(12))

    guidelines = [
        "Attach a title page to your assignment. Use the previous page as your cover sheet and ensure that you fill in the details accurately.",
        "All assignment work should be prepared using word processing software and should be organized and clearly structured.",
    ]
    for g in guidelines:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(g)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(3)

    add_paragraph(doc, "Word Processing Guidance", bold=True, font_size=Pt(13),
                  space_after=Pt(6), space_before=Pt(6))

    wp_guidelines = [
        "Use a font type that makes it easy for your examiner to read, for example, Arial or Times New Roman.",
        "Use 1.5-line line spacing. Left-justify all paragraphs.",
        "Ensure that all headings are consistent in terms of size and font style.",
        "Use the footer function to insert your Student USN, Name, Unit Code, Unit Title, and Page Number on each page.",
        "Use the spell check and grammar check functions to check your assignment.",
    ]
    for g in wp_guidelines:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(g)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(3)

    add_paragraph(doc, "Assignment Submission", bold=True, font_size=Pt(13),
                  space_after=Pt(6), space_before=Pt(6))

    sub_guidelines = [
        "Ensure that you give yourself enough time to complete the assignment by the due date. You must take responsibility for managing your own time effectively.",
        "Submit an electronic version uploaded to the student learning management system on or before the stated deadline.",
        "Submissions received up to 72 hours after the deadline will be penalised in line with JAIN College's Late Submission Policy.",
        "If you are unable to hand in your assignment on time and have valid reasons, such as illness, you may apply (in writing) for an extension to the due date. Mitigating circumstances must be submitted with documentary evidence.",
        "Non-submission of work without a valid reason will lead to an automatic REFERRAL. You will then be asked to complete an alternative assignment.",
    ]
    for g in sub_guidelines:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(g)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(3)

    # ── Declaration ──
    add_paragraph(doc, "Declaration of Authenticity", bold=True, font_size=Pt(13),
                  space_after=Pt(6), space_before=Pt(12))

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("Exercise meticulous care when incorporating the work or ideas of others into your assignment. Ensure that all sources are properly cited using the Harvard referencing system.")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING

    # Page break before Statement of Originality
    doc.add_page_break()

    add_paragraph(doc, "Statement of Originality and Student Declaration",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(13),
                  space_after=Pt(12), space_before=Pt(12))

    declarations = [
        "I hereby declare that I fully understand the concept of plagiarism, which includes the use of another individual's work without attribution or the use of one's own previously submitted work without acknowledgment.",
        "I know that plagiarism is a punishable offence because it constitutes theft, and the consequences of plagiarism as outlined in JAIN College's policies.",
        "I understand JAIN College's plagiarism, ghost writing and copying policy.",
        "I declare that all work presented by me for every aspect of this assignment is my own, and where I have made use of another person's work, it has been properly acknowledged according to the Harvard referencing convention.",
        "I acknowledge that the attachment of this document, signed or not, constitutes my agreement to the above.",
        "I understand that my assignment will not be considered as submitted if this declaration is not included.",
    ]
    for d in declarations:
        add_paragraph(doc, d, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))

    add_paragraph(doc, "", space_after=Pt(24))
    add_paragraph(doc, "Student's Signature: ……………………………\t\t\tDate: ………………",
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))


# ═══════════════════════════════════════════════════════════════
# Markdown → DOCX Content Parser
# ═══════════════════════════════════════════════════════════════

def parse_markdown_content(md_text):
    """
    Parse markdown text and return structured content blocks.
    Each block is a dict with type and content.
    """
    blocks = []
    lines = md_text.split('\n')
    i = 0
    
    # Skip header metadata (everything before first ---)
    # Find start of actual content
    content_started = False
    dash_count = 0
    start_i = 0
    
    for idx, line in enumerate(lines):
        if line.strip() == '---':
            dash_count += 1
            if dash_count >= 2:  # After second ---
                start_i = idx + 1
                break
    
    # If we couldn't find the pattern, look for first # TASK
    if start_i == 0:
        for idx, line in enumerate(lines):
            if line.strip().startswith('# TASK') or line.strip().startswith('# References'):
                start_i = idx
                break
    
    i = start_i
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip final checklist
        if stripped.startswith('**FINAL CHECKLIST'):
            break
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            i += 1
            continue
            
        # Skip metadata lines
        if stripped.startswith('**Student:**') or stripped.startswith('**USN:**') or stripped.startswith('**Note:**'):
            i += 1
            continue
        
        # Horizontal rules (skip)
        if stripped == '---':
            i += 1
            continue
        
        # Empty lines
        if not stripped:
            i += 1
            continue
        
        # Code blocks (like Maslow's pyramid) - capture as-is for special handling
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({'type': 'code', 'content': '\n'.join(code_lines)})
            continue
        
        # Tables
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            rows = []
            for tl in table_lines:
                # Skip separator rows
                if re.match(r'^\|[\s\-:|]+\|$', tl):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows.append(cells)
            
            if rows:
                blocks.append({'type': 'table', 'rows': rows})
            continue
        
        # Headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            blocks.append({'type': f'heading{level}', 'content': text})
            i += 1
            continue
        
        # Bullet lists (- or * prefix)
        if re.match(r'^[-*]\s', stripped):
            items = []
            while i < len(lines):
                l = lines[i].strip()
                if re.match(r'^[-*]\s', l):
                    items.append(l[2:].strip())
                    i += 1
                elif l.startswith('  ') and items:  # continuation line
                    items[-1] += ' ' + l.strip()
                    i += 1
                elif not l:  # empty line, might continue
                    # peek ahead
                    if i + 1 < len(lines) and re.match(r'^[-*]\s', lines[i+1].strip()):
                        i += 1
                        continue
                    else:
                        break
                else:
                    break
            blocks.append({'type': 'bullet_list', 'items': items})
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < len(lines):
                l = lines[i].strip()
                if re.match(r'^\d+\.\s', l):
                    text = re.sub(r'^\d+\.\s', '', l)
                    items.append(text)
                    i += 1
                elif l.startswith('  ') and items:
                    items[-1] += ' ' + l.strip()
                    i += 1
                elif not l:
                    if i + 1 < len(lines) and re.match(r'^\d+\.\s', lines[i+1].strip()):
                        i += 1
                        continue
                    else:
                        break
                else:
                    break
            blocks.append({'type': 'numbered_list', 'items': items})
            continue
        
        # Regular paragraph (collect continuation lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                break
            if l.startswith('#') or l.startswith('|') or l.startswith('```') or l == '---':
                break
            if re.match(r'^[-*]\s', l) or re.match(r'^\d+\.\s', l):
                break
            para_lines.append(l)
            i += 1
        
        text = ' '.join(para_lines)
        blocks.append({'type': 'paragraph', 'content': text})
    
    return blocks


def add_formatted_run(paragraph, text, font_name=FONT_NAME, font_size=Pt(11),
                      base_bold=False, base_italic=False):
    """Add text with inline markdown formatting (**bold**, *italic*)."""
    # Process bold and italic markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            clean = part[3:-3]
            run = paragraph.add_run(clean)
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            clean = part[2:-2]
            run = paragraph.add_run(clean)
            run.bold = True
            run.italic = base_italic
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            clean = part[1:-1]
            run = paragraph.add_run(clean)
            run.italic = True
            run.bold = base_bold
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
        
        run.font.name = font_name
        run.font.size = font_size


def render_blocks_to_doc(doc, blocks):
    """Render parsed content blocks into the document."""
    for block in blocks:
        btype = block['type']
        
        if btype == 'heading1':
            # Main task headers (TASK A, TASK B, References)
            text = block['content']
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 1']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = HEADING1_SIZE
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            run.bold = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'heading2':
            text = block['content']
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 2']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = HEADING2_SIZE
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            run.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'heading3':
            text = block['content']
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 3']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = HEADING3_SIZE
            run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            run.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'heading4':
            text = block['content']
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_run(p, text, font_size=Pt(11), base_bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = LINE_SPACING

        elif btype == 'paragraph':
            text = block['content']
            # Skip notes about diagrams
            if text.startswith('[NOTE:') or text.startswith('**[NOTE:'):
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_run(p, text, font_size=Pt(11))
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'bullet_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_formatted_run(p, item, font_size=Pt(11))
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'numbered_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Number')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_formatted_run(p, item, font_size=Pt(11))
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = LINE_SPACING
        
        elif btype == 'table':
            rows = block['rows']
            if rows:
                table = add_table_with_borders(doc, rows, header_row=True, font_size=Pt(10))
                doc.add_paragraph("")  # spacing after table
        
        elif btype == 'code':
            # For Maslow's pyramid, add a note about the diagram
            content = block['content']
            if 'SELF' in content or 'ACTUAL' in content or 'PHYSIOLOGICAL' in content:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("[Maslow's Hierarchy of Needs Pyramid Diagram]")
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.paragraph_format.space_after = Pt(6)
                
                # Add the pyramid as a simple table representation
                pyramid_data = [
                    ["Level", "Need", "Workplace Application"],
                    ["5 (Top)", "Self-Actualisation", "Challenging work, creativity, autonomy"],
                    ["4", "Esteem", "Recognition, achievement, meaningful titles"],
                    ["3", "Social/Belonging", "Teamwork, social events, supportive culture"],
                    ["2", "Safety & Security", "Job security, pensions, safe workplace"],
                    ["1 (Base)", "Physiological", "Fair wages, breaks, comfortable conditions"],
                ]
                add_table_with_borders(doc, pyramid_data, header_row=True, font_size=Pt(10))
                doc.add_paragraph("")
            else:
                # Generic code block - render as indented text
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)


# ═══════════════════════════════════════════════════════════════
# Main Document Builder
# ═══════════════════════════════════════════════════════════════

def build_document(unit_key):
    """Build a complete assignment document for the given unit."""
    unit = UNITS[unit_key]
    print(f"\n{'='*60}")
    print(f"Building: {unit['title']}")
    print(f"{'='*60}")
    
    doc = Document()
    
    # ── Set default styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING
    
    # Set Heading styles
    for level, size in [(1, HEADING1_SIZE), (2, HEADING2_SIZE), (3, HEADING3_SIZE)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT_NAME
        hs.font.size = size
        hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        hs.font.bold = True
    
    # ── Setup page and footer ──
    section = doc.sections[0]
    footer_text = f"{STUDENT_USN}_KareemSchultz_{unit['footer_code']}_{unit['footer_title']}"
    setup_footer(section, footer_text)
    
    # ── 1. Cover Page ──
    build_cover_page(doc, unit)
    doc.add_page_break()
    
    # ── 2. Guidelines ──
    build_guidelines(doc)
    doc.add_page_break()
    
    # ── 3. Assignment Content ──
    # Add main heading
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(unit['heading'])
    run.font.name = FONT_NAME
    run.font.size = HEADING1_SIZE
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    
    # Read and parse markdown content
    print(f"  Reading draft: {unit['draft_path']}")
    with open(unit['draft_path'], 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Split content and references
    if '# References' in md_content:
        parts = md_content.split('# References', 1)
        main_content = parts[0]
        ref_content = parts[1].split('**FINAL CHECKLIST')[0] if '**FINAL CHECKLIST' in parts[1] else parts[1]
    else:
        main_content = md_content
        ref_content = ""
    
    # Parse and render main content
    blocks = parse_markdown_content(main_content)
    print(f"  Parsed {len(blocks)} content blocks")
    render_blocks_to_doc(doc, blocks)
    
    # ── 4. References ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("References")
    run.font.name = FONT_NAME
    run.font.size = HEADING1_SIZE
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    
    # Parse references
    ref_lines = [l.strip() for l in ref_content.strip().split('\n') if l.strip() and not l.strip().startswith('---')]
    for ref_line in ref_lines:
        if ref_line.startswith('- [ ]') or ref_line.startswith('**FINAL'):
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Apply hanging indent for Harvard style
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = LINE_SPACING
        
        add_formatted_run(p, ref_line, font_size=Pt(11))
    
    # ── Save ──
    output_path = os.path.join(OUTPUT_DIR, unit['filename'])
    doc.save(output_path)
    print(f"  ✓ Saved: {output_path}")
    
    # Verify
    verify_doc = Document(output_path)
    print(f"  ✓ Verification: {len(verify_doc.paragraphs)} paragraphs, {len(verify_doc.tables)} tables")
    
    return output_path


# ═══════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("SCQF Level 6 Assignment Document Builder")
    print("=" * 60)
    print(f"Student: {STUDENT_NAME} ({STUDENT_USN})")
    print(f"College: {COLLEGE}")
    print(f"Output: {OUTPUT_DIR}")
    
    results = []
    for unit_key in UNITS:
        try:
            path = build_document(unit_key)
            results.append((unit_key, path, "SUCCESS"))
        except Exception as e:
            print(f"  ✗ ERROR: {e}")
            import traceback
            traceback.print_exc()
            results.append((unit_key, None, f"FAILED: {e}"))
    
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    for unit_key, path, status in results:
        unit = UNITS[unit_key]
        print(f"  {unit['code']} - {unit['title']}: {status}")
        if path:
            size = os.path.getsize(path)
            print(f"    File: {path} ({size:,} bytes)")
    print("\nDone!")
