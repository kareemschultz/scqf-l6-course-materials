"""
Generate Final Assignment Documents
Student: KAREEM SCHULTZ
USN: 252IFCBR0596
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Student Details
STUDENT_NAME = "KAREEM SCHULTZ"
STUDENT_USN = "252IFCBR0596"
SUBMISSION_DATE = "January 2026"

# Base paths
BASE_PATH = r"C:\Users\admin\Documents\SCQF L6 - Jain University Work\Course Matrix and Syllabus"

def set_footer(doc, footer_text):
    """Add footer to all sections of the document"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = footer_text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with proper formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Add space after table
    return table

def create_cover_page(doc, unit_code, unit_title):
    """Create a professional cover page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("JAIN UNIVERSITY")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'

    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SCQF Level 6 Foundation Diploma")
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()

    # Unit Info
    unit_para = doc.add_paragraph()
    unit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = unit_para.add_run(f"{unit_code}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    unit_title_para = doc.add_paragraph()
    unit_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = unit_title_para.add_run(unit_title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Student Details Table
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    labels = ["Student Name:", "Student USN:", "Submission Date:"]
    values = [STUDENT_NAME, STUDENT_USN, SUBMISSION_DATE]

    for i, (label, value) in enumerate(zip(labels, values)):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

    doc.add_page_break()

def create_declaration_page(doc):
    """Create the Statement of Originality and Student Declaration"""
    add_heading(doc, "Statement of Originality and Student Declaration", level=1)

    doc.add_paragraph()

    declaration_text = """I confirm that this assignment is my own work and has not been submitted for any other qualification. All sources of information have been acknowledged and referenced using the Harvard referencing system.

I understand that plagiarism and collusion are serious academic offences and may result in disciplinary action.

I have read and understood the assessment brief and have addressed all the required learning outcomes."""

    add_paragraph(doc, declaration_text)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature area
    sig = doc.add_paragraph()
    sig.add_run("Student Signature: ").bold = True
    sig.add_run("_" * 40)

    doc.add_paragraph()

    date = doc.add_paragraph()
    date.add_run("Date: ").bold = True
    date.add_run("_" * 20)

    doc.add_page_break()

def generate_he9e46():
    """Generate HE9E 46 - Contemporary Business Issues assignment"""
    print("Generating HE9E 46 assignment...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Cover page
    create_cover_page(doc, "HE9E 46", "Contemporary Business Issues")

    # Declaration
    create_declaration_page(doc)

    # TASK A
    add_heading(doc, "TASK A: Features and Characteristics of SMEs (50 marks)", level=1)

    # A1
    add_heading(doc, "A1: Definition of an SME (5 marks)", level=2)
    add_paragraph(doc, """So what exactly counts as an SME? The European Commission sets out pretty clear guidelines based on employee numbers and financial size. According to their 2003 recommendation, a business is considered an SME if it falls into one of these categories:""")

    add_table(doc,
        ["Category", "Employees", "Annual Turnover", "OR", "Balance Sheet Total"],
        [
            ["Medium-sized", "Less than 250", "Up to €50 million", "OR", "Up to €43 million"],
            ["Small", "Less than 50", "Up to €10 million", "OR", "Up to €10 million"],
            ["Micro", "Less than 10", "Up to €2 million", "OR", "Up to €2 million"]
        ])

    add_paragraph(doc, """The employee number is a hard limit, but businesses only need to meet either the turnover OR the balance sheet requirement - not both. What surprised me when I looked into this was just how dominant SMEs are in the economy. They make up roughly 99% of all businesses in the EU and create about two out of every three jobs (European Commission, 2003). That's a massive chunk of the economy.""")

    # A2
    add_heading(doc, "A2: Three Key Features of SMEs (15 marks)", level=2)

    add_heading(doc, "Feature 1: Flexibility and Agility", level=3)
    add_paragraph(doc, """If there's one thing that sets small businesses apart, it's how quickly they can change direction. I've seen this firsthand - big companies have layers of management and endless approval processes, while a small business owner can make a decision on Monday and implement it by Tuesday.

The COVID pandemic really showed this off. Local restaurants near me were offering delivery within days of lockdowns starting, while some of the big chains were still figuring out their strategy weeks later. With fewer employees, simpler structures, and owners who are actually there every day, small businesses can pivot fast.

That said, this can cut both ways. Quick decisions made without proper thinking can backfire badly. The flexibility that saves a business one day might sink it the next if the owner makes a rash call.""")

    add_heading(doc, "Feature 2: Independent Ownership and Management", level=3)
    add_paragraph(doc, """Most SMEs are run by the people who own them - whether that's a sole trader, a family, or a small group of partners. There's no distant board of directors or shareholders to answer to.

Think about a local bakery. If the owner hears customers asking for gluten-free options, they could add them to the menu by next week. In a large chain, that same decision would need market research, committee meetings, and probably months of planning before anything changed.

This close connection between ownership and management means the business often reflects who the owner is - their personality, their values, their quirks. The flip side? Everything depends on that one person. If they get ill, lose motivation, or just make bad choices, the whole business suffers. There's no backup.""")

    add_heading(doc, "Feature 3: Limited Resources but High Innovation Potential", level=3)
    add_paragraph(doc, """Let's be honest - SMEs usually don't have much money to throw around. Smaller budgets, fewer staff, less fancy equipment. On paper, that looks like a disadvantage.

But here's something interesting: constraints often force creativity. When you can't outspend your competition, you have to outthink them. Some of the most innovative products have come from scrappy startups working out of garages, not from corporate R&D labs. They found clever solutions to problems that bigger companies overlooked - or just didn't bother with.

Of course, limited resources also mean less cushion when things go wrong. A big company can survive a few bad months by dipping into reserves. A small business operating on thin margins might not make it through the same rough patch.""")

    # A3
    add_heading(doc, "A3: Characteristics at Each Business Life Cycle Stage (15 marks)", level=2)
    add_paragraph(doc, "Every business goes through stages - a bit like growing up, really. Understanding where you are helps prepare for what's coming next.")

    add_heading(doc, "Stage 1: Start-up", level=3)
    add_paragraph(doc, """This is where it all begins - turning an idea into an actual business. The focus here is pretty simple: prove that someone will actually pay for what you're selling.

What it looks like: Sales are tiny or non-existent at first. You're almost certainly losing money due to setup costs. Cash flow is everything - loads of startups fail just because they run out of money, even if the idea was good. The structure is basic - maybe just the founder and a couple of helpers. Everyone does a bit of everything. The failure rate is brutal - around 90% of startups don't make it.

Zomato is a good example. Back in 2008-2010, it was just a simple website listing restaurant menus. Nothing like the delivery giant it became later. Every big company started somewhere small.""")

    add_heading(doc, "Stage 2: Growth", level=3)
    add_paragraph(doc, """Once you've proved people want what you're selling, things start moving fast. This stage is exciting but honestly pretty stressful.

What it looks like: Sales shoot up as more customers find you. Profit starts appearing. The challenge becomes scaling up without quality falling apart. You need more people, which means actual HR stuff - contracts, training, management. The informal "we're all mates" culture starts needing more structure. You might need outside funding to keep up with demand.

The tricky part? What worked with five employees completely breaks when there are fifty. The systems, the communication, the way decisions get made - all of it needs to change.""")

    add_heading(doc, "Stage 3: Maturity", level=3)
    add_paragraph(doc, """A mature business has found its place in the market. Growth slows down, but things become more predictable.

What it looks like: Steady revenue, loyal customers. Operations run smoothly and efficiently. The main threat is competitors trying to steal your customers. There's a real risk of getting comfortable and missing changes in the market. Focus shifts from growing to protecting what you've built. Cash generation is usually strong.

Amul fits here - they've been operating successfully for decades. But even mature businesses can't just coast. They need to keep adapting or risk becoming irrelevant.""")

    add_heading(doc, "Stage 4: Renewal", level=3)
    add_paragraph(doc, """Instead of accepting decline, some businesses reinvent themselves. This takes guts and usually a fair bit of disruption.

What it looks like: Heavy investment in new products or technologies. Might mean entering completely new markets. Often requires fresh leadership or at least fresh thinking. Risky, but sometimes the only alternative to slowly dying. Company culture might need a major shake-up.

LEGO is my favourite example of this. In the early 2000s they were nearly bankrupt - too many product lines, losing focus. They stripped back to basics, refocused on what they were actually good at, and came back stronger than ever.""")

    add_heading(doc, "Stage 5: Decline", level=3)
    add_paragraph(doc, """When a business fails to adapt to what customers want or how the market is changing, it starts going downhill.

What it looks like: Sales and profits drop. The competitive edge that made the business successful fades. Products feel outdated. Cost-cutting becomes necessary. Hard decisions about whether to try a turnaround, sell up, or close.

Nokia's mobile phone story is the classic example. They were the biggest phone manufacturer in the world, but completely missed the smartphone shift. Apple and Samsung ate their lunch while they were still making better and better versions of phones nobody wanted anymore.""")

    # A4
    add_heading(doc, "A4: Objectives at Each Life Cycle Stage (15 marks)", level=2)
    add_paragraph(doc, "What a business tries to achieve changes as it moves through these stages. The priorities of a startup are completely different from those of a mature company.")

    add_heading(doc, "Start-up Stage Objectives", level=3)
    add_paragraph(doc, """1. Survival - This sounds dramatic, but it's the reality. Just getting through the first year is an achievement.

2. Finding Product-Market Fit - Proving that customers actually want what you're selling and will pay enough for it to be viable.

3. Getting Noticed - Building awareness on a shoestring budget. Most startups can't afford big marketing campaigns.

4. Keeping the Lights On - Making sure there's enough cash to pay the bills until revenue picks up.""")

    add_heading(doc, "Growth Stage Objectives", level=3)
    add_paragraph(doc, """1. Grabbing Market Share - Getting customers before competitors do.

2. Boosting Revenue - Capitalising on momentum while it lasts.

3. Building the Team - Hiring people and creating systems that can handle bigger volumes.

4. Reaching Consistent Profit - Moving from occasional good months to reliable earnings.""")

    add_heading(doc, "Maturity Stage Objectives", level=3)
    add_paragraph(doc, """1. Maximising Profit - Squeezing the best returns from established operations.

2. Defending Position - Stopping competitors from poaching your customers.

3. Improving Efficiency - Finding ways to do things cheaper without sacrificing quality.

4. Finding New Revenue Streams - Reducing dependence on products that might not last forever.""")

    add_heading(doc, "Renewal Stage Objectives", level=3)
    add_paragraph(doc, """1. Driving Innovation - Developing new offerings to replace ones that are fading.

2. Adapting to Change - Responding to shifts in technology or customer preferences.

3. Transforming the Organisation - Making the cultural and structural changes needed to compete in a new way.

4. Setting Up Long-term Growth - Building foundations for the next chapter of success.""")

    add_heading(doc, "Decline Stage Objectives", level=3)
    add_paragraph(doc, """1. Limiting Losses - Cutting costs to reduce the damage.

2. Protecting Cash - Keeping enough money available to have options.

3. Making the Hard Decisions - Figuring out whether to attempt a revival, sell, or shut down.

4. Communicating Honestly - Keeping employees, suppliers, and customers informed about what's happening.""")

    doc.add_page_break()

    # TASK B
    add_heading(doc, "TASK B: SME Business Strategies (50 marks)", level=1)

    # B1
    add_heading(doc, "B1: Internal and External Challenges Facing SMEs (10 marks)", level=2)
    add_paragraph(doc, "Running a small business means dealing with challenges from all directions - some you can control, others you can't.")

    add_heading(doc, "Internal Challenges", level=3)

    add_paragraph(doc, "Challenge 1: Cash Flow Problems", bold=True)
    add_paragraph(doc, """If I had to pick the single biggest killer of small businesses, it's cash flow. Research suggests about 82% of business failures link back to running out of money (NetSuite, 2024). Not failing because the idea was bad - failing because the money ran out before the idea could work.

Unlike big corporations sitting on cash reserves and easy access to credit, SMEs often operate on razor-thin margins. They struggle to pay suppliers and staff during quiet periods, take advantage of opportunities when they pop up, and survive unexpected costs or economic rough patches.

Banks want collateral and track records that new businesses simply don't have. And plenty of owners make things worse by mixing personal and business finances, making it nearly impossible to see how the business is actually doing.""")

    add_paragraph(doc, "Challenge 2: Skills Gaps and Owner Limitations", bold=True)
    add_paragraph(doc, """Here's something that doesn't get talked about enough: most SME owners are experts in what they sell, not in running a business. A brilliant chef might open a restaurant without knowing much about accounting, marketing, or managing people.

This leads to financial decisions made on gut feeling rather than proper analysis, marketing that wastes precious budget on stuff that doesn't work, trouble finding and keeping good employees, and everything depending on the owner for every single decision.

What works when you can personally watch everything falls apart once the business grows beyond one person's reach.""")

    add_heading(doc, "External Challenges", level=3)

    add_paragraph(doc, "Challenge 1: Competing Against Giants", bold=True)
    add_paragraph(doc, """Small businesses often find themselves up against competitors with vastly more resources. Imagine running a local grocery shop while Reliance Fresh opens down the road, or trying to compete with Amazon online.

Bigger competitors can offer lower prices because they buy in bulk, outspend you on advertising a hundred to one, take losses in one area while staying profitable overall, and offer more choice and convenience.

And with globalisation, it's not just local competition anymore. International companies can enter markets that used to be protected by geography.""")

    add_paragraph(doc, "Challenge 2: Regulations and Economic Conditions", bold=True)
    add_paragraph(doc, """There's a whole category of headaches that small businesses have no control over. Regulations - taxes, employment law, health and safety - eat up time and money that could be spent actually running the business.

When India introduced GST, loads of small businesses struggled with the new compliance requirements. The burden falls proportionately harder on small operations than big ones who can afford compliance teams.

Economic conditions add another layer. Inflation pushes up costs for materials and wages, while customers become more price-sensitive. Big companies have bargaining power; small ones often just have to absorb the hit.""")

    # B2
    add_heading(doc, "B2: Role of Four Functional Activities Across the Life Cycle (20 marks)", level=2)
    add_paragraph(doc, "Every business depends on four core areas: Operations, Marketing, HR, and Finance. What each area focuses on shifts as the business grows and changes.")

    add_heading(doc, "Operations Function", level=3)
    add_paragraph(doc, """Operations is about actually producing and delivering whatever you sell. It's the engine that keeps things running.

At Start-up: The focus is getting something - anything - out the door. It doesn't need to be perfect or efficient; it needs to work. A new restaurant is figuring out its recipes and kitchen flow through trial and error.

During Growth: Now you need to produce more without quality dropping. What worked for ten customers a day doesn't work for a hundred. Processes need documenting and standardising.

At Maturity: The focus shifts to doing things better, faster, cheaper. Lean principles, continuous improvement, squeezing waste out of the system.

In Decline: Cut back to match reduced demand. Keep only what's essential running.""")

    add_heading(doc, "Marketing Function", level=3)
    add_paragraph(doc, """Marketing identifies what customers want and tells them how you can provide it.

At Start-up: Working with minimal budget, trying to get noticed. Word of mouth, social media, anything free or cheap. The goal is reaching those early customers who'll take a chance on something new.

During Growth: More aggressive. The budget increases. Campaigns get more sophisticated. The aim is capturing market share while momentum is there.

At Maturity: Keep existing customers happy. Build loyalty. Fight off competitors trying to lure them away.

In Decline: Budgets get slashed. Focus on the most valuable remaining customers.""")

    add_heading(doc, "Human Resources Function", level=3)
    add_paragraph(doc, """HR deals with people - hiring, training, keeping them motivated.

At Start-up: Usually informal or basically non-existent. The founder does the hiring. Everyone pitches in on everything. Culture develops naturally from whoever's there.

During Growth: Suddenly you're hiring fast. You need actual policies, clear job descriptions, training programmes. The informal approach stops working.

At Maturity: Focus on keeping good people, developing specialists, planning for when key people leave or retire.

In Decline: The difficult job of managing layoffs while trying to hold onto the people you can't afford to lose.""")

    add_heading(doc, "Finance Function", level=3)
    add_paragraph(doc, """Finance manages the money - where it comes from, where it goes, whether there's enough.

At Start-up: Finding funding and watching every penny. How long can we last before we run out? Every expense gets questioned.

During Growth: May need more capital to fund expansion. Systems need upgrading to handle more transactions and complexity.

At Maturity: Maximising returns. Decisions about dividends, investments, maybe acquisitions.

In Decline: Preserving cash becomes critical. May involve selling assets, cutting costs, or planning for the end.""")

    add_heading(doc, "How They Work Together", level=3)
    add_paragraph(doc, """These four areas don't work in isolation. Marketing makes promises that Operations has to keep. HR provides the people who make everything happen. Finance funds it all and measures whether it's working.

At different stages, different functions lead. Start-ups are often operations-focused (build the product). Growth businesses are marketing-driven (get customers). Mature businesses focus on finance (maximise returns). Understanding this helps owners allocate their limited resources where they'll matter most.""")

    # B3
    add_heading(doc, "B3: Analysis of Two Business Strategies (20 marks)", level=2)
    add_paragraph(doc, "SMEs can't compete with large corporations on every front - they simply don't have the resources. They need to be strategic about where they focus. Two approaches that work particularly well for smaller businesses are differentiation and niche focus.")

    add_heading(doc, "Strategy 1: Differentiation Strategy", level=3)
    add_paragraph(doc, """Differentiation means being distinctly different from competitors in ways customers care about. Instead of competing on price, you compete on being unique.

What This Looks Like in Practice: A differentiating business might offer genuinely better quality that justifies charging more, customer service that goes beyond what's expected, features competitors don't have, a brand personality customers connect with emotionally, or personalisation that big companies can't match.

What Happens When It Works:

The Good: Better margins because customers pay more for something they value. Loyal customers who are buying something they can't easily get elsewhere. Protection from competitors because it's hard to copy what makes you unique. Strong brand because you become known for something specific.

The Risks: Costs more because maintaining uniqueness requires ongoing investment. Copycats appear because success attracts imitators. Tastes change and what customers value today might not matter tomorrow. Smaller market because not everyone pays premium prices.

Real Example: Apple is the obvious big example - not the cheapest, but people pay extra for the design and experience. For something smaller-scale, think about a local café that uses only local, organic ingredients and creates an atmosphere no chain could replicate. They're not trying to be Starbucks; they're being something Starbucks can't be.""")

    add_heading(doc, "Strategy 2: Niche/Focus Strategy", level=3)
    add_paragraph(doc, """A niche strategy means concentrating on serving one specific type of customer really well, rather than trying to appeal to everyone.

What This Looks Like in Practice: A niche business might focus on a specific area (best pizza in this particular neighbourhood), a specific group (clothing for tall women), a specialist product (only vintage mechanical watches), or an underserved need (gluten-free bakery items).

What Happens When It Works:

The Good: Deep customer knowledge because you really understand what these specific customers need. Strong relationships because specialist service creates loyal fans. Less competition because big companies often ignore small niches as not worth their time. Efficient focus because all resources go to one thing instead of being spread thin. Potential dominance because you can own your little corner of the market.

The Risks: Limited size because the niche might be too small for significant growth. Vulnerability because if the niche disappears, so does your business. Over-dependence because all eggs in one basket is risky. Attracting attention because if you're too successful, bigger players might notice.

Real Example: Razer built its whole business around gaming accessories - keyboards, mice, laptops designed for gamers. They didn't try to compete in the general PC market; they became the specialist that serious gamers trust. On a smaller scale, a law firm that only helps startups with registration becomes the go-to for every new entrepreneur in town.""")

    add_heading(doc, "Why These Strategies Suit SMEs", level=3)
    add_paragraph(doc, """Both approaches play to what small businesses do well: Flexibility to adapt quickly to what specific customers want, personal touch for building genuine relationships, passion from owner-operators who really know and care about their field, and speed for making decisions fast without corporate bureaucracy.

The key is committing fully. As business strategist Michael Porter pointed out, companies that try to do everything end up "stuck in the middle" - not cheap enough to win on price, not special enough to charge premium prices. Small businesses succeed by being excellent at something specific.""")

    doc.add_page_break()

    # References
    add_heading(doc, "References", level=1)

    refs = [
        "European Commission (2003) 'Commission Recommendation of 6 May 2003 concerning the definition of micro, small and medium-sized enterprises', Official Journal of the European Union, L 124, pp. 36-41.",
        "",
        "Corporate Finance Institute (n.d.) 'Business Life Cycle - Understanding the 5 Different Stages'. Available at: https://corporatefinanceinstitute.com/resources/valuation/business-life-cycle/ (Accessed: 20 January 2026).",
        "",
        "NetSuite (2024) '10 Top Financial Challenges for Small Businesses and How to Overcome Them'. Available at: https://www.netsuite.com/portal/resource/articles/business-strategy/small-business-financial-challenges.shtml (Accessed: 20 January 2026).",
        "",
        "Porter, M.E. (1985) Competitive Advantage: Creating and Sustaining Superior Performance. New York: Free Press.",
        "",
        "JAIN Online (2025) 'HE9E 46 Contemporary Business Issues Week 1 Lecture Slides'. Bangalore: JAIN University."
    ]

    for ref in refs:
        if ref:
            add_paragraph(doc, ref)

    # Set footer
    footer_text = f"{STUDENT_USN}_KareemSchultz_HE9E 46_Contemporary Business Issues"
    set_footer(doc, footer_text)

    # Save
    output_path = os.path.join(BASE_PATH, "My_Assignments", "HE9E_46_Contemporary_Business", "Final", "HE9E_46_Final_Assignment.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def generate_j229_76():
    """Generate J229 76 - Understanding Business assignment"""
    print("Generating J229 76 assignment...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Cover page
    create_cover_page(doc, "J229 76", "Understanding Business")

    # Declaration
    create_declaration_page(doc)

    # TASK A
    add_heading(doc, "TASK A: Large Organisations (50 marks)", level=1)

    # A1
    add_heading(doc, "A1: Comparing Features of Three Large Organisations from Different Sectors (10 marks)", level=2)
    add_paragraph(doc, "Large organisations don't all work the same way. Depending on which sector they're in - private, public, or third sector - they have completely different purposes, ownership structures, and ways of operating. Looking at three very different organisations helps show these differences clearly.")

    add_heading(doc, "Tata Motors (Private Sector)", level=3)
    add_paragraph(doc, """Tata Motors is one of India's biggest car manufacturers, established back in 1945 and based in Mumbai. As a private sector company, its main job is making money for the people who own it - its shareholders.

What defines Tata Motors: Who owns it - Shareholders, through the holding company Tata Sons. You can buy and sell Tata shares on Indian stock exchanges. Where the money comes from - Selling cars, issuing shares, bank loans, and reinvesting profits. Who calls the shots - A board of directors representing shareholders. What drives decisions - Ultimately, it's about returns for investors. How big - Over 58,000 employees globally, and they own Jaguar Land Rover too.""")

    add_heading(doc, "NHS Scotland (Public Sector)", level=3)
    add_paragraph(doc, """NHS Scotland is the publicly-funded healthcare system for Scotland, set up by the National Health Service (Scotland) Act back in 1947.

What defines NHS Scotland: Who owns it - The government, on behalf of all Scottish citizens. Where the money comes from - Entirely from taxes - when you go to the doctor or hospital, you don't pay at the time. Who calls the shots - NHS Boards that answer to Scottish Ministers. What drives decisions - Providing healthcare, not making profit. How big - Around 160,000 staff, £17 billion annual budget, 14 regional boards.""")

    add_heading(doc, "Oxfam (Third Sector)", level=3)
    add_paragraph(doc, """Oxfam is an international charity that started in Oxford in 1942, now working in over 90 countries around the world.

What defines Oxfam: Who owns it - Nobody, really - it's run by trustees on behalf of the people it helps. Where the money comes from - Donations, charity shops, grants, fundraising. Who calls the shots - Trustees following the charitable objectives. What drives decisions - Ending poverty - not profit, not politics. How big - 20+ member organisations, reached 14.3 million people last year.""")

    add_heading(doc, "Comparison Summary", level=3)
    add_table(doc,
        ["Feature", "Tata Motors", "NHS Scotland", "Oxfam"],
        [
            ["Sector", "Private", "Public", "Third"],
            ["Purpose", "Make profit", "Deliver services", "Social good"],
            ["Ownership", "Shareholders", "Government", "Trustees"],
            ["Funding", "Sales, shares", "Taxation", "Donations"],
            ["Accountable to", "Shareholders", "Ministers/citizens", "Beneficiaries"]
        ])

    add_paragraph(doc, """What strikes me is how the same activity - running a large organisation - looks completely different depending on why it exists. Tata's decisions ultimately come back to "will this make money?", NHS Scotland asks "will this help patients?", and Oxfam considers "will this reduce poverty?" """)

    # A2
    add_heading(doc, "A2: Comparing Organisational Types and Growth Methods (10 marks)", level=2)

    add_heading(doc, "Organisational Types", level=3)
    add_paragraph(doc, """Tata Motors is a Public Limited Company (PLC). That means anyone can buy shares on the stock exchange. This gives access to huge amounts of capital - if investors believe in the company, they'll put money in. The trade-off is transparency. PLCs have to publish their accounts and face scrutiny from shareholders who want returns.

NHS Scotland is what's called a statutory body - it exists because Parliament said it should. The National Health Service (Scotland) Act 1947 literally created it. Its structure isn't designed for commercial efficiency; it's designed by law to deliver healthcare.

Oxfam is registered as a charitable foundation. Unlike companies, charities can't take profits out. Any money left over has to go back into the mission. This changes everything about how decisions get made.""")

    add_heading(doc, "How They've Grown", level=3)
    add_paragraph(doc, """Tata Motors - Growing by Buying: The biggest example is when Tata bought Jaguar Land Rover from Ford in 2008 for $2.3 billion. Overnight, Tata went from being an Indian car company to owning two of Britain's most famous car brands. They didn't have to build those brands from scratch - they just bought them. More recently, in 2024, they split the company into separate commercial and passenger vehicle divisions, letting each focus on its own market.

NHS Scotland - Growing Through Politics: NHS Scotland doesn't really "grow" in the business sense. It expands when the Scottish Government decides to put more money in. New hospitals, more staff, expanded services - all of this happens through budget decisions made by politicians based on what the population needs.

Oxfam - Growing Through Partnership: Oxfam started as one committee in Oxford during World War II. Now it's a confederation of 20+ independent national organisations working together under the Oxfam International umbrella. This model lets them expand into new countries while respecting local knowledge.""")

    # A3
    add_heading(doc, "A3: Objectives of Four Large Organisations (15 marks)", level=2)

    add_heading(doc, "1. Reliance Jio", level=3)
    add_paragraph(doc, """Main Objective: Connecting Every Indian to the Internet

When Jio launched in 2016, internet data in India was expensive and patchy. Jio's founder, Mukesh Ambani, had a vision: internet access should be like water or electricity - a basic utility everyone can afford.

Why This Matters: Jio completely disrupted the telecom market. They made data so cheap that competitors had to merge or exit. Now Jio has over 465 million subscribers. But it's bigger than just business success. Cheap internet has brought online education, digital banking, and access to information to people who never had it before. A farmer in a remote village can now check market prices on their phone. That's transformative.""")

    add_heading(doc, "2. Tata Group", level=3)
    add_paragraph(doc, """Main Objective: Building the Nation Through Business

The Tata Group has an unusual setup - the majority of shares in Tata Sons (the holding company) are owned by charitable trusts. So when Tata companies make profit, much of it eventually goes to social causes.

Why This Matters: This objective goes back generations. The Tata name in India means something - it's synonymous with integrity and nation-building. They built India's first steel plant, first airline, first five-star hotel. People trust Tata. They'll buy a Tata car partly because of what Tata represents.""")

    add_heading(doc, "3. Amul", level=3)
    add_paragraph(doc, """Main Objective: Fair Prices for Farmers

Amul is a cooperative - owned by the dairy farmers who supply the milk. Their objective is to ensure farmers get a fair deal rather than being exploited by middlemen.

Why This Matters: The numbers are striking: Amul returns 85% of every rupee back to farmers. The global average is about 33%. This model helped trigger the "White Revolution" that made India the world's largest milk producer. Around 3.6 million farmers are now members.""")

    add_heading(doc, "4. Unilever", level=3)
    add_paragraph(doc, """Main Objective: Sustainable Growth

Unilever says it wants to grow while reducing its environmental footprint. Targets include net zero emissions by 2039, 100% recyclable packaging by 2025, and regenerative farming on 1 million hectares.

Why This Matters: Customers increasingly care about environmental impact - especially younger consumers. If Unilever ignores this, competitors who do take it seriously will steal market share. There's also risk management: climate change threatens supply chains.""")

    # A4
    add_heading(doc, "A4: Three Internal Structures with Justification (15 marks)", level=2)

    add_heading(doc, "1. Flat Organisational Structure", level=3)
    add_paragraph(doc, """What it is: Few layers of management. The person at the top isn't that far from the person at the bottom. Communication flows more directly.

Example: Tech startups often use this - developers might report directly to founders.

Why organisations use it: Quick decisions without layers of approval. Better communication that doesn't get distorted. Motivated staff who feel more involved. Lower costs with fewer managers to pay.

The downsides: Gets messy as the company grows. Managers overwhelmed with too many direct reports. Less clear career paths.

Works best for: Smaller companies, creative industries, anywhere speed and innovation matter more than strict control.""")

    add_heading(doc, "2. Tall/Hierarchical Structure", level=3)
    add_paragraph(doc, """What it is: Many management levels. Clear chain of command from top to bottom. Each manager supervises a small number of people.

Example: The military, big banks, traditional manufacturing companies.

Why organisations use it: Crystal clear accountability where everyone knows exactly who they report to. Visible career ladder for progression. Specialisation where each level can focus on their specific area. Tight control for senior management oversight.

The downsides: Slow decisions needing multiple approvals. Messages get garbled passing through layers. Departments become silos that don't talk to each other.

Works best for: Large organisations where control and compliance matter - banking, healthcare, government.""")

    add_heading(doc, "3. Matrix Structure", level=3)
    add_paragraph(doc, """What it is: Combines functional departments (marketing, finance, HR) with project or product teams. People might report to two bosses.

Example: Consulting firms, construction companies, multinationals working on complex projects.

Why organisations use it: Resource sharing where specialists can work on multiple projects. Flexibility for teams to form and dissolve as needed. Knowledge spreads as people bring expertise from their department. Customer focus through organising around specific clients or products.

The downsides: Two bosses means potential for conflicting priorities. Power struggles between functional and project managers. Needs really good communication to work properly.

Works best for: Complex organisations juggling multiple projects that need diverse expertise.""")

    doc.add_page_break()

    # TASK B
    add_heading(doc, "TASK B: Business Environment (50 marks)", level=1)

    # B1
    add_heading(doc, "B1: Impact of Two Internal Factors (15 marks)", level=2)
    add_paragraph(doc, "Internal factors are things within the organisation's control that affect how well it performs. Corporate culture and management decisions are two of the biggest.")

    add_heading(doc, "Positive Example: Toyota's Corporate Culture", level=3)
    add_paragraph(doc, """Toyota has been one of the most successful car companies for over 70 years. A big part of that comes down to how they do things - what they call "The Toyota Way."

Two core ideas drive everything: Respect for People - actually valuing employees and partners. Continuous Improvement (Kaizen) - never being satisfied, always looking for better ways.

How this creates success:

Quality obsession: At Toyota, any worker on the production line can stop everything if they spot a defect. In most factories, stopping production costs money and workers would get in trouble for it. At Toyota, they'd get in trouble for NOT stopping the line. Problems get fixed immediately.

Eliminating waste: The Toyota Production System includes "Just-in-Time" manufacturing - producing only what's needed when it's needed. This reduces inventory costs and improves efficiency.

Everyone's ideas matter: Because workers feel respected and their ideas are valued, they actually think about improvements. Small continuous improvements add up to huge gains over time.

Result: Decades of reliability reputation, efficient operations, and workers who genuinely care about the company's success.""")

    add_heading(doc, "Negative Example: Boeing's Management Decisions", level=3)
    add_paragraph(doc, """The Boeing 737 MAX disaster is a brutal example of what happens when internal decisions go wrong. Two crashes in 2018 and 2019 killed 346 people. This wasn't bad luck - it was the result of choices made inside Boeing.

What went wrong:

Culture flipped from engineering to finance: Boeing used to be obsessed with engineering excellence. But former CEOs deliberately changed this. One actually said he wanted to run Boeing "like a business rather than a great engineering firm." Profit became the priority. Safety became a cost to be minimised.

Cutting corners on the 737 MAX: To compete with Airbus, Boeing rushed the 737 MAX to market. Rather than design a new aircraft, they modified an existing design and added software called MCAS to compensate for handling problems. To save money on pilot training, Boeing didn't fully explain this new system.

Ignoring warnings: Boeing's own engineers raised concerns about MCAS. Test pilots experienced problems in simulators. But management dismissed the warnings because addressing them would delay the project and cost money.

The consequences: 346 people dead. Every 737 MAX grounded worldwide. $2.5 billion settlement with the US government. $87 billion wiped off shareholder value. Reputation possibly destroyed forever.

The cost-cutting that was supposed to protect profits ended up costing far more than doing things properly ever would have.""")

    # B2
    add_heading(doc, "B2: PESTEC Analysis - Impact of Two External Factors (20 marks)", level=2)
    add_paragraph(doc, "External factors are things outside the organisation's control that still affect it. PESTEC looks at Political, Economic, Social, Technological, Environmental, and Competitive factors.")

    add_heading(doc, "Netflix - Technological Factors", level=3)
    add_paragraph(doc, """Netflix wouldn't exist without technology - and technology continues to shape everything they do.

How technology helps:

Streaming itself: Netflix started mailing DVDs. The shift to streaming in 2007 changed entertainment forever. But this was only possible because internet speeds improved. Netflix now reaches 260+ million subscribers worldwide because the technology caught up with the vision.

AI and personalisation: Netflix uses machine learning to analyse viewing habits and suggest content. The recommendation engine reportedly influences 80% of content watched on Netflix. It keeps people subscribed.

Their own infrastructure: Netflix built a global content delivery network called Open Connect, placing servers close to users around the world. This means less buffering and better quality than competitors.

Technology challenges: Netflix faces competition from Disney+, Amazon Prime, Apple TV+ - all enabled by the same streaming technology Netflix pioneered. TikTok and short-form video platforms are changing how younger people consume content entirely.""")

    add_heading(doc, "McDonald's - Social Factors", level=3)
    add_paragraph(doc, """What people think, believe, and care about massively impacts a consumer business like McDonald's.

Health consciousness: People increasingly know that fast food isn't exactly healthy. McDonald's has had to respond - adding salads, fruit options, nutritional information. They've reduced portion sizes and reformulated recipes to lower salt, sugar, and fat content in some markets.

Ethical concerns: Modern consumers want to know where ingredients come from, how workers are treated, and what happens to packaging. McDonald's has invested in sustainable sourcing and committed to recyclable packaging by 2025.

Cultural adaptation: Operating in 100+ countries, McDonald's must adapt to local food cultures. In India, where many people don't eat beef, you won't find a Big Mac. Instead there's the McAloo Tikki (potato burger). In Japan, teriyaki burgers.

The bottom line: Social factors have pushed McDonald's to innovate menus, be more transparent about sourcing, invest in sustainability, and research local markets carefully.""")

    # B3
    add_heading(doc, "B3: Two Stakeholder Conflicts of Interest (15 marks)", level=2)
    add_paragraph(doc, "Different groups have different interests in any organisation. When those interests clash, someone has to give way or compromise.")

    add_heading(doc, "Conflict 1: Shareholders vs Employees", level=3)
    add_paragraph(doc, """What's the conflict? Shareholders own the company and want returns on their investment. One way to increase profits is to cut costs - and wages are often the biggest cost. Employees want the opposite - job security, fair pay, good conditions.

A real example: During COVID, British Airways cut thousands of jobs while parent company IAG kept paying dividends to shareholders. Workers argued this was deeply unfair.

How it might be resolved: Companies can try to balance both groups fairly. Long-term thinking recognises that demotivated workers hurt performance eventually. Employee ownership through giving workers shares aligns their interests with shareholders. Honest communication about difficult trade-offs helps maintain trust.""")

    add_heading(doc, "Conflict 2: Business Expansion vs Local Community", level=3)
    add_paragraph(doc, """What's the conflict? Businesses want to grow - building new facilities benefits shareholders and employees. But local residents may have concerns about traffic, pollution, housing costs, and changes to neighbourhood character.

A real example: When Amazon announced plans for its second headquarters, cities competed aggressively. But when Amazon chose New York's Long Island City, local residents and politicians pushed back hard about housing costs, strained transport, and whether tax incentives were fair. The opposition got so intense that Amazon pulled out entirely.

How it might be resolved: Actually listen through genuine consultation before plans are finalised. Mitigate impacts by investing in infrastructure and local hiring. Compromise by modifying plans to address specific concerns. Share benefits so the local area actually gains from the development.""")

    doc.add_page_break()

    # References
    add_heading(doc, "References", level=1)

    refs = [
        "Amul (n.d.) About Us - The Amul Model. Available at: https://amul.com/m/about-us (Accessed: 23 January 2026).",
        "",
        "Harvard Business School (2024) 'Why Boeing's Problems with the 737 MAX Began More Than 25 Years Ago', Working Knowledge. Available at: https://www.library.hbs.edu/working-knowledge/ (Accessed: 23 January 2026).",
        "",
        "NHS Scotland (n.d.) About NHS Scotland. Available at: https://www.scot.nhs.uk/about-nhs-scotland/ (Accessed: 23 January 2026).",
        "",
        "Oxfam International (n.d.) How We Are Organized. Available at: https://www.oxfam.org/en/what-we-do/about/how-we-are-organized (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'Netflix PESTEL/PESTLE Analysis'. Available at: https://panmore.com/netflix-pestel-pestle-analysis-recommendations-case-study (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'McDonald's PESTEL/PESTLE Analysis'. Available at: https://panmore.com/mcdonalds-pestel-pestle-analysis-recommendations (Accessed: 23 January 2026).",
        "",
        "Panmore Institute (n.d.) 'Toyota's Organizational Culture'. Available at: https://panmore.com/toyota-organizational-culture-characteristics-analysis (Accessed: 23 January 2026).",
        "",
        "Reliance Jio (2024) About Us. Available at: https://www.jio.com (Accessed: 23 January 2026).",
        "",
        "Tata Motors (2024) Annual Report 2024. Mumbai: Tata Motors Limited.",
        "",
        "Unilever (2025) Sustainability Goals. Available at: https://www.unilever.com/sustainability/ (Accessed: 23 January 2026).",
        "",
        "Wikipedia (2026) 'The Toyota Way'. Available at: https://en.wikipedia.org/wiki/The_Toyota_Way (Accessed: 23 January 2026)."
    ]

    for ref in refs:
        if ref:
            add_paragraph(doc, ref)

    # Set footer
    footer_text = f"{STUDENT_USN}_KareemSchultz_J229 76_Understanding Business"
    set_footer(doc, footer_text)

    # Save
    output_path = os.path.join(BASE_PATH, "My_Assignments", "J229_76_Understanding_Business", "Final", "J229_76_Final_Assignment.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def generate_j22a_76():
    """Generate J22A 76 - Management of People and Finance assignment"""
    print("Generating J22A 76 assignment...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Cover page
    create_cover_page(doc, "J22A 76", "Management of People and Finance")

    # Declaration
    create_declaration_page(doc)

    # TASK A
    add_heading(doc, "TASK A: Management of People (60 marks)", level=1)

    # A1
    add_heading(doc, "A1: Three HRM Approaches (15 marks)", level=2)
    add_paragraph(doc, "Human Resource Management is basically about getting the people side of business right. Hire the right people, help them do their jobs well, and keep them growing. Three key approaches are recruitment, training, and development.")

    add_heading(doc, "Recruitment", level=3)
    add_paragraph(doc, """Recruitment is finding and hiring people for jobs. Sounds simple, but getting it wrong is expensive - not just the wasted salary, but lost productivity, training costs down the drain, and sometimes damage to team morale.

How the process typically works:
1. Job Analysis - Work out what the role actually involves and what skills are needed
2. Job Description - Write down the duties, responsibilities, and working conditions
3. Person Specification - Define what qualifications, skills, and experience you're looking for
4. Advertising - Get the word out through job websites, social media, internal notice boards, agencies
5. Shortlisting - Go through applications and pick who to interview
6. Selection - Interviews, tests, whatever else helps you choose
7. Offer and Onboarding - Make the offer and help the new person settle in

Internal vs External: Promoting from within is cheaper and faster, plus it shows existing staff they can progress. The downside is you're not getting fresh ideas. Hiring from outside brings new perspectives but costs more, takes longer, and carries more risk.""")

    add_heading(doc, "Training", level=3)
    add_paragraph(doc, """Training is about helping employees do their current job better. It costs money upfront, but pays off through better performance and fewer mistakes.

Different types:
- Induction Training: For new starters - showing them around, explaining policies, introducing colleagues
- On-the-Job Training: Learning by doing, usually with someone experienced guiding you
- Off-the-Job Training: Courses, workshops, conferences away from normal work

Why bother with training: People do their jobs better and faster. Fewer mistakes, better quality. Fewer accidents. Staff feel valued and are more likely to stay.""")

    add_heading(doc, "Development", level=3)
    add_paragraph(doc, """Training is about today. Development is about tomorrow - preparing people for future roles.

How it happens:
- Mentoring: Pairing someone junior with someone senior who shares experience and advice
- Coaching: One-to-one support focusing on specific skills
- Job Rotation: Moving people through different roles to broaden their experience
- Management Programmes: Structured training preparing high-potential people for leadership

Why development matters: If you don't develop your best people, they'll leave for somewhere that will. Development also means you've got people ready to step up when senior roles open.""")

    # A2
    add_heading(doc, "A2: Maslow's Hierarchy of Needs (15 marks)", level=2)
    add_paragraph(doc, "Abraham Maslow's Hierarchy of Needs has been around since 1943, and it's still one of the most useful ways to think about what motivates people. The basic idea: we have different levels of needs, and we're more focused on lower levels until those are satisfied.")

    add_heading(doc, "Level 1: Physiological Needs", level=3)
    add_paragraph(doc, """The basics of survival - food, water, shelter, sleep. In work terms, this translates to earning enough to live on.

In the workplace: Someone on minimum wage is primarily thinking about paying rent and buying food. They might put up with a terrible job because they need the income.

What employers can do: Pay fair wages. Give proper breaks. Keep the workplace at a reasonable temperature.""")

    add_heading(doc, "Level 2: Safety and Security Needs", level=3)
    add_paragraph(doc, """Once basic needs are covered, we start wanting stability and protection from harm.

In the workplace: Someone hearing redundancy rumours will be anxious and distracted. They can't focus on doing great work when they're worried about losing their job.

What employers can do: Offer permanent contracts where possible. Provide pension schemes, health insurance. Maintain a safe working environment.""")

    add_heading(doc, "Level 3: Social/Belonging Needs", level=3)
    add_paragraph(doc, """Humans are social. We need connection, friendship, feeling like we belong.

In the workplace: Someone working remotely who never sees colleagues can feel isolated and disconnected. They might have good pay and job security but still be unhappy.

What employers can do: Create opportunities for teamwork. Organise social events. Build a supportive, friendly culture.""")

    add_heading(doc, "Level 4: Esteem Needs", level=3)
    add_paragraph(doc, """We want to feel respected - both self-respect and recognition from others.

In the workplace: Someone who consistently delivers great work but never hears acknowledgment will eventually burn out or leave. Recognition matters.

What employers can do: Recognition programmes. Meaningful job titles. Praise from managers. Giving people responsibility.""")

    add_heading(doc, "Level 5: Self-Actualisation", level=3)
    add_paragraph(doc, """The highest level - reaching your full potential, doing meaningful work.

In the workplace: A senior manager who's financially secure and respected might still feel unfulfilled if their work isn't challenging or meaningful.

What employers can do: Provide challenging assignments. Give autonomy. Allow creativity. Connect work to a larger purpose.

Using Maslow in Practice: Different people are at different levels. A graduate might focus on levels 1 and 2 (paying debts, getting stable). A senior employee might be motivated by levels 4 and 5 (recognition, challenges). Money isn't the only motivator.""")

    # A3
    add_heading(doc, "A3: Five Types of Industrial Action (15 marks)", level=2)
    add_paragraph(doc, "Industrial action is what happens when employees and employers can't agree on something important - pay, conditions, job security - and workers take steps to put pressure on the employer.")

    add_heading(doc, "1. Strike", level=3)
    add_paragraph(doc, """A strike means workers completely stop working. No work, no production. It's the nuclear option.

How it works: Employees don't come to work. Often they'll set up picket lines outside.

Impact: Production stops dead. Orders can't be fulfilled. Revenue disappears. Customers may go to competitors. But workers lose their wages too, so strikes are usually a last resort.

Example: Rail strikes in the UK during 2022-23 caused massive travel disruption.""")

    add_heading(doc, "2. Work-to-Rule", level=3)
    add_paragraph(doc, """Workers do exactly what their contract says - nothing more, nothing less. Every single rule gets followed to the letter.

How it works: Instead of using common sense to keep things moving, workers insist on following every procedure precisely.

Impact: Major slowdowns without technically breaking any rules. Hard for employers to discipline. Exposes how much organisations rely on workers going above and beyond.

Example: Airport security following every procedure exactly, creating huge queues.""")

    add_heading(doc, "3. Go-Slow", level=3)
    add_paragraph(doc, """Workers deliberately work at a reduced pace.

How it works: Everything takes longer. More careful over every task. Moving slower.

Impact: Output drops, targets get missed. Hard to prove it's deliberate. Workers still get paid normal wages.

Example: Factory workers taking twice as long over each assembly, halving daily production.""")

    add_heading(doc, "4. Overtime Ban", level=3)
    add_paragraph(doc, """Workers refuse to work any hours beyond what their contract requires.

How it works: When the clock hits end of shift, everyone goes home regardless of workload.

Impact: Especially damaging during busy periods. Deadlines get missed. May need expensive temporary staff.

Example: NHS nurses refusing overtime during winter when hospitals are struggling.""")

    add_heading(doc, "5. Sit-In", level=3)
    add_paragraph(doc, """Workers occupy the workplace but refuse to work.

How it works: Employees show up and stay, but don't do their jobs.

Impact: Production stops completely. Prevents employers bringing in replacement workers. Gets lots of media attention.

Example: Factory workers occupying a plant that's been announced for closure.

Overall Effects: Industrial action hurts organisations through lost money, damaged customer relationships, poisoned workplace relations, management distraction, and reputation damage. But it also gives workers power when negotiations have failed.""")

    # A4
    add_heading(doc, "A4: Impact of Employment Legislation (15 marks)", level=2)
    add_paragraph(doc, "Organisations have to follow laws that protect employees. Three big ones are Health and Safety, Equality, and Data Protection.")

    add_heading(doc, "Health and Safety at Work Act 1974", level=3)
    add_paragraph(doc, """This law says employers must take reasonable steps to protect the health and safety of their workers.

What employers must do: Provide safe equipment. Ensure materials are handled safely. Give proper training. Maintain safe working environments.

Impact on organisations: Costs for safety equipment, training, risk assessments. But fewer accidents means lower insurance costs, fewer compensation claims, and less lost working time. Non-compliance can mean prosecution, fines, even prison.""")

    add_heading(doc, "Equality Act 2010", level=3)
    add_paragraph(doc, """This protects people from discrimination based on protected characteristics: age, disability, gender reassignment, marriage, pregnancy, race, religion, sex, and sexual orientation.

What employers must do: No direct or indirect discrimination. Make reasonable adjustments for disabled employees. Pay equally for equal work. Prevent harassment.

Impact on organisations: Recruitment, policies, and culture must all be fair and inclusive. Beyond avoiding lawsuits, diverse workplaces often perform better. Tribunal awards for discrimination have no upper limit.""")

    add_heading(doc, "Data Protection Act 2018 / UK GDPR", level=3)
    add_paragraph(doc, """This controls how organisations handle personal information.

Key principles: Process data lawfully and fairly. Collect only for specific purposes. Keep it accurate. Don't keep it longer than needed. Process securely.

Impact on organisations: HR must be careful about what employee data they collect and store. Staff can request access to their data. Breaches must be reported within 72 hours. Fines up to £17.5 million or 4% of global turnover.""")

    doc.add_page_break()

    # TASK B
    add_heading(doc, "TASK B: Management of Finance (40 marks)", level=1)

    # B1
    add_heading(doc, "B1: Three Sources of Finance (10 marks)", level=2)
    add_paragraph(doc, "Businesses need money to start, operate, and grow. Where that money comes from matters a lot.")

    add_heading(doc, "1. Retained Profit", level=3)
    add_paragraph(doc, """Retained profit is money the business made and kept instead of paying it out to owners.

How it works: Business makes £100,000 profit, pays £40,000 in dividends. The £60,000 left over is retained profit.

Why choose it: Free money with no interest payments. Keep ownership without giving away shares. No permission needed. Use it however you want.

The catch: Only works if the business is already profitable. Shareholders might grumble about lower dividends.""")

    add_heading(doc, "2. Bank Loan", level=3)
    add_paragraph(doc, """Borrow a set amount from a bank, pay it back over time with interest.

How it works: Apply for a loan, bank assesses the risk, if approved you get the money. Make regular repayments until it's paid off.

Why choose it: Large amounts available. Predictable repayments make budgeting easier. Keep ownership. Interest is tax-deductible.

The catch: Interest adds to the cost. Repayments required whether business is doing well or not. May need collateral.""")

    add_heading(doc, "3. Share Capital", level=3)
    add_paragraph(doc, """Selling shares - pieces of ownership in the company - to investors.

How it works: Company issues new shares and sells them. Investors become part-owners with a share of future profits.

Why choose it: No repayment required. Big amounts possible. Shared risk - if the business fails, investors lose their investment. Investors often bring useful knowledge.

The catch: You're giving away ownership. Existing owners' share shrinks. Investors expect dividends. Can lose control if investors get enough shares.""")

    # B2
    add_heading(doc, "B2: Five Purposes of Financial Statements (15 marks)", level=2)
    add_paragraph(doc, "Financial statements are formal records of what's happening with a business's money.")

    add_heading(doc, "1. Helping Investment Decisions", level=3)
    add_paragraph(doc, "People thinking about investing need information. Is this business profitable? Is it growing? How does it compare to alternatives? Without reliable financial statements, investors would be guessing.")

    add_heading(doc, "2. Informing Lending Decisions", level=3)
    add_paragraph(doc, "When a business wants a loan, the bank needs to assess whether they'll get their money back. Financial statements show if the business can afford repayments and what could serve as security.")

    add_heading(doc, "3. Enabling Business Planning", level=3)
    add_paragraph(doc, "Managers use financial statements to run the business - comparing actual performance to budget, identifying trends, making decisions about pricing and investment.")

    add_heading(doc, "4. Assessing Financial Health", level=3)
    add_paragraph(doc, "Financial statements show what the business owns, what it owes, whether it can pay its debts, and whether the trend is positive or negative.")

    add_heading(doc, "5. Meeting Legal and Tax Requirements", level=3)
    add_paragraph(doc, "Limited companies must prepare and file financial statements by law. These are used for calculating tax, informing shareholders, and allowing auditors to verify accuracy.")

    # B3
    add_heading(doc, "B3: Five Accounting Ratios and Their Limitations (15 marks)", level=2)
    add_paragraph(doc, "Accounting ratios turn raw financial numbers into meaningful measures you can compare over time or against other businesses.")

    add_heading(doc, "1. Gross Profit Margin", level=3)
    add_paragraph(doc, """Formula: (Gross Profit ÷ Revenue) × 100

What it tells you: How efficiently the business produces or sources what it sells. Higher means more money left after direct costs.

Example: £500,000 revenue, £200,000 gross profit = 40% margin.

What's good: Varies by industry. Supermarkets run on 2-5%. Software companies might hit 80%+.""")

    add_heading(doc, "2. Net Profit Margin", level=3)
    add_paragraph(doc, """Formula: (Net Profit ÷ Revenue) × 100

What it tells you: Overall profitability after ALL costs - production costs, rent, wages, interest, tax.

Example: £500,000 revenue, £50,000 net profit = 10% margin.

Why it matters: Comparing with gross margin shows how much overheads are eating into profits.""")

    add_heading(doc, "3. Current Ratio", level=3)
    add_paragraph(doc, """Formula: Current Assets ÷ Current Liabilities

What it tells you: Can the business pay what it owes in the next 12 months?

Example: £150,000 assets, £100,000 liabilities = 1.5:1 ratio.

What's good: Around 2:1 is often considered healthy. Too low suggests cash flow problems.""")

    add_heading(doc, "4. Acid Test (Quick Ratio)", level=3)
    add_paragraph(doc, """Formula: (Current Assets - Inventory) ÷ Current Liabilities

What it tells you: Same as current ratio, but strips out inventory which might be hard to sell quickly.

Example: £150,000 assets, £60,000 inventory, £100,000 liabilities = 0.9:1.

What's good: 1:1 is often the minimum comfortable level.""")

    add_heading(doc, "5. Return on Capital Employed (ROCE)", level=3)
    add_paragraph(doc, """Formula: (Operating Profit ÷ Capital Employed) × 100

What it tells you: How effectively the business uses its capital to generate profit.

Example: £80,000 profit, £400,000 capital = 20% ROCE.

What's good: Should exceed the cost of borrowing. If paying 10% interest but only making 8% on capital, you're destroying value.""")

    add_heading(doc, "Limitations of Accounting Ratios", level=3)
    add_paragraph(doc, """1. Looking backwards: Ratios are based on historical data. They tell you what happened, not what will happen.

2. Apples and oranges: Different companies use different accounting methods, making comparison misleading.

3. Industry differences: What's good in one industry is terrible in another.

4. Snapshot in time: Financial statements capture one specific moment that might not be typical.

5. Window dressing: Companies can manipulate figures before reporting to make ratios look better.

6. Missing the non-financial: Ratios don't capture customer satisfaction, employee morale, or brand strength.

7. Context is everything: A falling ratio isn't necessarily bad (might reflect investment). A rising ratio isn't necessarily good (might reflect harmful cost-cutting).""")

    doc.add_page_break()

    # References
    add_heading(doc, "References", level=1)

    refs = [
        "ACAS (n.d.) 'Strikes and Industrial Action'. Available at: https://www.acas.org.uk/strikes-and-industrial-action (Accessed: 23 January 2026).",
        "",
        "BDC (n.d.) 'What are Financial Statements?'. Available at: https://www.bdc.ca/en/articles-tools/entrepreneur-toolkit/templates-business-guides/glossary/financial-statements (Accessed: 23 January 2026).",
        "",
        "CIPD (n.d.) 'Data Protection and GDPR in the Workplace'. Available at: https://www.cipd.org/uk/knowledge/factsheets/data-protection-factsheet/ (Accessed: 23 January 2026).",
        "",
        "Corporate Finance Institute (n.d.) 'Financial Ratios'. Available at: https://corporatefinanceinstitute.com/resources/accounting/financial-ratios/ (Accessed: 23 January 2026).",
        "",
        "Health and Safety Executive (n.d.) 'Health and Safety at Work Act 1974'. Available at: https://www.hse.gov.uk (Accessed: 23 January 2026).",
        "",
        "Simply Psychology (n.d.) 'Maslow's Hierarchy of Needs'. Available at: https://www.simplypsychology.org/maslow.html (Accessed: 23 January 2026).",
        "",
        "Vaia (n.d.) 'Sources of Finance'. Available at: https://www.vaia.com/en-us/explanations/business-studies/financial-performance/sources-of-finance/ (Accessed: 23 January 2026)."
    ]

    for ref in refs:
        if ref:
            add_paragraph(doc, ref)

    # Set footer
    footer_text = f"{STUDENT_USN}_KareemSchultz_J22A 76_Management of People and Finance"
    set_footer(doc, footer_text)

    # Save
    output_path = os.path.join(BASE_PATH, "My_Assignments", "J22A_76_Management_People_Finance", "Final", "J22A_76_Final_Assignment.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print("Generating Final Assignment Documents")
    print(f"Student: {STUDENT_NAME}")
    print(f"USN: {STUDENT_USN}")
    print("=" * 60)
    print()

    # Generate all three assignments
    he9e46_path = generate_he9e46()
    j229_path = generate_j229_76()
    j22a_path = generate_j22a_76()

    print()
    print("=" * 60)
    print("All assignments generated successfully!")
    print("=" * 60)
    print()
    print("Files created:")
    print(f"  1. {he9e46_path}")
    print(f"  2. {j229_path}")
    print(f"  3. {j22a_path}")
    print()
    print("REMINDER: Please sign the declaration in each document before submission!")
