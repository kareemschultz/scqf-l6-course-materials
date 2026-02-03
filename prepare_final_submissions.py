"""
Prepare Final Assignment Submissions
Copies official JAIN College templates and fills in student details.
Content needs to be manually copied from the updated draft files.
"""

import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt

# Student Details
STUDENT_NAME = "KAREEM SCHULTZ"
STUDENT_USN = "252IFCBR0596"
COLLEGE = "JAIN College"

# Base paths
BASE_DIR = Path(r"C:\Users\admin\Documents\SCQF L6 - Jain University Work\Course Matrix and Syllabus")

# Template paths (using the ones in the main directory)
TEMPLATES = {
    "HE9E_46": BASE_DIR / "HE9E 46_Contemporary Business Issues_Student Assessment Template.docx",
    "J229_76": BASE_DIR / "J229 76_Understanding Business._Student Assessment Template.docx",
    "J22A_76": BASE_DIR / "J22A 76_Management of People and Finance_Student_Assessment Template.docx"
}

# Output paths
OUTPUT_DIR = BASE_DIR / "Final_Submissions"

OUTPUTS = {
    "HE9E_46": OUTPUT_DIR / f"{STUDENT_USN}_KareemSchultz_HE9E_46_Contemporary_Business_Issues.docx",
    "J229_76": OUTPUT_DIR / f"{STUDENT_USN}_KareemSchultz_J229_76_Understanding_Business.docx",
    "J22A_76": OUTPUT_DIR / f"{STUDENT_USN}_KareemSchultz_J22A_76_Management_People_Finance.docx"
}

def fill_student_details(doc, usn, name, college):
    """
    Try to find and fill student details in the document.
    This looks for common placeholder patterns in the template.
    """
    # Common text patterns to look for and replace
    replacements = {
        "Student USN:": f"Student USN: {usn}",
        "Student Name:": f"Student Name: {name}",
        "College Name & Site:": f"College Name & Site: {college}",
        "College Name & Site: JAIN College": f"College Name & Site: {college}",
    }

    changes_made = 0

    # Search through paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text and new_text not in para.text:
                # Check if it's just the label without value
                if para.text.strip() == old_text or para.text.strip() == old_text.rstrip(':'):
                    para.text = new_text
                    changes_made += 1

    # Search through tables (cover pages often use tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    # Fill in USN field
                    if "Student USN" in text and usn not in text:
                        if text == "Student USN:" or text == "Student USN":
                            para.text = f"Student USN: {usn}"
                            changes_made += 1
                    # Fill in Name field
                    if "Student Name" in text and name not in text:
                        para.text = f"Student Name: {name}"
                        changes_made += 1

    return changes_made


def prepare_submission(unit_code):
    """Copy template and prepare for submission."""
    template_path = TEMPLATES[unit_code]
    output_path = OUTPUTS[unit_code]

    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}")
        return False

    # Create output directory if needed
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Copy the template
    shutil.copy(template_path, output_path)
    print(f"Copied template: {template_path.name}")

    # Open and try to fill in student details
    try:
        doc = Document(output_path)
        changes = fill_student_details(doc, STUDENT_USN, STUDENT_NAME, COLLEGE)
        doc.save(output_path)
        print(f"  -> Saved to: {output_path.name}")
        if changes > 0:
            print(f"  -> Filled in {changes} student detail fields")
        else:
            print(f"  -> Note: Couldn't auto-fill student details - please fill manually")
    except Exception as e:
        print(f"  -> Warning: Could not modify document: {e}")
        print(f"  -> Template copied but needs manual editing")

    return True


def main():
    print("=" * 60)
    print("PREPARING FINAL ASSIGNMENT SUBMISSIONS")
    print(f"Student: {STUDENT_NAME} ({STUDENT_USN})")
    print("=" * 60)
    print()

    # Process each assignment
    success_count = 0
    for unit_code in TEMPLATES.keys():
        print(f"\nProcessing {unit_code}...")
        if prepare_submission(unit_code):
            success_count += 1

    print()
    print("=" * 60)
    print(f"COMPLETED: {success_count}/{len(TEMPLATES)} templates prepared")
    print(f"Output folder: {OUTPUT_DIR}")
    print("=" * 60)
    print()
    print("NEXT STEPS:")
    print("1. Open each .docx file in Word")
    print("2. Fill in the cover page details (USN, Name, Date)")
    print("3. Copy content from the corresponding Draft_v2_Humanized.md file")
    print("4. Add Maslow's pyramid diagram to J22A 76 (required!)")
    print("5. Add footer to every page: StudentUSN_Name_UnitCode_UnitTitle_PageNumber")
    print("6. Sign the declaration page")
    print("7. Final read-through")
    print("8. Submit to LMS before deadline")
    print()
    print("DRAFT FILES LOCATION:")
    print("  - HE9E_46: My_Assignments/HE9E_46_Contemporary_Business/Drafts/HE9E_46_Draft_v2_Humanized.md")
    print("  - J229_76: My_Assignments/J229_76_Understanding_Business/Drafts/J229_76_Draft_v2_Humanized.md")
    print("  - J22A_76: My_Assignments/J22A_76_Management_People_Finance/Drafts/J22A_76_Draft_v2_Humanized.md")


if __name__ == "__main__":
    main()
