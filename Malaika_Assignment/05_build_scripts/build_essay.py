"""
Builds Malaika's MGMT268 Assignment #1 as a formatted Word document.
Topic: Why the Behavioral Approach is necessary beyond the Rational Approach in HRM Job Design.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import datetime

OUTPUT = Path(__file__).parent / "Malaika_MGMT268_Assessment1.docx"

ESSAY = {
    "title": "MGMT268: Individual Discussion – Assessment #1",
    "subtitle": "The Behavioral Approach to Job Design: A Necessary Complement to Scientific Management",
    "student": "Malaika",
    "course": "MGMT268 – Management of Human Resources",
    "date": "April 2026",

    "sections": [
        {
            "heading": "Introduction",
            "body": (
                "Organizations are complex human systems in which performance depends not only on the "
                "logical arrangement of tasks but also on the motivations, attitudes, and psychological "
                "states of the people performing those tasks. The question of how jobs should be designed "
                "is therefore one of the most fundamental questions in Human Resource Management (HRM). "
                "Two major approaches have emerged over the decades: the Rational Approach, rooted in "
                "Frederick Winslow Taylor's Scientific Management, and the Behavioral Approach, grounded "
                "in human relations theory and the Job Characteristics Model developed by Hackman and Oldham.\n\n"

                "The Rational Approach treats job design as an engineering problem — breaking work down "
                "into its simplest, most efficient components. While this produces measurable efficiency "
                "gains in the short term, it fails to account for the human element of work. This discussion "
                "argues that the Behavioral Approach is not merely useful but essential in managing human "
                "resources effectively. When viewed against the full scope of HRM theory covered in Topics "
                "1 through 4 of this course, it becomes clear that efficiency without motivation, "
                "satisfaction, and meaning ultimately undermines the very productivity it seeks to achieve. "
                "Both approaches are necessary, but the Behavioral Approach addresses the critical human "
                "dimensions that Scientific Management ignores."
            )
        },
        {
            "heading": "The Rational Approach: Strengths and Origins",
            "body": (
                "The Rational Approach to job design is grounded in the principles of Scientific Management, "
                "developed by Frederick Winslow Taylor and published in The Principles of Scientific "
                "Management (1909). Taylor's central premise was that by simplifying and optimizing a job, "
                "productivity can be increased significantly. Through time and motion studies, he determined "
                "that although a job can be done in many ways, there is one optimal way to achieve maximum "
                "productivity (Dessler, 1997). This approach emphasized standardization, task simplification, "
                "and the use of incentives to motivate workers to complete tasks as quickly as possible.\n\n"

                "As described in the course materials, the Rational Approach 'emphasizes the scientific "
                "study of a series of tasks and designs the job so that each job includes as few tasks as "
                "possible and can therefore be performed by any worker' (Topic 4 Notes). A real-world "
                "illustration is General Motors, whose competitive strategy of overall cost leadership "
                "uses assembly line workers with very narrowly designed tasks (Topic 4 Notes). Each worker "
                "performs a single, repetitive, highly optimized task. The logic is sound from an "
                "engineering standpoint: the simpler the task, the faster it can be performed, the less "
                "training required, and the lower the cost per unit.\n\n"

                "Scientific Management also gave rise to several HRM practices that remain relevant today, "
                "including job analysis and design, compensation using incentives as motivation, recruitment "
                "and selection, and the need for training and development (Topic 1 Notes). In this sense, "
                "the Rational Approach laid the groundwork for the modern HRM function. Its contributions "
                "cannot be dismissed; they were transformative in their time and continue to inform "
                "operational efficiency in manufacturing and process-driven industries."
            )
        },
        {
            "heading": "Why the Rational Approach Alone is Insufficient",
            "body": (
                "Despite its contributions, Scientific Management was described as obsolete by the 1920s "
                "due to failure caused by poor implementation and, critically, low worker morale (Topic 1 "
                "Notes). This is a telling historical verdict. The Rational Approach treats the worker as "
                "an extension of the machine — a unit of production rather than a human being with needs, "
                "aspirations, and psychological states. This attitude was captured vividly in the language "
                "of the era: workers were literally referred to as 'hands' — farm hands, factory hands, "
                "deck hands — as in the phrase 'all hands on deck.' They were not persons; they were "
                "appendages of the production process (Evolution of HRM, Video Lecture).\n\n"

                "Before Taylor's Scientific Management, workplaces operated under what historians call "
                "the 'Drive System,' where supervisors forced productivity through threats, strict "
                "monitoring, and humiliation — at times pointing to the factory gates to remind workers "
                "of the crowds of unemployed people waiting to take their jobs (Evolution of HRM, Video "
                "Lecture). Taylor replaced brute coercion with scientific precision, but did not "
                "fundamentally change the underlying assumption: that workers were incapable of "
                "thinking, that managers held all knowledge, and that workers merely needed to be "
                "directed. As illustrated in the video lecture, Taylor famously spent four months "
                "studying workers to determine that the optimal shovel load was exactly 21 pounds — "
                "not 20, not 22, but 21 pounds. In this system, managers did all the thinking and "
                "workers did all the doing. The human being was removed from the equation entirely.\n\n"

                "The consequence was inevitable: when workers are confined to narrow, repetitive tasks "
                "with no autonomy, no variety, and no sense of contribution, the result is boredom, "
                "disengagement, high absenteeism, and high turnover. These outcomes are precisely the "
                "opposite of what efficient job design is supposed to achieve. The Rational Approach "
                "focused entirely on the technical conditions of work while completely ignoring the "
                "human conditions (Evolution of HRM, Video Lecture). Managers must get both right.\n\n"

                "The landmark Hawthorne Studies of 1927, conducted at the Western Electric plant by "
                "Elton Mayo and Fritz Roethlisberger, delivered the empirical proof. Originally "
                "commissioned to examine the impact of lighting on productivity — a classically "
                "rational inquiry — the studies discovered something far more significant: the "
                "attention paid to workers caused their productivity to increase (Topic 1 Notes). "
                "Social factors in the work environment have a significant effect on worker "
                "productivity. How workers feel matters as much as how tasks are arranged.\n\n"

                "This gave rise to the Human Relations movement of the 1930s, the Labour Movement's "
                "demands for dignity and fair conditions, and ultimately Behavioural Science — all of "
                "which converged on the same conclusion: the purely Rational Approach to managing "
                "people is insufficient (Topic 1 Notes; Anthony et al., 1996). The modern HRM "
                "function, which replaced Personnel Management in the 1980s, placed great emphasis "
                "on performance, job satisfaction, worker involvement and commitment, and rewards "
                "based on individual and group performance (Topic 1 Notes). Personnel Management "
                "was largely administrative — a 'hire and fire' operation. HRM recognizes that "
                "people are a firm's most important asset and must be valued, empowered, and "
                "committed (Evolution of HRM, Video Lecture). This paradigm shift is the "
                "institutional acknowledgment that the Rational Approach alone is not enough."
            )
        },
        {
            "heading": "The Behavioral Approach: Theory and Framework",
            "body": (
                "The Behavioral Approach to job design examines a job's core dimensions — those aspects "
                "of the job that lead to meaningful work, responsibility for outcomes, and knowledge of "
                "results — and builds jobs around these components (Topic 4 Notes). According to Higgins "
                "(1994), four major sets of factors influence the outcome of any given job: environmental "
                "factors (PEST, industry, market), organizational factors (mission, goals, leadership "
                "style), individual characteristics of the job holder (personality, needs, attitudes, "
                "values, motivation), and the primary job dimensions themselves.\n\n"

                "The most influential framework within the Behavioral Approach is the Job Characteristics "
                "Model developed by Hackman and Oldham. This model identifies five core job characteristics "
                "that must be built into a job's design:\n\n"

                "1. Skill Variety: the degree to which the job requires employees to perform a wide range "
                "of operations or use a variety of procedures. This directly addresses the monotony "
                "created by the Rational Approach's task simplification.\n\n"

                "2. Task Identity: the extent to which employees complete a whole piece of work and can "
                "clearly identify the results of their efforts. Assembly line workers who tighten one "
                "bolt have no task identity; they cannot see their contribution in the final product.\n\n"

                "3. Task Significance: the extent to which an employee perceives a significant impact on "
                "others as a result of their work. Workers who understand that their work matters are "
                "more engaged and motivated.\n\n"

                "4. Autonomy: the extent to which employees have freedom, independence, and a major say "
                "in scheduling, selecting equipment, and deciding procedures. Scientific Management "
                "removed autonomy entirely by specifying the 'one best way' for every task.\n\n"

                "5. Feedback: the degree to which employees receive information about how well they are "
                "performing. Without feedback, workers cannot improve, learn, or feel a sense of "
                "accomplishment (Topic 4 Notes).\n\n"

                "These five characteristics, when present in a job, generate three critical psychological "
                "states: Experienced Meaningfulness (the degree to which employees perceive work as "
                "valuable and worthwhile), Experienced Responsibility (the degree to which employees feel "
                "accountable for outcomes), and Knowledge of Results (the degree to which employees "
                "understand how well they are performing) (Topic 4 Notes, Lecture). When these "
                "psychological states are activated, the model predicts highly positive personal and "
                "work outcomes: high internal motivation, high job satisfaction, high work quality, and "
                "low absenteeism and turnover (Topic 4 Lecture). These outcomes are exactly what any "
                "organization managing human resources seeks to achieve — and they cannot be produced "
                "by task efficiency alone."
            )
        },
        {
            "heading": "Practical Tools of the Behavioral Approach",
            "body": (
                "The Behavioral Approach is not merely theoretical. HR managers have concrete tools at "
                "their disposal to redesign jobs in ways that address employees' psychological needs:\n\n"

                "Job Rotation involves moving an employee from one job to another, predominantly at the "
                "same level. This breaks the routine of doing the same job repeatedly, reduces boredom, "
                "and improves skill variety (Topic 4 Notes). While it does not fundamentally enrich the "
                "job, it is a practical first step in addressing the psychological damage done by "
                "prolonged, repetitive task performance.\n\n"

                "Job Enlargement refers to increasing the scope of a job by extending the range of its "
                "duties and tasks within the same level. The goal is to include different activities "
                "rather than simply adding more workload. Job enlargement can improve task identity "
                "and task significance — employees begin to see and own a larger portion of the "
                "production process (Topic 4 Notes).\n\n"

                "Job Enrichment is the most comprehensive of the three interventions. It adds depth to "
                "the job as well as breadth, providing the job holder with more autonomy, responsibility, "
                "and feedback (Topic 4 Notes). Enrichment directly targets the highest-order "
                "psychological needs of workers — the need to feel in control of their work and "
                "accountable for their outcomes. Volvo, for instance, organizes production around "
                "work teams rather than individual assembly line workers, giving teams collective "
                "ownership of a complete vehicle (Topic 4 Notes). This stands in direct contrast to "
                "General Motors' rational approach and represents a deliberate organizational choice "
                "to invest in behavioral job design to compete on product quality and differentiation.\n\n"

                "The fact that both General Motors and Volvo produce cars — the same commodity — yet "
                "use fundamentally different job design philosophies illustrates that the choice between "
                "rational and behavioral approaches is not a matter of industry but of competitive "
                "strategy and organizational values. Organizations that compete on cost may lean on "
                "the Rational Approach; those that compete on quality, innovation, and differentiation "
                "require the Behavioral Approach."
            )
        },
        {
            "heading": "The Behavioral Approach in the Broader HRM Context (Topics 1–4)",
            "body": (
                "The necessity of the Behavioral Approach becomes even clearer when viewed within the "
                "full scope of the HRM function as covered in Topics 1 through 4. HRM is defined as "
                "'the policies, practices and systems that influence employees' behaviour, attitudes and "
                "performance' (Noe et al., 2006). The key word here is behaviour — and behaviour is "
                "driven by psychological states, not just task efficiency.\n\n"

                "Job analysis, as covered in Topic 4, underpins virtually every HRM activity: recruitment, "
                "selection, performance appraisal, training and development, compensation, and career "
                "planning. A job designed purely along rational lines produces a job description focused "
                "only on tasks, duties, and responsibilities. A job designed with the Behavioral Approach "
                "in mind produces a richer job specification that accounts for the knowledge, skills, "
                "abilities, and other characteristics (KSAOs) needed — including personality traits, "
                "achievement motivation, and growth need strength (Topic 4 Notes).\n\n"

                "Growth Need Strength — the strength of a person's need for personal accomplishment, "
                "learning, and development — is a key moderating variable in the Hackman-Oldham model "
                "(Topic 4 Lecture). Not every employee will respond equally to an enriched job; those "
                "with high growth need strength will thrive, while those with low growth need strength "
                "may prefer simpler tasks. This means effective HR management requires assessing "
                "individual differences — a fundamentally behavioral, not rational, activity.\n\n"

                "Furthermore, from Topic 2 and 3, HR Strategy and Planning requires that the HR function "
                "align the design of jobs with the strategic goals of the organization. A strategy built "
                "around product differentiation, innovation, or customer service excellence demands a "
                "workforce that is motivated, engaged, and intellectually invested — qualities that "
                "emerge from behavioral job design, not from narrowly specified task performance. "
                "Strategic HRM is impossible without the Behavioral Approach."
            )
        },
        {
            "heading": "Conclusion",
            "body": (
                "The Rational Approach to job design, rooted in Taylor's Scientific Management, made "
                "an undeniable contribution to the development of modern organizations and the HRM "
                "function. Its emphasis on efficiency, standardization, and the systematic study of "
                "work processes established a foundation upon which contemporary HR practices were built. "
                "However, the historical record is unambiguous: Scientific Management was effectively "
                "obsolete by the 1920s, undermined not by its logic but by its failure to account for "
                "the human being doing the work.\n\n"

                "The Behavioral Approach does not reject efficiency — it redefines it. True organizational "
                "efficiency is not achieved by designing the fastest task; it is achieved by designing "
                "jobs that workers are motivated to perform well, day after day, with commitment, quality, "
                "and low turnover. The Hackman-Oldham Job Characteristics Model provides a rigorous, "
                "empirically grounded framework for doing exactly that. Job rotation, job enlargement, "
                "and job enrichment give HR managers the practical tools to translate the theory into "
                "action.\n\n"

                "The Rational Approach answers the question: 'What is the most efficient way to perform "
                "this task?' The Behavioral Approach answers the deeper question: 'How do we design this "
                "job so that the person performing it remains motivated, satisfied, and committed to "
                "performing it well?' Managing human resources means managing human beings — with all "
                "their psychological complexity, individual differences, and need for meaning. The "
                "Behavioral Approach is not a supplement to the Rational Approach; it is a correction "
                "to its fundamental blind spot. Both are necessary, but in the management of human "
                "resources, the Behavioral Approach is indispensable."
            )
        },
        {
            "heading": "References",
            "body": (
                "Anthony, W. P., Perrewe, P. L., & Kacmar, K. M. (1996). Strategic Human Resource "
                "Management. Harcourt Brace.\n\n"
                "Dessler, G. (1997). Human Resource Management (7th ed.). Prentice Hall.\n\n"
                "Hackman, J. R., & Oldham, G. R. (1976). Motivation through the design of work: Test "
                "of a theory. Organizational Behavior and Human Performance, 16(2), 250–279.\n\n"
                "Higgins, J. M. (1994). The Management Challenge (2nd ed.). Macmillan.\n\n"
                "Noe, R. A., Hollenbeck, J. R., Gerhart, B., & Wright, P. M. (1996). Human Resource "
                "Management: Gaining a Competitive Advantage. Irwin.\n\n"
                "Noe, R. A., Hollenbeck, J. R., Gerhart, B., & Wright, P. M. (2006). Human Resource "
                "Management: Gaining a Competitive Advantage (5th ed.). McGraw-Hill.\n\n"
                "Taylor, F. W. (1909). The Principles of Scientific Management. Harper & Brothers.\n\n"
                "MGMT268 Video Lecture. (2025). Evolution of HRM [Video]. UWI Global Campus, "
                "Topic 1. https://www.youtube.com/watch?v=Kxc8KceOb14\n\n"
                "MGMT268 Video Lecture. (2025). Job Analysis [Video]. UWI Global Campus, "
                "Topic 4. https://www.youtube.com/watch?v=oas5n1nFHQQ\n\n"
                "MGMT268 Video Lecture. (2025). Job Design [Video]. UWI Global Campus, "
                "Topic 4. https://www.youtube.com/watch?v=uUG-Z5sg2UM"
            )
        }
    ]
}


def add_paragraph(doc, text, style='Normal', bold=False, size=12, space_before=0, space_after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title block
    add_paragraph(doc, ESSAY['title'], bold=True, size=14, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, ESSAY['subtitle'], bold=False, size=12, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, ESSAY['course'], size=11, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, ESSAY['date'], size=11, space_after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    for section in ESSAY['sections']:
        heading = doc.add_heading(section['heading'], level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        # Handle body with paragraphs separated by \n\n
        paras = section['body'].split('\n\n')
        for para_text in paras:
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.first_line_indent = Inches(0.3)
                run = p.add_run(para_text.strip())
                run.font.size = Pt(12)

    doc.save(str(OUTPUT))
    print(f"Essay saved: {OUTPUT}")

    # Word count estimate
    total = sum(len(s['body'].split()) for s in ESSAY['sections'])
    print(f"Estimated word count: ~{total} words")


if __name__ == '__main__':
    build_doc()
