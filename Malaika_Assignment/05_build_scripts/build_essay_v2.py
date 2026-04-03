"""Malaika MGMT268 Assessment #1 -- Human voice rewrite + diagram embedded."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUTPUT       = Path(__file__).parent / "Malaika_MGMT268_Assessment1_FINAL.docx"
DIAGRAM      = Path(__file__).parent / "hackman_oldham_model.png"
COMPARISON   = Path(__file__).parent / "comparison_table.png"
GNS_DIAGRAM  = Path(__file__).parent / "gns_diagram.png"
DIALOGUE_IMG = Path(__file__).parent / "hr_dialogue_illustration.png"
ANIM_FRAME   = Path(__file__).parent / "animation_frame.png"
VIDEO_PATH   = (Path(__file__).parent /
                "hr_animation/media/videos/main/1080p60/HRMDialogueScene.mp4")

STUDENT_NAME    = "Malaika Abdul-Jabar"
STUDENT_ID      = "320122505"
COURSE_CODE     = "MGMT268"
COURSE_TITLE    = "Human Resource Management"
LECTURER        = "Mr. K. Francis"
SUBMISSION_DATE = "April 2026"

# ---------------------------------------------------------------------------
# Section content (plain ASCII -- no curly quotes / em-dashes)
# ---------------------------------------------------------------------------

SECTIONS = [
("Introduction",
"""When I first looked at this question, I honestly thought the rational approach covered everything. Jobs are broken into tasks, tasks are optimised, and workers perform them -- what else is there? But as we worked through Topics 1 to 4, I started to see a much bigger picture. The truth is, organisations are not just systems of tasks -- they are made up of people, and people are far more complex than any efficiency model can fully capture.

This discussion argues that while the rational approach to job design -- rooted in Frederick Taylor's Scientific Management -- laid important groundwork for how we think about work, it is simply not sufficient on its own. The behavioural approach corrects a fundamental flaw in the rational model: it actually accounts for the human being doing the job. Based on everything covered from Topic 1 through to Topic 4, my position is that both approaches are needed, but if you strip out the behavioural side, you will eventually see your workforce disengage, underperform, and walk out the door."""),

("The Rational Approach -- What It Gets Right",
"""To be fair to Scientific Management, it genuinely transformed how work was organised. Frederick Winslow Taylor published The Principles of Scientific Management in 1909, and his central idea was straightforward: find the one best way to do every job, break it down into the smallest possible tasks, and train workers to execute those tasks as efficiently as possible (Dessler, 1997; Topic 1 Notes).

What stood out in the Topic 1 video lecture was a story about Taylor that really illustrates this mindset -- he apparently spent four whole months studying workers just to determine that the optimal shovel load was exactly 21 pounds. Not 20, not 22, but 21 (MGMT268 Video, Evolution of HRM, Topic 1). That level of precision captures the spirit of the rational approach perfectly: everything is measured, everything is optimised, and there is one scientifically correct way to do everything.

And it worked -- at least in certain environments. General Motors is a real example from our Topic 4 notes. Their competitive strategy is cost leadership, so assembly line workers perform very narrowly defined, repetitive tasks. It keeps production fast and costs low (Topic 4 Notes). From a purely operational standpoint, that makes sense. The rational approach also gave us job analysis, job descriptions, task specifications, and competency-based recruitment -- all of which are still core HRM practices today (Topic 1 Notes)."""),

("Why the Rational Approach Is Not Enough -- The Human Problem",
"""Here is where things start to fall apart for the purely rational model. The course notes state directly that Scientific Management was considered largely obsolete by the 1920s, due to failure caused by poor implementation and, critically, low worker morale (Topic 1 Notes). That is a significant verdict.

The problem is not the logic -- it is the assumption behind it. The rational approach treats workers as resources, no different from the machines they operate. The Topic 1 video lecture described this in a way that stuck with me: workers were literally called "hands" -- farm hands, factory hands, deck hands -- they were not seen as people, just as appendages of the production process (MGMT268 Video, Evolution of HRM, Topic 1). Before Taylor, workplaces ran on what the video calls the "Drive System" -- supervisors would use threats and intimidation to force output, sometimes literally pointing to the factory gates to remind workers of the unemployed crowds outside waiting to take their jobs (MGMT268 Video, Evolution of HRM, Topic 1). Taylor replaced that with scientific precision, but the underlying attitude -- that workers only needed to be directed and paid -- remained the same.

What happens when you treat people that way over time? The same thing happening in the scenario this discussion started with: employees disengage, productivity drops, absenteeism rises, and good workers leave. The Hawthorne Studies of 1927 at the Western Electric plant, conducted by Elton Mayo and Fritz Roethlisberger, gave us the first real scientific evidence of something that now seems obvious. The experiment was originally designed to study the effect of lighting on productivity -- a classically rational question. What it found instead changed management theory: the attention paid to workers, the simple act of acknowledging them as human beings, increased their productivity (Topic 1 Notes).

As the video lecture puts it: getting the technical conditions right is not enough -- you have to get the human conditions right too (MGMT268 Video, Evolution of HRM, Topic 1). That is the gap the behavioural approach fills. See Figure 2 below for a side-by-side comparison of the two approaches."""),

("The Behavioural Approach -- Designing for Human Beings",
"""The behavioural approach does not reject efficiency -- it adds the layer the rational model completely ignores: the psychological dimension of work. As defined in our Topic 4 notes, this approach examines a job's core dimensions -- the aspects that lead to meaningful work, a sense of responsibility for outcomes, and knowledge of results (Topic 4 Notes).

The most important framework here is the Hackman and Oldham Job Characteristics Model (1976). The model identifies five core characteristics that every well-designed job should have:

[NUMBERED_LIST]
1. Skill Variety -- the job should require a range of different skills, directly countering the monotony created by task simplification.
2. Task Identity -- the worker should complete a whole piece of work and be able to see their contribution in the final output. An assembly line worker who tightens one bolt all day has zero task identity.
3. Task Significance -- the work should feel like it matters and has an impact on others. Purpose is a powerful motivator.
4. Autonomy -- the worker should have genuine freedom and control over how they do the job. Scientific Management removed this entirely by prescribing the "one best way" for every task.
5. Feedback -- the worker should receive clear, ongoing information about how well they are performing (Topic 4 Notes; Hackman & Oldham, 1976).
[/NUMBERED_LIST]

When these five characteristics are present, they activate three critical psychological states: the work feels meaningful, the worker feels personally responsible for outcomes, and they have knowledge of results. The personal and work outcomes that follow are exactly what any HR manager wants: high internal motivation, high job satisfaction, high work quality, and low absenteeism and turnover (Topic 4 Lecture). The rational approach has no mechanism for producing these outcomes. See Figure 1 below for a visual summary of the model."""),

("Practical Tools -- Putting the Behavioural Approach Into Practice",
"""The behavioural approach is not just theory. HR managers have three practical tools for redesigning jobs to address employees' psychological needs:

Job Rotation involves moving employees between roles at the same level on a regular basis. It reduces boredom and builds skill variety without fundamentally restructuring the organisation (Topic 4 Notes).

Job Enlargement expands the horizontal scope of a role by adding tasks at the same level, giving workers a broader view of the work and improving task identity and significance (Topic 4 Notes).

Job Enrichment is the most impactful of the three. It adds depth by giving workers more autonomy, responsibility, and feedback -- not just more tasks, but genuine ownership of the work (Topic 4 Notes).

The contrast between General Motors and Volvo from our Topic 4 notes is a perfect illustration. Both companies make cars. GM uses rational, narrowly defined task design for cost efficiency. Volvo, competing on quality and differentiation, organises production around work teams who collectively own a larger portion of the manufacturing process. Same product, fundamentally different job design philosophy, and different competitive outcomes (Topic 4 Notes). The choice between approaches is not just an HR issue -- it is a strategic one."""),

("Connecting It All Back to Topics 1-4",
"""Looking at Topics 1 through 4 together, the case for the behavioural approach only gets stronger.

From Topic 1, the entire evolution of HRM -- from early personnel management to the modern HR function -- was driven by the recognition that people are an organisation's most important asset (Noe et al., 2006). Personnel management was largely administrative -- hiring, filing, firing. HRM, by contrast, is strategic and people-centred. That shift happened because the purely rational view of workers proved insufficient (MGMT268 Video, Evolution of HRM, Topic 1).

From Topics 2 and 3, HR Strategy and Planning requires that job design aligns with organisational goals. A strategy built on quality, innovation, or customer service demands workers who are motivated and genuinely invested in the work. That kind of workforce does not emerge from narrowly defined, top-down task engineering. It requires the behavioural approach.

From Topic 4, the Hackman-Oldham model also introduces Growth Need Strength -- the degree to which a person needs personal accomplishment, learning, and development -- as a moderating variable (Topic 4 Lecture). Not every employee will respond the same way to an enriched job, which is precisely why HR managers need behavioural tools: to understand individuals and design or assign roles that actually engage them (Topic 4 Notes). See Figure 3 below for a visual representation of the full model including Growth Need Strength."""),

("Conclusion",
"""The rational approach to job design is not wrong -- it is incomplete. Taylor's Scientific Management gave us the tools to analyse and optimise work, and those contributions are still relevant today in industries where cost and speed are the primary competitive levers. But the model forgot that the person doing the work is a human being, not a machine.

The Hawthorne Studies showed us in 1927 that how workers feel matters. Decades of research since then have reinforced that finding. The Hackman-Oldham Job Characteristics Model gives us a clear, evidence-based framework for designing jobs that produce what every organisation actually needs: motivated, satisfied, high-performing employees who show up and stay.

What I found most compelling working through this topic was realising that the rational and behavioural approaches are not really opposites -- they address different dimensions of the same problem. Taylor asks: what is the most efficient way to structure this task? Hackman and Oldham ask: how do we make this a job that a person actually wants to do well? In the long run, the second question matters more. You cannot manage human resources without managing the human beings those resources belong to.

This is why the behavioural approach is not an optional add-on -- it is a necessary correction. An organisation that only applies the rational approach will optimise its processes while slowly hollowing out its workforce. The evidence from our course -- the Hawthorne findings, the GM versus Volvo comparison, the research behind the Job Characteristics Model -- all points in the same direction: people perform at their best when the work itself gives them a reason to. That is not a soft insight. That is the foundation of effective human resource management."""),

("[APPENDIX]",
"""Appendix A: Dialogue Illustration -- The Rational vs Behavioural Debate in HRM

The illustration below (Figure 4) is a hand-drawn style dialogue between a Manager and an HR Manager, depicting the key arguments covered in this discussion. It was produced as a supplementary visual aid using Excalidraw.

[DIALOGUE_IMG]

Appendix B: Companion Animation

A short animated video (HRMDialogueScene.mp4, approximately 2.5 minutes) accompanies this submission. The animation uses dialogue between the same characters to walk through the evolution of HRM thinking from Scientific Management through to the Hackman-Oldham Job Characteristics Model, drawing directly on course content from Topics 1 to 4.

[ANIM_FRAME]

The animation is included in the submission folder and can be played with any standard media player."""),

("References",
"""Anthony, W. P., Perrewe, P. L., & Kacmar, K. M. (1996). Strategic Human Resource Management. Harcourt Brace.

Dessler, G. (1997). Human Resource Management (7th ed.). Prentice Hall.

Hackman, J. R., & Oldham, G. R. (1976). Motivation through the design of work: Test of a theory. Organizational Behavior and Human Performance, 16(2), 250-279.

Higgins, J. M. (1994). The Management Challenge (2nd ed.). Macmillan.

MGMT268 Video Lecture. (2025). Evolution of HRM [Video lecture]. UWI Global Campus, Topic 1. https://www.youtube.com/watch?v=Kxc8KceOb14

MGMT268 Video Lecture. (2025). Job Analysis [Video lecture]. UWI Global Campus, Topic 4. https://www.youtube.com/watch?v=oas5n1nFHQQ

MGMT268 Video Lecture. (2025). Job Design [Video lecture]. UWI Global Campus, Topic 4. https://www.youtube.com/watch?v=uUG-Z5sg2UM

Noe, R. A., Hollenbeck, J. R., Gerhart, B., & Wright, P. M. (1996). Human Resource Management: Gaining a Competitive Advantage. Irwin.

Noe, R. A., Hollenbeck, J. R., Gerhart, B., & Wright, P. M. (2006). Human Resource Management: Gaining a Competitive Advantage (5th ed.). McGraw-Hill.

Taylor, F. W. (1909). The Principles of Scientific Management. Harper & Brothers."""),
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_header_footer(doc):
    """Add a simple header (course info) and footer (student info + page number)."""
    for section in doc.sections:
        # ----- HEADER -----
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(f"{COURSE_CODE} | {COURSE_TITLE}  --  {STUDENT_NAME}  |  ID: {STUDENT_ID}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # ----- FOOTER -----
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_left = fp.add_run(f"{STUDENT_ID} | {STUDENT_NAME} | {COURSE_CODE} | Page ")
        run_left.font.size = Pt(9)

        # Insert page number field
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run_left._r.append(fld)

        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run_left._r.append(instrText)

        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run_left._r.append(fld2)


def add_apa_reference(doc, text):
    """Add a single APA reference with 0.5-inch hanging indent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # Hanging indent: left = 0.5 in, first_line_indent = -0.5 in
    p.paragraph_format.left_indent       = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(text.strip())
    r.font.size = Pt(12)


def add_numbered_characteristics(doc, items):
    """Add the 5 JCM characteristics as a proper numbered list."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after       = Pt(6)
        p.paragraph_format.left_indent       = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(0)
        # Bold the term, normal the explanation
        parts = item.split(' -- ', 1)
        r1 = p.add_run(parts[0])
        r1.bold = True
        r1.font.size = Pt(12)
        if len(parts) > 1:
            r2 = p.add_run(' -- ' + parts[1])
            r2.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # ---- COVER PAGE ----
    for txt, sz, bold in [
        ("MGMT268 -- Human Resource Management", 14, True),
        ("Individual Discussion -- Assessment #1", 13, True),
        ("", 6, False),
        ("Why Do We Need the Behavioural Approach", 13, False),
        ("When Managing Human Resources in Organisations?", 13, False),
        ("Is the Rational Approach Sufficient?", 12, False),
        ("", 6, False),
        (f"Student Name:   {STUDENT_NAME}", 12, False),
        (f"Student ID:        {STUDENT_ID}", 12, False),
        (f"Email:               malaika.abduljabar@my.open.uwi.edu", 11, False),
        (f"Lecturer:           {LECTURER}", 12, False),
        (f"Date:                {SUBMISSION_DATE}", 12, False),
        ("", 6, False),
        ("Word count: approximately 2,000 words (excluding references)", 10, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(sz)

    # Page break after cover
    doc.add_page_break()

    # ---- HEADER / FOOTER ----
    add_header_footer(doc)

    # figures inserted after specific sections
    POST_FIGURES = {
        "Why the Rational Approach Is Not Enough -- The Human Problem": (
            COMPARISON, "Figure 2: Scientific Management vs the Behavioural Approach -- A Comparative Overview", 6.2),
        "Connecting It All Back to Topics 1-4": (
            GNS_DIAGRAM, "Figure 3: Hackman & Oldham's Job Characteristics Model including Growth Need Strength (1976)", 6.2),
    }

    def add_figure(img_path, caption, width=6.2):
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(width))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(14)
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(10)

    def add_normal_para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.3)
        r = p.add_run(text)
        r.font.size = Pt(12)

    # ---- BODY ----
    diagram_added = False

    for heading, body in SECTIONS:

        # ── Appendix ──────────────────────────────────────────────────────────
        if heading == "[APPENDIX]":
            doc.add_page_break()
            h = doc.add_heading("Appendices", level=1)
            h.paragraph_format.space_before = Pt(14)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

            for para in body.strip().split('\n\n'):
                para = para.strip()
                if not para:
                    continue
                if para == "[DIALOGUE_IMG]":
                    add_figure(DIALOGUE_IMG,
                               "Figure 4: Illustration -- The Rational vs Behavioural Debate in HRM "
                               "(created with Excalidraw)", 6.2)
                elif para == "[ANIM_FRAME]":
                    add_figure(ANIM_FRAME,
                               "Figure 5: Storyboard frame from companion animation HRMDialogueScene.mp4 "
                               "(T-chart scene, ~1:06)", 5.8)
                else:
                    add_normal_para(para)
            continue

        # ── Section heading ───────────────────────────────────────────────────
        h = doc.add_heading(heading, level=1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(6)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        # ── References ────────────────────────────────────────────────────────
        if heading == "References":
            for ref_line in body.strip().split('\n\n'):
                ref_line = ref_line.strip()
                if ref_line:
                    add_apa_reference(doc, ref_line)
            continue

        # ── Numbered characteristics list ─────────────────────────────────────
        if '[NUMBERED_LIST]' in body:
            pre, rest  = body.split('[NUMBERED_LIST]', 1)
            list_block, post = rest.split('[/NUMBERED_LIST]', 1)

            for para in pre.strip().split('\n\n'):
                para = para.strip()
                if para:
                    add_normal_para(para)

            items = []
            for line in list_block.strip().split('\n'):
                line = line.strip()
                if line and line[0].isdigit() and '. ' in line:
                    items.append(line.split('. ', 1)[1])
            add_numbered_characteristics(doc, items)

            for para in post.strip().split('\n\n'):
                para = para.strip()
                if para:
                    add_normal_para(para)
                    if 'Figure 1' in para and DIAGRAM.exists() and not diagram_added:
                        diagram_added = True
                        add_figure(DIAGRAM,
                                   "Figure 1: Hackman & Oldham's Job Characteristics Model (1976)")
            continue

        # ── Normal paragraphs ─────────────────────────────────────────────────
        for para in body.strip().split('\n\n'):
            para = para.strip()
            if para:
                add_normal_para(para)

        # Insert figure after section if mapped
        if heading in POST_FIGURES:
            img, cap, w = POST_FIGURES[heading]
            add_figure(img, cap, w)

    doc.save(str(OUTPUT))
    words = sum(len(s.split()) for _, s in SECTIONS)
    print(f"Saved: {OUTPUT}")
    print(f"Est. word count: ~{words}")


if __name__ == '__main__':
    build()
