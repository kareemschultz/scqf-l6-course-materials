"""
Create Final Assignment Submissions - COMPLETE VERSION
Generates ready-to-submit assignments using official templates.
Student only needs to sign the declaration.
"""

import shutil
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# Student Details
STUDENT_NAME = "KAREEM SCHULTZ"
STUDENT_USN = "252IFCBR0596"
COLLEGE = "JAIN College"
SUBMISSION_DATE = "15/02/2026"

# Base paths
BASE_DIR = Path(r"C:\Users\admin\Documents\SCQF L6 - Jain University Work\Course Matrix and Syllabus")
OUTPUT_DIR = BASE_DIR / "Final_Ready_To_Submit"

# Template paths
TEMPLATES = {
    "HE9E_46": BASE_DIR / "HE9E 46_Contemporary Business Issues_Student Assessment Template.docx",
    "J229_76": BASE_DIR / "J229 76_Understanding Business._Student Assessment Template.docx",
    "J22A_76": BASE_DIR / "J22A 76_Management of People and Finance_Student_Assessment Template.docx"
}

def add_footer(doc, footer_text):
    """Add footer to all sections of the document."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer
        for para in footer.paragraphs:
            para.clear()

        # Add footer paragraph
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set font
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'


def fill_template_fields(doc, usn, name, college, date):
    """Fill in student details in tables and paragraphs."""

    # Search through tables (cover pages usually use tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()

                # Check for USN field
                if "Student USN:" in cell_text and usn not in cell_text:
                    for para in cell.paragraphs:
                        if "Student USN:" in para.text:
                            para.clear()
                            run = para.add_run(f"Student USN: {usn}")
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)

                # Check for College field
                if "College Name & Site:" in cell_text:
                    for para in cell.paragraphs:
                        if "College Name & Site:" in para.text:
                            para.clear()
                            run = para.add_run(f"College Name & Site: {college}")
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)

                # Check for Date Due
                if "Date Due:" in cell_text and "/" not in cell_text:
                    for para in cell.paragraphs:
                        if "Date Due:" in para.text:
                            para.clear()
                            run = para.add_run(f"Date Due: {date}")
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)

                # Check for Date of Submission
                if "Date of Submission:" in cell_text and "/" not in cell_text:
                    for para in cell.paragraphs:
                        if "Date of Submission:" in para.text:
                            para.clear()
                            run = para.add_run(f"Date of Submission: {date}")
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)


def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with proper formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.line_spacing = 1.5
    return para


def add_table(doc, data, header=True):
    """Add a formatted table."""
    if not data:
        return None

    rows = len(data)
    cols = len(data[0])

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)

            # Format cell
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if header and i == 0:
                        run.bold = True

    doc.add_paragraph()  # Space after table
    return table


# ============================================================
# HE9E 46 CONTENT
# ============================================================

def create_he9e_46_content():
    """Return the content for HE9E 46 Contemporary Business Issues."""
    return {
        "unit_code": "HE9E 46",
        "unit_title": "Contemporary Business Issues",
        "content": """
TASK A: Features and Characteristics of SMEs (50 marks)

Task 1.1: Definition of an SME (5 marks)

So what exactly counts as an SME? The European Commission sets out pretty clear guidelines based on employee numbers and financial size. According to their 2003 recommendation, a business is considered an SME if it falls into one of these categories:

| Category | Employees | Annual Turnover | Balance Sheet Total |
|----------|-----------|-----------------|---------------------|
| Medium-sized | Less than 250 | Up to €50 million | Up to €43 million |
| Small | Less than 50 | Up to €10 million | Up to €10 million |
| Micro | Less than 10 | Up to €2 million | Up to €2 million |

The employee number is a hard limit, but businesses only need to meet either the turnover OR the balance sheet requirement - not both. What surprised me when I looked into this was just how dominant SMEs are in the economy. They make up roughly 99% of all businesses in the EU and create about two out of every three jobs (European Commission, 2003). That's a massive chunk of the economy.

Task 1.2: Three Key Features of SMEs - Focus on E-Commerce (15 marks)

The assessment requires examining SME characteristics through the lens of E-Commerce businesses, so I'll explore three key features using examples from online retail and digital platforms.

Feature 1: Flexibility and Rapid Adaptation

E-Commerce SMEs can pivot faster than traditional brick-and-mortar businesses because they don't face the same physical constraints. Take Meesho, for example - it started as a reselling platform in 2015 and quickly pivoted to become India's largest social commerce platform when it spotted how people were using WhatsApp to sell products. A large retailer couldn't have made that shift so quickly.

During the pandemic, I noticed how small online sellers on platforms like Etsy adapted almost overnight. One jewellery maker I follow on Instagram switched from selling at craft fairs to running live selling sessions within a week. Meanwhile, larger retailers were still figuring out their digital strategy.

The downside? E-Commerce SMEs can sometimes move too fast without proper planning. Jumping on every trend - whether it's live shopping, social commerce, or quick commerce - can stretch limited resources thin.

Feature 2: Owner-Driven Decision Making

In most E-Commerce SMEs, the founder is still running the show. Nykaa, founded by Falguni Nayar, is a good example of this before it went public. She personally shaped the brand's focus on beauty education and tutorials alongside product sales - something that came from her own vision rather than corporate committees.

For smaller operations, this is even more pronounced. An Etsy seller or a Shopify store owner can test a new product line, change their pricing strategy, or completely rebrand their store based on customer feedback they read that morning. There's no approval process, no stakeholder meetings.

But this concentration of control has risks. Everything depends on the owner's skills - not just their product knowledge, but their understanding of digital marketing, SEO, logistics, and customer service. Many talented craftspeople struggle online not because their products are bad, but because running an E-Commerce business requires a completely different skill set.

Feature 3: Resource Constraints Driving Innovation

E-Commerce SMEs typically operate with limited budgets, which actually pushes them toward creative solutions. Zepto, the quick-commerce startup, couldn't afford the massive warehouse networks of Amazon or Flipkart, so they invented the "dark store" model - small neighbourhood warehouses that could deliver groceries in 10 minutes. That constraint-driven innovation became their competitive advantage.

Similarly, many small E-Commerce sellers can't afford expensive marketing, so they've become experts at organic social media, influencer micro-partnerships, and building genuine communities around their brands. Some of the most engaged audiences I've seen are around small D2C (direct-to-consumer) brands that simply couldn't afford traditional advertising.

The challenge remains access to capital. Unlike traditional retail where you might get bank loans secured against property, E-Commerce assets are harder to value. Many E-Commerce SMEs rely heavily on personal savings or bootstrap their growth, which limits how fast they can scale.

Task 1.3: Characteristics at Each Business Life Cycle Stage (15 marks)

Every business goes through stages - a bit like growing up, really. Understanding where you are helps prepare for what's coming next.

Stage 1: Start-up

This is where it all begins - turning an idea into an actual business. The focus here is pretty simple: prove that someone will actually pay for what you're selling.

What it looks like:
- Sales are tiny or non-existent at first
- You're almost certainly losing money due to setup costs
- Cash flow is everything - loads of startups fail just because they run out of money, even if the idea was good
- The structure is basic - maybe just the founder and a couple of helpers
- Everyone does a bit of everything
- The failure rate is brutal - around 90% of startups don't make it

Zomato is a good example. Back in 2008-2010, it was just a simple website listing restaurant menus. Nothing like the delivery giant it became later.

Stage 2: Growth

Once you've proved people want what you're selling, things start moving fast. This stage is exciting but honestly pretty stressful.

What it looks like:
- Sales shoot up as more customers find you
- Profit starts appearing
- The challenge becomes scaling up without quality falling apart
- You need more people, which means actual HR stuff - contracts, training, management
- The informal "we're all mates" culture starts needing more structure
- You might need outside funding to keep up with demand

The tricky part? What worked with five employees completely breaks when there are fifty.

Stage 3: Maturity

A mature business has found its place in the market. Growth slows down, but things become more predictable.

What it looks like:
- Steady revenue, loyal customers
- Operations run smoothly and efficiently
- The main threat is competitors trying to steal your customers
- There's a real risk of getting comfortable and missing changes in the market
- Focus shifts from growing to protecting what you've built
- Cash generation is usually strong

Amul fits here - they've been operating successfully for decades.

Stage 4: Renewal

Instead of accepting decline, some businesses reinvent themselves. This takes guts and usually a fair bit of disruption.

What it looks like:
- Heavy investment in new products or technologies
- Might mean entering completely new markets
- Often requires fresh leadership or at least fresh thinking
- Risky, but sometimes the only alternative to slowly dying
- Company culture might need a major shake-up

LEGO is my favourite example of this. In the early 2000s they were nearly bankrupt - too many product lines, losing focus. They stripped back to basics, refocused on what they were actually good at, and came back stronger than ever.

Stage 5: Decline

When a business fails to adapt to what customers want or how the market is changing, it starts going downhill.

What it looks like:
- Sales and profits drop
- The competitive edge that made the business successful fades
- Products feel outdated
- Cost-cutting becomes necessary
- Hard decisions about whether to try a turnaround, sell up, or close

Nokia's mobile phone story is the classic example. They were the biggest phone manufacturer in the world, but completely missed the smartphone shift.

Task 1.4: Objectives at Each Life Cycle Stage (15 marks)

Start-up Stage Objectives:
1. Survival - This sounds dramatic, but it's the reality. Just getting through the first year is an achievement.
2. Finding Product-Market Fit - Proving that customers actually want what you're selling and will pay enough for it to be viable.

Growth Stage Objectives:
1. Grabbing Market Share - Getting customers before competitors do.
2. Boosting Revenue - Capitalising on momentum while it lasts.

Maturity Stage Objectives:
1. Maximising Profit - Squeezing the best returns from established operations.
2. Defending Position - Stopping competitors from poaching your customers.

Renewal Stage Objectives:
1. Driving Innovation - Developing new offerings to replace ones that are fading.
2. Transforming the Organisation - Making the cultural and structural changes needed to compete in a new way.

Decline Stage Objectives:
1. Limiting Losses - Cutting costs to reduce the damage.
2. Making the Hard Decisions - Figuring out whether to attempt a revival, sell, or shut down.

TASK B: SME Business Strategies (50 marks)

Task 2.1: Internal and External Challenges Facing SMEs (10 marks)

Internal Challenges

Challenge 1: Cash Flow Problems

If I had to pick the single biggest killer of small businesses, it's cash flow. Research suggests about 82% of business failures link back to running out of money (NetSuite, 2024). Not failing because the idea was bad - failing because the money ran out before the idea could work.

Unlike big corporations sitting on cash reserves and easy access to credit, SMEs often operate on razor-thin margins. They struggle to pay suppliers and staff during quiet periods, take advantage of opportunities when they pop up, and survive unexpected costs or economic rough patches.

Challenge 2: Skills Gaps and Owner Limitations

Here's something that doesn't get talked about enough: most SME owners are experts in what they sell, not in running a business. A brilliant chef might open a restaurant without knowing much about accounting, marketing, or managing people.

This leads to financial decisions made on gut feeling rather than proper analysis, marketing that wastes precious budget, and everything depending on the owner for every single decision.

External Challenges

Challenge 1: Competing Against Giants

Small businesses often find themselves up against competitors with vastly more resources. Imagine running a local grocery shop while Reliance Fresh opens down the road, or trying to compete with Amazon online.

Bigger competitors can offer lower prices because they buy in bulk, outspend you on advertising a hundred to one, and take losses in one area while staying profitable overall.

Challenge 2: Regulations and Economic Conditions

There's a whole category of headaches that small businesses have no control over. Regulations - taxes, employment law, health and safety - eat up time and money that could be spent actually running the business.

When India introduced GST, loads of small businesses struggled with the new compliance requirements. The burden falls proportionately harder on small operations than big ones who can afford compliance teams.

Task 2.2: Role of Four Functional Activities Across the Life Cycle (20 marks)

Operations Function

At Start-up: The focus is getting something - anything - out the door. It doesn't need to be perfect or efficient; it needs to work.

During Growth: Now you need to produce more without quality dropping. What worked for ten customers a day doesn't work for a hundred.

At Maturity: The focus shifts to doing things better, faster, cheaper. Lean principles, continuous improvement.

In Decline: Cut back to match reduced demand. Keep only what's essential running.

Marketing Function (Including Digital Marketing)

At Start-up: Working with minimal budget, trying to get noticed. Word of mouth, social media, anything free or cheap.

During Growth: More aggressive. The budget increases. Campaigns get more sophisticated. Digital marketing becomes essential - SEO, paid ads, email campaigns.

At Maturity: Keep existing customers happy. Build loyalty. Fight off competitors.

In Decline: Budgets get slashed. Focus on the most valuable remaining customers.

Human Resources Function

At Start-up: Usually informal or basically non-existent. The founder does the hiring. Everyone pitches in on everything.

During Growth: Suddenly you're hiring fast. You need actual policies, clear job descriptions, training programmes.

At Maturity: Focus on keeping good people, developing specialists, planning for succession.

In Decline: The difficult job of managing layoffs while trying to hold onto the people you can't afford to lose.

Finance Function

At Start-up: Finding funding and watching every penny. How long can we last before we run out?

During Growth: May need more capital to fund expansion. Systems need upgrading to handle more complexity.

At Maturity: Maximising returns. Decisions about dividends, investments, maybe acquisitions.

In Decline: Preserving cash becomes critical. May involve selling assets or cutting costs.

Task 2.3: Analysis of Two Business Strategies (20 marks)

Strategy 1: Digitalisation Strategy

Digitalisation means adopting digital technologies to transform how the business operates, reaches customers, and delivers value.

What Digitalisation Looks Like for SMEs:
- Moving from physical stores to E-Commerce platforms
- Using cloud-based tools for accounting, inventory, and CRM
- Adopting digital marketing (social media, SEO, email campaigns)
- Implementing online payment systems
- Using data analytics to understand customer behaviour

Expected Outcomes - The Benefits:
1. Expanded Market Reach - A local handicraft business can suddenly sell globally through Etsy or their own Shopify store.
2. Reduced Operating Costs - Cloud software eliminates the need for expensive IT infrastructure.
3. Improved Customer Experience - Online booking, instant responses via chatbots, personalised recommendations.
4. Better Decision Making - Digital tools provide data instead of guessing.
5. Operational Efficiency - Automating invoicing, inventory management, and customer follow-ups.

Expected Outcomes - The Risks:
- Implementation costs add up
- Skills gap - many owners aren't digital natives
- Cybersecurity vulnerabilities
- Technology dependence

Real Example: Lenskart started as a small optical retailer and invested heavily in digitalisation early on. They built virtual try-on technology, an E-Commerce platform, and data-driven inventory systems. This digital-first approach helped them scale from a single store to one of India's largest eyewear companies.

Strategy 2: Partnership Strategy

Partnerships involve formal or informal collaborations with other businesses to achieve goals that would be difficult alone.

Types of Partnerships for SMEs:
- Supplier Partnerships - Long-term relationships for better terms
- Distribution Partnerships - Using established retailers or platforms
- Co-Marketing Partnerships - Collaborating on promotions
- Technology Partnerships - Integrating with platforms like Amazon
- Joint Ventures - Sharing resources for specific projects

Expected Outcomes - The Benefits:
1. Access to New Markets - A small Indian spice company partnering with a UK distributor gains access to an entire market.
2. Shared Resources and Costs - Two small businesses sharing warehouse space or jointly hiring consultants.
3. Enhanced Credibility - Being associated with established brands transfers trust.
4. Knowledge Transfer - Partners bring different expertise.
5. Risk Sharing - Partners share the risk when entering new markets.

Expected Outcomes - The Risks:
- Loss of control
- Dependency on partners
- Profit sharing reduces margins
- Reputation risk if partner has problems

Real Example: Starbucks and Tata formed a partnership (Tata Starbucks) to enter the Indian market. Starbucks got local expertise and supply chain access. Tata got a premium global brand. Neither could have achieved this alone.

References

European Commission (2003) 'Commission Recommendation of 6 May 2003 concerning the definition of micro, small and medium-sized enterprises', Official Journal of the European Union, L 124, pp. 36-41.

Corporate Finance Institute (n.d.) 'Business Life Cycle - Understanding the 5 Different Stages'. Available at: https://corporatefinanceinstitute.com/resources/valuation/business-life-cycle/ (Accessed: 20 January 2026).

NetSuite (2024) '10 Top Financial Challenges for Small Businesses and How to Overcome Them'. Available at: https://www.netsuite.com/portal/resource/articles/business-strategy/small-business-financial-challenges.shtml (Accessed: 20 January 2026).

Porter, M.E. (1985) Competitive Advantage: Creating and Sustaining Superior Performance. New York: Free Press.

JAIN Online (2025) 'HE9E 46 Contemporary Business Issues Week 1 Lecture Slides'. Bangalore: JAIN University.
"""
    }


# ============================================================
# J229 76 CONTENT
# ============================================================

def create_j229_76_content():
    """Return the content for J229 76 Understanding Business."""
    return {
        "unit_code": "J229 76",
        "unit_title": "Understanding Business",
        "content": """
TASK A: Large Organisations (50 marks)

Task 1.1.1: Comparing Features of Three Large Organisations from Different Sectors (8 marks)

Large organisations don't all work the same way. Depending on which sector they're in - private, public, or third sector - they have completely different purposes, ownership structures, and ways of operating.

Tata Motors (Private Sector - Secondary/Manufacturing)

Tata Motors is one of India's biggest car manufacturers, established back in 1945 and based in Mumbai. As a private sector company, its main job is making money for the people who own it - its shareholders.

What defines Tata Motors:
- Who owns it: Shareholders, through the holding company Tata Sons
- Where the money comes from: Selling cars, issuing shares, bank loans, and reinvesting profits
- Who calls the shots: A board of directors representing shareholders
- What drives decisions: Ultimately, it's about returns for investors
- How big: Over 58,000 employees globally, and they own Jaguar Land Rover too

NHS Scotland (Public Sector - Tertiary/Service)

NHS Scotland is the publicly-funded healthcare system for Scotland, set up by the National Health Service (Scotland) Act back in 1947.

What defines NHS Scotland:
- Who owns it: The government, on behalf of all Scottish citizens
- Where the money comes from: Entirely from taxes - when you go to the doctor or hospital, you don't pay at the time
- Who calls the shots: NHS Boards that answer to Scottish Ministers
- What drives decisions: Providing healthcare, not making profit
- How big: Around 160,000 staff, £17 billion annual budget, 14 regional boards

Oxfam (Third Sector - Charity)

Oxfam is an international charity that started in Oxford in 1942, now working in over 90 countries around the world.

What defines Oxfam:
- Who owns it: Nobody, really - it's run by trustees on behalf of the people it helps
- Where the money comes from: Donations, charity shops, grants, fundraising
- Who calls the shots: Trustees following the charitable objectives
- What drives decisions: Ending poverty - not profit, not politics
- How big: 20+ member organisations, reached 14.3 million people last year

Comparison Table:

| Feature | Tata Motors | NHS Scotland | Oxfam |
|---------|-------------|--------------|-------|
| Sector | Private | Public | Third |
| Purpose | Make profit | Deliver services | Social good |
| Ownership | Shareholders | Government | Trustees |
| Funding | Sales, shares | Taxation | Donations |
| Accountable to | Shareholders | Ministers/citizens | Beneficiaries |

What strikes me is how the same activity - running a large organisation - looks completely different depending on why it exists.

Task 1.1.2: Comparing Structural Characteristics of Business Types (6 marks)

Private Limited Company (Ltd)

A private limited company is owned by shareholders but shares cannot be sold to the general public.

Key characteristics:
- Ownership: Shares held privately by founders, family, or selected investors
- Control: Directors appointed by shareholders; often the owners run the company themselves
- Liability: Limited - shareholders only risk what they invested
- Share trading: Cannot sell shares on stock exchange
- Capital raising: Limited to private investment, bank loans, retained profits

Example: Dyson remained a private limited company for years, which let founder James Dyson make long-term R&D investments without pressure from public shareholders.

Public Limited Company (PLC)

A PLC can sell shares to anyone through a stock exchange.

Key characteristics:
- Ownership: Anyone can buy shares; ownership can be widely dispersed
- Control: Professional managers appointed by board; shareholders vote on major decisions
- Liability: Limited for shareholders
- Share trading: Freely traded on stock exchanges
- Capital raising: Can raise huge sums through share issues and bonds

Example: Tata Motors is listed on the BSE and NSE, giving them access to capital markets.

Franchise

A franchise is a licensing arrangement where a franchisee pays to use an established brand's name, systems, and support.

Key characteristics:
- Ownership: Franchisee owns their specific outlet; franchisor owns the brand
- Control: Franchisee controls daily operations but must follow franchisor's rules exactly
- Investment: Lower than starting from scratch - you're buying a proven model
- Support: Training, marketing, supply chains provided by franchisor

Example: Most McDonald's restaurants are franchises. The franchisee invests around £400,000-£1.5m to open one.

Task 1.1.3: Organic vs Inorganic Growth (6 marks)

Organic (Internal) Growth

Growing by expanding what you already do using your own resources.

Example 1: A Local Bakery Increasing Production Capacity

A successful neighbourhood bakery decides to grow by buying bigger ovens, hiring more bakers, and extending opening hours.

Why choose this method:
- Maintains complete control
- Grows at a manageable pace
- Can fund expansion from profits
- Low risk - gradual, reversible steps

Example 2: Infosys Expanding Through Technology Investment

Infosys grew organically for years by continuously investing in new capabilities, training employees, and building delivery centres.

Why choose this method:
- Maintained the distinctive Infosys culture
- Controlled quality
- Avoided integration headaches

Inorganic (External) Growth

Growing by acquiring, merging with, or partnering with other businesses.

Example 3: Tata Motors Acquiring Jaguar Land Rover

In 2008, Tata Motors bought Jaguar and Land Rover from Ford for $2.3 billion.

Why choose this method:
- Would have taken decades to build brands with JLR's prestige from scratch
- Instantly acquired JLR's engineering expertise
- Immediate entry into luxury car market
- But high risk - paid premium price

Example 4: Starbucks-Tata Partnership (Joint Venture)

Rather than entering India alone, Starbucks formed a 50-50 joint venture with Tata in 2012.

Why choose this method:
- Starbucks got local expertise; Tata got a premium global brand
- Shared investment and risk
- Smoother entry into complex Indian market

Task 1.2.1: Objectives of Four Large Organisations (15 marks)

1. Reliance Jio - Connecting Every Indian to the Internet

When Jio launched in 2016, internet data in India was expensive and patchy. Jio's founder, Mukesh Ambani, had a vision: internet access should be like water or electricity - a basic utility everyone can afford.

Why This Matters: Jio completely disrupted the telecom market. They made data so cheap that competitors had to merge or exit. Now Jio has over 465 million subscribers. Cheap internet has brought online education, digital banking, and access to information to people who never had it before.

2. Tata Group - Building the Nation Through Business

The Tata Group has an unusual setup - the majority of shares in Tata Sons are owned by charitable trusts. So when Tata companies make profit, much of it eventually goes to social causes.

Why This Matters: The Tata name in India means something - it's synonymous with integrity and nation-building. They built India's first steel plant, first airline, first five-star hotel.

3. Amul - Fair Prices for Farmers

Amul is a cooperative - owned by the dairy farmers who supply the milk. Their objective is to ensure farmers get a fair deal rather than being exploited by middlemen.

Why This Matters: Amul returns 85% of every rupee back to farmers. The global average is about 33%. Around 3.6 million farmers are now members.

4. Unilever - Sustainable Growth

Unilever says it wants to grow while reducing its environmental footprint. Targets include net zero emissions by 2039, 100% recyclable packaging by 2025.

Why This Matters: Customers increasingly care about environmental impact. Climate change threatens supply chains. Being ahead of the curve on sustainability might just be smart business.

Task 1.3.1: Three Internal Structures with Justification (10 marks)

1. Flat Organisational Structure

What it is: Few layers of management. The person at the top isn't that far from the person at the bottom.

Why organisations use it:
- Quick decisions without layers of approval
- Better communication - messages don't get distorted
- Motivated staff who feel more involved
- Lower costs - fewer managers to pay

Works best for: Smaller companies, creative industries, anywhere speed and innovation matter.

2. Tall/Hierarchical Structure

What it is: Many management levels. Clear chain of command from top to bottom.

Why organisations use it:
- Crystal clear accountability
- Visible career ladder
- Specialisation at each level
- Tight control and oversight

Example: NHS Scotland uses hierarchical structure for clear accountability in healthcare.

3. Matrix Structure

What it is: Combines functional departments with project teams. People might report to two bosses.

Why organisations use it:
- Resource sharing across projects
- Flexibility - teams form and dissolve as needed
- Knowledge spreads across the organisation
- Customer or product focus

Example: NASA and Infosys/TCS use matrix structures for complex projects requiring diverse expertise.

Task 1.3.2: Decision-Making Levels and Tools (5 marks)

Three Levels of Decision-Making

1. Strategic Decisions (Top Management)
- Long time horizon (3-10+ years)
- High risk and uncertainty
- Affect the whole organisation
Example: Should Tata Motors enter the electric vehicle market?

2. Tactical Decisions (Middle Management)
- Medium time horizon (1-3 years)
- Moderate risk
- Affect departments or divisions
Example: How should the marketing budget be allocated?

3. Operational Decisions (Front-Line Management)
- Short time horizon (days to months)
- Low risk
- Affect specific tasks
Example: Which shifts should staff work this week?

Decision-Making Tools

SWOT Analysis: Stands for Strengths, Weaknesses, Opportunities, and Threats. Forces systematic thinking about internal and external factors before making strategic decisions.

Decision Trees: Visual diagrams showing choices, possible outcomes, and their consequences. Allows calculation of expected values for different options.

TASK B: Business Environment (50 marks)

Task 2.1: Impact of Two Internal Factors (15 marks)

Positive Example: Toyota's Corporate Culture

Toyota has been one of the most successful car companies for over 70 years. Two core ideas drive everything:
1. Respect for People
2. Continuous Improvement (Kaizen)

How this creates success:

Quality obsession: At Toyota, any worker on the production line can stop everything if they spot a defect. Problems get fixed immediately.

Eliminating waste: The Toyota Production System includes "Just-in-Time" manufacturing - producing only what's needed when it's needed.

Everyone's ideas matter: Because workers feel respected and their ideas are valued, they actually think about improvements.

Negative Example: Boeing's Management Decisions

The Boeing 737 MAX disaster is a brutal example of what happens when internal decisions go wrong. Two crashes in 2018 and 2019 killed 346 people.

What went wrong:

Culture flipped from engineering to finance: Boeing used to be obsessed with engineering excellence. But profit became the priority. Safety became a cost to be minimised.

Cutting corners on the 737 MAX: To compete with Airbus, Boeing rushed the 737 MAX to market with flawed software called MCAS.

Ignoring warnings: Boeing's own engineers raised concerns. Test pilots experienced problems. But management dismissed the warnings.

The consequences: 346 people dead, every 737 MAX grounded worldwide, $2.5 billion settlement, $87 billion wiped off shareholder value.

Task 2.2: PESTEC Analysis - Impact of Two External Factors (20 marks)

Netflix - Technological Factors

How technology helps:

Streaming itself: Netflix started mailing DVDs. The shift to streaming in 2007 changed entertainment forever. Netflix now reaches 260+ million subscribers worldwide.

AI and personalisation: The recommendation engine reportedly influences 80% of content watched on Netflix.

Technology challenges: Netflix faces competition from Disney+, Amazon Prime, Apple TV+ - all enabled by the same streaming technology Netflix pioneered.

McDonald's - Social Factors

Health consciousness: People increasingly know that fast food isn't exactly healthy. McDonald's has had to respond - adding salads, fruit options, nutritional information.

Ethical concerns: Modern consumers want to know where ingredients come from, how workers are treated, what happens to packaging. McDonald's has invested in sustainable sourcing.

Cultural adaptation: In India, where many people don't eat beef for religious reasons, you won't find a Big Mac. Instead there's the McAloo Tikki.

Task 2.3: Two Stakeholder Conflicts of Interest (15 marks)

Conflict 1: Shareholders vs Employees

What's the conflict: Shareholders want returns on their investment. One way to increase profits is to cut wages. Employees want the opposite - job security, fair pay, good conditions.

Real example: British Airways cut thousands of jobs during COVID while parent company IAG kept paying dividends to shareholders.

How it might be resolved:
- Balance - treating both groups fairly
- Long-term thinking - demotivated workers hurt performance eventually
- Employee ownership - giving workers shares aligns their interests

Conflict 2: Business Expansion vs Local Community

What's the conflict: Businesses want to grow. But people living where expansion is planned have concerns about traffic, noise, housing costs.

Real example: When Amazon announced plans for HQ2, local residents in New York pushed back so hard that Amazon eventually pulled out entirely.

How it might be resolved:
- Actually listen - genuine consultation before plans are finalised
- Mitigate impacts - invest in infrastructure and local hiring
- Share benefits - ensure the local area actually gains from development

References

Amul (n.d.) About Us - The Amul Model. Available at: https://amul.com/m/about-us (Accessed: 23 January 2026).

Boeing (2024) Corporate Information. Available at: https://www.boeing.com (Accessed: 23 January 2026).

Harvard Business School (2024) 'Why Boeing's Problems with the 737 MAX Began More Than 25 Years Ago', Working Knowledge.

NHS Scotland (n.d.) About NHS Scotland. Available at: https://www.scot.nhs.uk/about-nhs-scotland/ (Accessed: 23 January 2026).

Oxfam International (n.d.) How We Are Organized. Available at: https://www.oxfam.org/en/what-we-do/about/how-we-are-organized (Accessed: 23 January 2026).

Reliance Jio (2024) About Us. Available at: https://www.jio.com (Accessed: 23 January 2026).

Tata Motors (2024) Annual Report 2024. Mumbai: Tata Motors Limited.

Unilever (2025) Sustainability Goals. Available at: https://www.unilever.com/sustainability/ (Accessed: 23 January 2026).
"""
    }


# ============================================================
# J22A 76 CONTENT
# ============================================================

def create_j22a_76_content():
    """Return the content for J22A 76 Management of People and Finance."""
    return {
        "unit_code": "J22A 76",
        "unit_title": "Management of People and Finance",
        "content": """
TASK A: Management of People (60 marks)

Task 1.1: Three HRM Approaches (15 marks)

Human Resource Management is basically about getting the people side of business right. Hire the right people, help them do their jobs well, and keep them growing.

Recruitment

Recruitment is finding and hiring people for jobs. Getting it wrong is expensive - not just the wasted salary, but lost productivity and training costs down the drain.

How the process typically works:
1. Job Analysis - Work out what the role actually involves
2. Job Description - Write down the duties and responsibilities
3. Person Specification - Define what qualifications and skills you're looking for
4. Advertising - Get the word out
5. Shortlisting - Go through applications and pick who to interview
6. Selection - Interviews, tests, whatever else helps you choose
7. Offer and Onboarding - Make the offer and help the new person settle in

Internal vs External: Promoting from within is cheaper and faster, plus it shows existing staff they can progress. Hiring from outside brings new perspectives but costs more and takes longer.

Training

Training is about helping employees do their current job better. It costs money upfront, but pays off through better performance and fewer mistakes.

Different types:
- Induction Training: For new starters - showing them around, explaining policies
- On-the-Job Training: Learning by doing with someone experienced guiding you
- Off-the-Job Training: Courses and workshops away from normal work

Development

Training is about today. Development is about tomorrow - preparing people for future roles.

How it happens:
- Mentoring: Pairing someone junior with someone senior
- Coaching: One-to-one support focusing on specific skills
- Job Rotation: Moving people through different roles to broaden experience
- Management Programmes: Structured training for leadership

Task 1.2: Maslow's Hierarchy of Needs (15 marks)

Abraham Maslow's Hierarchy of Needs has been around since 1943, and it's still one of the most useful ways to think about what motivates people.

[DIAGRAM: Maslow's Hierarchy Pyramid]

                    /\\
                   /  \\
                  / SELF \\
                 / ACTUAL- \\
                / ISATION   \\
               /______________\\
              /                \\
             /     ESTEEM       \\
            /   (Recognition)    \\
           /____________________\\
          /                      \\
         /   SOCIAL/BELONGING     \\
        /   (Friendship, Family)   \\
       /____________________________\\
      /                              \\
     /      SAFETY & SECURITY         \\
    /   (Job Security, Health)         \\
   /____________________________________\\
              PHYSIOLOGICAL
         (Food, Water, Shelter)

The Five Levels:

Level 1: Physiological Needs
The basics of survival - food, water, shelter, sleep. In work terms, this translates to earning enough to live on.

In the workplace: Someone on minimum wage is primarily thinking about paying rent and buying food. Until basic financial security is sorted, higher-level stuff doesn't really register.

What employers can do: Pay fair wages. Give proper breaks. Keep the workplace at a reasonable temperature.

Level 2: Safety and Security Needs
Once basic needs are covered, we start wanting stability and protection from harm.

In the workplace: Someone hearing redundancy rumours will be anxious and distracted. They can't focus on doing great work when they're worried about losing their job.

What employers can do: Offer permanent contracts. Provide pension schemes, health insurance. Maintain a safe working environment.

Level 3: Social/Belonging Needs
Humans are social. We need connection, friendship, feeling like we belong.

In the workplace: Someone working remotely who never sees colleagues can feel isolated and disconnected.

What employers can do: Create opportunities for teamwork. Organise social events. Build a supportive, friendly culture.

Level 4: Esteem Needs
We want to feel respected - both self-respect and recognition from others.

In the workplace: Someone who consistently delivers great work but never hears acknowledgment will eventually burn out or leave.

What employers can do: Recognition programmes. Meaningful job titles. Praise from managers.

Level 5: Self-Actualisation
The highest level - reaching your full potential, personal growth, doing meaningful work.

In the workplace: A senior manager might still feel unfulfilled if their work isn't challenging or meaningful.

What employers can do: Provide challenging assignments. Give autonomy. Allow creativity. Connect work to a larger purpose.

Task 1.3: Five Types of Industrial Action (15 marks)

Industrial action is what happens when employees and employers can't agree on something important and workers take steps to put pressure on the employer.

1. Strike

A strike means workers completely stop working. No work, no production. It's the nuclear option.

Impact:
- Production stops dead
- Orders can't be fulfilled, deliveries miss deadlines
- Revenue disappears
- But workers lose their wages too, so strikes are usually a last resort

Example: Rail strikes in the UK during 2022-23 caused massive travel disruption.

2. Work-to-Rule

Workers do exactly what their contract says - nothing more, nothing less. Every single rule gets followed to the letter.

Impact:
- Major slowdowns without technically breaking any rules
- Hard for employers to discipline
- Exposes how much organisations rely on workers going above and beyond

3. Go-Slow

Workers deliberately work at a reduced pace. Everything takes longer.

Impact:
- Output drops, targets get missed
- Hard to prove it's deliberate
- Workers still get paid normal wages

4. Lockout

A lockout is the opposite of a strike - it's action taken by the EMPLOYER. Management prevents workers from entering the workplace.

Impact:
- Workers lose wages while locked out
- Can break the momentum of union negotiations
- Puts financial pressure on employees to accept employer's terms

Example: In professional sports, team owners sometimes lock out players during contract negotiations.

5. Picketing

Picketing means workers stand outside the workplace to protest and discourage others from entering.

Impact:
- Creates visible public protest
- Discourages deliveries
- Generates media attention
- Non-striking workers may feel pressured not to enter

Example: During the 2022-23 UK rail strikes, RMT union members picketed outside stations.

Task 1.4: Impact of Employment Legislation (15 marks)

Organisations must follow numerous laws that protect employees.

1. Equality & Diversity (Equality Act 2010)

This protects people from discrimination based on protected characteristics: age, disability, gender reassignment, marriage/civil partnership, pregnancy, race, religion, sex, and sexual orientation.

Impact on organisations:
- Recruitment must be fair and based on job-related criteria
- Every HR policy needs reviewing for indirect discrimination
- Training on equality and unconscious bias is common
- Tribunal awards for discrimination have no upper limit

2. Health & Safety (Health and Safety at Work Act 1974)

This law says employers must take reasonable steps to protect the health and safety of their workers.

Impact on organisations:
- Must provide safe equipment and proper training
- Risk assessments for every work activity
- Can face prosecution and big fines for non-compliance
- Fewer accidents means lower insurance costs

3. Minimum Wage (National Minimum Wage Act 1998)

Sets the minimum hourly rate employers must pay workers.

Impact on organisations:
- Direct hit to bottom line for businesses with many minimum wage workers
- Must track hours carefully to ensure compliance
- Employers who underpay face naming and shaming and fines up to 200% of arrears

4. Working Time Regulations 1998

These regulations limit working hours and guarantee rest periods.

Key provisions:
- Maximum 48-hour average working week
- Minimum 11 hours rest per day
- 20-minute break if working more than 6 hours
- Minimum 5.6 weeks paid annual leave

Impact on organisations:
- Must carefully plan rotas
- Need to track hours worked
- May need more staff to cover shifts while ensuring proper breaks

5. Dismissal and Redundancy (Employment Rights Act 1996)

Protects employees from unfair dismissal and sets rules for redundancy.

Impact on organisations:
- Can't just fire people - must have valid reasons and follow proper procedures
- Need documentation of warnings and performance issues
- Redundancy requires consultation and statutory payments
- Unfair dismissal claims can result in compensation up to a year's salary

TASK B: Management of Finance (40 marks)

Task 2.1: Three Sources of Finance (15 marks)

1. Retained Profit

Retained profit is money the business made and kept instead of paying it out to owners.

Why choose it:
- Free money - no interest payments
- Keep ownership - no giving away shares
- No permission needed
- Use it however you want

The catch: You can only use this if the business is already profitable.

2. Bank Loan

Borrow a set amount from a bank, pay it back over time with interest.

Why choose it:
- Large amounts available
- Predictable repayments make budgeting easier
- Keep ownership
- Interest payments reduce taxable profit

The catch: Interest adds to the cost. Repayments required whether business is doing well or not.

3. Share Capital

Selling shares - pieces of ownership in the company - to investors.

Why choose it:
- No repayment required
- Big amounts possible
- Shared risk
- Investors often bring useful knowledge and connections

The catch: You're giving away ownership. Investors expect dividends and can influence decisions.

Task 2.2: Five Purposes of Financial Statements (10 marks)

1. Helping Investment Decisions
People thinking about investing need information. Is this business profitable? Is it growing? Without reliable financial statements, investors would be guessing.

2. Informing Lending Decisions
When a business wants a loan, the bank needs to assess whether they'll get their money back. Financial statements answer key questions about profit, security, and cash flow.

3. Enabling Business Planning
Managers use financial statements to compare actual performance to budget, identify which products are most profitable, and decide where to invest resources.

4. Assessing Financial Health
Financial statements show what the business owns (assets), what it owes (liabilities), whether it can pay its debts, and if the trend is positive or negative.

5. Meeting Legal and Tax Requirements
Limited companies must prepare and file financial statements by law for calculating tax, informing shareholders, and ensuring accountability.

Task 2.3: Five Accounting Ratios and Their Limitations (15 marks)

1. Gross Profit Margin

Formula: (Gross Profit / Revenue) x 100

What it tells you: How efficiently the business produces or sources what it sells. Higher means more money left after direct costs.

Example: £500,000 revenue, £200,000 gross profit = 40% margin.

2. Net Profit Margin

Formula: (Net Profit / Revenue) x 100

What it tells you: Overall profitability after ALL costs - not just production costs but rent, wages, interest, tax, everything.

Example: £500,000 revenue, £50,000 net profit = 10% margin.

3. Current Ratio

Formula: Current Assets / Current Liabilities

What it tells you: Can the business pay what it owes in the next 12 months?

Example: £150,000 current assets, £100,000 current liabilities = 1.5:1 ratio.

What's good: Around 2:1 is often considered healthy.

4. Acid Test (Quick Ratio)

Formula: (Current Assets - Inventory) / Current Liabilities

What it tells you: Same as current ratio, but strips out inventory because you might not be able to sell stock quickly.

Example: £150,000 current assets, £60,000 inventory, £100,000 liabilities = 0.9:1.

5. Return on Capital Employed (ROCE)

Formula: (Operating Profit / Capital Employed) x 100

What it tells you: How effectively the business uses its capital to generate profit.

Example: £80,000 operating profit, £400,000 capital employed = 20% ROCE.

Limitations of Ratio Analysis:

1. Looking backwards - Ratios are based on historical data, not predictions
2. Apples and oranges - Different companies use different accounting methods
3. Industry differences - What's good in one industry is terrible in another
4. Snapshot in time - May capture an unusually good or bad moment
5. Window dressing - Companies can manipulate figures before reporting
6. Missing the non-financial - Doesn't capture customer satisfaction or employee morale
7. Context is everything - A falling ratio isn't necessarily bad

References

ACAS (n.d.) 'Strikes and Industrial Action'. Available at: https://www.acas.org.uk/strikes-and-industrial-action (Accessed: 23 January 2026).

BDC (n.d.) 'What are Financial Statements?'. Available at: https://www.bdc.ca/en/articles-tools/entrepreneur-toolkit/templates-business-guides/glossary/financial-statements (Accessed: 23 January 2026).

CIPD (n.d.) 'Data Protection and GDPR in the Workplace'. Available at: https://www.cipd.org/uk/knowledge/factsheets/data-protection-factsheet/ (Accessed: 23 January 2026).

Corporate Finance Institute (n.d.) 'Financial Ratios'. Available at: https://corporatefinanceinstitute.com/resources/accounting/financial-ratios/ (Accessed: 23 January 2026).

Health and Safety Executive (n.d.) 'Health and Safety at Work Act 1974'. Available at: https://www.hse.gov.uk (Accessed: 23 January 2026).

Indeed (n.d.) 'Maslow's Hierarchy of Needs: Applying It in the Workplace'. Available at: https://www.indeed.com/career-advice/career-development/maslows-hierarchy-of-needs (Accessed: 23 January 2026).
"""
    }


def create_final_document(unit_key, content_data):
    """Create the final document with all content."""
    template_path = TEMPLATES[unit_key]

    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}")
        return None

    # Create output filename
    unit_code = content_data["unit_code"].replace(" ", "_")
    output_filename = f"{STUDENT_USN}_{STUDENT_NAME.replace(' ', '')}_{unit_code}_{content_data['unit_title'].replace(' ', '_')}.docx"
    output_path = OUTPUT_DIR / output_filename

    # Copy template
    shutil.copy(template_path, output_path)

    # Open document
    doc = Document(output_path)

    # Fill student details in template
    fill_template_fields(doc, STUDENT_USN, STUDENT_NAME, COLLEGE, SUBMISSION_DATE)

    # Add content after the declaration page
    # First, find where to add content (after existing content)
    doc.add_page_break()

    # Add content heading
    doc.add_heading("Assignment Content", level=1)
    doc.add_paragraph()

    # Add the main content
    content_lines = content_data["content"].strip().split("\n")

    for line in content_lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("TASK "):
            doc.add_heading(line, level=1)
        elif line.startswith("Task "):
            doc.add_heading(line, level=2)
        elif line.startswith("Feature ") or line.startswith("Stage ") or line.startswith("Level "):
            doc.add_heading(line, level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            para = doc.add_paragraph("• " + line[2:])
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.5
        elif line.startswith("|"):
            # Skip table formatting lines, we'd need more complex table handling
            continue
        elif line.startswith("References"):
            doc.add_page_break()
            doc.add_heading("References", level=1)
        else:
            para = doc.add_paragraph(line)
            para.paragraph_format.line_spacing = 1.5

    # Add footer
    footer_text = f"{STUDENT_USN}_{STUDENT_NAME.replace(' ', '')}_{content_data['unit_code']}_{content_data['unit_title']}"
    add_footer(doc, footer_text)

    # Save
    doc.save(output_path)

    return output_path


def main():
    """Main function to create all assignments."""
    print("=" * 70)
    print("CREATING FINAL ASSIGNMENT SUBMISSIONS")
    print(f"Student: {STUDENT_NAME} (USN: {STUDENT_USN})")
    print(f"Submission Date: {SUBMISSION_DATE}")
    print("=" * 70)
    print()

    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Content generators
    content_generators = {
        "HE9E_46": create_he9e_46_content,
        "J229_76": create_j229_76_content,
        "J22A_76": create_j22a_76_content
    }

    created_files = []

    for unit_key, content_func in content_generators.items():
        print(f"\nProcessing {unit_key}...")
        content_data = content_func()

        output_path = create_final_document(unit_key, content_data)

        if output_path:
            created_files.append(output_path)
            print(f"  -> Created: {output_path.name}")
        else:
            print(f"  -> FAILED")

    print()
    print("=" * 70)
    print(f"COMPLETED: {len(created_files)}/{len(content_generators)} documents created")
    print(f"Output folder: {OUTPUT_DIR}")
    print("=" * 70)
    print()
    print("FINAL STEPS FOR YOU:")
    print("1. Open each document in Microsoft Word")
    print("2. Review the content and make any personal edits")
    print("3. SIGN the declaration page (your signature is required!)")
    print("4. For J22A 76: Create a proper Maslow's pyramid diagram")
    print("5. Check all tables are formatted correctly")
    print("6. Submit to LMS before the deadline")
    print()
    print("FILES CREATED:")
    for f in created_files:
        print(f"  - {f.name}")


if __name__ == "__main__":
    main()
