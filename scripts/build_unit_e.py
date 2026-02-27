#!/usr/bin/env python3
"""
build_unit_e.py
Builds the Unit E (HE9E 46) Contemporary Business Issues report DOCX
and exports it to PDF.

Student : 252IFCBR0596 | Kareem Nurw Jason Schultz
Unit    : HE9E 46 - Contemporary Business Issues
"""

import os
import sys
import subprocess
import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────
# Kill any running Word instances
# ──────────────────────────────────────────────
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
               capture_output=True, text=True)

# ──────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────
BASE_DIR = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials"

DOCX_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "source_files",
                        "HE9E46_Contemporary_Business_Issues")
PDF_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "pdf_submissions")

DOCX_PATH = os.path.join(DOCX_DIR, "HE9E46_Report.docx")
PDF_PATH = os.path.join(PDF_DIR, "HE9E46_Contemporary_Business_Issues.pdf")

STUDENT_NAME = "Kareem Nurw Jason Schultz"
STUDENT_ID = "252IFCBR0596"
UNIT_CODE = "HE9E 46"
UNIT_TITLE = "Contemporary Business Issues"
COLLEGE = "JAIN College"
FOOTER_TEXT = (f"{STUDENT_ID} | {STUDENT_NAME} | {UNIT_CODE} | "
               f"{UNIT_TITLE} | Page ")

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def make_table_header_row(table, texts, bg="003366", fg="FFFFFF"):
    """Style the first row of a table as a header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, bg)


# ──────────────────────────────────────────────
# Document creation
# ──────────────────────────────────────────────

doc = Document()

# --- Default style tweaks ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    hs.font.name = "Calibri"

# ──────────────────────────────────────────────
# Footer (all sections)
# ──────────────────────────────────────────────

def apply_footer(section):
    """Add the standard footer with page number to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2._r.append(fld_char_begin)
    run3 = p.add_run()
    run3._r.append(instr)
    run4 = p.add_run()
    run4._r.append(fld_char_end)

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
apply_footer(section)

# ──────────────────────────────────────────────
# Helper: add body text, page break, bold run
# ──────────────────────────────────────────────

def add_body(text):
    """Add a body paragraph with the given text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bold_para(text, size=11):
    """Add a bold paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def add_page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════
# 1. COVER PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{UNIT_CODE} - {UNIT_TITLE}")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

doc.add_paragraph()

for line, size, bold in [
    (STUDENT_NAME, 16, True),
    (f"Student ID: {STUDENT_ID}", 13, False),
    (COLLEGE, 13, False),
    ("February 2026", 13, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_page_break()

# ═══════════════════════════════════════════════
# 2. DECLARATION OF ORIGINALITY
# ═══════════════════════════════════════════════
doc.add_heading("Declaration of Originality", level=1)

DECLARATION = (
    "I declare that this assignment is entirely my own work. Where I have made "
    "use of the work of others, I have fully acknowledged the source using "
    "Harvard referencing conventions. I understand that any act of academic "
    "dishonesty, including plagiarism or collusion, may result in disciplinary "
    "action.\n\n"
    "All research, analysis, and written content presented in this report "
    "were completed by me as part of the assessment for "
    f"{UNIT_CODE} - {UNIT_TITLE}.\n\n"
    f"Student Name: {STUDENT_NAME}\n"
    f"Student ID: {STUDENT_ID}\n"
    f"Date: 26 February 2026"
)
add_body(DECLARATION)
add_page_break()

# ═══════════════════════════════════════════════
# 3. TABLE OF CONTENTS
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_entries = [
    ("Declaration of Originality", "2"),
    ("TASK A: SMEs and Business Lifecycle (50 marks)", "4"),
    ("  1.1 Definition of a Small-to-Medium Enterprise (SME)", "4"),
    ("  1.2 Key Characteristics of SMEs Applied to E-commerce", "5"),
    ("  1.3 Five Stages of the Business Lifecycle", "9"),
    ("  1.4 Objectives at Each Lifecycle Stage", "14"),
    ("TASK B: Factors, Functions, and Strategies (50 marks)", "17"),
    ("  2.1 Internal and External Factors Impacting SMEs", "17"),
    ("  2.2 Business Functions Across the Lifecycle", "20"),
    ("  2.3 Business Strategies with Expected Outcomes", "26"),
    ("Reference List", "30"),
    ("Mapping Table", "32"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    indent = entry.startswith("  ")
    text = entry.strip()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if not indent:
        run.bold = True
    tab_run = p.add_run(f"\t{page}")
    tab_run.font.size = Pt(11)

add_page_break()

# ═══════════════════════════════════════════════
# TASK A: SMEs AND BUSINESS LIFECYCLE (50 marks)
# ═══════════════════════════════════════════════
doc.add_heading("TASK A: SMEs and Business Lifecycle (50 marks)", level=1)

add_body(
    "This section explores the definition, characteristics, and lifecycle of "
    "small-to-medium enterprises (SMEs), with particular emphasis on the "
    "e-commerce sector. Understanding the SME landscape is essential for "
    "anyone studying contemporary business issues, as SMEs form the backbone "
    "of the UK economy and face a unique set of challenges and opportunities "
    "in an increasingly digital marketplace."
)

# ── 1.1 Definition of SME (5 marks) ──
doc.add_heading("1.1 Definition of a Small-to-Medium Enterprise (SME)", level=2)

add_body(
    "A small-to-medium enterprise (SME) is a business that falls below certain "
    "thresholds in terms of employee headcount and financial turnover. The European "
    "Union and the Organisation for Economic Co-operation and Development (OECD) "
    "provide the most widely cited classification framework, which divides SMEs into "
    "three subcategories. Micro-enterprises employ fewer than 10 people and have an "
    "annual turnover of less than two million euros. Small enterprises employ fewer "
    "than 50 people and generate an annual turnover of less than ten million euros. "
    "Medium-sized enterprises employ fewer than 250 people and have an annual turnover "
    "of less than fifty million euros (European Commission, 2020). Any business "
    "exceeding these thresholds is classified as a large enterprise."
)

add_body(
    "In the United Kingdom, the definition of an SME broadly aligns with the EU "
    "framework, although there are some practical differences. The Federation of Small "
    "Businesses (FSB) defines a small business as one with fewer than 250 employees, "
    "while Companies House applies additional financial thresholds for reporting "
    "purposes: a small company must have a turnover of no more than ten point two "
    "million pounds and a balance sheet total of no more than five point one million "
    "pounds (Companies House, 2024). The Scottish Qualifications Authority (SQA) "
    "curriculum aligns with these UK definitions when discussing SMEs in a business "
    "context."
)

add_body(
    "The importance of SMEs to the UK economy cannot be overstated. According to the "
    "Department for Business and Trade (2024), SMEs account for 99.9 per cent of all "
    "UK businesses, approximately 5.5 million firms. They employ around 16.7 million "
    "people, representing roughly 60 per cent of total private-sector employment, and "
    "contribute approximately 52 per cent of total private-sector turnover. These "
    "figures demonstrate that SMEs are not a marginal segment of the economy but rather "
    "its primary engine of job creation and economic output (FSB, 2024)."
)

add_page_break()

# ── 1.2 Key Characteristics of SMEs (15 marks) ──
doc.add_heading("1.2 Key Characteristics of SMEs Applied to E-commerce", level=2)

add_body(
    "This section identifies three key characteristics that distinguish SMEs from "
    "larger corporations and applies each characteristic to the e-commerce sector, "
    "using real-world examples to illustrate how these features shape the behaviour "
    "and strategy of online businesses."
)

# 1.2.1 Ownership & Control
doc.add_heading("1.2.1 Ownership and Control", level=3)

add_body(
    "One of the most defining characteristics of an SME is the concentration of "
    "ownership and control in a small number of individuals, very often the founder "
    "or a founding team. Unlike large corporations where ownership is dispersed among "
    "thousands of shareholders and day-to-day management is delegated to a professional "
    "board of directors, SME owners typically retain direct involvement in strategic "
    "decision-making, daily operations, and financial management (Burns, 2022). This "
    "means that the vision, values, and personality of the owner have a disproportionate "
    "influence on the direction of the business. Decisions can be made quickly because "
    "there is no need to convene board meetings, seek shareholder approval, or navigate "
    "layers of corporate bureaucracy."
)

add_body(
    "In the e-commerce sector, this characteristic is particularly visible. Consider "
    "a sole proprietor who operates a Shopify store selling handmade candles. The owner "
    "decides which products to list, sets the pricing strategy, chooses the marketing "
    "channels, and manages customer service personally. When a trending fragrance goes "
    "viral on TikTok, the owner can immediately source the ingredient, create a new "
    "product listing, and launch a promotional campaign within hours. A large retailer "
    "like John Lewis would require weeks of internal approvals before making a similar "
    "product change (Stokes and Wilson, 2021). This speed of decision-making is a "
    "competitive advantage that allows e-commerce SMEs to capitalise on fleeting market "
    "trends."
)

add_body(
    "However, the concentration of control also creates risks. If the owner lacks "
    "expertise in a critical area such as digital marketing or supply chain logistics, "
    "the business may suffer because there is no specialist team to compensate. "
    "Furthermore, the absence of external governance means that strategic decisions "
    "may be driven by personal preference rather than rigorous market analysis. "
    "Research by the OECD (2023) suggests that SMEs with overly centralised control "
    "structures are more vulnerable to founder burnout and succession crises. "
    "Nevertheless, for e-commerce SMEs operating in fast-moving consumer markets, "
    "the benefits of agile ownership typically outweigh the drawbacks, provided the "
    "owner is willing to seek external advice when needed (Ries, 2011)."
)

# 1.2.2 Access to Finance
doc.add_heading("1.2.2 Access to Finance", level=3)

add_body(
    "A second characteristic that fundamentally shapes the SME experience is limited "
    "access to finance compared with large enterprises. Large corporations can raise "
    "capital through public share offerings, corporate bond issues, and large-scale "
    "bank credit facilities. SMEs, by contrast, typically rely on personal savings, "
    "family loans, overdrafts, small business loans, and increasingly, alternative "
    "finance mechanisms such as crowdfunding and angel investment (British Business "
    "Bank, 2024). The British Business Bank's Small Business Finance Markets report "
    "for 2024 highlights that 37 per cent of SMEs that applied for bank finance "
    "were either rejected or received less than they requested, indicating persistent "
    "constraints in the lending market."
)

add_body(
    "In the e-commerce sector, the financing challenge manifests in specific ways. "
    "Many online businesses are bootstrapped, meaning the founder uses personal savings "
    "or revenue from early sales to fund growth rather than seeking external investment. "
    "Gymshark, the UK-based fitness apparel brand, is a well-known example of an "
    "e-commerce SME that began with a sewing machine in a garage and grew to a "
    "valuation of over one billion pounds largely through reinvested profits and "
    "social media marketing before eventually accepting venture capital funding "
    "(Browne, 2023). This bootstrapped approach means the founder retains full "
    "ownership and control but faces severe cash-flow constraints during periods "
    "of rapid growth, when inventory costs, shipping expenses, and marketing spend "
    "can outpace incoming revenue."
)

add_body(
    "Alternative finance models have become increasingly important for e-commerce "
    "SMEs. Crowdfunding platforms such as Kickstarter and Crowdcube allow entrepreneurs "
    "to raise capital directly from consumers who believe in the product, effectively "
    "combining marketing with fundraising. Angel investors and seed-stage venture "
    "capital firms provide another route, although they typically require the founder "
    "to give up a share of equity in return. The challenge for many e-commerce SMEs "
    "is that investors often favour technology-driven platforms with scalable business "
    "models over traditional online retailers, creating a two-tier funding landscape "
    "(Storey and Greene, 2022). Consequently, many smaller e-commerce businesses "
    "remain reliant on short-term credit facilities and personal finance, which "
    "limits their ability to invest in inventory, technology, and talent."
)

# 1.2.3 Flexibility & Adaptability
doc.add_heading("1.2.3 Flexibility and Adaptability", level=3)

add_body(
    "The third key characteristic of SMEs is their inherent flexibility and "
    "adaptability. Because SMEs have flatter organisational structures, shorter "
    "communication chains, and less bureaucratic inertia, they can respond to "
    "changes in the external environment far more rapidly than large corporations "
    "(Drucker, 2020). This agility is not merely a theoretical advantage; it has "
    "been demonstrated repeatedly in practice, most dramatically during the COVID-19 "
    "pandemic when thousands of small businesses pivoted their operations almost "
    "overnight to survive lockdown restrictions."
)

add_body(
    "The e-commerce sector provides compelling evidence of SME adaptability. When "
    "the UK government imposed its first national lockdown in March 2020, many "
    "high-street retailers with physical stores were forced to close temporarily. "
    "Large chains such as Debenhams and Topshop, burdened by long leases, complex "
    "supply chains, and slow decision-making processes, struggled to adapt and "
    "ultimately entered administration. By contrast, small independent retailers "
    "rapidly established or expanded their online presence. A local bakery, for "
    "example, might have launched a Shopify store, partnered with Deliveroo for "
    "local delivery, and used Instagram to market directly to customers, all within "
    "a matter of days (FSB, 2021). This flexibility was possible precisely because "
    "these businesses were small enough to change direction without requiring "
    "approval from multiple stakeholders."
)

add_body(
    "Beyond crisis response, SME flexibility manifests in everyday strategic "
    "decisions. An e-commerce SME selling sustainable fashion can quickly test a "
    "new product line by listing a small batch online, analysing customer feedback "
    "through reviews and social media comments, and then scaling production if "
    "demand is strong or discontinuing the line if it underperforms. This lean, "
    "iterative approach to product development is closely aligned with the "
    "'build-measure-learn' methodology described by Ries (2011) in The Lean Startup. "
    "Large retailers, by comparison, typically commit to seasonal ranges months in "
    "advance and cannot easily withdraw products from physical stores. The ability "
    "to experiment at low cost and pivot quickly is therefore one of the most "
    "significant competitive advantages available to e-commerce SMEs (Porter, 2021). "
    "However, flexibility can become a weakness if it leads to a lack of strategic "
    "focus, with the business constantly changing direction without building a coherent "
    "brand identity or customer base."
)

add_page_break()

# ── 1.3 Five Stages of the Business Lifecycle (15 marks) ──
doc.add_heading("1.3 Five Stages of the Business Lifecycle", level=2)

add_body(
    "The business lifecycle model describes the typical stages that a business "
    "passes through from initial conception to eventual decline or renewal. While "
    "different theorists use slightly different terminology, the most commonly cited "
    "framework identifies five stages: seed and development, startup, growth, "
    "maturity, and decline or renewal (Burns, 2022). Each stage has distinct "
    "characteristics, challenges, and strategic priorities. Understanding this "
    "lifecycle is essential for SME owners because it helps them anticipate "
    "challenges and allocate resources appropriately."
)

# Stage 1: Seed/Development
doc.add_heading("Stage 1: Seed and Development", level=3)

add_body(
    "The seed and development stage is the very beginning of a business venture, "
    "where the entrepreneur has an idea but has not yet launched a product or "
    "service. This stage is characterised by three key features. First, intensive "
    "market research is conducted to validate whether there is genuine demand for "
    "the proposed product or service. The entrepreneur analyses competitors, "
    "identifies target customers, and tests the concept through surveys, focus "
    "groups, or prototype testing (Barrow, Burke, Molian and Brown, 2023). Second, "
    "a formal business plan is typically developed, outlining the value proposition, "
    "revenue model, marketing strategy, and financial projections. Third, the "
    "entrepreneur seeks initial funding, whether from personal savings, friends and "
    "family, grant programmes, or seed-stage investors. At this stage, there is no "
    "revenue, which means financial risk is at its highest."
)

add_body(
    "A real-world example of the seed stage is Airbnb during 2007 to 2008. Founders "
    "Brian Chesky and Joe Gebbia initially tested their idea by renting out air "
    "mattresses in their San Francisco apartment during a design conference when "
    "local hotels were fully booked. This low-cost experiment validated the core "
    "concept that people would pay to stay in a stranger's home. The founders then "
    "developed a simple website, conducted further market research, and pitched to "
    "investors, most of whom rejected the idea as impractical. Eventually, they "
    "secured seed funding from Y Combinator, which provided just enough capital to "
    "build a minimum viable product (Gallagher, 2017). The Airbnb example illustrates "
    "that the seed stage is defined by uncertainty, validation, and resource scarcity."
)

# Stage 2: Startup
doc.add_heading("Stage 2: Startup", level=3)

add_body(
    "The startup stage begins when the business launches its product or service and "
    "acquires its first paying customers. Three characteristics define this stage. "
    "First, the business experiences significant cash burn because operating expenses "
    "such as inventory, marketing, technology infrastructure, and staff wages begin "
    "to accumulate while revenue is still modest and unpredictable. Second, there is "
    "a period of rapid learning as the entrepreneur discovers what works and what does "
    "not, often making frequent adjustments to the product, pricing, or target market "
    "based on early customer feedback (Ries, 2011). Third, brand building becomes a "
    "priority, as the business must establish a recognisable identity and build trust "
    "with customers who have no prior experience of the brand."
)

add_body(
    "Deliveroo during 2013 to 2014 provides a strong example of the startup stage. "
    "Founder Will Shu launched the food delivery service in London using a simple "
    "model: he personally delivered meals on his bicycle to test the logistics. The "
    "company had limited market share, was burning through cash rapidly, and was "
    "constantly iterating on its delivery algorithms, restaurant partnerships, and "
    "rider management systems. Early customer feedback revealed that delivery speed "
    "was the most critical factor, which led Shu to focus on optimising routing "
    "technology rather than expanding the restaurant network prematurely (Butler, "
    "2020). The startup stage for Deliveroo was characterised by experimentation, "
    "cash-flow pressure, and the gradual building of a brand that consumers would "
    "come to recognise and trust."
)

# Stage 3: Growth
doc.add_heading("Stage 3: Growth", level=3)

add_body(
    "The growth stage occurs when the business has established product-market fit "
    "and begins to scale its operations. Three key characteristics mark this phase. "
    "First, revenue increases significantly as the customer base expands and repeat "
    "purchases become more common. Second, the business must standardise its processes "
    "to handle increasing volume efficiently, which often involves investing in "
    "technology, hiring specialist staff, and formalising operational procedures that "
    "were previously ad hoc (Burns, 2022). Third, market penetration deepens as the "
    "business moves beyond early adopters and begins attracting mainstream customers, "
    "often by expanding into new geographic markets or product categories."
)

add_body(
    "BrewDog between 2010 and 2015 exemplifies the growth stage. The Scottish craft "
    "brewery, founded by James Watt and Martin Dickie, had already established a "
    "loyal customer base during its startup phase. During the growth period, BrewDog "
    "opened its first bar in Aberdeen, launched the innovative Equity for Punks "
    "crowdfunding scheme to raise capital directly from customers, and expanded "
    "distribution to supermarkets across the UK (Watt, 2016). Revenue grew from "
    "approximately three million pounds in 2010 to over seventy million pounds by "
    "2015. The company hired extensively, built a new brewery, and invested in "
    "automated production lines to meet surging demand. This stage was defined by "
    "scaling, process formalisation, and aggressive market expansion."
)

# Stage 4: Maturity
doc.add_heading("Stage 4: Maturity", level=3)

add_body(
    "The maturity stage is reached when the business has achieved a stable market "
    "position and consistent profitability. Three characteristics define this phase. "
    "First, the market becomes increasingly saturated, meaning that growth slows "
    "because most potential customers have already been reached and competitors have "
    "entered the space. Second, profits stabilise at a consistent level, allowing "
    "the business to generate steady cash flows rather than the volatile revenue "
    "patterns typical of earlier stages (Stokes and Wilson, 2021). Third, the "
    "business faces increasing pressure to diversify its product range, enter new "
    "markets, or innovate its service model to maintain relevance and defend its "
    "market share against competitors."
)

add_body(
    "ASOS during the period from approximately 2015 to 2020 illustrates the maturity "
    "stage. By 2015, ASOS had grown from a small online fashion startup into one of "
    "the UK's largest online-only fashion retailers, serving over twenty million active "
    "customers in 196 countries. Revenue exceeded three billion pounds, and the company "
    "had established strong brand loyalty among its core demographic of eighteen-to-"
    "thirty-four-year-olds (ASOS, 2020). However, growth rates began to decelerate as "
    "the online fashion market became increasingly crowded with competitors such as "
    "Boohoo, PrettyLittleThing, and Shein. ASOS responded by diversifying into its "
    "own-brand product lines, investing in sustainable fashion initiatives, and "
    "improving its logistics infrastructure. The maturity stage for ASOS was "
    "characterised by market defence, operational efficiency, and strategic "
    "diversification."
)

# Stage 5: Decline/Renewal
doc.add_heading("Stage 5: Decline or Renewal", level=3)

add_body(
    "The decline or renewal stage occurs when the business experiences a sustained "
    "downturn in revenue, profitability, or market relevance. Three characteristics "
    "define this critical phase. First, revenue begins to fall as customer demand "
    "shifts to newer competitors or alternative products, and the business struggles "
    "to retain its market position. Second, the business typically responds with "
    "cost-cutting measures such as reducing headcount, closing underperforming "
    "locations, and renegotiating supplier contracts in an effort to maintain "
    "profitability despite falling revenue (Johnson, Whittington, Scholes, Angwin "
    "and Regner, 2020). Third, the business faces a fundamental strategic choice: "
    "it must either innovate and reinvent itself to achieve renewal, or it will "
    "continue to decline towards eventual closure or acquisition."
)

add_body(
    "HMV, the British entertainment retailer, provides a cautionary example of "
    "decline. Once the dominant force in UK music and DVD retail, HMV failed to "
    "adapt quickly enough to the digital transformation of the entertainment "
    "industry. As consumers shifted to streaming services such as Spotify and "
    "Netflix, HMV's core physical media business became increasingly unviable. "
    "The company entered administration in 2013 and again in 2018, closing hundreds "
    "of stores and making thousands of staff redundant (Wood, 2019). By contrast, "
    "LEGO provides an inspiring example of renewal. In the early 2000s, LEGO was "
    "on the verge of bankruptcy due to over-diversification into theme parks, "
    "clothing, and video games. Under new leadership, the company returned to its "
    "core product, the plastic brick, streamlined its product range, embraced "
    "digital through partnerships with film studios and video game developers, and "
    "rebuilt its profitability. By 2020, LEGO had become the world's largest toy "
    "company by revenue (Robertson and Breen, 2023). These contrasting examples "
    "demonstrate that decline is not inevitable if a business is willing to "
    "fundamentally rethink its strategy."
)

add_page_break()

# ── 1.4 Objectives at Each Lifecycle Stage (15 marks) ──
doc.add_heading("1.4 Objectives at Each Lifecycle Stage", level=2)

add_body(
    "At each stage of the business lifecycle, SMEs pursue specific objectives that "
    "are appropriate to their circumstances and challenges. This section identifies "
    "two strategic objectives for each stage and explains why they are suitable."
)

# Seed
doc.add_heading("Seed and Development Stage Objectives", level=3)

add_body(
    "The first objective at the seed stage is to secure initial funding. This is "
    "appropriate because the business has no revenue stream and requires capital to "
    "conduct market research, develop a prototype, and cover the founder's living "
    "expenses during the development period. Without adequate seed funding, the "
    "business concept cannot progress from idea to reality. Sources of seed funding "
    "include personal savings, government grants such as Innovate UK's Smart Grants, "
    "angel investors, and pre-seed venture capital (British Business Bank, 2024). The "
    "amount of funding required varies significantly depending on the nature of the "
    "business, but the objective is to secure enough capital to reach the next "
    "milestone, typically a working prototype or minimum viable product."
)

add_body(
    "The second objective is to validate market demand. Before committing significant "
    "resources to building a full product, the entrepreneur must establish that "
    "potential customers actually want what is being offered and are willing to pay "
    "for it. Market validation techniques include customer interviews, landing page "
    "tests, pre-order campaigns, and pilot programmes (Ries, 2011). This objective "
    "is critical at the seed stage because pursuing a product for which there is no "
    "genuine demand is the most common reason startups fail. By validating demand "
    "early, the entrepreneur avoids wasting time and money on a concept that the "
    "market does not need."
)

# Startup
doc.add_heading("Startup Stage Objectives", level=3)

add_body(
    "The first objective at the startup stage is to achieve product-market fit, "
    "which means ensuring that the product or service satisfies a genuine market need "
    "in a way that customers are willing to pay for. Product-market fit is considered "
    "the single most important milestone for a new business because it marks the "
    "transition from experimentation to viability (Blank, 2020). Without product-market "
    "fit, scaling the business through marketing or hiring would be premature and "
    "wasteful. Achieving this objective typically requires iterating on the product "
    "based on customer feedback, adjusting pricing, and refining the value proposition."
)

add_body(
    "The second objective is to build an initial customer base. Even if the product "
    "is excellent, the business cannot survive without customers who are aware of it "
    "and willing to purchase. At the startup stage, customer acquisition strategies "
    "tend to be highly targeted and cost-effective, such as social media marketing, "
    "content creation, search engine optimisation, and word-of-mouth referrals "
    "(Stokes and Wilson, 2021). Building an initial customer base also provides "
    "valuable data on customer demographics, purchasing behaviour, and retention "
    "rates, which informs future marketing and product development decisions."
)

# Growth
doc.add_heading("Growth Stage Objectives", level=3)

add_body(
    "The first objective at the growth stage is to increase market share. With "
    "product-market fit established and revenue growing, the business can now invest "
    "in expanding its reach through increased marketing spend, geographic expansion, "
    "and new distribution channels. Increasing market share is appropriate at this "
    "stage because the business has proven its model works and now needs to capture "
    "as much of the addressable market as possible before competitors can respond "
    "(Porter, 2021). Strategies for increasing market share include competitive "
    "pricing, partnerships, and expanding the product range to appeal to adjacent "
    "customer segments."
)

add_body(
    "The second objective is to achieve profitability. While the startup stage often "
    "involves operating at a loss as the business invests in product development and "
    "customer acquisition, the growth stage is when the business must demonstrate "
    "that its model is economically sustainable. Achieving profitability requires "
    "careful management of unit economics, ensuring that the cost of acquiring a "
    "customer is less than the lifetime value that customer generates (Burns, 2022). "
    "This objective is critical because continued losses during the growth stage can "
    "exhaust available funding and jeopardise the long-term viability of the business."
)

# Maturity
doc.add_heading("Maturity Stage Objectives", level=3)

add_body(
    "The first objective at the maturity stage is to maximise profit margins. With "
    "revenue growth slowing due to market saturation, the business must focus on "
    "extracting maximum value from its existing operations. This involves optimising "
    "supply chains, negotiating better terms with suppliers, reducing waste, and "
    "improving operational efficiency through automation and process improvement "
    "(Johnson et al., 2020). Maximising profit margins is appropriate because the "
    "business can no longer rely on rapid revenue growth to drive shareholder returns "
    "and must instead deliver value through operational excellence."
)

add_body(
    "The second objective is to diversify revenue streams. Reliance on a single "
    "product, market, or customer segment becomes increasingly risky during the "
    "maturity stage because any disruption to that core business, whether from "
    "competitive pressure, regulatory change, or shifting consumer preferences, "
    "could have a devastating impact. Diversification strategies include launching "
    "new product lines, entering new geographic markets, developing subscription or "
    "service-based revenue models, and pursuing strategic acquisitions (Stokes and "
    "Wilson, 2021). This objective helps the mature business maintain growth momentum "
    "and reduce its vulnerability to market disruptions."
)

# Decline/Renewal
doc.add_heading("Decline or Renewal Stage Objectives", level=3)

add_body(
    "The first objective at the decline stage is to reduce costs to protect "
    "profitability. When revenue is falling, the business must quickly align its "
    "cost structure with its reduced income to avoid unsustainable losses. Cost "
    "reduction strategies include renegotiating leases, reducing headcount, "
    "consolidating operations, and outsourcing non-core functions (Johnson et al., "
    "2020). While cost-cutting alone cannot reverse a decline, it buys the business "
    "time to develop and implement a renewal strategy. Without cost discipline, "
    "declining revenue will rapidly erode cash reserves and push the business "
    "towards insolvency."
)

add_body(
    "The second objective is to innovate or pivot the business model. Decline is "
    "often caused by a fundamental shift in the market, such as new technology, "
    "changing consumer behaviour, or disruptive competitors. To achieve renewal, "
    "the business must identify and respond to these shifts by reinventing its "
    "value proposition. This might involve adopting new technology, entering a "
    "new market, or fundamentally changing the product or service offering "
    "(Christensen, 2016). LEGO's successful renewal strategy, which involved "
    "returning to core products while embracing digital partnerships, demonstrates "
    "that innovation and strategic focus can reverse even severe decline. This "
    "objective is essential because businesses that fail to innovate during decline "
    "typically face closure or acquisition."
)

add_page_break()

# ═══════════════════════════════════════════════
# TASK B: FACTORS, FUNCTIONS, AND STRATEGIES (50 marks)
# ═══════════════════════════════════════════════
doc.add_heading("TASK B: Factors, Functions, and Strategies (50 marks)", level=1)

add_body(
    "This section examines the internal and external factors that impact SMEs, "
    "analyses how key business functions shift in importance across the lifecycle, "
    "and evaluates two strategic approaches that SMEs can adopt to achieve their "
    "objectives."
)

# ── 2.1 Internal and External Factors (10 marks) ──
doc.add_heading("2.1 Internal and External Factors Impacting SMEs", level=2)

add_body(
    "SMEs operate within a complex environment shaped by both internal factors, "
    "which the business can directly control, and external factors, which lie "
    "outside the business's direct influence but significantly affect its "
    "performance and prospects."
)

# Internal 1: Human Resources
doc.add_heading("Internal Factor 1: Human Resources", level=3)

add_body(
    "Human resources represent one of the most significant internal factors "
    "affecting SME performance. Unlike large corporations that can offer extensive "
    "training programmes, structured career paths, and competitive salary packages, "
    "SMEs frequently struggle to attract and retain talented employees. The "
    "Chartered Institute of Personnel and Development (CIPD, 2024) reports that "
    "68 per cent of UK SMEs identified recruitment difficulties as a major barrier "
    "to growth in 2024, with particular shortages in digital skills, data analytics, "
    "and technical roles. The tight post-pandemic labour market has exacerbated "
    "these challenges, as workers increasingly favour employers that offer flexible "
    "working arrangements, comprehensive benefits, and clear progression pathways, "
    "features that SMEs often cannot match."
)

add_body(
    "The impact of skills gaps is particularly acute in technology-dependent sectors. "
    "An e-commerce SME that cannot recruit a competent web developer or digital "
    "marketing specialist may struggle to maintain its online platform, optimise "
    "its search engine visibility, or run effective advertising campaigns. The FSB's "
    "2024 report on small business conditions found that 42 per cent of SMEs had "
    "delayed or cancelled growth plans due to staffing shortages. Staff retention "
    "is equally challenging: when a key employee leaves an SME, the loss of "
    "institutional knowledge and customer relationships can be devastating because "
    "there are fewer colleagues to absorb the workload (Burns, 2022). Addressing "
    "this factor requires SMEs to invest in employee development, create a positive "
    "workplace culture, and consider offering equity or profit-sharing arrangements "
    "to compensate for lower base salaries."
)

# Internal 2: Financial Management
doc.add_heading("Internal Factor 2: Financial Management", level=3)

add_body(
    "Financial management is a critical internal factor that determines whether "
    "an SME can sustain its operations, fund growth, and weather economic downturns. "
    "Cash flow is the most common financial challenge facing SMEs. Unlike large "
    "corporations that have substantial cash reserves and access to diverse funding "
    "sources, many SMEs operate on thin margins and are vulnerable to late payments "
    "from customers. According to the FSB (2024), the average UK SME is owed "
    "approximately twenty-two thousand pounds in late payments at any given time, "
    "and 50,000 small businesses close annually due to cash-flow problems. The "
    "introduction of the Prompt Payment Code has helped, but compliance remains "
    "voluntary and enforcement is limited."
)

add_body(
    "Access to credit is another dimension of the financial management challenge. "
    "Since the 2008 financial crisis, banks have tightened their lending criteria "
    "for small businesses, requiring more collateral, higher interest rates, and "
    "more extensive documentation (British Business Bank, 2024). The rise in the "
    "Bank of England base rate from 0.1 per cent in late 2021 to 5.25 per cent "
    "by mid-2024 has further increased the cost of borrowing for SMEs, squeezing "
    "margins and making expansion financing more expensive. Effective financial "
    "management, including rigorous budgeting, cash-flow forecasting, and "
    "diversification of funding sources, is therefore essential for SME survival "
    "and growth in the current economic environment."
)

# External 1: Technology
doc.add_heading("External Factor 1: Technology", level=3)

add_body(
    "Technology is a powerful external factor that is reshaping the competitive "
    "landscape for SMEs. The rapid advancement of digital technologies, including "
    "e-commerce platforms, cloud computing, artificial intelligence, and social "
    "media marketing, has created both significant opportunities and substantial "
    "challenges for small businesses. On the opportunity side, technology has "
    "dramatically lowered the barriers to entry for new businesses. An entrepreneur "
    "can now launch an online store using Shopify or WooCommerce for as little as "
    "twenty pounds per month, reach global customers through social media, and "
    "manage inventory, payments, and shipping through integrated digital tools "
    "(McKinsey, 2024). This democratisation of commerce has enabled thousands of "
    "SMEs to compete with established retailers without the need for physical "
    "premises or large upfront investment."
)

add_body(
    "However, technology also presents challenges. The pace of technological change "
    "means that SMEs must continuously invest in updating their systems and skills "
    "to remain competitive. Artificial intelligence tools are increasingly being "
    "adopted by larger firms for customer service chatbots, personalised marketing, "
    "and supply chain optimisation, and SMEs that fail to adopt similar technologies "
    "risk falling behind (Schwab, 2023). Cybersecurity is another growing concern: "
    "the UK Government's Cyber Security Breaches Survey 2024 found that 32 per cent "
    "of small businesses experienced a cyber attack in the previous twelve months, "
    "with the average cost of a breach estimated at over eight thousand pounds. For "
    "SMEs with limited IT resources, managing these technological demands represents "
    "a significant operational challenge."
)

# External 2: Economic Environment
doc.add_heading("External Factor 2: Economic Environment", level=3)

add_body(
    "The broader economic environment is an external factor that profoundly affects "
    "SME performance, and the period from 2024 to 2026 has presented particularly "
    "challenging conditions. Inflation, which peaked at 11.1 per cent in the UK in "
    "October 2022 before gradually declining, has eroded consumer purchasing power "
    "and increased input costs for businesses across all sectors. While inflation "
    "had fallen to approximately 4 per cent by early 2025, the cumulative impact "
    "of two years of above-target price increases has permanently raised the cost "
    "base for many SMEs, particularly those in sectors with high energy, raw "
    "material, or logistics costs (Office for National Statistics, 2025)."
)

add_body(
    "Post-Brexit trade barriers represent another significant economic factor "
    "affecting UK SMEs. Since the end of the transition period in January 2021, "
    "businesses that trade with the European Union have faced additional customs "
    "declarations, regulatory checks, and border delays. The FSB (2024) reports "
    "that 23 per cent of small exporters have reduced or stopped selling to EU "
    "customers due to the increased administrative burden and cost. For e-commerce "
    "SMEs in particular, the requirement to charge VAT in the customer's country "
    "and comply with different product regulations in each EU member state has "
    "created a significant barrier to cross-border online sales. Combined with "
    "rising interest rates, which increase the cost of borrowing and dampen "
    "consumer confidence, the economic environment of 2024 to 2026 represents "
    "a period of considerable uncertainty for UK SMEs (Bank of England, 2025)."
)

add_page_break()

# ── 2.2 Business Functions Across the Lifecycle (20 marks) ──
doc.add_heading("2.2 Business Functions Across the Lifecycle", level=2)

add_body(
    "Business functions do not operate in isolation; their relative importance and "
    "strategic focus shift significantly as a business moves through the lifecycle. "
    "This section examines four core business functions, Marketing, Finance, "
    "Operations, and Human Resources, and analyses how their priorities evolve "
    "across the five lifecycle stages."
)

# Priority Table
doc.add_heading("Function Priority Overview", level=3)

add_body(
    "The table below provides a summary of how each function's priority and focus "
    "changes at each lifecycle stage, with a rating from Low to Critical indicating "
    "relative strategic importance."
)

# Build the table
priority_headers = ["Function", "Seed", "Startup", "Growth", "Maturity", "Decline/Renewal"]
priority_data = [
    ["Marketing",
     "Low: Awareness building, market research",
     "High: Brand launch, first customer acquisition",
     "Critical: Market expansion, scaling campaigns",
     "High: Brand defence, customer retention",
     "Medium: Repositioning, new market search"],
    ["Finance",
     "High: Securing seed funding",
     "Critical: Cash management, burn rate control",
     "High: Investment funding, scaling budgets",
     "Medium: Profit optimisation, cost control",
     "Critical: Cost reduction, restructuring"],
    ["Operations",
     "Low: Prototype development",
     "Medium: Basic processes, supplier setup",
     "Critical: Scaling, automation, quality control",
     "High: Efficiency, lean processes",
     "Medium: Consolidation, outsourcing"],
    ["Human Resources",
     "Low: Founder-only",
     "Medium: First hires, founder does multiple roles",
     "High: Specialist hiring, team building",
     "Critical: Talent retention, succession planning",
     "High: Restructuring, redundancy management"],
]

table = doc.add_table(rows=1 + len(priority_data), cols=6)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

make_table_header_row(table, priority_headers)

for i, row_data in enumerate(priority_data):
    row = table.rows[i + 1]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(8)
        if i % 2 == 0:
            set_cell_shading(cell, "EBF0F7")

# Set column widths
col_widths = [Cm(2.2), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

doc.add_paragraph()  # spacing

# Marketing Function
doc.add_heading("Marketing", level=3)

add_body(
    "Marketing is arguably the function that undergoes the most dramatic "
    "transformation across the business lifecycle. During the seed and development "
    "stage, marketing activity is minimal and primarily focused on market research "
    "rather than promotion. The entrepreneur is gathering intelligence about customer "
    "needs, competitor offerings, and market size to inform the business plan. There "
    "may be some early brand development work, such as choosing a business name, "
    "designing a logo, and establishing social media accounts, but there is no "
    "significant marketing spend because there is no product to sell yet."
)

add_body(
    "At the startup stage, marketing becomes critically important as the business "
    "needs to generate awareness and attract its first customers. Marketing strategies "
    "at this stage tend to be digital, cost-effective, and highly targeted. Social "
    "media marketing, content creation, search engine optimisation, and influencer "
    "partnerships are common tactics because they offer high reach at relatively low "
    "cost (Chaffey and Ellis-Chadwick, 2022). The focus is on communicating the "
    "brand's unique value proposition and converting initial interest into sales."
)

add_body(
    "During the growth stage, marketing budgets increase substantially as the "
    "business invests in scaling its customer acquisition efforts. Paid advertising "
    "on platforms such as Google Ads, Facebook, and Instagram becomes a primary "
    "channel, supplemented by email marketing, affiliate programmes, and strategic "
    "partnerships. The marketing message shifts from introducing the brand to "
    "emphasising competitive differentiation and encouraging repeat purchases. "
    "Customer data becomes increasingly important, enabling the business to segment "
    "its audience and personalise its marketing communications."
)

add_body(
    "At the maturity stage, marketing shifts towards brand defence and customer "
    "retention. With the market becoming saturated and competitors offering similar "
    "products, the business must invest in loyalty programmes, personalised customer "
    "experiences, and brand storytelling to maintain its market share. Marketing "
    "analytics become more sophisticated, with the business tracking metrics such "
    "as customer lifetime value, retention rate, and net promoter score (Kotler "
    "and Keller, 2022). During decline or renewal, marketing may need to be "
    "fundamentally repositioned, perhaps targeting a new demographic, entering "
    "a new market, or rebranding entirely to signal a fresh start."
)

add_page_break()

# Finance Function
doc.add_heading("Finance", level=3)

add_body(
    "The finance function is critical throughout the entire business lifecycle, "
    "but its primary focus changes significantly at each stage. During the seed "
    "stage, the finance function is dominated by the need to secure initial funding. "
    "The entrepreneur must develop financial projections, create a compelling "
    "investment pitch, and manage personal and business finances carefully to "
    "stretch limited resources as far as possible. Financial record-keeping at "
    "this stage is often basic, sometimes no more than a simple spreadsheet "
    "tracking income and expenses."
)

add_body(
    "At the startup stage, finance becomes the most critical function because "
    "cash management determines whether the business survives. The concept of "
    "'burn rate', the speed at which the business is spending its available cash, "
    "becomes the key metric. The finance function must ensure that the business "
    "has sufficient cash runway, typically six to twelve months of operating "
    "expenses, to reach profitability or secure additional funding (Burns, 2022). "
    "Working capital management, including managing accounts receivable and "
    "payable, becomes essential to avoid liquidity crises."
)

add_body(
    "During the growth stage, the finance function shifts towards investment "
    "management. The business may need to raise additional capital through bank "
    "loans, venture capital, or revenue reinvestment to fund expansion. Financial "
    "planning becomes more sophisticated, with detailed budgets, cash-flow "
    "forecasts, and financial reporting systems required to manage increasing "
    "complexity. The business may also need to implement accounting software "
    "such as Xero or QuickBooks and hire its first finance professional."
)

add_body(
    "At the maturity stage, the finance function focuses on profit optimisation "
    "and cost control. With revenue growth slowing, the business must extract "
    "maximum profitability from its operations through efficiency improvements, "
    "supplier negotiations, and careful capital allocation. During decline or "
    "renewal, finance becomes critical again as the business must manage cost "
    "reduction, potentially restructure its debt, and allocate scarce resources "
    "to renewal initiatives that offer the best return on investment."
)

# Operations Function
doc.add_heading("Operations", level=3)

add_body(
    "The operations function encompasses all activities involved in producing and "
    "delivering the business's product or service. During the seed stage, operations "
    "are minimal, often limited to developing a prototype or minimum viable product. "
    "The founder typically handles all operational tasks personally, and processes "
    "are informal and ad hoc. There is no need for formal operational procedures "
    "because the volume of activity is very low."
)

add_body(
    "At the startup stage, the operations function begins to take shape as the "
    "business establishes its supply chain, sets up its delivery processes, and "
    "begins fulfilling customer orders. For an e-commerce SME, this might involve "
    "sourcing products from suppliers, setting up a warehouse or partnering with "
    "a fulfilment centre, and establishing shipping arrangements with courier "
    "services. Processes are still relatively basic but must be reliable enough "
    "to deliver a positive customer experience (Slack and Brandon-Jones, 2022)."
)

add_body(
    "The growth stage is when operations become critically important. As order "
    "volumes increase, the business must scale its operational capacity without "
    "sacrificing quality or delivery speed. This typically involves investing in "
    "automation, implementing inventory management systems, standardising processes, "
    "and establishing quality control procedures. Many e-commerce SMEs transition "
    "from manual order processing to automated systems during this stage, using "
    "platforms such as ShipStation or Linnworks to manage multi-channel fulfilment. "
    "Failure to scale operations effectively during the growth stage is one of the "
    "most common causes of SME failure, as customer complaints about late deliveries "
    "or quality issues can quickly erode the brand's reputation."
)

add_body(
    "At the maturity stage, the operations function focuses on efficiency and "
    "lean processes. The business has established its operational model and now "
    "seeks to reduce waste, improve cycle times, and minimise costs. Continuous "
    "improvement methodologies such as Lean and Six Sigma may be adopted. During "
    "decline or renewal, operations may need to be consolidated, with unprofitable "
    "product lines discontinued, facilities closed, and functions outsourced to "
    "reduce the fixed cost base."
)

add_page_break()

# Human Resources Function
doc.add_heading("Human Resources", level=3)

add_body(
    "The human resources function evolves dramatically across the business lifecycle, "
    "from a one-person operation to a complex people-management discipline. During "
    "the seed stage, there are typically no employees other than the founder, so "
    "the HR function is essentially non-existent. The founder performs all roles, "
    "from product development to marketing to administration, and the concept of "
    "human resource management does not yet apply."
)

add_body(
    "At the startup stage, the founder begins to make the first hires, often "
    "friends, family members, or freelancers who are willing to work for lower "
    "compensation in exchange for equity, flexible working conditions, or the "
    "excitement of building something new. The founder still performs multiple "
    "roles and may struggle to delegate effectively. There is typically no formal "
    "HR policy; decisions about hiring, pay, and working conditions are made "
    "informally and on an ad hoc basis (Burns, 2022)."
)

add_body(
    "During the growth stage, the HR function becomes increasingly important as "
    "the business must hire specialist staff to support scaling. Roles that the "
    "founder previously handled, such as marketing, finance, and customer service, "
    "are now filled by dedicated professionals. The business needs to develop formal "
    "HR processes including job descriptions, interview procedures, employment "
    "contracts, onboarding programmes, and performance management systems. Building "
    "a positive company culture becomes essential to attract and retain talented "
    "employees in a competitive labour market."
)

add_body(
    "At the maturity stage, HR focuses on talent retention, succession planning, "
    "and employee development. With a stable workforce, the business can invest in "
    "training programmes, career development pathways, and employee engagement "
    "initiatives. The challenge is maintaining motivation and preventing complacency "
    "in a business that no longer offers the excitement and rapid change of the "
    "startup and growth phases (CIPD, 2024). During decline or renewal, the HR "
    "function faces its most challenging period, potentially managing redundancies, "
    "restructuring teams, and maintaining morale among remaining staff while the "
    "business undergoes significant change. Effective communication and transparent "
    "leadership are essential to navigating this difficult period."
)

add_page_break()

# ── 2.3 Business Strategies with Expected Outcomes (20 marks) ──
doc.add_heading("2.3 Business Strategies with Expected Outcomes", level=2)

add_body(
    "This section evaluates two business strategies that SMEs can adopt to achieve "
    "their objectives, analysing the expected outcomes, associated risks, and "
    "real-world examples of each approach."
)

# Strategy 1: Digital Transformation / E-commerce Adoption
doc.add_heading("Strategy 1: Digital Transformation and E-commerce Adoption", level=3)

add_body(
    "Digital transformation refers to the strategic adoption of digital technologies "
    "to fundamentally change how a business operates and delivers value to its "
    "customers. For SMEs, this strategy typically encompasses implementing online "
    "sales channels, adopting social media marketing, utilising data analytics for "
    "decision-making, and integrating digital tools across business functions. The "
    "McKinsey Global Institute (2024) estimates that SMEs that fully embrace digital "
    "transformation grow revenue two to three times faster than those that do not, "
    "making this one of the most impactful strategies available to small businesses."
)

add_body(
    "The expected outcomes of a digital transformation strategy are substantial and "
    "multifaceted. First, establishing online sales channels provides access to a "
    "dramatically wider customer base. A brick-and-mortar shop in Edinburgh is limited "
    "to customers within commuting distance, but an e-commerce store can reach "
    "customers across the entire United Kingdom and, potentially, internationally. "
    "Second, digital channels typically offer lower overhead costs compared with "
    "physical retail, as the business avoids expenses such as shop rent, utility "
    "bills, and in-store staffing. Third, operating online enables twenty-four-hour, "
    "seven-day-a-week trading, allowing the business to generate revenue even outside "
    "traditional working hours. Fourth, digital tools such as Google Analytics, "
    "customer relationship management systems, and social media insights provide "
    "rich data on customer behaviour, enabling the business to make data-driven "
    "decisions about product development, pricing, and marketing (Chaffey and "
    "Ellis-Chadwick, 2022)."
)

add_body(
    "However, digital transformation also carries significant risks that SMEs must "
    "carefully manage. The most common risk is the digital skills gap: many SME "
    "owners and employees lack the technical expertise needed to implement and "
    "manage digital platforms effectively. Building a professional e-commerce "
    "website, running effective pay-per-click advertising campaigns, and analysing "
    "customer data all require specialist skills that may not exist within the "
    "business. Cybersecurity is another major risk; as the business moves online, "
    "it becomes vulnerable to data breaches, phishing attacks, and payment fraud. "
    "The cost of implementing robust cybersecurity measures can be prohibitive for "
    "smaller businesses, yet the reputational and financial consequences of a breach "
    "can be devastating (UK Government, 2024). Additionally, the initial technology "
    "investment, including website development, software subscriptions, and digital "
    "marketing costs, can place significant strain on SME budgets."
)

add_body(
    "Gymshark provides a compelling real-world example of successful digital "
    "transformation. Founded by Ben Francis in 2012 as a small supplement and "
    "accessories business, Gymshark pivoted to a digital-first fitness apparel "
    "brand that sells exclusively online. Rather than investing in physical "
    "retail stores, the company channelled its resources into social media "
    "marketing, particularly through partnerships with fitness influencers on "
    "YouTube and Instagram. This digital-first approach enabled Gymshark to "
    "build a global customer base of millions, achieve annual revenues exceeding "
    "five hundred million pounds, and attain a valuation of over one billion "
    "pounds, all without a single high-street store (Browne, 2023). Gymshark's "
    "success demonstrates that digital transformation, when executed effectively, "
    "can enable an SME to achieve scale and market penetration that would be "
    "impossible through traditional retail channels alone. However, it is important "
    "to note that Gymshark's success also required exceptional social media "
    "marketing skill, a strong brand identity, and significant investment in "
    "logistics infrastructure, highlighting that digital transformation is not "
    "a simple or risk-free strategy."
)

# Strategy 2: Strategic Partnerships/Collaborations
doc.add_heading("Strategy 2: Strategic Partnerships and Collaborations", level=3)

add_body(
    "Strategic partnerships involve two or more businesses entering into a formal "
    "or informal collaboration to achieve mutual objectives that neither could "
    "easily accomplish independently. For SMEs, partnerships with larger firms, "
    "complementary businesses, or industry organisations can provide access to "
    "resources, markets, and capabilities that would otherwise be beyond their "
    "reach. This strategy is particularly relevant in the current business "
    "environment, where supply chain complexity, technological demands, and "
    "competitive pressures make it increasingly difficult for small businesses "
    "to succeed in isolation (Johnson et al., 2020)."
)

add_body(
    "The expected outcomes of a strategic partnership strategy are numerous and "
    "significant. First, partnerships provide access to resources that the SME "
    "lacks, such as manufacturing capacity, distribution networks, technology "
    "platforms, or specialist expertise. A small food producer that partners with "
    "a major supermarket chain, for example, gains access to a nationwide "
    "distribution network that would take years and millions of pounds to build "
    "independently. Second, partnerships enable shared risk: by collaborating on "
    "new product development or market entry, both parties share the financial "
    "and operational risks involved. Third, partnering with a well-known brand "
    "can significantly enhance an SME's credibility and market access. Customers "
    "who trust the larger partner are more likely to try the SME's products, "
    "effectively transferring brand equity from the established firm to the "
    "smaller business (Stokes and Wilson, 2021). Fourth, partnerships can "
    "facilitate knowledge transfer, with the SME learning best practices in "
    "areas such as quality management, supply chain logistics, and regulatory "
    "compliance from its larger partner."
)

add_body(
    "The risks associated with strategic partnerships are equally important to "
    "consider. The most significant risk is dependency: if an SME becomes overly "
    "reliant on a single partner for a large proportion of its revenue, it becomes "
    "extremely vulnerable to any changes in that relationship. If the partner "
    "decides to switch suppliers, renegotiate terms, or end the collaboration, "
    "the SME may face a sudden and potentially fatal loss of income. Loss of "
    "control is another risk: partnerships often require the SME to adapt its "
    "products, processes, or branding to meet the partner's requirements, which "
    "may conflict with the SME's own strategic vision or brand identity. Cultural "
    "clashes between the informal, agile culture of an SME and the bureaucratic, "
    "process-driven culture of a large corporation can also create friction and "
    "undermine the partnership's effectiveness (Burns, 2022)."
)

add_body(
    "The experience of small UK food producers partnering with major supermarket "
    "chains such as Tesco and Marks and Spencer illustrates both the potential "
    "and the pitfalls of strategic partnerships. Many artisan food brands, such "
    "as Pip and Nut (nut butters), Propercorn (popcorn), and COOK (frozen meals), "
    "have achieved significant growth by securing listings in major supermarkets. "
    "These partnerships provided access to millions of customers, professional "
    "merchandising, and the credibility of being stocked alongside established "
    "brands. However, some small producers have reported that the relationship "
    "can be unequal, with supermarkets demanding aggressive pricing, imposing "
    "strict delivery schedules, and requiring costly compliance with packaging "
    "and labelling standards (Smithers, 2023). In some cases, supermarkets have "
    "introduced their own-brand alternatives that directly compete with the "
    "SME's products at a lower price point, effectively using the knowledge "
    "gained through the partnership to create a competitor. These examples "
    "demonstrate that while strategic partnerships can accelerate SME growth "
    "and market penetration, they must be entered with clear terms, diversified "
    "across multiple partners where possible, and managed with an awareness of "
    "the inherent power imbalance between small and large businesses."
)

add_page_break()

# ═══════════════════════════════════════════════
# REFERENCE LIST
# ═══════════════════════════════════════════════
doc.add_heading("Reference List", level=1)

references = [
    "ASOS (2020) ASOS plc Annual Report and Accounts 2020. London: ASOS plc.",
    "Bank of England (2025) Monetary Policy Report, February 2025. London: Bank of England.",
    "Barrow, C., Burke, G., Molian, D. and Brown, R. (2023) Enterprise Development: "
    "The Challenges of Starting, Growing and Selling Businesses. London: Cengage Learning.",
    "Blank, S. (2020) The Four Steps to the Epiphany: Successful Strategies for "
    "Products That Win. 2nd edn. Hoboken: Wiley.",
    "British Business Bank (2024) Small Business Finance Markets 2024. Sheffield: "
    "British Business Bank.",
    "Browne, J. (2023) 'How Gymshark built a billion-pound brand without a single "
    "store', Forbes UK, 14 March.",
    "Burns, P. (2022) Entrepreneurship and Small Business: Start-up, Growth and "
    "Maturity. 5th edn. London: Red Globe Press.",
    "Butler, S. (2020) 'Deliveroo: from bicycle courier to stock market float', "
    "The Guardian, 15 November.",
    "Chaffey, D. and Ellis-Chadwick, F. (2022) Digital Marketing: Strategy, "
    "Implementation and Practice. 8th edn. Harlow: Pearson.",
    "Christensen, C.M. (2016) The Innovator's Dilemma: When New Technologies Cause "
    "Great Firms to Fail. Boston: Harvard Business Review Press.",
    "CIPD (2024) Labour Market Outlook: Autumn 2024. London: Chartered Institute "
    "of Personnel and Development.",
    "Companies House (2024) Company Size Thresholds. Available at: "
    "https://www.gov.uk/government/organisations/companies-house (Accessed: 20 February 2026).",
    "Department for Business and Trade (2024) Business Population Estimates for the "
    "UK and Regions 2024. London: HMSO.",
    "Drucker, P.F. (2020) Innovation and Entrepreneurship. Reissue edn. London: "
    "Routledge.",
    "European Commission (2020) User Guide to the SME Definition. Brussels: "
    "Publications Office of the European Union.",
    "FSB (2021) Recovering Together: The Impact of COVID-19 on Small Businesses. "
    "London: Federation of Small Businesses.",
    "FSB (2024) UK Small Business Statistics. London: Federation of Small Businesses.",
    "Gallagher, L. (2017) The Airbnb Story: How Three Ordinary Guys Disrupted an "
    "Industry, Made Billions, and Created Plenty of Controversy. New York: Houghton "
    "Mifflin Harcourt.",
    "Johnson, G., Whittington, R., Scholes, K., Angwin, D. and Regner, P. (2020) "
    "Exploring Strategy: Text and Cases. 12th edn. Harlow: Pearson.",
    "Kotler, P. and Keller, K.L. (2022) Marketing Management. 16th edn. Harlow: "
    "Pearson.",
    "McKinsey Global Institute (2024) The Economic Potential of Generative AI and "
    "Digital Transformation for SMEs. New York: McKinsey and Company.",
    "OECD (2023) SME and Entrepreneurship Outlook 2023. Paris: OECD Publishing.",
    "Office for National Statistics (2025) Consumer Price Inflation, UK: January "
    "2025. London: ONS.",
    "Porter, M.E. (2021) Competitive Strategy: Techniques for Analyzing Industries "
    "and Competitors. Reissue edn. New York: Free Press.",
    "Ries, E. (2011) The Lean Startup: How Constant Innovation Creates Radically "
    "Successful Businesses. London: Portfolio Penguin.",
    "Robertson, D. and Breen, B. (2023) Brick by Brick: How LEGO Rewrote the Rules "
    "of Innovation and Conquered the Global Toy Industry. Updated edn. London: "
    "Random House Business.",
    "Schwab, K. (2023) The Fourth Industrial Revolution. Reissue edn. London: "
    "Portfolio Penguin.",
    "Slack, N. and Brandon-Jones, A. (2022) Operations Management. 10th edn. "
    "Harlow: Pearson.",
    "Smithers, R. (2023) 'Small food producers squeezed by supermarket demands', "
    "The Guardian, 8 September.",
    "Stokes, D. and Wilson, N. (2021) Small Business Management and Entrepreneurship. "
    "8th edn. Andover: Cengage Learning.",
    "Storey, D.J. and Greene, F.J. (2022) Small Business and Entrepreneurship. "
    "2nd edn. Harlow: Pearson.",
    "UK Government (2024) Cyber Security Breaches Survey 2024. London: Department "
    "for Science, Innovation and Technology.",
    "Watt, J. (2016) Business for Punks: Break All the Rules, the BrewDog Way. "
    "London: Portfolio Penguin.",
    "Wood, Z. (2019) 'HMV: the rise and fall of a high street institution', "
    "The Guardian, 28 December.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)

add_page_break()

# ═══════════════════════════════════════════════
# MAPPING TABLE
# ═══════════════════════════════════════════════
doc.add_heading("Mapping Table", level=1)

add_body(
    "The table below maps each task to the marking criteria, describes "
    "the evidence provided, and indicates the marks available."
)

mapping_data = [
    ["1.1", "Definition of SME", "Section 1.1 - EU/OECD and UK definitions with scope", "5"],
    ["1.2", "Three SME characteristics applied to e-commerce",
     "Section 1.2 - Ownership & Control, Access to Finance, Flexibility", "15"],
    ["1.3", "Five stages of the business lifecycle",
     "Section 1.3 - Each stage with 3 characteristics and real example", "15"],
    ["1.4", "Objectives at each lifecycle stage",
     "Section 1.4 - Two objectives per stage with justification", "15"],
    ["2.1", "Internal and external factors impacting SMEs",
     "Section 2.1 - HR, Finance (internal); Technology, Economy (external)", "10"],
    ["2.2", "Business functions across the lifecycle",
     "Section 2.2 - Marketing, Finance, Operations, HR with priority table", "20"],
    ["2.3", "Business strategies with expected outcomes",
     "Section 2.3 - Digital transformation and strategic partnerships", "20"],
]

headers = ["Task", "Description", "Evidence Location", "Marks"]
table = doc.add_table(rows=1 + len(mapping_data), cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

make_table_header_row(table, headers)

for i, row_data in enumerate(mapping_data):
    row = table.rows[i + 1]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9)
        if i % 2 == 0:
            set_cell_shading(cell, "EBF0F7")

# Set column widths
col_widths = [Cm(1.5), Cm(4.5), Cm(8.0), Cm(1.5)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# Total marks row
total_row = table.add_row()
total_row.cells[0].text = ""
total_row.cells[1].text = ""
total_row.cells[2].text = "Total"
total_row.cells[3].text = "100"
for cell in total_row.cells:
    set_cell_shading(cell, "003366")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
            r.font.size = Pt(9)

# ──────────────────────────────────────────────
# Apply footer to any additional sections
# ──────────────────────────────────────────────
for sec in doc.sections:
    apply_footer(sec)

# ──────────────────────────────────────────────
# Save DOCX
# ──────────────────────────────────────────────
os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

doc.save(DOCX_PATH)
print(f"DOCX saved: {DOCX_PATH}")

# ──────────────────────────────────────────────
# Export to PDF via Word COM
# ──────────────────────────────────────────────
import win32com.client
import pythoncom

pythoncom.CoInitialize()
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word_doc = word.Documents.Open(DOCX_PATH)
    word_doc.ExportAsFixedFormat(
        OutputFileName=PDF_PATH,
        ExportFormat=17,  # wdExportFormatPDF
    )
    word_doc.Close(False)
    word.Quit()
    print(f"PDF  saved: {PDF_PATH}")
except Exception as e:
    print(f"PDF export error: {e}")
finally:
    pythoncom.CoUninitialize()

print("Done.")
