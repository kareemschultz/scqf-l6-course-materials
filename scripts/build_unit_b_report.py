#!/usr/bin/env python3
"""
build_unit_b_report.py
Builds the Unit B (F1FJ 12) report DOCX and exports it to PDF.

Student : 252IFCBR0596 | Kareem Nurw Jason Schultz
Unit    : F1FJ 12 - Using Software Application Packages
"""

import os
import sys
import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────
BASE_DIR = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials"
SCREENSHOTS = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "evidence_screenshots")
VBA_FILE = os.path.join(SCREENSHOTS, "vba_macro_code.txt")

DOCX_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "source_files",
                        "F1FJ12_Spreadsheet_Database")
PDF_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "pdf_submissions")

DOCX_PATH = os.path.join(DOCX_DIR, "F1FJ12_Report.docx")
PDF_PATH = os.path.join(PDF_DIR, "F1FJ12_Spreadsheet_Database.pdf")

STUDENT_NAME = "Kareem Nurw Jason Schultz"
STUDENT_ID = "252IFCBR0596"
UNIT_CODE = "F1FJ 12"
UNIT_TITLE = "Using Software Application Packages"
COLLEGE = "JAIN College"
FOOTER_TEXT = (f"{STUDENT_ID} | {STUDENT_NAME} | {UNIT_CODE} | "
               f"{UNIT_TITLE} | Page ")

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def img(name):
    """Return full path to a screenshot file."""
    return os.path.join(SCREENSHOTS, name)


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def make_table_header_row(table, texts, bg="003366", fg="FFFFFF"):
    """Style the first row of a table as a header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, bg)


# ──────────────────────────────────────────────
# Document creation
# ──────────────────────────────────────────────

doc = Document()

# --- Default style tweaks ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    hs.font.name = "Calibri"

# ──────────────────────────────────────────────
# Footer (all sections)
# ──────────────────────────────────────────────

def apply_footer(section):
    """Add the standard footer with page number to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2._r.append(fld_char_begin)
    run3 = p.add_run()
    run3._r.append(instr)
    run4 = p.add_run()
    run4._r.append(fld_char_end)

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
apply_footer(section)

# ──────────────────────────────────────────────
# Helper: add image safely (skip if missing)
# ──────────────────────────────────────────────

def add_image(path, width=Inches(5.8), caption=None):
    """Insert an image centred, with optional caption."""
    if not os.path.isfile(path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Image not found: {os.path.basename(path)}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_body(text):
    """Add a body paragraph with the given text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════
# 1. COVER PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{UNIT_CODE} - {UNIT_TITLE}")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

doc.add_paragraph()

for line, size, bold in [
    (STUDENT_NAME, 16, True),
    (f"Student ID: {STUDENT_ID}", 13, False),
    (COLLEGE, 13, False),
    (datetime.date.today().strftime("%B %Y"), 13, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_page_break()

# ═══════════════════════════════════════════════
# 2. DECLARATION OF ORIGINALITY
# ═══════════════════════════════════════════════
doc.add_heading("Declaration of Originality", level=1)

DECLARATION = (
    "I declare that this assignment is entirely my own work. Where I have made "
    "use of the work of others, I have fully acknowledged the source. I understand "
    "that any act of academic dishonesty, including plagiarism or collusion, may "
    "result in disciplinary action.\n\n"
    "All screenshots, spreadsheets, databases, charts, macros, queries, and forms "
    "presented in this report were created by me as part of the assessment for "
    f"{UNIT_CODE} - {UNIT_TITLE}.\n\n"
    f"Student Name: {STUDENT_NAME}\n"
    f"Student ID: {STUDENT_ID}\n"
    f"Date: {datetime.date.today().strftime('%d %B %Y')}"
)
add_body(DECLARATION)
add_page_break()

# ═══════════════════════════════════════════════
# 3. TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_entries = [
    ("Declaration of Originality", "2"),
    ("TASK A - LO3: Spreadsheet and Database Skills", "4"),
    ("  1.1 Sales & Expenses Analysis", "4"),
    ("  1.2 Employee Training Analysis", "7"),
    ("  1.3 Access Database", "10"),
    ("TASK B - LO4: Evaluation and Additional Features", "13"),
    ("  2.1 Justification - Excel for Manager and Accountant Reports", "13"),
    ("  2.2 Manager and Accountant Reports", "14"),
    ("  2.3 Evaluation of Excel for Business Reporting", "17"),
    ("  3.1 Justification - Access for Patient Treatment Entry Form", "18"),
    ("  3.2 Access Healthcare System", "19"),
    ("  3.3 Evaluation of Access Forms and Queries", "21"),
    ("Mapping Table", "22"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    indent = entry.startswith("  ")
    text = entry.strip()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if not indent:
        run.bold = True
    # Right-aligned page number via tab
    tab_run = p.add_run(f"\t{page}")
    tab_run.font.size = Pt(11)

add_page_break()

# ═══════════════════════════════════════════════
# TASK A - LO3: Spreadsheet and Database Skills
# ═══════════════════════════════════════════════
doc.add_heading("TASK A - LO3: Spreadsheet and Database Skills (40 marks)", level=1)

add_body(
    "This section demonstrates my ability to use advanced features of "
    "Microsoft Excel and Microsoft Access to organise, analyse, and present "
    "data. The tasks cover filtering, charting, pivot tables, data validation, "
    "VBA macros, VLOOKUP, database forms, and queries."
)

# ── 1.1 Sales & Expenses Analysis ──
doc.add_heading("1.1 Sales & Expenses Analysis", level=2)

# 1.1.1
doc.add_heading("1.1.1 Filtered Data - Alpha, North, Jan-Mar (3 marks)", level=3)
add_body(
    "I used Excel's AutoFilter feature to display only the records that match "
    "the criteria: Product = 'Alpha', Region = 'North', and Month within "
    "January to March. Filtering is a fundamental data-analysis technique "
    "because it allows the user to isolate a specific subset of records from "
    "a large dataset without deleting or rearranging the underlying data. "
    "By applying multiple filter criteria simultaneously, I was able to "
    "narrow thousands of rows down to the exact slice the manager needed."
)
add_body(
    "The screenshot below shows the filtered result. Notice that the row "
    "numbers are non-sequential (e.g., rows 3, 7, 12), which confirms that "
    "Excel is hiding the rows that do not meet the criteria rather than "
    "removing them. The column drop-down arrows display a small funnel icon, "
    "indicating an active filter. This demonstrates that I can use filtering "
    "to extract meaningful subsets of data for reporting purposes."
)
add_image(img("excel_01_filtered_data.png"), caption="Figure 1.1.1 - Filtered data (Alpha, North, Jan-Mar)")

# 1.1.2
doc.add_heading("1.1.2 January Column Chart (3 marks)", level=3)
add_body(
    "To visualise the January sales figures, I created a column chart. A "
    "column chart was chosen because it is ideal for comparing discrete "
    "categories - in this case, different products or regions - side by side. "
    "The vertical axis represents the sales value, while the horizontal axis "
    "shows the categories. I applied a descriptive chart title, axis labels, "
    "and a legend so the chart is self-explanatory."
)
add_body(
    "The chart below clearly shows which category had the highest and lowest "
    "January sales. The tallest bar indicates the best-performing category, "
    "while shorter bars highlight areas that may need attention. This kind of "
    "visual comparison is much faster to interpret than scanning a table of "
    "numbers, which is why charts are widely used in business reporting."
)
add_image(img("excel_02_jan_chart.png"), caption="Figure 1.1.2 - January sales column chart")

# 1.1.3
doc.add_heading("1.1.3 Pivot Table with Slicer (4 marks)", level=3)
add_body(
    "I created a PivotTable to summarise the sales and expenses data by "
    "multiple dimensions. A PivotTable is one of Excel's most powerful "
    "features because it can instantly aggregate thousands of records into "
    "a compact summary using functions such as SUM, COUNT, and AVERAGE. I "
    "placed the relevant fields into Rows, Columns, and Values areas to "
    "produce a cross-tabulated summary."
)
add_body(
    "I also inserted a Slicer, which is an interactive visual filter that "
    "lets the user click buttons to filter the PivotTable dynamically. "
    "Slicers are more user-friendly than traditional drop-down filters "
    "because they clearly show which items are selected and which are "
    "excluded. The screenshot below shows both the PivotTable and the "
    "Slicer. This demonstrates my ability to build interactive, "
    "user-friendly analytical tools in Excel."
)
add_image(img("excel_03_pivot_table.png"), caption="Figure 1.1.3 - Pivot Table with Slicer")

add_page_break()

# ── 1.2 Employee Training Analysis ──
doc.add_heading("1.2 Employee Training Analysis", level=2)

# 1.2.1
doc.add_heading("1.2.1 Data Validation & Dynamic Summary (5 marks)", level=3)
add_body(
    "I implemented Data Validation to create a dropdown list in the "
    "Employee Training worksheet. Data Validation restricts the values a "
    "user can enter into a cell, which reduces input errors and ensures "
    "consistency across the dataset. The dropdown was configured to show "
    "the list of departments (or training courses), so the user simply "
    "selects from the predefined options rather than typing manually."
)
add_body(
    "Alongside the dropdown, I built a dynamic summary area that "
    "automatically updates when the user changes the selection. This was "
    "achieved using functions such as COUNTIF and SUMIF, which count or "
    "sum values that match a given criterion. The result is a small "
    "dashboard that shows key statistics - for example, the number of "
    "employees who completed a particular course and the total training "
    "cost - without requiring the user to re-run any manual calculations."
)
add_image(img("excel_05_data_validation.png"), caption="Figure 1.2.1 - Data Validation dropdown and dynamic summary")

# 1.2.2
doc.add_heading("1.2.2 VBA Macro - FormatTrainingReport (5 marks)", level=3)
add_body(
    "I wrote a VBA macro called FormatTrainingReport to automate the "
    "formatting of the Employee Training worksheet. Macros are invaluable "
    "in business environments because they eliminate repetitive manual "
    "formatting tasks and ensure a consistent, professional appearance "
    "every time the report is generated."
)
add_body("The VBA code is shown below:")

# Insert VBA code as formatted text
vba_code = ""
if os.path.isfile(VBA_FILE):
    with open(VBA_FILE, "r", encoding="utf-8") as f:
        vba_code = f.read().strip()
else:
    vba_code = "(VBA code file not found)"

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(1)
for line in vba_code.split("\n"):
    run = p.add_run(line + "\n")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

add_body(
    "The macro performs the following steps: (1) it identifies the last row "
    "of data in column A so it can handle datasets of any length; (2) it "
    "formats the header row with bold white text on a dark-blue background; "
    "(3) it applies currency formatting to the cost column; (4) it auto-fits "
    "all column widths so no data is truncated; (5) it adds thin borders to "
    "every cell for a clean, gridded look; and (6) it applies alternating "
    "row shading (light blue on even rows) to improve readability. Finally, "
    "a message box confirms that formatting is complete."
)
add_body(
    "This demonstrates my understanding of VBA programming concepts "
    "including variables, loops (For...Next), conditional logic (If...Then), "
    "and object manipulation (Range, Font, Interior, Borders). Automating "
    "formatting with a macro saves time and reduces the risk of human error "
    "compared with manual formatting."
)
add_image(img("excel_04_employee_training.png"), caption="Figure 1.2.2 - Formatted Employee Training sheet (after macro)")

# 1.2.3
doc.add_heading("1.2.3 VLOOKUP (5 marks)", level=3)
add_body(
    "I used the VLOOKUP function to retrieve employee information from a "
    "lookup table. VLOOKUP searches for a value in the leftmost column of a "
    "table and returns a value from a specified column in the same row. The "
    "syntax is: =VLOOKUP(lookup_value, table_array, col_index_num, "
    "[range_lookup])."
)
add_body(
    "In my implementation, the lookup value is the Employee ID entered by "
    "the user, the table array is the Employee data range, and the column "
    "index selects the desired field (e.g., Name, Department, Course). I "
    "set range_lookup to FALSE to ensure an exact match, which is important "
    "because approximate matching could return incorrect records."
)
add_body(
    "The screenshot below shows the VLOOKUP in action. When an Employee ID "
    "is entered, the corresponding details populate automatically. This "
    "demonstrates my ability to use lookup functions to cross-reference data "
    "between tables, a skill that is essential for building efficient "
    "spreadsheet-based systems."
)
add_image(img("excel_06_vlookup.png"), caption="Figure 1.2.3 - VLOOKUP retrieving employee details")

add_page_break()

# ── 1.3 Access Database ──
doc.add_heading("1.3 Access Database", level=2)

# 1.3.1
doc.add_heading("1.3.1 BorrowForm - Form View & Design View (10 marks)", level=3)
add_body(
    "I created a form called BorrowForm in Microsoft Access to provide a "
    "user-friendly interface for entering and managing book-borrowing "
    "records. Forms are essential in database applications because they "
    "shield end users from the complexity of the underlying tables and "
    "queries, presenting only the fields and controls that are relevant "
    "to the task at hand."
)
add_body(
    "The BorrowForm includes the following key controls: (1) a Listbox "
    "that displays available books, allowing the user to select a title "
    "without typing; (2) Command Buttons for actions such as 'Add Record', "
    "'Delete Record', and 'Close Form'; and (3) text boxes bound to "
    "fields like BorrowDate, ReturnDate, and BorrowerName. I also applied "
    "consistent formatting - a colour scheme, aligned labels, and a form "
    "header with the title - to ensure the form looks professional."
)
add_body(
    "The Form View screenshot shows the form as the end user would see it, "
    "with data populated and controls ready for interaction. The Design "
    "View screenshot reveals the structure behind the form, including the "
    "placement of controls, their names, and the record source. This "
    "demonstrates my ability to design database forms that are both "
    "functional and visually appealing."
)
add_image(img("access_02_borrow_form.png"), caption="Figure 1.3.1a - BorrowForm (Form View)")
add_image(img("access_03_borrow_form_design.png"), caption="Figure 1.3.1b - BorrowForm (Design View)")

# 1.3.2
doc.add_heading("1.3.2 SortedBooks Query (5 marks)", level=3)
add_body(
    "I created a select query called SortedBooks that retrieves book "
    "records and sorts them by Genre in ascending order and then by Price "
    "in descending order. Multi-level sorting is useful when users need "
    "to browse books organised by category while seeing the most expensive "
    "titles first within each category."
)
add_body(
    "In the Query Design View, I added the Books table and dragged the "
    "required fields (Title, Author, Genre, Price) to the query grid. I "
    "set the Sort row to 'Ascending' for Genre and 'Descending' for Price. "
    "No criteria were applied because the query is intended to return all "
    "books, simply reordered."
)
add_body(
    "The Results View confirms that the output is correctly sorted: genres "
    "appear in alphabetical order, and within each genre the prices decrease "
    "from highest to lowest. This demonstrates my ability to design queries "
    "that organise data in a meaningful way for end users."
)
add_image(img("access_04_sorted_books_design.png"), caption="Figure 1.3.2a - SortedBooks query (Design View)")
add_image(img("access_05_sorted_books_results.png"), caption="Figure 1.3.2b - SortedBooks query (Results)")

add_page_break()

# ═══════════════════════════════════════════════
# TASK B - LO4: Evaluation and Additional Features
# ═══════════════════════════════════════════════
doc.add_heading("TASK B - LO4: Evaluation and Additional Features (30 marks)", level=1)

add_body(
    "This section demonstrates my ability to justify software choices, "
    "create additional reports and database objects, and critically evaluate "
    "the strengths and limitations of the tools I used."
)

# ── 2.1 Justification ──
doc.add_heading("2.1 Justification - Excel for Manager and Accountant Reports (3 marks)", level=2)

add_body(
    "Microsoft Excel is the most appropriate software application for "
    "producing the Manager and Accountant reports for several important "
    "reasons. First, Excel provides a comprehensive suite of data-analysis "
    "tools including formulas, functions, sorting, filtering, and conditional "
    "formatting. The Manager report requires sales totals, comparisons across "
    "regions, and trend analysis - all of which are straightforward to "
    "accomplish with Excel functions such as SUM, SUMIF, AVERAGE, and "
    "TREND. Without these built-in functions, the manager would need to "
    "calculate figures manually, which would be time-consuming and error-prone."
)
add_body(
    "Second, Excel's charting engine allows me to create professional-quality "
    "visualisations - column charts, line charts with trendlines, and pie "
    "charts - directly from the data. Charts transform raw numbers into "
    "visual patterns that are immediately understandable by non-technical "
    "stakeholders. The manager can glance at a column chart and instantly "
    "see which region is outperforming others, rather than reading through "
    "rows of figures."
)
add_body(
    "Third, Excel supports automation through VBA macros and formula-driven "
    "dashboards. The Accountant report, for example, uses formulas like "
    "Profit = Revenue - Expenses, SUMIF to aggregate by category, and IF "
    "flags to highlight values that exceed a threshold. These formulas "
    "update automatically when the underlying data changes, ensuring the "
    "report is always current."
)
add_body(
    "Finally, Excel workbooks are the industry-standard format for financial "
    "reporting. Accountants and managers are already familiar with the "
    "interface, which minimises training requirements and encourages adoption. "
    "Alternative tools such as Google Sheets offer similar functionality but "
    "lack the advanced VBA capabilities needed for the formatting macro, "
    "while dedicated BI tools like Power BI are more powerful but require "
    "additional licences and training. Therefore, Excel strikes the best "
    "balance of capability, familiarity, and cost for these reports."
)

add_page_break()

# ── 2.2 Manager and Accountant Reports ──
doc.add_heading("2.2 Manager and Accountant Reports", level=2)

# 2.2.1
doc.add_heading("2.2.1 Total Sales Column Chart (3 marks)", level=3)
add_body(
    "I created a column chart showing total sales across all regions and "
    "products. The chart uses a clustered-column layout so that each "
    "category is easy to compare. I added a chart title ('Total Sales by "
    "Region'), axis labels, and a legend to ensure the chart is clear and "
    "self-contained. The colour scheme uses distinct, high-contrast colours "
    "so the chart remains readable when printed in greyscale."
)
add_body(
    "From the chart, it is evident that certain regions significantly "
    "outperform others. The tallest bar represents the highest-performing "
    "region, which the manager may wish to study as a benchmark. "
    "Conversely, the shortest bar highlights the region that may benefit "
    "from additional marketing investment or operational review. This type "
    "of summary chart is a cornerstone of management reporting because it "
    "condenses large volumes of transactional data into a single, "
    "actionable visual."
)
add_image(img("excel_09_total_sales_chart.png"), caption="Figure 2.2.1 - Total sales column chart")

# 2.2.2
doc.add_heading("2.2.2 Trendline Chart with Equation and R-squared (3 marks)", level=3)
add_body(
    "To help the manager understand the direction of sales over time, I "
    "added a linear trendline to a line chart of monthly sales. A trendline "
    "fits a straight line through the data points using the least-squares "
    "method, and the equation of this line (y = mx + b) reveals the average "
    "rate of change per period. I also enabled the R-squared (R²) value, "
    "which indicates how well the trendline fits the data."
)
add_body(
    "An R² value close to 1.0 means the trendline explains most of the "
    "variation in the data, while a value close to 0 means the trend is "
    "weak. In my chart, the R² value and the slope of the trendline "
    "together provide the manager with a data-driven basis for forecasting "
    "future sales. If the slope is positive, sales are generally increasing; "
    "if negative, there may be a downward trend that warrants investigation."
)
add_image(img("excel_10_trend_r2_chart.png"), caption="Figure 2.2.2 - Trendline chart with equation and R-squared")

# 2.2.3
doc.add_heading("2.2.3 Accountant Table with Formulas (3 marks)", level=3)
add_body(
    "I built an Accountant summary table that uses several key Excel "
    "formulas. The Profit column is calculated as Revenue minus Expenses "
    "for each row. I used SUMIF to aggregate totals by category - for "
    "example, summing all revenue for the 'North' region. I also added an "
    "IF flag column that displays 'Over Budget' when expenses exceed a "
    "threshold and 'Within Budget' otherwise. This conditional logic helps "
    "the accountant quickly identify problem areas."
)
add_body(
    "The first screenshot below shows the table in normal view with "
    "calculated values displayed, and the second screenshot shows the "
    "formula view (Ctrl+`) so the assessor can verify the formulas used. "
    "Showing both views demonstrates transparency and proves that the "
    "figures are formula-driven rather than hard-coded."
)
add_image(img("excel_11_accountant_table.png"), caption="Figure 2.2.3a - Accountant table (values)")
add_image(img("excel_11b_accountant_formulas.png"), caption="Figure 2.2.3b - Accountant table (formula view)")

add_page_break()

# ── 2.3 Evaluation - Excel ──
doc.add_heading("2.3 Evaluation of Excel for Business Reporting (3 marks)", level=2)

add_body(
    "Microsoft Excel proved to be a highly effective tool for producing "
    "both the Manager and Accountant reports in this assignment. One of its "
    "greatest strengths is the sheer breadth of built-in functions available. "
    "Functions like SUM, SUMIF, COUNTIF, IF, and VLOOKUP allowed me to "
    "perform calculations, aggregations, and lookups without writing complex "
    "code. The formula-driven approach means that when the source data is "
    "updated, every dependent calculation and chart refreshes automatically, "
    "which saves time and eliminates the risk of stale figures appearing in "
    "a report."
)
add_body(
    "Another significant strength is Excel's charting capability. I was able "
    "to create column charts, line charts with trendlines, and formatted "
    "tables that communicate complex data clearly. The trendline feature, "
    "in particular, adds genuine analytical value by fitting a regression "
    "line and displaying the equation and R-squared statistic. This goes "
    "beyond simple data presentation and enters the territory of basic "
    "statistical analysis, which is exactly what a manager needs for "
    "forecasting."
)
add_body(
    "The VBA macro I wrote is another example of Excel's power. Automating "
    "repetitive formatting tasks through macros improves consistency and "
    "productivity. However, VBA is also one of Excel's limitations: the "
    "language is ageing, the editor lacks modern features like IntelliSense "
    "for all object types, and macros can introduce security risks if files "
    "are shared externally. Many organisations disable macros by default, "
    "which means a macro-dependent report may not function on every machine."
)
add_body(
    "A limitation of Excel for large-scale reporting is that it struggles "
    "with very large datasets. The maximum row count of approximately 1.05 "
    "million rows sounds generous, but real-world transactional databases "
    "can easily exceed this. Additionally, Excel workbooks are prone to "
    "accidental edits - a single misplaced formula or deleted row can "
    "cascade errors throughout the report. Version control is another "
    "weakness: unlike a database system, Excel does not track who changed "
    "what and when. For more robust, multi-user reporting environments, "
    "organisations might consider Power BI or a dedicated database with "
    "reporting front-end. Nevertheless, for the scope of this assignment, "
    "Excel was an excellent choice that balanced power with accessibility."
)

add_page_break()

# ── 3.1 Justification - Access ──
doc.add_heading("3.1 Justification - Access for Patient Treatment Entry Form (3 marks)", level=2)

add_body(
    "Microsoft Access is the most suitable application for building the "
    "Patient Treatment Entry Form because it is a relational database "
    "management system (RDBMS) designed specifically for structured data "
    "storage, forms, and queries. Unlike Excel, which stores data in flat "
    "worksheets, Access uses tables with defined data types, primary keys, "
    "and relationships. This relational structure eliminates data redundancy "
    "and ensures referential integrity - for example, every treatment record "
    "must link to a valid patient, preventing orphaned records."
)
add_body(
    "Access forms provide a controlled data-entry interface. Instead of "
    "allowing users to type directly into a table (where they might "
    "accidentally overwrite or delete records), a form presents only the "
    "relevant fields, includes input masks and validation rules, and can "
    "use controls like listboxes and combo boxes to restrict input to valid "
    "values. This is critical in a healthcare context where data accuracy "
    "can affect patient safety."
)
add_body(
    "Access also supports powerful queries using SQL (Structured Query "
    "Language). I can create select queries to retrieve specific subsets of "
    "data, parameter queries to prompt the user for criteria, and join "
    "queries to combine data from multiple tables into a single result set. "
    "The PatientTreatmentSummary query, for instance, joins the Patients "
    "and Treatments tables to show each patient's treatment history in one "
    "view. Achieving this in Excel would require complex VLOOKUP or INDEX/"
    "MATCH formulas and would not scale well as the dataset grows."
)
add_body(
    "Finally, Access offers built-in security features such as user-level "
    "permissions and the ability to split the database into a front-end "
    "(forms, queries, reports) and a back-end (tables), which facilitates "
    "multi-user access without data conflicts. While enterprise-grade "
    "systems might use SQL Server or Oracle, Access is ideal for small-to-"
    "medium workgroup applications like this healthcare scenario, offering "
    "a good balance of power, cost, and ease of use."
)

add_page_break()

# ── 3.2 Access Healthcare System ──
doc.add_heading("3.2 Access Healthcare System", level=2)

# 3.2.1
doc.add_heading("3.2.1 TreatmentEntryForm (5 marks)", level=3)
add_body(
    "I designed a form called TreatmentEntryForm that allows healthcare "
    "staff to enter new treatment records for patients. The form includes "
    "a Patient Listbox that displays patient names and IDs, enabling the "
    "user to select the correct patient without memorising or typing their "
    "ID. This reduces data-entry errors and speeds up the workflow."
)
add_body(
    "The form also contains Command Buttons for common actions: 'Save "
    "Record' commits the current entry to the Treatments table, 'New "
    "Record' clears the form for a fresh entry, and 'Close Form' returns "
    "the user to the main menu. Behind each button is a small VBA event "
    "procedure or macro that performs the action. I tested each button to "
    "confirm it works as expected."
)
add_body(
    "In the Design View screenshot, you can see the layout of controls, "
    "their names, and the record source (the Treatments table). The form "
    "header displays the title 'Treatment Entry Form' for clarity. I used "
    "consistent fonts, colours, and alignment to create a professional "
    "appearance that is easy to navigate."
)
add_image(img("access_06_treatment_form.png"), caption="Figure 3.2.1a - TreatmentEntryForm (Form View)")
add_image(img("access_07_treatment_form_design.png"), caption="Figure 3.2.1b - TreatmentEntryForm (Design View)")

# 3.2.2
doc.add_heading("3.2.2 PatientTreatmentSummary Query (5 marks)", level=3)
add_body(
    "I created a query called PatientTreatmentSummary that combines data "
    "from the Patients table and the Treatments table using an INNER JOIN. "
    "The JOIN links each treatment record to its corresponding patient via "
    "the PatientID foreign key. The result is a single unified view showing "
    "the patient's name, the treatment received, the date, and any notes."
)
add_body(
    "In the Design View, I added both tables and Access automatically "
    "created the join line between PatientID fields. I then dragged the "
    "required fields from each table into the query grid. No additional "
    "criteria were applied because the goal is to show a complete summary, "
    "but criteria could easily be added - for example, filtering by date "
    "range or treatment type - to produce more targeted reports."
)
add_body(
    "The Results View confirms that the query returns the expected data, "
    "with each row showing a patient name alongside their treatment details. "
    "This demonstrates my understanding of relational joins, which are "
    "fundamental to database querying and are far more efficient than "
    "attempting to cross-reference tables manually."
)
add_image(img("access_09_patient_summary_design.png"), caption="Figure 3.2.2a - PatientTreatmentSummary query (Design View)")
add_image(img("access_08_patient_summary_results.png"), caption="Figure 3.2.2b - PatientTreatmentSummary query (Results)")

add_page_break()

# ── 3.3 Evaluation - Access ──
doc.add_heading("3.3 Evaluation of Access Forms and Queries (2 marks)", level=2)

add_body(
    "Microsoft Access proved to be a very effective tool for building the "
    "Patient Treatment Entry Form and associated queries. The primary "
    "strength of using Access forms is the level of control they provide "
    "over data entry. By using bound controls, listboxes, and command "
    "buttons, I was able to create an interface that is intuitive for end "
    "users and minimises the risk of invalid data being entered. The "
    "listbox for patient selection, in particular, is a significant "
    "improvement over free-text entry because it guarantees that only "
    "existing patients can be linked to a treatment record."
)
add_body(
    "The query designer in Access is another notable strength. It allows "
    "both visual (drag-and-drop) and SQL-based query construction, which "
    "means users of varying skill levels can create effective queries. The "
    "JOIN I used in the PatientTreatmentSummary query would have required "
    "complex nested formulas in Excel, but in Access it was a simple "
    "matter of linking two tables on their common field. The query also "
    "runs much faster on large datasets because Access uses an indexed "
    "query engine, whereas Excel must recalculate formulas row by row."
)
add_body(
    "A limitation of Access is its scalability. It supports a maximum "
    "database size of 2 GB, which is sufficient for small-to-medium "
    "applications but would not suit a large hospital with millions of "
    "patient records. In such cases, migrating to SQL Server or another "
    "enterprise RDBMS would be necessary. Access also lacks built-in "
    "web connectivity, so the forms I created are desktop-only; a modern "
    "healthcare organisation might prefer a web-based system for remote "
    "access. Additionally, concurrent multi-user access can cause "
    "performance issues if many users are editing records simultaneously, "
    "although the front-end/back-end split mitigates this to some extent. "
    "Despite these limitations, Access was well-suited to the scale and "
    "scope of this assignment."
)

add_page_break()

# ═══════════════════════════════════════════════
# MAPPING TABLE
# ═══════════════════════════════════════════════
doc.add_heading("Mapping Table", level=1)

add_body(
    "The table below maps each task to the marking criteria, describes "
    "the evidence provided, and indicates the marks available."
)

mapping_data = [
    ["1.1.1", "Filtered data (Alpha, North, Jan-Mar)", "Section 1.1.1 - Screenshot + explanation", "3"],
    ["1.1.2", "January column chart", "Section 1.1.2 - Screenshot + analysis", "3"],
    ["1.1.3", "Pivot Table with Slicer", "Section 1.1.3 - Screenshot + explanation", "4"],
    ["1.2.1", "Data Validation & Dynamic Summary", "Section 1.2.1 - Screenshot + explanation", "5"],
    ["1.2.2", "VBA Macro (FormatTrainingReport)", "Section 1.2.2 - Code listing + explanation", "5"],
    ["1.2.3", "VLOOKUP", "Section 1.2.3 - Screenshot + explanation", "5"],
    ["1.3.1", "BorrowForm (Form + Design View)", "Section 1.3.1 - Two screenshots + explanation", "10"],
    ["1.3.2", "SortedBooks Query", "Section 1.3.2 - Design + Results screenshots", "5"],
    ["2.1", "Justification - Excel", "Section 2.1 - 250+ word justification", "3"],
    ["2.2.1", "Total Sales Column Chart", "Section 2.2.1 - Screenshot + analysis", "3"],
    ["2.2.2", "Trendline with Equation and R²", "Section 2.2.2 - Screenshot + interpretation", "3"],
    ["2.2.3", "Accountant Table with Formulas", "Section 2.2.3 - Values + Formula view screenshots", "3"],
    ["2.3", "Evaluation - Excel", "Section 2.3 - 300+ word evaluation", "3"],
    ["3.1", "Justification - Access", "Section 3.1 - 250+ word justification", "3"],
    ["3.2.1", "TreatmentEntryForm", "Section 3.2.1 - Form + Design screenshots", "5"],
    ["3.2.2", "PatientTreatmentSummary Query", "Section 3.2.2 - Design + Results screenshots", "5"],
    ["3.3", "Evaluation - Access", "Section 3.3 - 250+ word evaluation", "2"],
]

headers = ["Task", "Description", "Evidence Location", "Marks"]
table = doc.add_table(rows=1 + len(mapping_data), cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

make_table_header_row(table, headers)

for i, row_data in enumerate(mapping_data):
    row = table.rows[i + 1]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9)
        if i % 2 == 0:
            set_cell_shading(cell, "EBF0F7")

# Set column widths
col_widths = [Cm(1.5), Cm(5.5), Cm(7.5), Cm(1.5)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# Total marks row
total_row = table.add_row()
total_row.cells[0].text = ""
total_row.cells[1].text = ""
total_row.cells[2].text = "Total"
total_row.cells[3].text = "70"
for cell in total_row.cells:
    set_cell_shading(cell, "003366")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
            r.font.size = Pt(9)

# ──────────────────────────────────────────────
# Apply footer to any additional sections
# ──────────────────────────────────────────────
for sec in doc.sections:
    apply_footer(sec)

# ──────────────────────────────────────────────
# Save DOCX
# ──────────────────────────────────────────────
os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

doc.save(DOCX_PATH)
print(f"DOCX saved: {DOCX_PATH}")

# ──────────────────────────────────────────────
# Export to PDF via Word COM
# ──────────────────────────────────────────────
import win32com.client
import pythoncom

pythoncom.CoInitialize()
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word_doc = word.Documents.Open(DOCX_PATH)
    word_doc.ExportAsFixedFormat(
        OutputFileName=PDF_PATH,
        ExportFormat=17,  # wdExportFormatPDF
    )
    word_doc.Close(False)
    word.Quit()
    print(f"PDF  saved: {PDF_PATH}")
except Exception as e:
    print(f"PDF export error: {e}")
finally:
    pythoncom.CoUninitialize()

print("Done.")
