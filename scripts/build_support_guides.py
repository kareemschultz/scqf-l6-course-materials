"""Build Step-by-Step Excel and Access feature guides as DOCX then export to PDF."""
import subprocess, time, os

subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
time.sleep(1)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'SCQF_L6_SUPPORT_PACK', 'F1FJ12_Spreadsheet_Database'
)

def add_step(doc, num, text, substeps=None):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {num}: ')
    run.font.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    if substeps:
        for sub in substeps:
            sp = doc.add_paragraph(sub, style='List Bullet')
            for r in sp.runs:
                r.font.size = Pt(10)

def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0, 100, 0)

# ======================================================================
# EXCEL FEATURES GUIDE
# ======================================================================
doc = Document()
doc.add_heading('Step-by-Step Excel Features Guide', level=0)
doc.add_paragraph(
    'This guide explains HOW to use the Excel features required for the '
    'F1FJ12 assessment. It does not contain analysis or evaluation text. '
    'Follow these steps to build your own workbook.'
)

# --- Filtering ---
doc.add_heading('1. Filtering Data', level=1)
doc.add_paragraph(
    'Filtering lets you show only the rows that match specific criteria.'
)
add_step(doc, 1, 'Click any cell inside your data range.')
add_step(doc, 2, 'Go to Data tab > Filter (or press Ctrl + Shift + L).')
add_step(doc, 3, 'Click the dropdown arrow on the column header you want to filter.')
add_step(doc, 4, 'Untick "Select All", then tick only the values you want to show.')
add_step(doc, 5, 'Click OK. Only matching rows will be displayed.')
add_step(doc, 6, 'To remove the filter, click the dropdown again and select "Clear Filter".')
add_tip(doc, 'You can filter multiple columns at once. Apply filters one column at a time.')

# --- Column Chart ---
doc.add_heading('2. Creating a Column Chart', level=1)
add_step(doc, 1, 'Select the data range you want to chart (including headers).')
add_step(doc, 2, 'Go to Insert tab > Charts group > Column Chart.')
add_step(doc, 3, 'Choose "Clustered Column" (the first option).')
add_step(doc, 4, 'Click the chart to select it, then use the Chart Design tab to:', substeps=[
    'Add a chart title (click "Add Chart Element" > Chart Title)',
    'Add axis titles if needed',
    'Change the colour scheme using "Change Colors"',
])
add_step(doc, 5, 'Right-click the chart and select "Move Chart" to place it on a separate sheet if required.')
add_tip(doc, 'To compare two data series (e.g. Sales vs Expenses), include both columns in your selection.')

# --- Pivot Table ---
doc.add_heading('3. Creating a Pivot Table with Slicer', level=1)
doc.add_paragraph('A pivot table summarises large datasets by grouping and aggregating data.')
add_step(doc, 1, 'Click any cell inside your data range.')
add_step(doc, 2, 'Go to Insert tab > PivotTable.')
add_step(doc, 3, 'In the dialog box, choose "New Worksheet" and click OK.')
add_step(doc, 4, 'In the PivotTable Fields pane on the right:', substeps=[
    'Drag a text field (e.g. Product) to the Rows area',
    'Drag another text field (e.g. Region) to the Columns area',
    'Drag a number field (e.g. Sales) to the Values area',
])
add_step(doc, 5, 'To add a Slicer: click inside the pivot table, then go to PivotTable Analyze > Insert Slicer.')
add_step(doc, 6, 'Tick the field you want to filter by (e.g. Region) and click OK.')
add_step(doc, 7, 'Click items in the slicer to filter the pivot table interactively.')
add_tip(doc, 'Hold Ctrl to select multiple items in a slicer. Use the clear filter button (funnel icon) to reset.')

# --- Data Validation ---
doc.add_heading('4. Data Validation Dropdown', level=1)
add_step(doc, 1, 'Click the cell where you want the dropdown.')
add_step(doc, 2, 'Go to Data tab > Data Validation.')
add_step(doc, 3, 'Under "Allow", select "List".')
add_step(doc, 4, 'In the "Source" box, either:', substeps=[
    'Type the options separated by commas (e.g. Sales,IT,HR,Finance)',
    'Or select a range of cells containing the options',
])
add_step(doc, 5, 'Click OK. The cell now has a dropdown arrow.')
add_tip(doc, 'Use a named range for cleaner formulas: Formulas > Name Manager > New.')

# --- Dynamic Summary with SUMIFS / COUNTIFS ---
doc.add_heading('5. Dynamic Summary Formulas', level=1)
doc.add_paragraph(
    'These formulas calculate totals based on conditions, which makes them '
    'ideal for linking to a dropdown selection.'
)
doc.add_paragraph('SUMIFS syntax:')
p = doc.add_paragraph('=SUMIFS(sum_range, criteria_range1, criteria1, ...)')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
doc.add_paragraph('Example: =SUMIFS(E2:E100, C2:C100, G1)  where G1 is your dropdown cell.')
doc.add_paragraph('')
doc.add_paragraph('COUNTIFS syntax:')
p = doc.add_paragraph('=COUNTIFS(criteria_range1, criteria1, ...)')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
doc.add_paragraph('Example: =COUNTIFS(C2:C100, G1)  counts rows matching the dropdown value.')

# --- VBA Macro ---
doc.add_heading('6. Recording a VBA Macro', level=1)
add_step(doc, 1, 'Go to View tab > Macros > Record Macro.')
add_step(doc, 2, 'Give the macro a name (e.g. FormatReport). No spaces allowed in the name.')
add_step(doc, 3, 'Click OK to start recording.')
add_step(doc, 4, 'Perform the formatting actions you want the macro to repeat:', substeps=[
    'Select header row and apply Bold',
    'Apply currency format to the Cost column',
    'Auto-fit column widths (Home > Format > AutoFit Column Width)',
    'Apply a fill colour to the header row',
])
add_step(doc, 5, 'Go to View tab > Macros > Stop Recording.')
add_step(doc, 6, 'To view the code: press Alt + F11 to open the VBA Editor.')
add_step(doc, 7, 'To run the macro: View > Macros > select your macro > Run.')
add_tip(doc, 'Save the file as .xlsm (macro-enabled workbook), not .xlsx.')

# --- VLOOKUP ---
doc.add_heading('7. VLOOKUP Formula', level=1)
doc.add_paragraph('VLOOKUP searches the first column of a range and returns a value from another column.')
doc.add_paragraph('Syntax:')
p = doc.add_paragraph('=VLOOKUP(lookup_value, table_array, col_index_num, [range_lookup])')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
add_step(doc, 1, 'Identify your lookup value (the value you want to search for).')
add_step(doc, 2, 'Identify the table range (the lookup value must be in the FIRST column).')
add_step(doc, 3, 'Count which column number contains the value you want returned.')
add_step(doc, 4, 'Use FALSE for exact match (almost always what you want).')
doc.add_paragraph('Example: =VLOOKUP(A2, Sheet2!$A$1:$D$20, 3, FALSE)')
add_tip(doc, 'Use $ signs to lock the table range (absolute reference) so it does not shift when copied.')

# --- Line Chart with Trendline ---
doc.add_heading('8. Line Chart with Trendline, R-Squared, and Forecast', level=1)
add_step(doc, 1, 'Select your time-series data (e.g. months and total sales).')
add_step(doc, 2, 'Insert > Line Chart > Line with Markers.')
add_step(doc, 3, 'Click on the data line in the chart to select it.')
add_step(doc, 4, 'Right-click the line > Add Trendline.')
add_step(doc, 5, 'In the Trendline Options panel:', substeps=[
    'Select "Linear"',
    'Tick "Display Equation on chart"',
    'Tick "Display R-squared value on chart"',
    'Set "Forward" to 2 periods under Forecast',
])
add_step(doc, 6, 'Click Close. The trendline, equation, and R-squared value will appear on the chart.')
add_tip(doc, 'R-squared (R^2) ranges from 0 to 1. A value closer to 1 means the trendline fits the data well.')

# --- IF and SUMIF ---
doc.add_heading('9. IF and SUMIF Functions', level=1)
doc.add_paragraph('IF formula - tests a condition and returns different values:')
p = doc.add_paragraph('=IF(condition, value_if_true, value_if_false)')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
doc.add_paragraph('Example: =IF(D2<20000, "Below Target", "On Target")')
doc.add_paragraph('')
doc.add_paragraph('SUMIF formula - adds values that meet a condition:')
p = doc.add_paragraph('=SUMIF(range, criteria, [sum_range])')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
doc.add_paragraph('Example: =SUMIF(B2:B20, "Product X", D2:D20)')

# --- Formula View ---
doc.add_heading('10. Switching to Formula View', level=1)
add_step(doc, 1, 'Press Ctrl + ` (grave accent key, usually above Tab).')
add_step(doc, 2, 'All cells now show the actual formulas instead of calculated values.')
add_step(doc, 3, 'Widen columns if formulas are cut off (double-click column borders).')
add_step(doc, 4, 'Take your screenshot while in formula view.')
add_step(doc, 5, 'Press Ctrl + ` again to return to normal view.')

excel_docx = os.path.join(BASE, 'Step_By_Step_Excel_Features_Guide.docx')
doc.save(excel_docx)
print(f'Created {excel_docx}')

# ======================================================================
# ACCESS FEATURES GUIDE
# ======================================================================
doc = Document()
doc.add_heading('Step-by-Step Access Features Guide', level=0)
doc.add_paragraph(
    'This guide explains HOW to use the Access features required for the '
    'F1FJ12 assessment. Follow these steps to build your own database.'
)

# --- Creating Tables ---
doc.add_heading('1. Creating Tables', level=1)
add_step(doc, 1, 'Open Microsoft Access and create a new blank database.')
add_step(doc, 2, 'In the default Table1, switch to Design View (right-click tab > Design View).')
add_step(doc, 3, 'Enter your field names, data types, and descriptions:', substeps=[
    'AutoNumber for ID fields (automatically generates unique numbers)',
    'Short Text for names, titles, categories (set Field Size)',
    'Number (Long Integer) for foreign key fields',
    'Currency for money values',
    'Date/Time for dates',
    'Yes/No for true/false fields',
])
add_step(doc, 4, 'Set the primary key: right-click the ID field > Primary Key.')
add_step(doc, 5, 'Save the table with a descriptive name (e.g. Books, Members, Patients).')
add_step(doc, 6, 'Repeat for each table your brief requires.')

# --- Relationships ---
doc.add_heading('2. Creating Relationships', level=1)
add_step(doc, 1, 'Go to Database Tools tab > Relationships.')
add_step(doc, 2, 'If no tables appear, click Show Table and add all your tables.')
add_step(doc, 3, 'Drag the primary key field from one table onto the matching foreign key field in another.')
add_step(doc, 4, 'In the Edit Relationships dialog:', substeps=[
    'Verify the correct fields are shown',
    'Tick "Enforce Referential Integrity"',
    'Optionally tick "Cascade Update" and "Cascade Delete"',
])
add_step(doc, 5, 'Click Create. A line appears connecting the two tables.')
add_step(doc, 6, 'Repeat for all relationships.')
add_tip(doc, 'One-to-Many is the most common relationship type. The "one" side is the table with the primary key.')

# --- Forms ---
doc.add_heading('3. Building a Form', level=1)
add_step(doc, 1, 'Select the table you want to build a form for in the Navigation Pane.')
add_step(doc, 2, 'Go to Create tab > Form. This creates a basic automatic form.')
add_step(doc, 3, 'Switch to Design View to customise the layout.')
add_step(doc, 4, 'To add a Listbox (dropdown with values from another table):', substeps=[
    'In Design View, go to Form Design Tools > Design tab',
    'Click the List Box control in the Controls group',
    'Draw it on the form where you want it',
    'The wizard will guide you: choose "I want the list box to look up values in a table"',
    'Select the source table and the field(s) to display',
])
add_step(doc, 5, 'To add Command Buttons:', substeps=[
    'Click the Button control in the Controls group',
    'Draw it on the form',
    'The wizard offers common actions: "Add Record", "Delete Record", "Close Form", etc.',
    'Choose the action, pick a label or icon, and click Finish',
])
add_step(doc, 6, 'Switch to Form View to test your form.')
add_tip(doc, 'Name your controls clearly (e.g. lstPatientName, btnAddRecord) so your report can reference them.')

# --- Queries ---
doc.add_heading('4. Creating a Query', level=1)
doc.add_paragraph('A Select query retrieves specific data from one or more tables.')
add_step(doc, 1, 'Go to Create tab > Query Design.')
add_step(doc, 2, 'In the Show Table dialog, add the table(s) you need and close.')
add_step(doc, 3, 'Double-click the fields you want to include in the query results.')
add_step(doc, 4, 'To sort results: click the Sort row and choose Ascending or Descending.')
add_step(doc, 5, 'To filter: type criteria in the Criteria row (e.g. "Fiction" or >10).')
add_step(doc, 6, 'Click Run (the red ! icon) to see the results.')

doc.add_heading('5. Creating a Join Query (Multiple Tables)', level=1)
add_step(doc, 1, 'Go to Create tab > Query Design.')
add_step(doc, 2, 'Add both tables. If relationships exist, a join line appears automatically.')
add_step(doc, 3, 'Add fields from BOTH tables to the query grid.')
add_step(doc, 4, 'Set sort order and criteria as needed.')
add_step(doc, 5, 'Run the query. Results will show combined data from both tables.')
add_tip(doc, 'If no join line appears, drag the linking field from one table to the other to create the join manually.')

# --- Entering Sample Data ---
doc.add_heading('6. Entering Sample Data', level=1)
doc.add_paragraph(
    'Before taking screenshots, make sure your tables have enough sample data '
    '(at least 8-10 records per table) so that your forms and queries show '
    'meaningful results.'
)
add_step(doc, 1, 'Open each table in Datasheet View.')
add_step(doc, 2, 'Type sample data directly into the rows. Leave the AutoNumber ID field blank (it fills automatically).')
add_step(doc, 3, 'For foreign key fields, enter values that match existing primary keys in the related table.')
add_tip(doc, 'Enter data in the "one" side tables first (e.g. Books, Members) before entering data in junction tables (e.g. Borrowing).')

access_docx = os.path.join(BASE, 'Step_By_Step_Access_Features_Guide.docx')
doc.save(access_docx)
print(f'Created {access_docx}')

# ======================================================================
# Export both to PDF via Word COM
# ======================================================================
print('\nExporting to PDF...')
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
time.sleep(1)

import pythoncom
import win32com.client

for docx_path, pdf_name in [
    (excel_docx, 'Step_By_Step_Excel_Features_Guide.pdf'),
    (access_docx, 'Step_By_Step_Access_Features_Guide.pdf'),
]:
    pdf_path = os.path.join(BASE, pdf_name)
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = False
    try:
        d = word.Documents.Open(os.path.abspath(docx_path))
        d.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)
        d.Close(False)
        sz = os.path.getsize(pdf_path)
        print(f'Exported {pdf_name} ({sz:,} bytes)')
    finally:
        word.Quit()
        pythoncom.CoUninitialize()
    # Clean up intermediate DOCX
    os.remove(docx_path)
    print(f'Removed intermediate {os.path.basename(docx_path)}')

print('\nAll guides created successfully.')
