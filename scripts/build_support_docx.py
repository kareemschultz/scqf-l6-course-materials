"""Build all DOCX templates for the support pack."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'SCQF_L6_SUPPORT_PACK')

def style_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_placeholder(doc, text="[Write your explanation here...]"):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True

def add_table_with_headers(doc, headers, rows=3, placeholder="[...]"):
    table = doc.add_table(rows=1 + rows, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r in range(1, rows + 1):
        for c in range(len(headers)):
            table.rows[r].cells[c].text = placeholder
    return table

# ======================================================================
# 1. F1FJ12 - Sample Access Table Structure
# ======================================================================
doc = Document()
doc.add_heading('Access Database - Table Structure Reference', level=0)
doc.add_paragraph(
    'This document shows example table structures for an Access database. '
    'Use this as a guide to plan your own tables, fields, and data types. '
    'You must create your own tables with your own field names.'
)

style_heading(doc, 'Example Table 1: Books')
add_table_with_headers(doc,
    ['Field Name', 'Data Type', 'Description', 'Primary Key?'],
    rows=5, placeholder='')
t = doc.tables[-1]
data = [
    ('BookID', 'AutoNumber', 'Unique identifier for each book', 'Yes'),
    ('Title', 'Short Text (100)', 'Book title', 'No'),
    ('Author', 'Short Text (50)', 'Author full name', 'No'),
    ('Genre', 'Short Text (30)', 'Category (e.g. Fiction, Science)', 'No'),
    ('Price', 'Currency', 'Retail price', 'No'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

style_heading(doc, 'Example Table 2: Members')
add_table_with_headers(doc,
    ['Field Name', 'Data Type', 'Description', 'Primary Key?'],
    rows=5, placeholder='')
t = doc.tables[-1]
data = [
    ('MemberID', 'AutoNumber', 'Unique member identifier', 'Yes'),
    ('FirstName', 'Short Text (30)', 'Member first name', 'No'),
    ('Surname', 'Short Text (30)', 'Member surname', 'No'),
    ('Email', 'Short Text (80)', 'Contact email', 'No'),
    ('JoinDate', 'Date/Time', 'Date membership started', 'No'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

style_heading(doc, 'Example Table 3: Borrowing (Junction Table)')
add_table_with_headers(doc,
    ['Field Name', 'Data Type', 'Description', 'Foreign Key?'],
    rows=4, placeholder='')
t = doc.tables[-1]
data = [
    ('BorrowID', 'AutoNumber', 'Unique borrow record', 'No (PK)'),
    ('MemberID', 'Number (Long)', 'Links to Members table', 'Yes'),
    ('BookID', 'Number (Long)', 'Links to Books table', 'Yes'),
    ('BorrowDate', 'Date/Time', 'Date book was borrowed', 'No'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

style_heading(doc, 'Relationships')
doc.add_paragraph(
    'In Access, you create relationships between tables using the Relationships '
    'window (Database Tools > Relationships). Drag the primary key field from '
    'one table onto the matching foreign key field in another table.'
)
doc.add_paragraph('Typical relationships for the above tables:')
doc.add_paragraph('Members.MemberID (one) -> Borrowing.MemberID (many)', style='List Bullet')
doc.add_paragraph('Books.BookID (one) -> Borrowing.BookID (many)', style='List Bullet')
doc.add_paragraph(
    'Tick "Enforce Referential Integrity" when creating relationships to '
    'prevent orphan records.'
)

style_heading(doc, 'Your Own Tables')
doc.add_paragraph(
    'Your brief will specify what tables to create. Plan your tables on paper '
    'first: list the fields, decide on data types, identify the primary key, '
    'and work out which fields link between tables.'
)
add_placeholder(doc, '[Plan your table structures here before building in Access...]')

out = os.path.join(BASE, 'F1FJ12_Spreadsheet_Database', 'Sample_Access_Table_Structure.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 2. F1FE12 - Basic Welcome Pack Template
# ======================================================================
doc = Document()
doc.add_heading('[Your Conference Name Here]', level=0)
doc.add_paragraph('[Your Conference Subtitle]')
doc.add_paragraph('[Venue Name]')
doc.add_paragraph('[Date]')
doc.add_paragraph('')

style_heading(doc, 'Page 1: Table of Contents')
add_placeholder(doc,
    '[Insert an automatic Table of Contents here: References > Table of Contents. '
    'Make sure you have applied Heading 1 and Heading 2 styles to your section '
    'headings BEFORE inserting the TOC.]')

doc.add_page_break()
style_heading(doc, 'Page 2: Welcome and Conference Overview')
add_placeholder(doc, '[Insert your conference logo here with Tight or Square text wrapping.]')
add_placeholder(doc,
    '[Write a welcome paragraph introducing the conference, its purpose, '
    'and what delegates can expect from the day.]')

style_heading(doc, 'Conference Agenda', level=2)
doc.add_paragraph(
    '[Create a table with two columns: Time (3cm) and Session (10cm). '
    'Apply shading to the header row. Use single-line borders.]'
)
add_table_with_headers(doc, ['Time', 'Session'], rows=8, placeholder='[...]')

doc.add_page_break()
style_heading(doc, 'Page 3: Venue and Practical Information')
style_heading(doc, 'Venue Location', level=2)
add_placeholder(doc, '[Describe how to get to the venue, parking, public transport.]')
style_heading(doc, 'Wi-Fi Information', level=2)
add_placeholder(doc, '[Provide network name and password.]')
style_heading(doc, 'Contact Information', level=2)
add_placeholder(doc, '[Provide organiser contact details.]')

doc.add_page_break()
style_heading(doc, 'Page 4: Delegate Badges (Mail Merge)')
doc.add_paragraph(
    '[This page will be replaced by your mail merge output. Set up the badge '
    'template with merge fields like <<FirstName>> <<Surname>> <<Organisation>> '
    'and then finish the merge to generate one badge per delegate.]'
)

out = os.path.join(BASE, 'F1FE12_Word_Presentation', 'Basic_Welcome_Pack_Template.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 3. J22976 - Essay Structure Outline
# ======================================================================
doc = Document()
doc.add_heading('J229 76 - Understanding Business: Structure Outline', level=0)
doc.add_paragraph(
    'This is a blank structural outline. Fill in each section with your '
    'own research, analysis, and Harvard-referenced sources. Use your own '
    'chosen organisations throughout.'
)

doc.add_page_break()
style_heading(doc, 'TASK A')

style_heading(doc, '1.1.1 Comparison of Organisations', level=2)
doc.add_paragraph(
    '[Choose THREE organisations from DIFFERENT sectors (e.g. private, public, '
    'third sector). Compare them across: scale, structure, specialisation, '
    'regulation, and capital requirements.]'
)
add_placeholder(doc)

style_heading(doc, '1.1.2 Types of Business Ownership', level=2)
doc.add_paragraph('[Compare at least three types of ownership structure.]')
add_placeholder(doc)

style_heading(doc, '1.1.3 Organic vs Inorganic Growth', level=2)
doc.add_paragraph('[Explain both growth types. Provide real business examples for each.]')
add_placeholder(doc)

style_heading(doc, '1.2.1 Business Objectives', level=2)
doc.add_paragraph(
    '[Analyse objectives of your chosen organisations. Explain WHY each '
    'objective was important at the time it was adopted.]'
)
add_placeholder(doc)

style_heading(doc, '1.3.1 Organisational Structures', level=2)
doc.add_paragraph(
    '[Explain hierarchical, functional, and matrix structures. Apply each '
    'to a real organisation.]'
)
add_placeholder(doc)

style_heading(doc, '1.3.2 Decision-Making', level=2)
doc.add_paragraph(
    '[Explain strategic, tactical, and operational decisions. '
    'Include a SWOT analysis and decision tree example.]'
)
add_placeholder(doc)

style_heading(doc, 'TASK B')

style_heading(doc, '2.1 Internal Factors', level=2)
doc.add_paragraph('[Analyse internal factors affecting two organisations.]')
add_placeholder(doc)

style_heading(doc, '2.2 External Factors (PESTEC)', level=2)
doc.add_paragraph(
    '[Analyse external factors using the PESTEC framework: Political, Economic, '
    'Social, Technological, Environmental, Competitive.]'
)
add_placeholder(doc)

style_heading(doc, '2.3 Stakeholders', level=2)
doc.add_paragraph('[Identify stakeholders, their influence, and potential conflicts.]')
add_placeholder(doc)

style_heading(doc, 'Reference List')
add_placeholder(doc, '[List all Harvard references here in alphabetical order.]')

style_heading(doc, 'Mapping Table')
add_table_with_headers(doc,
    ['Task', 'Marking Criteria', 'Evidence / Page Ref'],
    rows=10, placeholder='[...]')

out = os.path.join(BASE, 'J22976_Understanding_Business', 'Essay_Structure_Outline.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 4. J22976 - Comparison Table Template
# ======================================================================
doc = Document()
doc.add_heading('Comparison Table Templates', level=0)
doc.add_paragraph(
    'Use these blank tables to organise your comparative analysis. '
    'Replace all placeholder text with your own research.'
)

style_heading(doc, 'Table 1: Organisation Comparison')
add_table_with_headers(doc,
    ['Criteria', 'Organisation 1\n[Name]', 'Organisation 2\n[Name]', 'Organisation 3\n[Name]'],
    rows=6, placeholder='[...]')
t = doc.tables[-1]
criteria = ['Sector', 'Scale (size)', 'Structure', 'Specialisation', 'Regulation', 'Capital']
for i, c in enumerate(criteria):
    t.rows[i+1].cells[0].text = c

doc.add_paragraph('')
style_heading(doc, 'Table 2: Ownership Type Comparison')
add_table_with_headers(doc,
    ['Feature', 'Type 1\n[e.g. Private Ltd]', 'Type 2\n[e.g. Public Ltd]', 'Type 3\n[e.g. Franchise]'],
    rows=5, placeholder='[...]')
t = doc.tables[-1]
features = ['Ownership', 'Liability', 'Decision-making', 'Raising capital', 'Example business']
for i, f in enumerate(features):
    t.rows[i+1].cells[0].text = f

doc.add_paragraph('')
style_heading(doc, 'Table 3: Stakeholder Analysis')
add_table_with_headers(doc,
    ['Stakeholder', 'Interest', 'Influence Level', 'Potential Conflict'],
    rows=6, placeholder='[...]')
t = doc.tables[-1]
stk = ['Shareholders', 'Employees', 'Customers', 'Government', 'Suppliers', 'Local Community']
for i, s in enumerate(stk):
    t.rows[i+1].cells[0].text = s

out = os.path.join(BASE, 'J22976_Understanding_Business', 'Comparison_Table_Template.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 5. J22A76 - Section Structure Template
# ======================================================================
doc = Document()
doc.add_heading('J22A 76 - Management of People and Finance: Structure', level=0)
doc.add_paragraph(
    'This template provides section headings and placeholders only. '
    'Write all content yourself using your own research and examples.'
)

doc.add_page_break()
style_heading(doc, 'TASK A: Human Resource Management')

style_heading(doc, '1.1 Three HRM Approaches', level=2)
doc.add_paragraph(
    '[Explain three different approaches to human resource management. '
    'For each approach, explain what it involves and how it contributes '
    'to effective HR management.]'
)
doc.add_paragraph('Approach 1: [Name]', style='List Number')
add_placeholder(doc)
doc.add_paragraph('Approach 2: [Name]', style='List Number')
add_placeholder(doc)
doc.add_paragraph('Approach 3: [Name]', style='List Number')
add_placeholder(doc)

style_heading(doc, "1.2 Maslow's Hierarchy of Needs", level=2)
doc.add_paragraph(
    '[Explain the theory. Include YOUR OWN diagram (not copied from the internet). '
    'Apply each level to a workplace context with examples.]'
)
add_placeholder(doc, '[Insert your Maslow diagram here.]')
doc.add_paragraph('Level 1 - Physiological Needs:', style='List Number')
add_placeholder(doc, '[Explain and give workplace example...]')
doc.add_paragraph('Level 2 - Safety Needs:', style='List Number')
add_placeholder(doc)
doc.add_paragraph('Level 3 - Social/Belonging Needs:', style='List Number')
add_placeholder(doc)
doc.add_paragraph('Level 4 - Esteem Needs:', style='List Number')
add_placeholder(doc)
doc.add_paragraph('Level 5 - Self-Actualisation:', style='List Number')
add_placeholder(doc)

style_heading(doc, '1.3 Industrial Action', level=2)
doc.add_paragraph('[Explain five forms of industrial action and their impact on businesses.]')
for form in ['Strike action', 'Go-slow', 'Work-to-rule', 'Lockout', 'Picketing']:
    doc.add_paragraph(f'{form}:', style='List Bullet')
    add_placeholder(doc)

style_heading(doc, '1.4 Employment Legislation', level=2)
doc.add_paragraph('[Explain the following areas of employment law:]')
for law in ['Equality and diversity legislation',
            'Health and Safety at Work Act',
            'National Minimum Wage / Living Wage',
            'Working Time Regulations',
            'Dismissal and Redundancy procedures']:
    doc.add_paragraph(f'{law}:', style='List Bullet')
    add_placeholder(doc)

doc.add_page_break()
style_heading(doc, 'TASK B: Finance')

style_heading(doc, '2.1 Sources of Finance', level=2)
doc.add_paragraph('[Explain three sources of finance - include both long-term and short-term.]')
add_placeholder(doc)

style_heading(doc, '2.2 Purposes of Financial Statements', level=2)
doc.add_paragraph('[Explain five purposes of financial statements for business decision-making.]')
add_placeholder(doc)

style_heading(doc, '2.3 Accounting Ratios', level=2)
doc.add_paragraph(
    '[Explain five accounting ratios. For each: state the formula, explain what it '
    'measures, and show how it helps analyse business performance.]'
)
add_placeholder(doc)

style_heading(doc, 'Reference List')
add_placeholder(doc, '[Harvard references in alphabetical order...]')

style_heading(doc, 'Mapping Table')
add_table_with_headers(doc,
    ['Task', 'Marking Criteria', 'Evidence / Page Ref'],
    rows=8, placeholder='[...]')

out = os.path.join(BASE, 'J22A76_Management_People_Finance', 'Section_Structure_Template.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 6. HE9E46 - Lifecycle Table Template
# ======================================================================
doc = Document()
doc.add_heading('Business Lifecycle Table Template', level=0)
doc.add_paragraph(
    'Complete this table with your own research. Identify characteristics '
    'and real business examples for each stage. Do not copy examples from '
    'classmates.'
)

style_heading(doc, 'Business Lifecycle Stages')
add_table_with_headers(doc,
    ['Stage', 'Characteristic 1', 'Characteristic 2', 'Characteristic 3', 'Real Example'],
    rows=5, placeholder='[...]')
t = doc.tables[-1]
stages = ['1. Seed / Start-up', '2. Growth', '3. Maturity', '4. Decline', '5. Renewal / Exit']
for i, s in enumerate(stages):
    t.rows[i+1].cells[0].text = s

doc.add_paragraph('')
style_heading(doc, 'Objectives Per Lifecycle Stage')
add_table_with_headers(doc,
    ['Stage', 'Objective 1', 'Objective 2'],
    rows=5, placeholder='[...]')
t = doc.tables[-1]
for i, s in enumerate(stages):
    t.rows[i+1].cells[0].text = s

out = os.path.join(BASE, 'HE9E46_Contemporary_Business_Issues', 'Lifecycle_Table_Template.docx')
doc.save(out)
print(f'Created {out}')

# ======================================================================
# 7. HE9E46 - Strategy Comparison Template
# ======================================================================
doc = Document()
doc.add_heading('Strategy Comparison Template', level=0)
doc.add_paragraph(
    'Use this template to structure your analysis of two business strategies. '
    'Choose strategies relevant to SMEs in your chosen sector.'
)

style_heading(doc, 'Strategy 1: [Name your first strategy]')
add_table_with_headers(doc,
    ['Aspect', 'Details'],
    rows=5, placeholder='[...]')
t = doc.tables[-1]
aspects = ['Strategy description', 'Why suitable for SMEs', 'Expected benefits',
           'Potential risks', 'Real-world example']
for i, a in enumerate(aspects):
    t.rows[i+1].cells[0].text = a

doc.add_paragraph('')
style_heading(doc, 'Strategy 2: [Name your second strategy]')
add_table_with_headers(doc,
    ['Aspect', 'Details'],
    rows=5, placeholder='[...]')
t = doc.tables[-1]
for i, a in enumerate(aspects):
    t.rows[i+1].cells[0].text = a

doc.add_paragraph('')
style_heading(doc, 'Comparison Summary')
add_placeholder(doc,
    '[Write a paragraph comparing the two strategies. Which is more suitable '
    'for your chosen sector? What are the trade-offs? Use evidence to support '
    'your evaluation.]')

out = os.path.join(BASE, 'HE9E46_Contemporary_Business_Issues', 'Strategy_Comparison_Template.docx')
doc.save(out)
print(f'Created {out}')

print('\nAll DOCX templates created successfully.')
