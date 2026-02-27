#!/usr/bin/env python3
"""
build_unit_d.py
Builds the Unit D (J22A 76) Management of People and Finance report DOCX
and exports it to PDF via Word COM automation.

Student : 252IFCBR0596 | Kareem Nurw Jason Schultz
Unit    : J22A 76 - Management of People and Finance
"""

import os
import sys
import subprocess
import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────
BASE_DIR = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials"

DOCX_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "source_files",
                        "J22A76_Management_People_Finance")
PDF_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "pdf_submissions")

DOCX_PATH = os.path.join(DOCX_DIR, "J22A76_Report.docx")
PDF_PATH = os.path.join(PDF_DIR, "J22A76_Management_People_Finance.pdf")

os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

STUDENT_NAME = "Kareem Nurw Jason Schultz"
STUDENT_ID = "252IFCBR0596"
UNIT_CODE = "J22A 76"
UNIT_TITLE = "Management of People and Finance"
COLLEGE = "JAIN College"
TOTAL_MARKS = 100
FOOTER_TEXT = (f"{STUDENT_ID} | {STUDENT_NAME} | {UNIT_CODE} | "
               f"{UNIT_TITLE} | Page ")

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def make_table_header_row(table, texts, bg="003366", fg="FFFFFF"):
    """Style the first row of a table as a header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, bg)


def apply_footer(section):
    """Add the standard footer with page number to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2._r.append(fld_char_begin)
    run3 = p.add_run()
    run3._r.append(instr)
    run4 = p.add_run()
    run4._r.append(fld_char_end)


# ──────────────────────────────────────────────
# Document creation
# ──────────────────────────────────────────────

doc = Document()

# --- Default style tweaks ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    hs.font.name = "Calibri"

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
apply_footer(section)


def add_body(text):
    """Add a body paragraph with the given text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_page_break():
    doc.add_page_break()


def add_bullet(text, level=0):
    """Add a bulleted list item."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_bold_body(bold_part, rest):
    """Add a paragraph with an initial bold run followed by normal text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(rest)
    return p


# ═══════════════════════════════════════════════
# 1. COVER PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{UNIT_CODE} - {UNIT_TITLE}")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

doc.add_paragraph()

for line, size, bold in [
    (STUDENT_NAME, 16, True),
    (f"Student ID: {STUDENT_ID}", 13, False),
    (COLLEGE, 14, True),
    (f"Total Marks: {TOTAL_MARKS}", 12, False),
    ("February 2026", 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

add_page_break()

# ═══════════════════════════════════════════════
# 2. DECLARATION OF ORIGINALITY
# ═══════════════════════════════════════════════
doc.add_heading("Declaration of Originality", level=1)
doc.add_paragraph()

add_body(
    "I, Kareem Nurw Jason Schultz (Student ID: 252IFCBR0596), hereby declare that "
    "this assignment for J22A 76 - Management of People and Finance is entirely my "
    "own work. I have not copied from any other student's work or from any other "
    "source except where due acknowledgement is made explicitly in the text, nor has "
    "any part been written for me by another person."
)

add_body(
    "I understand that plagiarism is a serious academic offence and that the penalties "
    "for submitting plagiarised work can include loss of marks, module failure, or "
    "further disciplinary action as determined by JAIN College's academic integrity policy."
)

add_body(
    "All sources of information have been properly cited and referenced using the "
    "Harvard referencing system. Where I have used direct quotations or closely "
    "paraphrased the work of others, this has been clearly indicated."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
r = p.add_run("Signed: ")
r.bold = True
p.add_run("Kareem Nurw Jason Schultz")

p = doc.add_paragraph()
r = p.add_run("Date: ")
r.bold = True
p.add_run("February 2026")

p = doc.add_paragraph()
r = p.add_run("Student ID: ")
r.bold = True
p.add_run(STUDENT_ID)

add_page_break()

# ═══════════════════════════════════════════════
# 3. TABLE OF CONTENTS
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
doc.add_paragraph()

toc_entries = [
    ("Declaration of Originality", "2"),
    ("Table of Contents", "3"),
    ("TASK A: Human Resource Management (60 marks)", ""),
    ("    Section 1.1 - Three HRM Approaches", "4"),
    ("    Section 1.2 - Maslow's Hierarchy of Needs", "8"),
    ("    Section 1.3 - Five Forms of Industrial Action", "11"),
    ("    Section 1.4 - Employment Legislation", "16"),
    ("TASK B: Finance (40 marks)", ""),
    ("    Section 2.1 - Three Sources of Finance", "21"),
    ("    Section 2.2 - Purposes of Financial Statements", "24"),
    ("    Section 2.3 - Five Accounting Ratios", "27"),
    ("Reference List", "31"),
    ("Mapping Table", "33"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if entry.startswith("    "):
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(entry.strip())
        run.font.size = Pt(11)
    elif entry.startswith("TASK"):
        run = p.add_run(entry)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    else:
        run = p.add_run(entry)
        run.font.size = Pt(11)
    if page:
        tab_run = p.add_run(f"  {'.' * 40}  {page}")
        tab_run.font.size = Pt(11)
        tab_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_page_break()

# ═══════════════════════════════════════════════
# TASK A: HUMAN RESOURCE MANAGEMENT (60 marks)
# ═══════════════════════════════════════════════

doc.add_heading("TASK A: Human Resource Management", level=1)
p = doc.add_paragraph()
r = p.add_run("Total: 60 marks")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

add_body(
    "This section examines the key functions of Human Resource Management (HRM) "
    "within modern organisations. HRM encompasses all aspects of managing people "
    "in the workplace, from recruitment through to departure, and plays a critical "
    "role in ensuring that organisations attract, develop, and retain talented "
    "employees (Armstrong and Taylor, 2020). The following sections explore three "
    "HRM approaches, motivational theory, industrial relations, and employment "
    "legislation in detail."
)

# ───────────────────────────────────────────────
# Section 1.1 - Three HRM Approaches (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 1.1 - Three HRM Approaches and How They Contribute "
                "to Effective HR Management", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Human Resource Management involves a range of strategic and operational "
    "approaches designed to maximise the contribution of employees to "
    "organisational success. This section examines three key HRM approaches: "
    "Recruitment and Selection, Training and Development, and Performance "
    "Management. Each approach is evaluated with real-world examples to "
    "demonstrate how they contribute to effective workforce management."
)

# --- 1.1.1 Recruitment & Selection ---
doc.add_heading("1.1.1 Recruitment and Selection", level=3)

add_body(
    "Recruitment and selection is one of the most fundamental HRM functions, "
    "as it determines the quality and suitability of the people who join an "
    "organisation. Recruitment refers to the process of attracting a pool of "
    "qualified candidates for a vacant position, while selection involves "
    "choosing the most appropriate candidate from that pool (CIPD, 2023). "
    "An effective recruitment and selection process typically begins with a "
    "thorough job analysis, which identifies the tasks, responsibilities, "
    "and requirements of the role. From this analysis, a job description is "
    "produced, outlining the duties and reporting relationships, along with "
    "a person specification, which details the essential and desirable "
    "qualifications, skills, experience, and personal attributes required "
    "of the successful candidate."
)

add_body(
    "The recruitment process can be conducted internally, through promotions "
    "or transfers, or externally, using methods such as online job boards, "
    "recruitment agencies, social media, and university careers services. "
    "For example, Google is well known for its rigorous recruitment strategy, "
    "which includes multiple rounds of structured interviews, coding "
    "challenges for technical roles, and behavioural assessments designed "
    "to evaluate cultural fit (Schmidt and Hunter, 1998). Google receives "
    "over three million applications per year but hires only a fraction, "
    "reflecting the company's commitment to selecting the very best talent "
    "(Bock, 2015)."
)

add_body(
    "Selection methods commonly include interviews (structured or unstructured), "
    "psychometric tests, skills assessments, assessment centres, and references. "
    "Assessment centres, in particular, are considered one of the most valid "
    "selection methods because they use a combination of exercises, such as "
    "group discussions, presentations, in-tray exercises, and role plays, to "
    "evaluate candidates against multiple competencies simultaneously "
    "(Robertson and Smith, 2001). The NHS, for example, uses assessment "
    "centres extensively when selecting senior managers and clinical leaders, "
    "ensuring that candidates demonstrate both technical competence and the "
    "interpersonal skills needed to lead diverse teams."
)

add_body(
    "Effective recruitment and selection contributes to HR management in several "
    "important ways. First, hiring the right person for the role reduces staff "
    "turnover, which is both costly and disruptive. The CIPD (2023) estimates "
    "that the average cost of replacing an employee is approximately 6,000 "
    "pounds when factoring in advertising, agency fees, interview time, "
    "training, and lost productivity. Second, good recruitment builds "
    "organisational capability by ensuring a diverse workforce with the skills "
    "needed to meet strategic objectives. Third, a fair and transparent "
    "recruitment process supports compliance with the Equality Act 2010, "
    "reducing the risk of discrimination claims and enhancing the "
    "organisation's reputation as an employer of choice."
)

# --- 1.1.2 Training & Development ---
doc.add_heading("1.1.2 Training and Development", level=3)

add_body(
    "Training and development is a core HRM function concerned with enhancing "
    "employees' knowledge, skills, and competencies to improve both individual "
    "performance and organisational effectiveness (Armstrong and Taylor, 2020). "
    "Training refers to planned activities designed to help employees acquire "
    "specific job-related skills, while development is a broader, longer-term "
    "process focused on personal and professional growth. Both are essential for "
    "maintaining a skilled, adaptable, and motivated workforce."
)

add_body(
    "There are several types of training commonly used in organisations. "
    "Induction training is provided to new employees to familiarise them with "
    "the organisation's culture, policies, health and safety procedures, and "
    "their specific role. A well-structured induction programme helps new "
    "starters settle in quickly and become productive sooner. On-the-job "
    "training involves learning whilst performing the actual work, often "
    "through coaching, mentoring, job rotation, or shadowing experienced "
    "colleagues. This approach is cost-effective and immediately relevant, "
    "as the employee learns in the real working environment. Off-the-job "
    "training takes place away from the workplace, through external courses, "
    "workshops, conferences, or online learning platforms. Although more "
    "expensive, off-the-job training provides access to specialist knowledge "
    "and allows employees to learn without the distractions of daily work."
)

add_body(
    "Continuing Professional Development (CPD) is an ongoing process where "
    "employees take responsibility for maintaining and updating their skills "
    "throughout their careers. Many professions, such as teaching, nursing, "
    "and accounting, require evidence of CPD as a condition of registration. "
    "For example, the John Lewis Partnership invests heavily in employee "
    "development, offering apprenticeships, management development programmes, "
    "and access to further education. As a partnership, John Lewis recognises "
    "that its employees (known as Partners) are also co-owners of the "
    "business, which creates a strong incentive to invest in their growth "
    "(John Lewis Partnership, 2023)."
)

add_body(
    "Training and development contributes to effective HR management by "
    "improving productivity, as employees who are properly trained make fewer "
    "mistakes and work more efficiently. It also enhances employee satisfaction "
    "and retention, since workers who feel that their employer is investing in "
    "their future are more likely to remain loyal. The CIPD (2023) reports "
    "that organisations with strong learning cultures experience 30 to 50 "
    "per cent higher retention rates than those that neglect training. "
    "Furthermore, training supports succession planning by developing internal "
    "talent to fill future leadership positions, reducing reliance on costly "
    "external recruitment."
)

# --- 1.1.3 Performance Management ---
doc.add_heading("1.1.3 Performance Management", level=3)

add_body(
    "Performance management is a systematic process that seeks to align "
    "individual employee goals with the broader strategic objectives of the "
    "organisation (Bach and Edwards, 2013). Rather than being a one-off annual "
    "event, effective performance management is a continuous cycle that involves "
    "setting clear expectations, monitoring progress, providing feedback, and "
    "supporting development. It is one of the most important HRM approaches "
    "because it directly links employee effort to organisational outcomes."
)

add_body(
    "The performance management process typically begins with goal setting, "
    "where managers and employees agree on specific, measurable, achievable, "
    "relevant, and time-bound (SMART) objectives for the review period. "
    "Performance is then monitored throughout the year through regular "
    "one-to-one meetings, informal check-ins, and the tracking of Key "
    "Performance Indicators (KPIs). KPIs are quantifiable metrics that "
    "measure how effectively an employee is achieving their objectives, "
    "such as sales targets, customer satisfaction scores, or project "
    "completion rates."
)

add_body(
    "Formal performance appraisals are typically conducted annually or "
    "biannually, providing an opportunity for a structured review of "
    "achievement. Many organisations now use 360-degree feedback, where "
    "performance input is gathered not just from the line manager but also "
    "from peers, subordinates, and sometimes customers. This provides a more "
    "rounded and balanced view of an employee's strengths and areas for "
    "improvement. For example, the NHS uses 360-degree feedback as part of "
    "its medical revalidation process, where doctors must demonstrate that "
    "they continue to meet professional standards (General Medical Council, "
    "2022)."
)

add_body(
    "Performance-related pay (PRP) is another tool used within performance "
    "management, linking financial rewards to the achievement of agreed "
    "targets. While PRP can be a powerful motivator, it must be implemented "
    "fairly and transparently to avoid perceptions of bias. Google, for "
    "instance, uses a combination of base salary, performance bonuses, and "
    "stock options to reward high performers, with performance ratings "
    "influencing the size of bonus payments (Bock, 2015)."
)

add_body(
    "Performance management contributes to effective HRM by ensuring that "
    "every employee understands what is expected of them and receives the "
    "support and feedback needed to achieve their goals. It enables early "
    "identification of underperformance so that appropriate interventions, "
    "such as additional training or coaching, can be put in place. It also "
    "supports talent identification by highlighting high-performing employees "
    "who may be suitable for promotion or leadership development programmes. "
    "Ultimately, performance management creates a culture of accountability "
    "and continuous improvement that benefits both employees and the "
    "organisation as a whole."
)

add_page_break()

# ───────────────────────────────────────────────
# Section 1.2 - Maslow's Hierarchy of Needs (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 1.2 - Maslow's Hierarchy of Needs", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Motivation is a central concern of HRM because the level of effort and "
    "commitment that employees bring to their work has a direct impact on "
    "organisational performance. One of the most widely recognised theories "
    "of motivation is Abraham Maslow's Hierarchy of Needs, first published "
    "in 1943. Maslow proposed that human needs are arranged in a hierarchical "
    "structure, often depicted as a pyramid, with the most basic physiological "
    "needs at the base and the need for self-actualisation at the apex. "
    "According to Maslow, individuals are motivated to satisfy lower-level "
    "needs before they can focus on higher-level ones (Maslow, 1943)."
)

# Pyramid diagram description
doc.add_heading("Diagram: Maslow's Hierarchy of Needs Pyramid", level=3)

# Build a table to visually represent the pyramid
pyramid_table = doc.add_table(rows=5, cols=1)
pyramid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pyramid_data = [
    ("Self-Actualisation", "E8D5B7"),
    ("Esteem Needs", "C9DAF8"),
    ("Belongingness and Love Needs", "D5E8D4"),
    ("Safety Needs", "FCE5CD"),
    ("Physiological Needs", "F4CCCC"),
]
for i, (label, colour) in enumerate(pyramid_data):
    cell = pyramid_table.rows[i].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    set_cell_shading(cell, colour)

add_body(
    "Figure 1: Maslow's Hierarchy of Needs, depicted as a five-level pyramid. "
    "The base represents physiological needs (most fundamental) and the peak "
    "represents self-actualisation (highest-order need)."
)

doc.add_heading("The Five Levels Explained", level=3)

add_bold_body(
    "Level 1 - Physiological Needs: ",
    "These are the most basic human needs required for survival, including "
    "food, water, shelter, warmth, and rest. In the workplace, physiological "
    "needs are addressed through fair wages that enable employees to afford "
    "housing, food, and basic living expenses. Some employers go further "
    "by providing on-site facilities. Google, for example, offers free meals, "
    "snacks, and beverages at all of its offices, directly satisfying "
    "physiological needs and removing barriers to productivity (Bock, 2015). "
    "Comfortable working conditions, adequate heating and lighting, and "
    "access to rest areas also fall within this category."
)

add_bold_body(
    "Level 2 - Safety Needs: ",
    "Once physiological needs are met, individuals seek security and "
    "stability. In the workplace, safety needs include job security, safe "
    "working conditions, health insurance, pension schemes, and protection "
    "from workplace hazards. The NHS is a strong example of an employer that "
    "addresses safety needs, as it offers permanent contracts, comprehensive "
    "pension schemes (the NHS Pension Scheme is one of the most generous in "
    "the UK), and strict compliance with health and safety regulations "
    "(NHS Employers, 2023). The Health and Safety at Work Act 1974 places "
    "a legal duty on all employers to ensure, so far as is reasonably "
    "practicable, the health, safety, and welfare of their employees, which "
    "directly supports this level of the hierarchy."
)

add_bold_body(
    "Level 3 - Belongingness and Love Needs: ",
    "This level relates to the human desire for social connection, "
    "friendship, and a sense of belonging. In the workplace, organisations "
    "can address these needs through team-building activities, social events, "
    "mentoring programmes, and inclusive workplace cultures. Effective "
    "communication between management and staff, opportunities for "
    "collaboration, and supportive line management all contribute to "
    "employees feeling that they are valued members of a team. Organisations "
    "that foster strong social bonds tend to experience lower absenteeism "
    "and higher levels of engagement (CIPD, 2023)."
)

add_bold_body(
    "Level 4 - Esteem Needs: ",
    "Esteem needs include the desire for recognition, respect, status, "
    "and a sense of achievement. Employers can address esteem needs through "
    "formal recognition schemes, such as employee of the month awards, "
    "performance bonuses, and public acknowledgement of achievements. "
    "Promotion opportunities, job titles, and increased responsibility also "
    "satisfy esteem needs. The John Lewis Partnership is a particularly "
    "relevant example, as its partnership model gives all employees a share "
    "in the profits and a voice in how the business is run. This ownership "
    "structure satisfies esteem needs by making employees feel respected "
    "and valued as stakeholders, not merely workers (John Lewis Partnership, "
    "2023)."
)

add_bold_body(
    "Level 5 - Self-Actualisation: ",
    "The highest level of the hierarchy represents the need to fulfil one's "
    "potential and achieve personal growth. In the workplace, self-actualisation "
    "can be supported through challenging work assignments, opportunities for "
    "creativity and innovation, leadership development programmes, and support "
    "for further education. Google's famous '20% time' policy, which "
    "historically allowed engineers to spend one day per week on personal "
    "projects, is a well-known example of how an employer can encourage "
    "self-actualisation by giving employees the freedom to pursue their "
    "passions (Bock, 2015). Products such as Gmail and Google News originated "
    "from this policy, demonstrating that supporting self-actualisation can "
    "also generate significant business value."
)

doc.add_heading("Evaluation of Maslow's Theory", level=3)

add_body(
    "Maslow's Hierarchy of Needs has several strengths as a motivational "
    "framework. It is intuitive and easy to understand, making it accessible "
    "to managers who may not have a background in psychology. It provides a "
    "useful structure for thinking about the different factors that motivate "
    "employees and highlights the importance of meeting basic needs before "
    "expecting higher-level engagement. Many organisations use the hierarchy, "
    "either explicitly or implicitly, when designing their reward and benefit "
    "packages, ensuring that they address needs at every level."
)

add_body(
    "However, the theory also has significant limitations. The rigid "
    "hierarchical structure has been criticised because it suggests that "
    "needs must be satisfied in a strict sequential order, which does not "
    "always reflect reality. For example, a poorly paid artist may prioritise "
    "self-actualisation over financial security, and many people pursue social "
    "belonging and esteem simultaneously rather than sequentially. Wahba and "
    "Bridwell (1976) conducted a comprehensive review of the research evidence "
    "and found little empirical support for the strict hierarchy. Furthermore, "
    "the theory is culturally biased, as it was developed in a Western context "
    "and may not apply equally to collectivist cultures where group needs take "
    "precedence over individual self-actualisation (Hofstede, 2001). Despite "
    "these limitations, Maslow's theory remains one of the most widely taught "
    "and referenced models in HRM and continues to provide a valuable starting "
    "point for understanding workplace motivation."
)

add_page_break()

# ───────────────────────────────────────────────
# Section 1.3 - Five Forms of Industrial Action (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 1.3 - Five Forms of Industrial Action and "
                "Their Impact", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Industrial action refers to measures taken by workers, usually organised "
    "through trade unions, to put pressure on their employer during a dispute "
    "over pay, conditions, or other workplace issues. Industrial action is a "
    "fundamental part of the employment relationship and is regulated by UK "
    "employment law, particularly the Trade Union and Labour Relations "
    "(Consolidation) Act 1992 and the Trade Union Act 2016. This section "
    "examines five forms of industrial action, evaluating the impact of each "
    "on both employers and employees."
)

# --- 1.3.1 Strikes ---
doc.add_heading("1.3.1 Strikes", level=3)

add_body(
    "A strike is the most visible and disruptive form of industrial action, "
    "involving the complete withdrawal of labour by workers. Employees refuse "
    "to attend work for a defined period, effectively halting production or "
    "service delivery. Strikes are typically used as a last resort after "
    "negotiations have broken down and are intended to demonstrate the "
    "collective power of the workforce."
)

add_body(
    "Strikes have a significant impact on employers. Production stops "
    "completely during the strike period, leading to lost revenue, missed "
    "deadlines, and potential breach of contracts with customers. For service "
    "organisations, strikes can cause severe disruption to users. The series "
    "of rail strikes in 2022 and 2023, organised by the RMT union in dispute "
    "over pay, job security, and working conditions, caused widespread "
    "disruption to commuters and businesses across the UK, with Network Rail "
    "estimating losses of tens of millions of pounds per strike day (BBC "
    "News, 2023). Public sector strikes, such as those by NHS nurses "
    "represented by the Royal College of Nursing in early 2023, raised "
    "concerns about patient safety and placed additional pressure on already "
    "stretched services."
)

add_body(
    "For employees, striking involves a loss of pay for the duration of the "
    "action, as employers are not obliged to pay workers who are on strike. "
    "This can cause financial hardship, particularly for lower-paid workers. "
    "There is also a risk that prolonged strikes may damage the relationship "
    "between workers and management, making it more difficult to reach a "
    "satisfactory resolution. However, strikes can also strengthen worker "
    "solidarity and demonstrate the importance of collective bargaining power."
)

add_body(
    "Under UK law, a strike is only lawful if it has been authorised by a "
    "properly conducted secret postal ballot, with a minimum turnout of 50 "
    "per cent, and if 14 days' notice has been given to the employer. The "
    "Trade Union Act 2016 introduced additional requirements for industrial "
    "action in 'important public services' such as health, education, and "
    "transport, where at least 40 per cent of those entitled to vote must "
    "have voted in favour. Workers who participate in lawful industrial "
    "action are protected from unfair dismissal for the first 12 weeks of "
    "the action (ACAS, 2023)."
)

# --- 1.3.2 Go-Slow ---
doc.add_heading("1.3.2 Go-Slow", level=3)

add_body(
    "A go-slow is a form of industrial action where workers continue to "
    "attend work but deliberately perform their duties at a reduced pace. "
    "The aim is to reduce productivity and output without the financial "
    "penalty of a full strike, as workers continue to receive their wages "
    "while still disrupting the employer's operations."
)

add_body(
    "The impact on employers can be substantial, particularly in "
    "manufacturing, logistics, or service environments where speed and "
    "efficiency are critical. Output falls, delivery times increase, and "
    "customer satisfaction may decline. For example, in the postal sector, "
    "go-slow actions by Royal Mail workers have historically led to delays "
    "in mail and parcel deliveries, affecting both personal and business "
    "customers (Communication Workers Union, 2022). Employers may find it "
    "difficult to address a go-slow because workers are technically still "
    "performing their duties, making it harder to take disciplinary action "
    "compared to a full strike."
)

add_body(
    "For employees, a go-slow is attractive because it allows them to "
    "exert pressure on the employer while still earning their regular pay. "
    "However, there are risks. If the go-slow is not officially sanctioned "
    "by the trade union or does not meet the legal requirements for "
    "industrial action, employees may face disciplinary action for failing "
    "to meet contractual performance standards. Prolonged go-slows can also "
    "strain relationships between workers and management and may lead to "
    "increased workplace tensions."
)

add_body(
    "Go-slows occupy a grey area in UK employment law. While there is no "
    "specific legislation that addresses go-slows directly, they may be "
    "considered a breach of the implied contractual duty to work with "
    "reasonable skill and care. If the go-slow is part of a lawful trade "
    "dispute and has been properly balloted, workers will receive the same "
    "protections as those participating in a strike (Lewis and Sargeant, "
    "2019)."
)

# --- 1.3.3 Work-to-Rule ---
doc.add_heading("1.3.3 Work-to-Rule", level=3)

add_body(
    "Work-to-rule is a form of industrial action in which workers strictly "
    "adhere to the precise terms of their employment contract and follow "
    "every workplace rule and procedure to the letter. In practice, most "
    "workplaces rely on employees exercising a degree of flexibility and "
    "goodwill, such as staying a few minutes late to complete a task, "
    "helping colleagues outside their immediate responsibilities, or using "
    "their initiative to solve problems. When workers adopt a work-to-rule "
    "approach, this flexibility disappears, and operations slow down "
    "significantly."
)

add_body(
    "The impact on employers is considerable because organisations depend "
    "on discretionary effort from employees. When workers refuse to do "
    "anything beyond the strict requirements of their contract, efficiency "
    "drops, decision-making slows, and bottlenecks emerge. In education, "
    "for example, teachers adopting a work-to-rule would refuse to attend "
    "after-school meetings that are not contractually required, stop "
    "volunteering for extracurricular activities, and strictly observe "
    "break times, leading to significant disruption to the running of "
    "the school (NASUWT, 2022)."
)

add_body(
    "For employees, work-to-rule is one of the least financially risky forms "
    "of industrial action because they continue to receive full pay while "
    "technically fulfilling their contractual obligations. It highlights the "
    "extent to which employers rely on the goodwill and discretionary effort "
    "of their workforce. However, it can create a tense and unpleasant "
    "working environment, and prolonged work-to-rule actions may damage "
    "team dynamics and morale."
)

add_body(
    "Legally, work-to-rule is difficult for employers to challenge because "
    "employees are, by definition, complying with the terms of their "
    "contract. However, if the action is organised as part of a trade "
    "dispute, it should follow the same ballot and notice requirements as "
    "other forms of industrial action to ensure legal protection for the "
    "workers involved (ACAS, 2023)."
)

# --- 1.3.4 Lockouts ---
doc.add_heading("1.3.4 Lockouts", level=3)

add_body(
    "A lockout is a form of industrial action initiated by the employer "
    "rather than the workers. During a lockout, the employer prevents "
    "workers from entering the workplace and carrying out their duties, "
    "effectively denying them the opportunity to work and earn wages. "
    "Lockouts are used by employers as a counter-measure during disputes, "
    "intended to put financial pressure on workers and their trade union "
    "to accept the employer's position."
)

add_body(
    "The impact on employers is that a lockout stops production or service "
    "delivery, similar to a strike. However, the employer retains control "
    "over the timing and duration of the action. Lockouts can be effective "
    "as a bargaining tool because they force workers to go without pay, "
    "which can be particularly pressuring for lower-paid employees. In the "
    "United States, lockouts have been used more frequently in industries "
    "such as professional sports and manufacturing, though they are less "
    "common in the UK. A notable international example was the 2012-2013 "
    "National Hockey League lockout in North America, where team owners "
    "locked out players for 113 days during a dispute over revenue sharing, "
    "resulting in significant financial losses for both sides."
)

add_body(
    "For employees, lockouts cause immediate financial hardship because "
    "they lose their wages for the duration of the action. Workers may "
    "struggle to meet financial commitments, and prolonged lockouts can "
    "create anxiety and uncertainty about job security. Trade unions may "
    "provide strike funds or hardship payments to support members during "
    "a lockout, but these are typically modest."
)

add_body(
    "In the UK, lockouts are subject to the same legal framework as other "
    "forms of industrial action. An employer who locks out workers during "
    "a lawful trade dispute is not required to pay wages for the period of "
    "the lockout. However, if the lockout is conducted in a manner that "
    "breaches the employment contract, such as failing to give adequate "
    "notice, it may give rise to claims for breach of contract or unfair "
    "dismissal (Lewis and Sargeant, 2019)."
)

# --- 1.3.5 Picketing ---
doc.add_heading("1.3.5 Picketing", level=3)

add_body(
    "Picketing involves workers gathering at or near the entrance to their "
    "workplace during a dispute to peacefully protest and persuade others "
    "not to cross the picket line. Picketing is often conducted alongside "
    "a strike and serves both a practical purpose, by discouraging other "
    "workers, suppliers, and customers from entering the premises, and a "
    "symbolic purpose, by drawing public attention to the dispute."
)

add_body(
    "The impact on employers can be significant, particularly if picketing "
    "is effective in preventing deliveries, discouraging customers, or "
    "generating negative media coverage. During the 2022-2023 Royal Mail "
    "strikes, picket lines outside sorting offices attracted considerable "
    "media attention and public sympathy, putting pressure on management "
    "to reach a settlement (Communication Workers Union, 2022). The "
    "reputational impact of picketing can be substantial, as images of "
    "workers on picket lines may damage the employer's brand and public "
    "perception."
)

add_body(
    "For employees, picketing is a way of demonstrating solidarity and "
    "making their voices heard. It can boost morale among strikers and "
    "strengthen the sense of collective purpose. However, picketing can "
    "also be physically demanding, particularly in poor weather conditions, "
    "and there is a risk of confrontation with those who wish to cross "
    "the picket line."
)

add_body(
    "UK law strictly regulates picketing to ensure it remains peaceful "
    "and lawful. Under the Trade Union and Labour Relations (Consolidation) "
    "Act 1992, picketing is only lawful if it takes place at or near the "
    "worker's own place of work, is conducted for the purpose of peacefully "
    "obtaining or communicating information, or peacefully persuading any "
    "person to work or not to work. There is a Code of Practice on "
    "Picketing issued by the Department for Business and Trade, which "
    "recommends that the number of pickets at any entrance should normally "
    "not exceed six. A trade union must appoint a picket supervisor who "
    "is familiar with the Code of Practice and who carries a letter of "
    "authorisation from the union. Secondary picketing, which involves "
    "picketing at a workplace other than one's own, is generally unlawful "
    "(ACAS, 2023)."
)

add_page_break()

# ───────────────────────────────────────────────
# Section 1.4 - Employment Legislation (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 1.4 - Employment Legislation", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Employment legislation provides the legal framework that governs the "
    "relationship between employers and employees in the UK. It establishes "
    "minimum standards for working conditions, protects employees from unfair "
    "treatment, and sets out the rights and obligations of both parties. "
    "This section examines five key areas of employment legislation, "
    "explaining what each law covers, the obligations it places on employers, "
    "the rights it grants to employees, and the consequences of non-compliance."
)

# --- Equality Act 2010 ---
doc.add_heading("1.4.1 Equality Act 2010", level=3)

add_body(
    "The Equality Act 2010 is the primary piece of UK legislation that "
    "protects individuals from discrimination in the workplace and in wider "
    "society. It replaced and consolidated previous anti-discrimination laws, "
    "including the Sex Discrimination Act 1975, the Race Relations Act 1976, "
    "and the Disability Discrimination Act 1995, creating a single, coherent "
    "legal framework."
)

add_body(
    "The Act identifies nine protected characteristics: age, disability, "
    "gender reassignment, marriage and civil partnership, pregnancy and "
    "maternity, race, religion or belief, sex, and sexual orientation. "
    "It prohibits several forms of discrimination, including direct "
    "discrimination (treating someone less favourably because of a protected "
    "characteristic), indirect discrimination (applying a policy or practice "
    "that puts people with a protected characteristic at a disadvantage), "
    "harassment (unwanted conduct related to a protected characteristic that "
    "creates an intimidating or hostile environment), and victimisation "
    "(treating someone unfairly because they have raised or supported a "
    "discrimination complaint)."
)

add_body(
    "Employers have a legal obligation to ensure that their recruitment "
    "processes, pay structures, promotion decisions, and workplace policies "
    "are free from discrimination. They must also make reasonable adjustments "
    "for disabled employees, such as providing accessible workstations, "
    "modified working hours, or assistive technology. Employers are vicariously "
    "liable for acts of discrimination committed by their employees in the "
    "course of employment, unless they can demonstrate that they took all "
    "reasonable steps to prevent it."
)

add_body(
    "Employees who believe they have been discriminated against can raise a "
    "grievance with their employer and, if unresolved, bring a claim to an "
    "employment tribunal. Remedies may include compensation (which is "
    "uncapped for discrimination claims), a declaration of the employee's "
    "rights, or a recommendation that the employer takes specific action to "
    "reduce the effect of the discrimination. Non-compliance with the "
    "Equality Act can result in significant financial penalties, reputational "
    "damage, and loss of employee trust (Equality and Human Rights Commission, "
    "2021)."
)

# --- Health & Safety at Work Act 1974 ---
doc.add_heading("1.4.2 Health and Safety at Work Act 1974", level=3)

add_body(
    "The Health and Safety at Work Act 1974 (HSWA) is the foundational "
    "piece of UK legislation governing workplace health and safety. It "
    "places a general duty on employers to ensure, so far as is reasonably "
    "practicable, the health, safety, and welfare of all their employees "
    "whilst at work. The Act also imposes duties on employees to take "
    "reasonable care of their own health and safety and that of others who "
    "may be affected by their acts or omissions."
)

add_body(
    "Employers are required to conduct thorough risk assessments to identify "
    "potential hazards in the workplace and implement appropriate control "
    "measures. They must provide employees with adequate information, "
    "instruction, training, and supervision to ensure safe working practices. "
    "Employers must also maintain the workplace in a safe condition, provide "
    "appropriate personal protective equipment (PPE) where necessary, and "
    "have procedures in place for dealing with emergencies."
)

add_body(
    "The Act is enforced by the Health and Safety Executive (HSE) and local "
    "authority environmental health departments. HSE inspectors have the "
    "power to enter workplaces without notice, carry out inspections, and "
    "take action where they find non-compliance. Enforcement actions include "
    "improvement notices (requiring the employer to remedy a contravention "
    "within a specified period), prohibition notices (requiring the employer "
    "to stop a dangerous activity immediately), and prosecution. Serious "
    "breaches can result in unlimited fines and, in cases where negligence "
    "leads to death, imprisonment under the Corporate Manslaughter and "
    "Corporate Homicide Act 2007 (HSE, 2023)."
)

add_body(
    "Employees have the right to a safe working environment, to receive "
    "necessary health and safety training, to be consulted on health and "
    "safety matters, and to stop work and leave an area if they believe they "
    "are in serious and imminent danger. They also have the right to raise "
    "health and safety concerns with their employer or the HSE without fear "
    "of dismissal or other detriment (ACAS, 2023)."
)

# --- National Minimum Wage Act 1998 ---
doc.add_heading("1.4.3 National Minimum Wage Act 1998", level=3)

add_body(
    "The National Minimum Wage Act 1998 established a legal minimum hourly "
    "rate of pay for workers in the UK. The National Living Wage (NLW), "
    "which applies to workers aged 21 and over (lowered from 23 in April "
    "2024), and the National Minimum Wage (NMW), which applies to younger "
    "workers and apprentices, are reviewed annually by the Low Pay Commission, "
    "an independent body that advises the government on appropriate rates."
)

add_body(
    "As of April 2024, the NLW is set at 11.44 pounds per hour for workers "
    "aged 21 and over. The NMW rates are 8.60 pounds per hour for workers "
    "aged 18 to 20, 6.40 pounds per hour for workers under 18, and 6.40 "
    "pounds per hour for apprentices. These rates are designed to protect "
    "workers from exploitation while ensuring that businesses can remain "
    "competitive."
)

add_body(
    "Employers are legally required to pay at least the minimum wage to all "
    "eligible workers. They must keep accurate records of hours worked and "
    "wages paid, and make these available for inspection if requested. "
    "Enforcement is carried out by HM Revenue and Customs (HMRC), which has "
    "the power to investigate complaints, carry out inspections, and issue "
    "notices of underpayment. Employers found to be paying below the minimum "
    "wage must pay arrears to the affected workers and may be issued with a "
    "financial penalty of up to 200 per cent of the arrears, up to a "
    "maximum of 20,000 pounds per worker. Persistent or deliberate non-"
    "compliance can result in criminal prosecution (Gov.uk, 2024)."
)

add_body(
    "The minimum wage has had a significant impact on UK businesses, "
    "particularly in low-pay sectors such as retail, hospitality, and social "
    "care. While it has lifted the earnings of millions of workers and "
    "reduced in-work poverty, some employers, particularly small businesses, "
    "argue that the rising minimum wage increases their labour costs and "
    "reduces their ability to create new jobs. Nonetheless, research by the "
    "Low Pay Commission (2023) suggests that the negative employment effects "
    "of the minimum wage have been minimal and that it has been broadly "
    "beneficial for both workers and the economy."
)

# --- Working Time Regulations 1998 ---
doc.add_heading("1.4.4 Working Time Regulations 1998", level=3)

add_body(
    "The Working Time Regulations 1998 (WTR) implement the European Working "
    "Time Directive into UK law and set limits on working hours, rest "
    "breaks, and annual leave entitlement. Although the UK has left the "
    "European Union, the Working Time Regulations remain in force as retained "
    "EU law and continue to provide important protections for UK workers."
)

add_body(
    "The key provisions of the WTR include a maximum average working week "
    "of 48 hours, calculated over a 17-week reference period. Workers have "
    "the right to opt out of the 48-hour limit voluntarily, but the opt-out "
    "must be in writing and can be withdrawn by the worker with seven days' "
    "notice. Adult workers are entitled to a rest break of at least 20 "
    "minutes if their working day is longer than six hours, a daily rest "
    "period of at least 11 consecutive hours between working days, and a "
    "weekly rest period of at least 24 uninterrupted hours in each seven-day "
    "period (or 48 hours in each 14-day period)."
)

add_body(
    "The WTR also provide for a minimum of 5.6 weeks' paid annual leave per "
    "year (28 days for a full-time worker), which can include public "
    "holidays. Employers are not permitted to replace the holiday entitlement "
    "with a payment in lieu, except when the worker's employment is "
    "terminated. Night workers are subject to additional protections, "
    "including a limit of an average of eight hours' work in any 24-hour "
    "period and the right to free health assessments."
)

add_body(
    "Failure to comply with the WTR can result in employment tribunal "
    "claims. Workers who are denied their rest breaks or holiday entitlement, "
    "or who are subjected to a detriment for refusing to work excessive "
    "hours, can bring a claim for compensation. Employers who fail to limit "
    "working hours or provide adequate rest breaks may also face enforcement "
    "action from the HSE (Gov.uk, 2024)."
)

# --- Dismissal & Redundancy ---
doc.add_heading("1.4.5 Dismissal and Redundancy", level=3)

add_body(
    "The law governing dismissal and redundancy is primarily contained in "
    "the Employment Rights Act 1996, which establishes the rights of "
    "employees in relation to termination of employment. An understanding "
    "of dismissal and redundancy law is essential for both employers and "
    "employees to ensure that terminations are handled fairly, lawfully, "
    "and with appropriate procedural safeguards."
)

add_body(
    "Fair dismissal requires the employer to have a potentially fair reason "
    "for the dismissal, which falls into one of five statutory categories: "
    "capability or qualifications, conduct, redundancy, contravention of a "
    "statutory duty or restriction, and 'some other substantial reason.' "
    "In addition to having a fair reason, the employer must follow a fair "
    "procedure, which typically involves investigating the issue, informing "
    "the employee of the allegations or circumstances, holding a formal "
    "meeting, allowing the employee to be accompanied, and providing a "
    "right of appeal. The ACAS Code of Practice on Disciplinary and "
    "Grievance Procedures provides detailed guidance on the steps that "
    "employers should follow (ACAS, 2023)."
)

add_body(
    "Unfair dismissal occurs when an employer terminates an employee's "
    "contract without a fair reason or without following a fair procedure. "
    "Some dismissals are automatically unfair, regardless of the reason, "
    "including dismissal for pregnancy, for making a protected disclosure "
    "(whistleblowing), for asserting a statutory right, or for taking part "
    "in lawful industrial action within the first 12 weeks. Employees with "
    "at least two years' continuous service have the right to bring a claim "
    "for unfair dismissal to an employment tribunal. If the claim is "
    "successful, the tribunal may order reinstatement, re-engagement, or "
    "compensation, which includes a basic award (calculated in the same "
    "way as a statutory redundancy payment) and a compensatory award "
    "(currently capped at the lower of 105,707 pounds or 52 weeks' pay)."
)

add_body(
    "Redundancy occurs when an employer dismisses employees because the "
    "business is closing, the workplace is closing, or the requirement for "
    "employees to carry out work of a particular kind has ceased or "
    "diminished. Employers are required to follow a fair redundancy "
    "procedure, which includes consulting with affected employees (and "
    "trade unions or employee representatives if 20 or more redundancies "
    "are proposed), applying fair and objective selection criteria, "
    "considering suitable alternative employment within the organisation, "
    "and allowing a right of appeal. Employees with at least two years' "
    "continuous service are entitled to a statutory redundancy payment, "
    "calculated based on age, length of service, and weekly pay (capped "
    "at a statutory maximum). Failure to follow a fair redundancy procedure "
    "can result in claims for unfair dismissal and additional compensation "
    "(Gov.uk, 2024)."
)

add_page_break()

# ═══════════════════════════════════════════════
# TASK B: FINANCE (40 marks)
# ═══════════════════════════════════════════════

doc.add_heading("TASK B: Finance", level=1)
p = doc.add_paragraph()
r = p.add_run("Total: 40 marks")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

add_body(
    "This section examines the financial aspects of business management, "
    "including the sources of finance available to organisations, the "
    "purposes of financial statements, and the use of accounting ratios "
    "to analyse business performance. A sound understanding of finance "
    "is essential for managers because financial decisions directly affect "
    "an organisation's ability to operate, grow, and achieve its strategic "
    "objectives (Atrill and McLaney, 2021)."
)

# ───────────────────────────────────────────────
# Section 2.1 - Sources of Finance (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 2.1 - Three Sources of Finance", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Businesses require finance for a variety of purposes, including "
    "starting up, day-to-day operations, expansion, and investment in new "
    "assets. The choice of finance depends on factors such as the amount "
    "needed, the duration, the cost, the level of risk, and the stage of "
    "the business. This section examines two long-term sources of finance "
    "(bank loans and retained profits) and one short-term source (overdraft), "
    "evaluating their suitability for different business scenarios."
)

# --- Bank Loans ---
doc.add_heading("2.1.1 Bank Loans (Long-Term)", level=3)

add_body(
    "A bank loan is a sum of money borrowed from a bank that is repaid "
    "over a fixed period, typically between one and 25 years, with interest. "
    "The borrower receives the full amount of the loan at the outset and "
    "makes regular repayments (usually monthly) that include both capital "
    "and interest. Bank loans can be secured (backed by an asset such as "
    "property) or unsecured (based on the borrower's creditworthiness)."
)

add_body(
    "There are several advantages of bank loans as a source of finance. "
    "First, the borrower receives the full amount immediately, which is "
    "useful for large, planned investments such as purchasing premises, "
    "equipment, or vehicles. Second, repayments are spread over a fixed "
    "period, making them predictable and easier to budget for. Third, the "
    "interest rate may be fixed, protecting the borrower from fluctuations "
    "in market rates. Fourth, the lender has no ownership stake in the "
    "business, so the borrower retains full control."
)

add_body(
    "However, there are also disadvantages. Interest must be paid regardless "
    "of whether the business is profitable, which can create a significant "
    "financial burden during difficult trading periods. Secured loans carry "
    "the risk that the pledged asset may be repossessed if the borrower "
    "defaults on repayments. Banks may also require the business to meet "
    "certain conditions (covenants), such as maintaining a minimum level "
    "of profitability or asset cover. For example, when Tesco undertook "
    "a major store expansion programme in the early 2000s, it used a "
    "combination of bank loans and retained profits to fund the "
    "construction of new supermarkets and distribution centres across the "
    "UK (Tesco PLC, 2023)."
)

add_body(
    "Bank loans are most appropriate for established businesses with a "
    "proven track record and sufficient assets to offer as security. They "
    "are particularly suitable for funding long-term capital investments "
    "where the benefits will be realised over several years, such as buying "
    "property or investing in major equipment. Start-up businesses may find "
    "it more difficult to obtain bank loans because they lack the trading "
    "history and assets that banks typically require as security."
)

# --- Retained Profits ---
doc.add_heading("2.1.2 Retained Profits (Long-Term)", level=3)

add_body(
    "Retained profits are the portion of a company's net profit that is "
    "not distributed to shareholders as dividends but is instead reinvested "
    "in the business. Retained profits are one of the most common and "
    "important sources of long-term finance, particularly for established "
    "businesses that have been trading profitably over a number of years."
)

add_body(
    "The main advantage of retained profits is that they are free of "
    "interest charges, unlike bank loans or other forms of borrowing. "
    "There are no repayment obligations, no interest costs, and no need "
    "to provide security. The business retains complete control over how "
    "the funds are used, and there is no dilution of ownership, as there "
    "would be with issuing new shares. Retained profits also signal "
    "financial health and stability, which can enhance the business's "
    "reputation with investors, creditors, and customers."
)

add_body(
    "However, retained profits have limitations. They are only available "
    "to businesses that are already profitable, which excludes start-ups "
    "and loss-making organisations. The amount available is limited to what "
    "the business has earned and not distributed, which may be insufficient "
    "for large-scale investments. There is also an opportunity cost, as "
    "profits that are retained cannot be paid to shareholders, which may "
    "make the company less attractive to investors who expect regular "
    "dividend income. For example, Apple Inc. famously accumulated over "
    "200 billion dollars in retained profits before eventually beginning "
    "to pay dividends and buy back shares in 2012, reflecting a strategic "
    "decision to reinvest profits into research, development, and new "
    "product lines (Apple Inc., 2023)."
)

add_body(
    "Retained profits are most appropriate for established, profitable "
    "businesses that wish to fund gradual expansion, invest in new "
    "equipment, or build up a financial reserve without taking on external "
    "debt. They are particularly suitable for businesses that value "
    "independence and wish to avoid the obligations and conditions that "
    "come with external finance."
)

# --- Overdraft ---
doc.add_heading("2.1.3 Bank Overdraft (Short-Term)", level=3)

add_body(
    "A bank overdraft is a short-term borrowing facility that allows a "
    "business to withdraw more money from its bank account than it "
    "currently holds, up to an agreed limit. Overdrafts are designed to "
    "help businesses manage short-term cash flow fluctuations, such as "
    "covering wages or supplier payments while waiting for customers to "
    "pay their invoices."
)

add_body(
    "The key advantage of an overdraft is its flexibility. The business "
    "only borrows what it needs, and interest is charged only on the "
    "amount overdrawn, not on the entire facility limit. It provides "
    "immediate access to funds without the need for a formal loan "
    "application, making it ideal for managing unexpected cash shortfalls "
    "or seasonal variations in income. For example, a retail business may "
    "use an overdraft facility during quieter trading months to cover "
    "expenses, knowing that cash flow will improve during peak periods "
    "such as the Christmas season."
)

add_body(
    "However, overdrafts have significant disadvantages. Interest rates "
    "on overdrafts are typically much higher than on bank loans, making "
    "them an expensive form of borrowing if used over extended periods. "
    "The facility is repayable on demand, which means the bank can "
    "withdraw the overdraft at any time, potentially leaving the business "
    "in a difficult financial position. Overdrafts are not suitable for "
    "funding long-term investments because of their short-term nature "
    "and high cost."
)

add_body(
    "Overdrafts are most appropriate for businesses of all sizes that "
    "need to manage short-term cash flow gaps. They are particularly "
    "useful for start-ups and small businesses that experience irregular "
    "income patterns. However, businesses that find themselves permanently "
    "overdrawn should consider converting to a longer-term loan to reduce "
    "interest costs and provide greater financial stability."
)

# --- Comparison Table ---
doc.add_heading("Comparison of Sources of Finance", level=3)

comp_table = doc.add_table(rows=4, cols=5)
comp_table.style = "Table Grid"
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_table_header_row(comp_table,
    ["Source", "Type", "Cost", "Best For", "Risk Level"])
comp_data = [
    ("Bank Loan", "Long-term", "Fixed/variable interest", "Established businesses; capital investment", "Medium"),
    ("Retained Profits", "Long-term", "No direct cost", "Profitable businesses; gradual expansion", "Low"),
    ("Bank Overdraft", "Short-term", "High variable interest", "All businesses; cash flow management", "Medium-High"),
]
for r_idx, row_data in enumerate(comp_data, start=1):
    for c_idx, text in enumerate(row_data):
        cell = comp_table.rows[r_idx].cells[c_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9)

add_page_break()

# ───────────────────────────────────────────────
# Section 2.2 - Purposes of Financial Statements (10 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 2.2 - Five Purposes of Financial Statements", level=2)
p = doc.add_paragraph()
r = p.add_run("(10 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Financial statements are formal records of a business's financial "
    "activities and position. They provide essential information to a wide "
    "range of stakeholders, including investors, managers, creditors, and "
    "HM Revenue and Customs (HMRC). The three main financial statements "
    "are the Income Statement (Profit and Loss Account), the Statement of "
    "Financial Position (Balance Sheet), and the Cash Flow Statement. This "
    "section first explains each statement and then examines five specific "
    "purposes that financial statements serve."
)

doc.add_heading("The Three Main Financial Statements", level=3)

add_bold_body(
    "Income Statement (Profit and Loss Account): ",
    "The income statement shows the revenue earned and the expenses "
    "incurred by a business over a specific period, typically one year. "
    "It calculates the gross profit (revenue minus cost of sales) and "
    "the net profit (gross profit minus operating expenses, interest, and "
    "tax). The income statement is a key indicator of profitability and "
    "shows whether the business is generating sufficient income to cover "
    "its costs and provide a return to its owners (Atrill and McLaney, "
    "2021)."
)

add_bold_body(
    "Statement of Financial Position (Balance Sheet): ",
    "The balance sheet provides a snapshot of a business's financial "
    "position at a specific point in time, showing what the business owns "
    "(assets), what it owes (liabilities), and the residual interest of "
    "the owners (equity). Assets are classified as non-current (long-term "
    "assets such as property and equipment) and current (short-term assets "
    "such as stock, debtors, and cash). Liabilities are similarly "
    "classified as non-current (long-term debts) and current (amounts "
    "due within one year). The fundamental accounting equation states that "
    "assets equal liabilities plus equity."
)

add_bold_body(
    "Cash Flow Statement: ",
    "The cash flow statement shows the actual cash inflows and outflows "
    "of a business over a specific period. It is divided into three "
    "sections: operating activities (cash generated from the core "
    "business), investing activities (cash used to buy or received from "
    "selling assets), and financing activities (cash from borrowing, "
    "share issues, or dividend payments). The cash flow statement is "
    "crucial because a profitable business can still fail if it runs "
    "out of cash, highlighting the important distinction between profit "
    "and cash flow."
)

doc.add_heading("Five Purposes of Financial Statements", level=3)

add_bold_body(
    "1. Assessing Profitability: ",
    "One of the primary purposes of financial statements is to enable "
    "stakeholders to assess the profitability of a business. The income "
    "statement shows whether the business is making a profit or a loss, "
    "and how that profit is distributed between cost of sales, operating "
    "expenses, and taxation. Investors use profitability information to "
    "determine whether the business is generating an adequate return on "
    "their investment, while managers use it to identify areas where costs "
    "can be reduced or revenue can be increased."
)

add_bold_body(
    "2. Evaluating Liquidity: ",
    "Financial statements allow stakeholders to assess the liquidity of "
    "a business, meaning its ability to meet its short-term financial "
    "obligations as they fall due. The balance sheet shows the relationship "
    "between current assets and current liabilities, and the cash flow "
    "statement reveals whether the business is generating sufficient cash "
    "to cover its day-to-day expenses. Creditors and suppliers are "
    "particularly interested in liquidity because it affects the "
    "likelihood that they will be paid on time."
)

add_bold_body(
    "3. Comparing Performance: ",
    "Financial statements enable comparison of business performance "
    "over time (trend analysis) and against competitors (benchmarking). "
    "By examining financial statements from successive years, managers "
    "can identify trends in revenue, profitability, and efficiency. "
    "Comparing the financial ratios of similar businesses within an "
    "industry helps investors and managers understand how well the "
    "business is performing relative to its peers and whether there "
    "are areas for improvement."
)

add_bold_body(
    "4. Tax Compliance: ",
    "Businesses are legally required to prepare financial statements "
    "for tax purposes. HMRC uses the information in the income statement "
    "to calculate the corporation tax (for companies) or income tax "
    "(for sole traders and partnerships) that is due. Accurate financial "
    "statements are essential for ensuring that the correct amount of "
    "tax is paid and for avoiding penalties for non-compliance. Companies "
    "are also required to file their financial statements with Companies "
    "House, where they become publicly available."
)

add_bold_body(
    "5. Securing Investment and Finance: ",
    "Financial statements are a critical tool for businesses seeking "
    "to raise finance from external sources. Banks and other lenders "
    "examine financial statements to assess the creditworthiness of a "
    "business before approving a loan. Investors review financial "
    "statements to evaluate the potential return on their investment and "
    "the level of risk involved. A business with strong, well-presented "
    "financial statements is more likely to secure favourable terms from "
    "lenders and attract investment from shareholders (Atrill and "
    "McLaney, 2021)."
)

add_page_break()

# ───────────────────────────────────────────────
# Section 2.3 - Five Accounting Ratios (15 marks)
# ───────────────────────────────────────────────
doc.add_heading("Section 2.3 - Five Accounting Ratios and How Each "
                "Analyses Performance", level=2)
p = doc.add_paragraph()
r = p.add_run("(15 marks)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_body(
    "Accounting ratios are mathematical calculations that use data from "
    "financial statements to analyse different aspects of a business's "
    "performance. They are widely used by managers, investors, and creditors "
    "because they provide a standardised way of measuring profitability, "
    "liquidity, and financial stability, enabling meaningful comparisons "
    "between businesses and over time. This section explains five key "
    "accounting ratios, including their formulae, interpretation, example "
    "calculations, and managerial applications."
)

# --- Gross Profit Margin ---
doc.add_heading("2.3.1 Gross Profit Margin", level=3)

add_body(
    "The gross profit margin measures the percentage of revenue that remains "
    "after deducting the cost of sales. It indicates how efficiently a "
    "business converts its revenue into gross profit, reflecting the "
    "effectiveness of its pricing strategy and production cost management."
)

p = doc.add_paragraph()
r = p.add_run("Formula: ")
r.bold = True
p.add_run("Gross Profit Margin = (Gross Profit / Revenue) x 100")

add_body(
    "Example calculation: A business has revenue of 500,000 pounds and cost "
    "of sales of 300,000 pounds. The gross profit is 500,000 minus 300,000 "
    "equals 200,000 pounds. The gross profit margin is (200,000 / 500,000) "
    "x 100 = 40 per cent."
)

add_body(
    "A higher gross profit margin indicates that the business is retaining "
    "a larger proportion of each pound of revenue as gross profit, which "
    "suggests effective cost control and a strong pricing strategy. A "
    "declining gross profit margin over time may indicate rising input costs, "
    "increased competition forcing price reductions, or inefficiencies in "
    "the production process. Typical gross profit margins vary significantly "
    "by industry, with software companies often achieving margins of 70 to "
    "90 per cent, while supermarkets typically operate on margins of 25 to "
    "30 per cent. Managers use this ratio to monitor production efficiency, "
    "evaluate supplier contracts, and make pricing decisions."
)

# --- Net Profit Margin ---
doc.add_heading("2.3.2 Net Profit Margin", level=3)

add_body(
    "The net profit margin measures the percentage of revenue that remains "
    "as net profit after all expenses, including operating costs, interest, "
    "and tax, have been deducted. It provides a comprehensive picture of "
    "overall profitability and how well the business controls its total "
    "cost base."
)

p = doc.add_paragraph()
r = p.add_run("Formula: ")
r.bold = True
p.add_run("Net Profit Margin = (Net Profit / Revenue) x 100")

add_body(
    "Example calculation: Using the same business, if operating expenses "
    "are 120,000 pounds, interest is 10,000 pounds, and tax is 14,000 "
    "pounds, the net profit is 200,000 minus 120,000 minus 10,000 minus "
    "14,000 equals 56,000 pounds. The net profit margin is (56,000 / "
    "500,000) x 100 = 11.2 per cent."
)

add_body(
    "A healthy net profit margin indicates that the business is managing "
    "all of its costs effectively, not just production costs. If the gross "
    "profit margin is strong but the net profit margin is weak, this "
    "suggests that operating expenses, such as administrative costs, rent, "
    "or marketing, are too high relative to revenue. Managers use the net "
    "profit margin to assess overall business efficiency, identify areas "
    "where overhead costs can be reduced, and evaluate whether the business "
    "is generating a satisfactory return for its owners."
)

# --- Current Ratio ---
doc.add_heading("2.3.3 Current Ratio", level=3)

add_body(
    "The current ratio is a liquidity ratio that measures a business's "
    "ability to pay its short-term obligations using its short-term assets. "
    "It compares current assets (cash, stock, and debtors) with current "
    "liabilities (amounts due within one year, such as trade creditors and "
    "short-term loans)."
)

p = doc.add_paragraph()
r = p.add_run("Formula: ")
r.bold = True
p.add_run("Current Ratio = Current Assets / Current Liabilities")

add_body(
    "Example calculation: A business has current assets of 150,000 pounds "
    "and current liabilities of 75,000 pounds. The current ratio is "
    "150,000 / 75,000 = 2:1."
)

add_body(
    "A current ratio of 2:1 is often cited as ideal, meaning the business "
    "has twice as many current assets as current liabilities, providing a "
    "comfortable margin of safety. A ratio below 1:1 indicates that the "
    "business may not have sufficient current assets to meet its short-term "
    "obligations, which could lead to cash flow problems and an inability "
    "to pay suppliers or employees on time. However, a very high current "
    "ratio (for example, 5:1) may suggest that the business is not using "
    "its assets efficiently, with too much cash sitting idle or excessive "
    "levels of stock. Managers use the current ratio to monitor liquidity, "
    "plan cash flow, and make decisions about stock levels, credit terms, "
    "and short-term borrowing."
)

# --- ROCE ---
doc.add_heading("2.3.4 Return on Capital Employed (ROCE)", level=3)

add_body(
    "Return on Capital Employed (ROCE) is a profitability ratio that "
    "measures how effectively a business uses its total capital (equity "
    "plus long-term borrowings) to generate profit. It is widely regarded "
    "as one of the most important financial ratios because it shows the "
    "return that the business generates for every pound of capital invested."
)

p = doc.add_paragraph()
r = p.add_run("Formula: ")
r.bold = True
p.add_run("ROCE = (Net Profit / Capital Employed) x 100")

p = doc.add_paragraph()
r = p.add_run("Where: ")
r.bold = True
p.add_run("Capital Employed = Total Assets - Current Liabilities "
          "(or Equity + Non-Current Liabilities)")

add_body(
    "Example calculation: A business has net profit of 56,000 pounds, "
    "total assets of 400,000 pounds, and current liabilities of 75,000 "
    "pounds. Capital employed is 400,000 minus 75,000 equals 325,000 "
    "pounds. ROCE is (56,000 / 325,000) x 100 = 17.2 per cent."
)

add_body(
    "A higher ROCE indicates that the business is generating more profit "
    "relative to the capital invested, which is attractive to investors. "
    "As a general guide, a ROCE above 15 per cent is considered good, while "
    "a ROCE below 10 per cent may suggest that the capital could be more "
    "productively invested elsewhere. Investors compare ROCE with the "
    "returns available from alternative investments, such as savings "
    "accounts or government bonds, to assess whether the level of risk "
    "is justified by the return. Managers use ROCE to evaluate the "
    "performance of different divisions or business units, to assess the "
    "viability of capital investment projects, and to benchmark performance "
    "against industry competitors."
)

# --- Debt-to-Equity Ratio ---
doc.add_heading("2.3.5 Debt-to-Equity Ratio", level=3)

add_body(
    "The debt-to-equity ratio measures the proportion of a business's "
    "financing that comes from debt (borrowed money) compared to equity "
    "(the owners' investment and retained profits). It is a key indicator "
    "of financial leverage and risk, showing the extent to which a business "
    "relies on external borrowing to fund its operations and growth."
)

p = doc.add_paragraph()
r = p.add_run("Formula: ")
r.bold = True
p.add_run("Debt-to-Equity Ratio = Total Debt / Total Equity")

add_body(
    "Example calculation: A business has total debt (including long-term "
    "loans and overdrafts) of 120,000 pounds and total equity of 200,000 "
    "pounds. The debt-to-equity ratio is 120,000 / 200,000 = 0.6, or "
    "0.6:1."
)

add_body(
    "A debt-to-equity ratio of 0.6:1 means that for every pound of equity, "
    "the business has 60 pence of debt, which is generally considered a "
    "moderate level of gearing. A lower ratio indicates that the business "
    "is primarily financed by equity, which reduces financial risk because "
    "there are fewer interest obligations and less exposure to changes in "
    "interest rates. A higher ratio indicates greater reliance on debt, "
    "which increases financial risk but may also amplify returns to "
    "shareholders through the effect of financial leverage. A ratio above "
    "2:1 is generally considered high risk, as it suggests the business "
    "is heavily indebted and may struggle to service its debt during periods "
    "of low profitability."
)

add_body(
    "Managers use the debt-to-equity ratio to make decisions about the "
    "optimal capital structure of the business, balancing the tax benefits "
    "and lower cost of debt against the increased financial risk. Creditors "
    "examine the ratio to assess the risk of lending, and investors use it "
    "to evaluate the risk profile of their investment. For example, utility "
    "companies typically have higher debt-to-equity ratios because their "
    "stable, regulated income streams can support higher levels of "
    "borrowing, while technology start-ups tend to have lower ratios "
    "because their volatile earnings make it harder to service debt "
    "(Atrill and McLaney, 2021)."
)

# --- Summary Ratio Table ---
doc.add_heading("Summary of Accounting Ratios", level=3)

ratio_table = doc.add_table(rows=6, cols=4)
ratio_table.style = "Table Grid"
ratio_table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_table_header_row(ratio_table,
    ["Ratio", "Formula", "Measures", "Ideal / Guideline"])
ratio_data = [
    ("Gross Profit Margin", "(GP / Revenue) x 100",
     "Production efficiency", "Higher is better; varies by industry"),
    ("Net Profit Margin", "(NP / Revenue) x 100",
     "Overall profitability", "Higher is better; typically 5-20%"),
    ("Current Ratio", "CA / CL",
     "Short-term liquidity", "Approx. 2:1"),
    ("ROCE", "(NP / CE) x 100",
     "Return on investment", "Above 15% is good"),
    ("Debt-to-Equity", "Total Debt / Total Equity",
     "Financial leverage", "Below 1:1 is conservative"),
]
for r_idx, row_data in enumerate(ratio_data, start=1):
    for c_idx, text in enumerate(row_data):
        cell = ratio_table.rows[r_idx].cells[c_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9)

add_page_break()

# ═══════════════════════════════════════════════
# REFERENCE LIST
# ═══════════════════════════════════════════════
doc.add_heading("Reference List", level=1)

references = [
    "ACAS (2023) Discipline and Grievances at Work: The ACAS Guide. London: "
    "Advisory, Conciliation and Arbitration Service. Available at: "
    "https://www.acas.org.uk (Accessed: 15 February 2026).",

    "Apple Inc. (2023) Annual Report 2023. Cupertino, CA: Apple Inc.",

    "Armstrong, M. and Taylor, S. (2020) Armstrong's Handbook of Human "
    "Resource Management Practice. 15th edn. London: Kogan Page.",

    "Atrill, P. and McLaney, E. (2021) Accounting and Finance for "
    "Non-Specialists. 12th edn. Harlow: Pearson Education.",

    "Bach, S. and Edwards, M. (2013) Managing Human Resources: Human "
    "Resource Management in Transition. 5th edn. Chichester: John Wiley "
    "and Sons.",

    "BBC News (2023) 'Rail strikes: How they have affected the UK.' BBC "
    "News, 5 January. Available at: https://www.bbc.co.uk/news/business-"
    "62634795 (Accessed: 15 February 2026).",

    "Bock, L. (2015) Work Rules! Insights from Inside Google That Will "
    "Transform How You Live and Lead. London: John Murray.",

    "CIPD (2023) Factsheet: Recruitment. London: Chartered Institute of "
    "Personnel and Development. Available at: https://www.cipd.org "
    "(Accessed: 15 February 2026).",

    "Communication Workers Union (2022) Royal Mail Dispute: CWU Position "
    "Statement. London: CWU.",

    "Equality and Human Rights Commission (2021) Equality Act 2010: "
    "Employment Statutory Code of Practice. London: EHRC.",

    "General Medical Council (2022) The Good Medical Practice Framework "
    "for Appraisal and Revalidation. London: GMC.",

    "Gov.uk (2024) National Minimum Wage and National Living Wage Rates. "
    "Available at: https://www.gov.uk/national-minimum-wage-rates "
    "(Accessed: 15 February 2026).",

    "Hofstede, G. (2001) Culture's Consequences: Comparing Values, "
    "Behaviors, Institutions and Organizations Across Nations. 2nd edn. "
    "Thousand Oaks, CA: Sage Publications.",

    "HSE (2023) Health and Safety at Work etc Act 1974: Guidance. "
    "Bootle: Health and Safety Executive.",

    "John Lewis Partnership (2023) Annual Report and Accounts 2023. "
    "London: John Lewis Partnership.",

    "Lewis, D. and Sargeant, M. (2019) Employment Law: The Essentials. "
    "15th edn. London: CIPD.",

    "Low Pay Commission (2023) National Minimum Wage: Low Pay Commission "
    "Report 2023. London: LPC.",

    "Maslow, A.H. (1943) 'A Theory of Human Motivation.' Psychological "
    "Review, 50(4), pp. 370-396.",

    "NASUWT (2022) Industrial Action: Guidance for Members. Birmingham: "
    "NASUWT The Teachers' Union.",

    "NHS Employers (2023) NHS Terms and Conditions of Service Handbook. "
    "Leeds: NHS Employers.",

    "Robertson, I.T. and Smith, M. (2001) 'Personnel Selection.' Journal "
    "of Occupational and Organizational Psychology, 74(4), pp. 441-472.",

    "Schmidt, F.L. and Hunter, J.E. (1998) 'The Validity and Utility of "
    "Selection Methods in Personnel Psychology.' Psychological Bulletin, "
    "124(2), pp. 262-274.",

    "Tesco PLC (2023) Annual Report and Financial Statements 2023. "
    "Welwyn Garden City: Tesco PLC.",

    "Wahba, M.A. and Bridwell, L.G. (1976) 'Maslow Reconsidered: A "
    "Review of Research on the Need Hierarchy Theory.' Organizational "
    "Behavior and Human Performance, 15(2), pp. 212-240.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.size = Pt(10)

add_page_break()

# ═══════════════════════════════════════════════
# MAPPING TABLE
# ═══════════════════════════════════════════════
doc.add_heading("Mapping Table", level=1)

add_body(
    "The following table maps each section of this report to the relevant "
    "outcomes and performance criteria of the J22A 76 unit specification."
)

map_table = doc.add_table(rows=8, cols=4)
map_table.style = "Table Grid"
map_table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_table_header_row(map_table,
    ["Section", "Topic", "Outcome", "Marks"])

mapping_data = [
    ("1.1", "Three HRM Approaches", "Outcome 1: Explain approaches to HRM", "15"),
    ("1.2", "Maslow's Hierarchy of Needs", "Outcome 1: Apply motivational theory", "15"),
    ("1.3", "Five Forms of Industrial Action", "Outcome 1: Industrial relations", "15"),
    ("1.4", "Employment Legislation", "Outcome 1: Employment law", "15"),
    ("2.1", "Three Sources of Finance", "Outcome 2: Sources of finance", "15"),
    ("2.2", "Purposes of Financial Statements", "Outcome 2: Financial statements", "10"),
    ("2.3", "Five Accounting Ratios", "Outcome 2: Analyse performance using ratios", "15"),
]

for r_idx, row_data in enumerate(mapping_data, start=1):
    for c_idx, text in enumerate(row_data):
        cell = map_table.rows[r_idx].cells[c_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Total marks row
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run("Total Marks: 100 (Task A: 60 + Task B: 40)")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

# ──────────────────────────────────────────────
# Save DOCX
# ──────────────────────────────────────────────
print(f"Saving DOCX to: {DOCX_PATH}")
doc.save(DOCX_PATH)
print(f"DOCX saved successfully: {os.path.getsize(DOCX_PATH):,} bytes")

# ──────────────────────────────────────────────
# Kill Word and export PDF via COM
# ──────────────────────────────────────────────
print("Killing any running Word instances...")
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
               capture_output=True, text=True)

print(f"Exporting PDF to: {PDF_PATH}")
import pythoncom
import win32com.client

pythoncom.CoInitialize()
try:
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    docx_abs = os.path.abspath(DOCX_PATH)
    pdf_abs = os.path.abspath(PDF_PATH)
    wdoc = word.Documents.Open(docx_abs)
    wdoc.ExportAsFixedFormat(pdf_abs, 17)  # 17 = wdExportFormatPDF
    wdoc.Close(False)
    word.Quit()
    print(f"PDF exported successfully: {os.path.getsize(PDF_PATH):,} bytes")
except Exception as e:
    print(f"PDF export error: {e}", file=sys.stderr)
    try:
        word.Quit()
    except:
        pass
finally:
    pythoncom.CoUninitialize()

print("Build complete.")
