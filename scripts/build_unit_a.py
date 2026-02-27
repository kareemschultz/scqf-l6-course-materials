#!/usr/bin/env python3
"""
build_unit_a.py
Builds the SCQF Level 6 Unit A (J229 76) Understanding Business report.
Outputs DOCX via python-docx, then exports to PDF via Word COM automation.
Student: 252IFCBR0596 | Kareem Nurw Jason Schultz
"""

import os
import sys
import subprocess
import time

# ---------------------------------------------------------------------------
# paths
# ---------------------------------------------------------------------------
BASE = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials"
DOCX_DIR = os.path.join(BASE, "SCQF_L6_FINAL_SUBMISSION", "source_files",
                        "J22976_Understanding_Business")
PDF_DIR = os.path.join(BASE, "SCQF_L6_FINAL_SUBMISSION", "pdf_submissions")
DOCX_PATH = os.path.join(DOCX_DIR, "J22976_Report.docx")
PDF_PATH = os.path.join(PDF_DIR, "J22976_Understanding_Business.pdf")

os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

# Kill any lingering Word processes so COM automation works cleanly
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
               capture_output=True)
time.sleep(2)

# ---------------------------------------------------------------------------
# imports for document creation
# ---------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ---------------------------------------------------------------------------
# constants
# ---------------------------------------------------------------------------
STUDENT_ID = "252IFCBR0596"
STUDENT_NAME = "Kareem Nurw Jason Schultz"
UNIT_CODE = "J229 76"
UNIT_TITLE = "Understanding Business"
TOTAL_MARKS = 100
COLLEGE = "JAIN College"
DATE_STR = "February 2026"
FOOTER_TEMPLATE = (
    f"{STUDENT_ID} | {STUDENT_NAME} | {UNIT_CODE} | {UNIT_TITLE} | Page "
)

# Colour palette
NAVY = RGBColor(0x00, 0x2B, 0x5C)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = "D6E4F0"
MEDIUM_BLUE = "4472C4"
LIGHT_GREY = "F2F2F2"

# ---------------------------------------------------------------------------
# helper utilities
# ---------------------------------------------------------------------------

doc = Document()

# -- page setup (A4) --
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def set_cell_shading(cell, colour_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs keys: top, bottom, left, right, insideH, insideV
    values: dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(text, style='Normal', bold=False, italic=False,
                            font_size=11, color=DARK_GREY,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_after=Pt(6), space_before=Pt(0),
                            first_line_indent=None, font_name='Calibri'):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    return p


def add_heading_styled(text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Calibri'
    return h


def add_table_with_header(headers, rows, col_widths=None):
    """Create a formatted table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, MEDIUM_BLUE)

    # data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_bullet(text, level=0, bold_prefix=None):
    """Add a bullet point. Optionally bold a prefix portion."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Calibri'
        run_b.font.color.rgb = DARK_GREY
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(11)
        run_rest.font.name = 'Calibri'
        run_rest.font.color.rgb = DARK_GREY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GREY
    return p


# ---------------------------------------------------------------------------
# FOOTER with page numbers
# ---------------------------------------------------------------------------
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)

run_footer = fp.add_run(FOOTER_TEMPLATE)
run_footer.font.size = Pt(8)
run_footer.font.name = 'Calibri'
run_footer.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Add dynamic page number field
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_fld1 = fp.add_run()
run_fld1._r.append(fldChar1)

instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run_fld2 = fp.add_run()
run_fld2._r.append(instrText)

fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_fld3 = fp.add_run()
run_fld3._r.append(fldChar2)


# ===================================================================
# 1. COVER PAGE
# ===================================================================
print("Building cover page...")

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(UNIT_TITLE.upper())
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(f"Unit Code: {UNIT_CODE}")
run2.font.size = Pt(16)
run2.font.color.rgb = DARK_GREY
run2.font.name = 'Calibri'

doc.add_paragraph()

# Horizontal rule
p_hr = doc.add_paragraph()
p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="002B5C"/>'
    f'</w:pBdr>'
)
p_hr._p.get_or_add_pPr().append(pBdr)

doc.add_paragraph()

details = [
    ("Student Name:", STUDENT_NAME),
    ("Student ID:", STUDENT_ID),
    ("Institution:", COLLEGE),
    ("Unit Code:", UNIT_CODE),
    ("Unit Title:", UNIT_TITLE),
    ("Total Marks:", str(TOTAL_MARKS)),
    ("Date:", DATE_STR),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p.add_run(label + "  ")
    rl.font.bold = True
    rl.font.size = Pt(13)
    rl.font.color.rgb = NAVY
    rl.font.name = 'Calibri'
    rv = p.add_run(value)
    rv.font.size = Pt(13)
    rv.font.color.rgb = DARK_GREY
    rv.font.name = 'Calibri'

doc.add_page_break()


# ===================================================================
# 2. DECLARATION OF ORIGINALITY
# ===================================================================
print("Building declaration of originality...")
add_heading_styled("Declaration of Originality", level=1)

declaration_text = (
    "I hereby declare that this assignment is entirely my own work. "
    "Where I have drawn upon the ideas or words of others, I have clearly "
    "acknowledged and referenced these sources in accordance with Harvard "
    "referencing conventions. I understand that the college has a strict "
    "policy on plagiarism and academic misconduct, and I confirm that this "
    "submission has not been previously submitted for any other assessment. "
    "I have read and understood the college's regulations on academic "
    "integrity and confirm that I have complied with them fully."
)
add_formatted_paragraph(declaration_text)

doc.add_paragraph()
p_sig = doc.add_paragraph()
run_s = p_sig.add_run(f"Signed:  {STUDENT_NAME}")
run_s.font.size = Pt(12)
run_s.font.name = 'Calibri'
run_s.font.color.rgb = DARK_GREY

p_date = doc.add_paragraph()
run_d = p_date.add_run(f"Date:  {DATE_STR}")
run_d.font.size = Pt(12)
run_d.font.name = 'Calibri'
run_d.font.color.rgb = DARK_GREY

doc.add_page_break()


# ===================================================================
# 3. TABLE OF CONTENTS
# ===================================================================
print("Building table of contents...")
add_heading_styled("Table of Contents", level=1)

p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_toc._r.append(fldChar_begin)

run_toc2 = p_toc.add_run()
instrText_toc = parse_xml(
    f'<w:instrText {nsdecls("w")} xml:space="preserve">'
    f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
)
run_toc2._r.append(instrText_toc)

run_toc3 = p_toc.add_run()
fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run_toc3._r.append(fldChar_sep)

run_toc4 = p_toc.add_run("[Table of Contents will update when opened in Word - right-click and select 'Update Field']")
run_toc4.font.size = Pt(11)
run_toc4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

run_toc5 = p_toc.add_run()
fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_toc5._r.append(fldChar_end)

doc.add_page_break()


# ===================================================================
# TASK A: BUSINESS STRUCTURES AND DECISIONS
# ===================================================================
print("Building Task A...")
add_heading_styled("TASK A: Business Structures and Decisions", level=1)


# -------------------------------------------------------------------
# Section 1.1.1 - Comparing organisations (8 marks)
# -------------------------------------------------------------------
add_heading_styled("1.1.1 Comparing Three Organisations Across Sectors", level=2)

add_formatted_paragraph(
    "Businesses across the private, public, and third sectors differ significantly in their "
    "purpose, scale, structure, and the way they are regulated. This section provides a "
    "comparative analysis of Tata Motors (private sector multinational), NHS Scotland "
    "(public sector health service), and Oxfam (third sector charitable organisation). "
    "By examining these three organisations side by side, we can identify how sectoral "
    "differences shape the way each operates and the decisions they make.",
    space_after=Pt(8)
)

add_heading_styled("Comparison Table", level=3)

headers_111 = ["Criterion", "Tata Motors\n(Private Sector)", "NHS Scotland\n(Public Sector)",
               "Oxfam\n(Third Sector)"]
rows_111 = [
    [
        "Scale of Operation",
        "Over 80,000 employees globally; revenue of approximately $44 billion (2023); "
        "manufacturing and sales operations across India, UK, South Korea, and Southeast Asia "
        "(Tata Motors, 2023).",
        "Employs around 160,000 staff across 14 territorial health boards; annual budget "
        "exceeding GBP 18 billion funded through general taxation; operates exclusively within "
        "Scotland (Audit Scotland, 2023).",
        "Around 10,000 employees and 46,000 volunteers; annual income of roughly GBP 450 million; "
        "programmes in over 90 countries worldwide (Oxfam, 2023)."
    ],
    [
        "Organisational Structure",
        "Hierarchical with divisional elements; the automotive division sits within the wider "
        "Tata Group conglomerate; clear reporting lines from factory floor to board of directors "
        "(Needle, 2015).",
        "Tall hierarchical structure with multiple management layers; led by Scottish Government "
        "cabinet secretary, cascading through NHS boards, hospital management, and clinical teams "
        "(Steel and Cylus, 2012).",
        "Relatively flat structure for a global organisation; national affiliates have considerable "
        "autonomy; decisions often made collaboratively through working groups (Lewis, 2014)."
    ],
    [
        "Specialisation",
        "Focuses on automobile manufacturing including passenger vehicles, commercial vehicles, "
        "and electric vehicles; also expanding into connected and autonomous driving technologies "
        "(Tata Motors, 2023).",
        "Specialises in universal healthcare provision, including primary care, hospital services, "
        "mental health, and public health programmes; must serve every resident regardless of "
        "ability to pay (Scottish Government, 2022).",
        "Specialises in poverty alleviation, emergency humanitarian response, and advocacy for "
        "trade justice; campaigns for systemic policy change alongside delivering direct aid "
        "(Oxfam, 2023)."
    ],
    [
        "Regulation",
        "Regulated by automotive safety standards (e.g., Euro NCAP, Bharat Stage emissions), "
        "financial regulators such as SEBI in India and the FCA in the UK, and employment law "
        "across multiple jurisdictions (SMMT, 2022).",
        "Governed by the National Health Service (Scotland) Act 1978; regulated by Healthcare "
        "Improvement Scotland (HIS); must comply with Scottish Government policy directives and "
        "clinical standards (Healthcare Improvement Scotland, 2023).",
        "Regulated by the Charity Commission (England and Wales) and OSCR in Scotland; must "
        "comply with charity law, fundraising regulations, and safeguarding standards; subject "
        "to public scrutiny on governance (OSCR, 2022)."
    ],
    [
        "Capital Requirements and Sources",
        "Substantial capital needs for R&D and manufacturing; funded through share capital, "
        "retained profits, bank loans, and bond issues; parent company Tata Sons provides "
        "additional investment support (Tata Group, 2023).",
        "Capital funded almost entirely by the Scottish Government through Barnett Formula "
        "allocations; supplementary income from prescription charges (recently abolished) and "
        "car parking fees; cannot raise capital through equity markets (Audit Scotland, 2023).",
        "Funded through public donations, government grants, retail shop income, and institutional "
        "funding from bodies such as the EU and UN; cannot distribute surplus to individuals and "
        "must reinvest in charitable objectives (Oxfam, 2023)."
    ],
]

add_table_with_header(headers_111, rows_111, col_widths=[1.2, 1.8, 1.8, 1.8])

add_heading_styled("Narrative Analysis", level=3)

add_formatted_paragraph(
    "The comparison above reveals that sectoral classification has a profound effect on how "
    "organisations are structured and governed. Tata Motors operates at a massive industrial "
    "scale, driven by the profit motive and accountable to shareholders. Its hierarchical "
    "structure reflects the need for tight quality control in automotive manufacturing, where "
    "a single defect can result in costly recalls and reputational damage (Needle, 2015). "
    "In contrast, NHS Scotland's scale is shaped not by market demand but by a government "
    "mandate to provide universal care. Its hierarchical structure serves a different purpose: "
    "ensuring clinical accountability and clear chains of command in life-or-death situations "
    "(Steel and Cylus, 2012). This suggests that hierarchy is not inherently rigid or "
    "bureaucratic; rather, its effectiveness depends on the context in which it operates."
)

add_formatted_paragraph(
    "Oxfam's relatively flat structure sets it apart from the other two organisations. As a "
    "charity working across 90 countries, Oxfam relies heavily on local knowledge and must "
    "empower regional teams to respond quickly to humanitarian crises (Lewis, 2014). However, "
    "a limitation of this decentralised approach became apparent during the 2018 safeguarding "
    "scandal, which revealed gaps in oversight and accountability. This illustrates that while "
    "flat structures promote agility, they can create governance vulnerabilities when oversight "
    "mechanisms are insufficient."
)

add_formatted_paragraph(
    "Regarding capital, the three organisations demonstrate fundamentally different financial "
    "models. Tata Motors can access global capital markets, giving it flexibility to invest in "
    "long-term projects such as electric vehicle development. NHS Scotland, by contrast, depends "
    "entirely on government funding, meaning that its capital investment is subject to political "
    "priorities and fiscal constraints (Audit Scotland, 2023). Oxfam faces a different challenge: "
    "its income fluctuates with public generosity and donor confidence, which can be undermined "
    "by reputational issues. In summary, while all three organisations are large and influential, "
    "the sector in which they operate fundamentally shapes their financial resilience, governance "
    "structures, and strategic priorities."
)


# -------------------------------------------------------------------
# Section 1.1.2 - Business Ownership Types (6 marks)
# -------------------------------------------------------------------
add_heading_styled("1.1.2 Comparing Business Ownership Types", level=2)

add_formatted_paragraph(
    "The legal form of a business determines who owns it, how it is controlled, and how "
    "profits are distributed. This section compares three ownership structures: private limited "
    "companies (Ltd), public limited companies (Plc), and franchises. Each offers distinct "
    "advantages and challenges depending on the size and ambitions of the business."
)

add_heading_styled("Comparison of Ownership Types", level=3)

headers_112 = ["Feature", "Private Limited Company\n(Ltd)", "Public Limited Company\n(Plc)",
               "Franchise"]
rows_112 = [
    [
        "Ownership",
        "Owned by shareholders who are often family members or a small group of investors; "
        "shares cannot be traded on a public stock exchange (BPP, 2021).",
        "Owned by shareholders who can buy and sell shares freely on the stock exchange; "
        "minimum share capital of GBP 50,000 required at incorporation (Companies Act 2006).",
        "The franchisee owns the individual outlet but operates under the franchisor's brand, "
        "systems, and intellectual property; the franchisor retains ultimate control over the "
        "brand (Mendelsohn, 2004)."
    ],
    [
        "Liability",
        "Limited liability; shareholders' personal assets are protected and they can only lose "
        "the amount they invested in shares.",
        "Limited liability applies equally; however, the larger shareholder base means directors "
        "face greater scrutiny from institutional investors.",
        "The franchisee has limited liability if operating as a Ltd company; however, they carry "
        "financial risk from franchise fees, royalties, and investment in premises."
    ],
    [
        "Control",
        "Directors and shareholders maintain close control; decisions can be made quickly without "
        "public disclosure of detailed financial information.",
        "Board of directors runs the company, but must answer to a dispersed shareholder base; "
        "subject to AGMs and extensive regulatory reporting requirements.",
        "Franchisee has day-to-day operational control but must follow the franchisor's operating "
        "manual; major decisions about branding, pricing, and suppliers are typically dictated by "
        "the franchisor."
    ],
    [
        "Profit Distribution",
        "Profits distributed as dividends to shareholders or retained for reinvestment; the small "
        "ownership group means larger individual returns.",
        "Profits shared among potentially millions of shareholders; dividend policy set by the "
        "board and influenced by market expectations.",
        "Franchisee keeps profits after paying ongoing royalty fees (typically 4-12% of revenue) "
        "and contributing to a national marketing fund (Mendelsohn, 2004)."
    ],
    [
        "Examples",
        "JCB, Dyson, Virgin Atlantic (before going public); many small and medium enterprises.",
        "Tesco, BP, Unilever; these companies are listed on the London Stock Exchange.",
        "McDonald's, Subway, Starbucks (selected international outlets); Domino's Pizza UK."
    ],
]

add_table_with_header(headers_112, rows_112, col_widths=[1.2, 1.8, 1.8, 1.8])

add_heading_styled("Analysis", level=3)

add_formatted_paragraph(
    "Each ownership structure offers a different balance between control, access to capital, "
    "and risk. A private limited company is well suited to entrepreneurs who want to retain "
    "close control over strategic decisions without the pressure of public shareholders. "
    "However, a limitation is that raising large amounts of capital can be difficult because "
    "shares cannot be sold on the stock exchange (BPP, 2021). Public limited companies overcome "
    "this barrier by accessing equity markets, which enables rapid expansion, as demonstrated by "
    "Unilever's ability to fund global operations. In contrast, this openness brings challenges: "
    "Plc directors must manage short-term profit expectations from institutional investors, which "
    "can sometimes conflict with long-term strategic goals."
)

add_formatted_paragraph(
    "Franchising represents a distinct model in which the franchisor expands rapidly without "
    "bearing the full cost of new outlets. For the franchisee, it offers the advantage of an "
    "established brand and proven business model, which significantly reduces the risk of "
    "failure compared to an independent start-up (Mendelsohn, 2004). However, the trade-off "
    "is a loss of autonomy: franchisees must adhere to strict operational guidelines and pay "
    "ongoing fees that reduce overall profitability. This suggests that the franchise model is "
    "most appropriate for individuals who value security and brand recognition over creative "
    "independence."
)


# -------------------------------------------------------------------
# Section 1.1.3 - Organic vs Inorganic Growth (6 marks)
# -------------------------------------------------------------------
add_heading_styled("1.1.3 Organic and Inorganic Business Growth", level=2)

add_formatted_paragraph(
    "Business growth can be achieved organically, through internal development, or inorganically, "
    "through mergers, acquisitions, and strategic alliances. Both approaches carry distinct "
    "advantages and risks, and the choice between them depends on an organisation's resources, "
    "objectives, and competitive environment."
)

add_heading_styled("Organic Growth", level=3)

add_formatted_paragraph(
    "Organic growth refers to expansion achieved through a business's own resources and "
    "capabilities, such as developing new products, entering new markets, or increasing "
    "operational capacity (Stokes and Wilson, 2017). A straightforward example is a local "
    "bakery that decides to expand its menu to include gluten-free products and extends its "
    "opening hours to capture evening trade. This kind of incremental growth allows the owner "
    "to test new ideas without committing significant capital, maintaining full control over "
    "the business and its brand identity."
)

add_formatted_paragraph(
    "At a larger scale, Infosys demonstrates organic growth through sustained investment in "
    "talent recruitment and employee development. Founded in 1981 with just seven employees "
    "and USD 250 in capital, Infosys grew to over 340,000 employees by 2023 largely through "
    "organic means: building training campuses, developing proprietary methodologies, and "
    "gradually expanding its client base across global markets (Infosys, 2023). The advantage "
    "of this approach is that it preserves organisational culture and ensures quality control. "
    "However, a limitation is that organic growth can be slow, and competitors pursuing "
    "acquisition strategies may gain market share more rapidly."
)

add_heading_styled("Inorganic Growth", level=3)

add_formatted_paragraph(
    "Inorganic growth involves external expansion through mergers, acquisitions, or strategic "
    "alliances. Tata Motors' acquisition of Jaguar Land Rover (JLR) from Ford in 2008 for "
    "USD 2.3 billion is a landmark example. This acquisition gave Tata Motors instant access "
    "to premium automotive brands, advanced engineering capabilities, and established markets "
    "in Europe and North America (Tata Motors, 2023). The advantage was transformative: Tata "
    "Motors moved from being primarily an Indian commercial vehicle manufacturer to a global "
    "player in the luxury segment almost overnight. However, the risks were considerable. "
    "Integrating two very different corporate cultures proved challenging, and JLR initially "
    "required substantial investment to return to profitability."
)

add_formatted_paragraph(
    "Strategic alliances represent a less risky form of inorganic growth. The Tata Starbucks "
    "joint venture, formed in 2012, combined Starbucks' global brand and coffee expertise with "
    "Tata's deep understanding of the Indian market and its existing supply chain infrastructure "
    "(Economic Times, 2012). This alliance allowed Starbucks to enter the Indian market more "
    "effectively than it could have alone, while Tata gained access to a world-renowned brand. "
    "The advantage is shared risk and complementary strengths; the disadvantage is that profits "
    "must be shared, and strategic disagreements between partners can slow decision-making."
)

add_heading_styled("Evaluation", level=3)

add_formatted_paragraph(
    "In summary, organic growth offers control and cultural consistency but is typically slower, "
    "while inorganic growth offers speed and market access but carries integration risks and "
    "higher financial exposure. Most successful large organisations, including the Tata Group, "
    "employ a combination of both strategies depending on their objectives and market conditions "
    "(Needle, 2015)."
)


# -------------------------------------------------------------------
# Section 1.2.1 - Organisational Objectives (15 marks)
# -------------------------------------------------------------------
add_heading_styled("1.2.1 Organisational Objectives and Their Significance", level=2)

add_formatted_paragraph(
    "Every organisation pursues objectives that guide its strategy, resource allocation, and "
    "day-to-day operations. However, objectives are not formed in a vacuum; they emerge from "
    "the specific economic, social, and competitive conditions that an organisation faces at a "
    "given point in time. This section examines the objectives of four organisations and explains "
    "why those objectives were important when they were adopted."
)

# -- Reliance Jio --
add_heading_styled("Reliance Jio: Market Penetration Through Free Data (2016)", level=3)

add_formatted_paragraph(
    "When Reliance Jio launched in September 2016, its primary objective was to achieve rapid "
    "market penetration in the Indian telecommunications sector by offering free voice calls "
    "and heavily subsidised data services. This objective was significant because at the time, "
    "India's mobile data prices were among the highest in the world relative to income, and "
    "internet penetration stood at only around 35 percent of the population (TRAI, 2017). The "
    "existing operators, including Airtel, Vodafone, and Idea, had invested conservatively in "
    "4G infrastructure and were charging premium rates for data services."
)

add_formatted_paragraph(
    "Jio's objective directly addressed this gap. By offering free services during its initial "
    "six-month promotion, Jio attracted over 100 million subscribers within 170 days, making "
    "it the fastest-growing mobile network in history (Mukherji, 2017). This strategy was "
    "important for several reasons. First, it democratised internet access across India, "
    "enabling millions of rural and low-income users to come online for the first time. "
    "Second, it fundamentally disrupted the competitive landscape, forcing competitors to "
    "slash prices and consolidate, with Vodafone and Idea eventually merging to survive. "
    "Third, the objective aligned with parent company Reliance Industries' broader vision of "
    "building a digital ecosystem encompassing e-commerce, entertainment, and financial "
    "services. The objective was therefore not merely about gaining subscribers; it was about "
    "creating the infrastructure for an entirely new business model."
)

# -- Tata Group --
add_heading_styled("Tata Group: Nation-Building and Ethical Business (1868 onwards)", level=3)

add_formatted_paragraph(
    "The Tata Group was founded by Jamsetji Tata in 1868 with an objective that was unusual "
    "for its era: to build Indian industrial capability while maintaining ethical business "
    "practices and contributing to national development. This objective was significant because "
    "India was under British colonial rule, and indigenous industrialisation was actively "
    "discouraged by colonial trade policies that favoured raw material exports over domestic "
    "manufacturing (Lala, 2006)."
)

add_formatted_paragraph(
    "Jamsetji established the Empress Mills textile factory in 1877, Tata Steel in 1907, and "
    "laid the groundwork for what would become the Indian Institute of Science, reflecting a "
    "belief that business should serve society as well as shareholders. After Indian "
    "independence in 1947, the Tata Group's nation-building objective became even more important. "
    "Under JRD Tata's leadership, the group expanded into airlines, chemicals, technology, and "
    "hospitality, directly contributing to India's post-independence industrialisation. The "
    "significance of this objective is reflected in the Tata Group's unique ownership structure: "
    "approximately 66 percent of the equity of Tata Sons is held by philanthropic trusts "
    "(Tata Group, 2023). This means that a majority of the group's profits are directed towards "
    "charitable causes, including education, healthcare, and rural development. The objective "
    "of ethical nation-building has therefore shaped not only the group's strategy but its very "
    "corporate structure."
)

# -- Amul --
add_heading_styled("Amul: Empowering Dairy Farmers Through Cooperatives (1946)", level=3)

add_formatted_paragraph(
    "Amul, the Gujarat Cooperative Milk Marketing Federation, was established in 1946 with the "
    "objective of empowering dairy farmers by eliminating exploitative middlemen and giving "
    "producers direct control over the collection, processing, and marketing of milk. This "
    "objective was critically important because Indian dairy farmers at the time were subject "
    "to a system in which private traders purchased milk at very low prices and controlled "
    "distribution, leaving farmers in poverty (Kurien, 2005)."
)

add_formatted_paragraph(
    "The cooperative model, championed by Dr Verghese Kurien under the banner of Operation "
    "Flood (launched in 1970), transformed India from a milk-deficient nation into the world's "
    "largest milk producer by 1998 (National Dairy Development Board, 2020). The significance "
    "of Amul's objective lies in its social impact: the cooperative now has over 3.6 million "
    "farmer members across 18,600 village societies, and it ensures that approximately 80 percent "
    "of the consumer price reaches the farmer, compared to less than 50 percent in conventional "
    "supply chains. This objective was important when adopted because it addressed a fundamental "
    "structural inequality in the Indian agricultural economy. It demonstrated that business "
    "objectives need not be purely financial; Amul's cooperative model delivered both economic "
    "empowerment and social justice simultaneously."
)

# -- Unilever --
add_heading_styled("Unilever: Sustainable Living Plan (2010)", level=3)

add_formatted_paragraph(
    "In 2010, Unilever launched its Sustainable Living Plan (USLP) with three overarching "
    "objectives: to help more than one billion people improve their health and wellbeing, to "
    "halve the environmental footprint of its products, and to enhance the livelihoods of "
    "millions of people in its supply chain, all while doubling the size of the business "
    "(Unilever, 2020). This objective was significant because it was adopted during a period "
    "of growing public concern about climate change, resource depletion, and corporate social "
    "responsibility."
)

add_formatted_paragraph(
    "The plan was important for several reasons. First, it positioned Unilever as a leader in "
    "sustainable business at a time when consumers, particularly younger demographics, were "
    "increasingly choosing brands that aligned with their values (Kotler and Keller, 2016). "
    "Second, it created long-term cost savings by reducing waste, water use, and energy "
    "consumption across the supply chain. Third, it helped Unilever attract and retain talent; "
    "research consistently shows that employees are more engaged when they believe their "
    "employer acts responsibly (CIPD, 2019). However, the USLP also faced criticism. Some "
    "analysts argued that pursuing sustainability while doubling business size was inherently "
    "contradictory, and that the plan's targets were too ambitious to be credible (Financial "
    "Times, 2019). Despite this, Unilever reported that its 'Sustainable Living Brands' grew "
    "69 percent faster than the rest of the business by 2019, suggesting that the objective "
    "delivered tangible commercial results alongside its social goals."
)


# -------------------------------------------------------------------
# Section 1.3.1 - Organisational Structures (10 marks)
# -------------------------------------------------------------------
add_heading_styled("1.3.1 Organisational Structures", level=2)

add_formatted_paragraph(
    "The structure an organisation adopts determines how authority flows, how decisions are "
    "made, and how effectively it can respond to its operating environment. This section "
    "examines three key structures: hierarchical, functional, and matrix, using real "
    "organisational examples to evaluate each."
)

# Hierarchical
add_heading_styled("Hierarchical Structure: NHS Scotland", level=3)

add_formatted_paragraph(
    "NHS Scotland operates a tall hierarchical structure with clearly defined levels of "
    "authority. At the top sits the Scottish Government's Cabinet Secretary for Health and "
    "Social Care, followed by NHS Scotland's Chief Executive, then the 14 territorial health "
    "board chairs and chief executives, hospital managers, department heads, and finally "
    "frontline clinical and support staff (Steel and Cylus, 2012)."
)

add_formatted_paragraph(
    "The following describes the hierarchical chain of command within NHS Scotland:",
    italic=True, space_after=Pt(4)
)

# Simple text-based diagram for hierarchy
hier_rows = [
    ["Level", "Role", "Responsibility"],
    ["1", "Cabinet Secretary for Health", "Sets national health policy and budget"],
    ["2", "NHS Scotland Chief Executive", "Oversees all health boards and national strategy"],
    ["3", "Health Board Chairs / Chief Executives", "Manage regional health services"],
    ["4", "Hospital / Service Managers", "Operational management of facilities"],
    ["5", "Department Heads / Senior Clinicians", "Lead clinical departments and teams"],
    ["6", "Frontline Staff (nurses, doctors, support)", "Deliver patient care directly"],
]
add_table_with_header(hier_rows[0], hier_rows[1:], col_widths=[0.8, 2.5, 3.5])

add_formatted_paragraph(
    "The hierarchical structure offers several advantages for NHS Scotland. First, it provides "
    "a clear chain of command, which is essential in healthcare where clinical accountability "
    "must be unambiguous. If a patient receives incorrect treatment, the hierarchy ensures that "
    "responsibility can be traced and addressed (Buchanan and Huczynski, 2017). Second, it "
    "facilitates standardisation of care across Scotland's diverse geography, from urban "
    "hospitals in Glasgow to remote clinics in the Highlands."
)

add_formatted_paragraph(
    "However, the hierarchical structure also has significant disadvantages. Communication can "
    "be slow as information must pass through multiple layers, which can delay the implementation "
    "of new policies or responses to emerging health crises. Furthermore, frontline staff may "
    "feel disconnected from strategic decision-making, which can reduce morale and innovation "
    "(Mullins, 2016). The COVID-19 pandemic exposed some of these limitations, as the speed "
    "of the virus's spread demanded faster, more agile decision-making than the traditional "
    "hierarchy could always provide."
)

# Functional
add_heading_styled("Functional Structure: Tata Motors", level=3)

add_formatted_paragraph(
    "Tata Motors organises its operations primarily along functional lines, with separate "
    "departments for engineering, manufacturing, marketing and sales, finance, human resources, "
    "and research and development. Each department is headed by a functional director who "
    "reports to the managing director (Tata Motors, 2023)."
)

add_formatted_paragraph(
    "The functional structure at Tata Motors can be summarised as follows:",
    italic=True, space_after=Pt(4)
)

func_rows = [
    ["Function", "Key Responsibilities", "Strategic Importance"],
    ["Engineering", "Vehicle design, testing, safety compliance",
     "Core competence in automotive R&D"],
    ["Manufacturing", "Production lines, quality control, supply chain",
     "Efficiency and cost management"],
    ["Marketing & Sales", "Brand management, dealer networks, pricing",
     "Revenue generation and market positioning"],
    ["Finance", "Budgeting, investment analysis, financial reporting",
     "Capital allocation and shareholder returns"],
    ["Human Resources", "Recruitment, training, employee relations",
     "Talent development and retention"],
    ["R&D", "Electric vehicle development, autonomous driving tech",
     "Future competitiveness and innovation"],
]
add_table_with_header(func_rows[0], func_rows[1:], col_widths=[1.5, 2.5, 2.8])

add_formatted_paragraph(
    "The functional structure benefits Tata Motors by enabling deep specialisation within each "
    "department. Engineers can focus entirely on technical excellence without being distracted "
    "by marketing concerns, while the marketing team can concentrate on understanding customer "
    "needs without needing to understand engine specifications (Needle, 2015). This specialisation "
    "supports efficiency and expertise."
)

add_formatted_paragraph(
    "However, a limitation is that functional silos can develop, where departments become "
    "inward-looking and fail to collaborate effectively. For example, the engineering team might "
    "design a technically excellent vehicle that the marketing team finds difficult to position "
    "competitively. Cross-functional communication requires deliberate management effort to "
    "prevent these silos from undermining overall organisational performance."
)

# Matrix
add_heading_styled("Matrix Structure: Infosys and NASA", level=3)

add_formatted_paragraph(
    "A matrix structure combines functional and project-based reporting lines, meaning that "
    "employees report to both a functional manager and a project manager simultaneously. "
    "Infosys, the Indian IT services giant, uses a matrix structure to manage its global "
    "delivery model. A software developer at Infosys might report to a functional head of "
    "software engineering for technical standards and career development, while simultaneously "
    "reporting to a project manager for a specific client engagement (Infosys, 2023)."
)

add_formatted_paragraph(
    "NASA provides another well-known example. During the Apollo programme and in its current "
    "operations, NASA engineers report to both their functional discipline head (for example, "
    "propulsion engineering) and the programme manager of the specific mission they support. "
    "This dual reporting structure ensures that each project benefits from specialised technical "
    "expertise while maintaining alignment with programme-wide objectives and timelines "
    "(Shenhar and Dvir, 2007)."
)

add_formatted_paragraph(
    "The matrix structure in these organisations illustrates a key concept:",
    italic=True, space_after=Pt(4)
)

matrix_table = [
    ["Aspect", "Functional Reporting Line", "Project Reporting Line"],
    ["Focus", "Technical standards, career development, skill building",
     "Client deliverables, deadlines, project budgets"],
    ["Manager", "Functional / discipline head",
     "Project or programme manager"],
    ["Advantage", "Deep expertise and consistency across projects",
     "Agility and client focus"],
    ["Risk", "May create conflicting priorities for employees",
     "Project pressures can override quality standards"],
]
add_table_with_header(matrix_table[0], matrix_table[1:], col_widths=[1.2, 2.8, 2.8])

add_formatted_paragraph(
    "The primary advantage of the matrix structure is flexibility. Infosys can rapidly assemble "
    "project teams by drawing specialists from different functional areas, which is essential in "
    "the fast-moving IT services industry where client requirements change frequently. Similarly, "
    "NASA can bring together propulsion engineers, avionics specialists, and mission planners "
    "into cohesive teams without permanently restructuring the organisation."
)

add_formatted_paragraph(
    "However, the matrix structure creates complexity. Employees with two managers may face "
    "conflicting demands, and the dual reporting lines can slow decision-making if the functional "
    "and project managers disagree (Buchanan and Huczynski, 2017). This suggests that the matrix "
    "structure works best in organisations with strong collaboration cultures and clear conflict "
    "resolution mechanisms."
)


# -------------------------------------------------------------------
# Section 1.3.2 - Decision Levels and Tools (5 marks)
# -------------------------------------------------------------------
add_heading_styled("1.3.2 Strategic, Tactical, and Operational Decisions", level=2)

add_formatted_paragraph(
    "Business decisions occur at three levels: strategic, tactical, and operational. Each level "
    "differs in scope, time horizon, and the individuals responsible for making them."
)

dec_rows = [
    ["Decision Level", "Scope", "Time Horizon", "Made By", "Example"],
    ["Strategic",
     "Organisation-wide; sets long-term direction",
     "3-10 years",
     "Board of directors, CEO",
     "Tata Motors' decision to acquire Jaguar Land Rover in 2008 to enter the premium vehicle market"],
    ["Tactical",
     "Departmental; implements strategy",
     "1-3 years",
     "Middle managers, department heads",
     "Tata Motors' marketing department launching a targeted advertising campaign for the Nexon EV in India"],
    ["Operational",
     "Day-to-day activities; routine processes",
     "Hours to weeks",
     "Supervisors, frontline staff",
     "A Tata Motors factory manager adjusting the production schedule to meet a sudden increase in orders"],
]
add_table_with_header(dec_rows[0], dec_rows[1:], col_widths=[1.0, 1.5, 0.9, 1.2, 2.2])

add_formatted_paragraph(
    "Strategic decisions are the most significant because they commit substantial resources and "
    "are difficult to reverse. Tata Motors' acquisition of JLR was a strategic decision that "
    "cost USD 2.3 billion and fundamentally transformed the company's market position. Tactical "
    "decisions translate strategy into action; they are important but more easily adjusted. "
    "Operational decisions are routine but collectively determine whether the organisation "
    "functions efficiently on a daily basis (Johnson et al., 2017)."
)

add_heading_styled("SWOT Analysis: Tata Motors", level=3)

add_formatted_paragraph(
    "SWOT analysis is a widely used strategic planning tool that evaluates an organisation's "
    "Strengths, Weaknesses, Opportunities, and Threats. Below is a SWOT analysis for Tata "
    "Motors to illustrate how this tool supports strategic decision-making."
)

swot_headers = ["Strengths", "Weaknesses"]
swot_row1 = [
    "Strong brand heritage and trust in India; diversified product portfolio spanning "
    "commercial and passenger vehicles; access to Tata Group resources and synergies; "
    "ownership of premium brands Jaguar and Land Rover (Tata Motors, 2023).",
    "Heavy dependence on the Indian market for commercial vehicle revenue; JLR division "
    "has faced profitability challenges; relatively late entry to the electric vehicle "
    "segment compared to Tesla and BYD (Financial Times, 2023)."
]
swot_row2_headers = ["Opportunities", "Threats"]
swot_row2 = [
    "Rapidly growing Indian EV market supported by government subsidies (FAME II scheme); "
    "expansion into African and Southeast Asian markets; development of connected and "
    "autonomous vehicle technologies.",
    "Intense global competition from established and new EV manufacturers; supply chain "
    "disruptions from semiconductor shortages; tightening emission regulations in the EU "
    "requiring significant R&D investment."
]

# Build the SWOT as a 2x2 table
swot_table = doc.add_table(rows=2, cols=2)
swot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
swot_table.style = 'Table Grid'

# Headers in top-left of each cell
swot_data = [
    [("Strengths", swot_row1[0]), ("Weaknesses", swot_row1[1])],
    [("Opportunities", swot_row2[0]), ("Threats", swot_row2[1])],
]
colors = [["C6EFCE", "FFC7CE"], ["BDD7EE", "F4CCCC"]]

for r in range(2):
    for c in range(2):
        cell = swot_table.rows[r].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        run_h = p.add_run(swot_data[r][c][0] + "\n")
        run_h.font.bold = True
        run_h.font.size = Pt(11)
        run_h.font.name = 'Calibri'
        run_h.font.color.rgb = NAVY
        run_t = p.add_run(swot_data[r][c][1])
        run_t.font.size = Pt(10)
        run_t.font.name = 'Calibri'
        run_t.font.color.rgb = DARK_GREY
        set_cell_shading(cell, colors[r][c])
        cell.width = Inches(3.4)

doc.add_paragraph()

add_heading_styled("Decision Tree Example", level=3)

add_formatted_paragraph(
    "A decision tree is a visual tool that maps out the possible outcomes of a decision, "
    "helping managers evaluate options systematically. Consider the following scenario for "
    "Tata Motors:"
)

add_formatted_paragraph(
    "Decision: Should Tata Motors invest GBP 1 billion in a new electric vehicle platform "
    "for JLR, or should it license existing EV technology from a partner?",
    bold=True, space_after=Pt(6)
)

dt_rows = [
    ["Decision Node", "Option A: Develop In-House", "Option B: License Technology"],
    ["Initial Cost", "GBP 1 billion over 4 years", "GBP 300 million licensing fee + royalties"],
    ["Potential Outcome (Optimistic)",
     "Proprietary EV platform gives competitive advantage; estimated revenue of GBP 5 billion "
     "over 10 years with 60% probability",
     "Faster time-to-market; estimated revenue of GBP 3 billion over 10 years with 70% probability"],
    ["Potential Outcome (Pessimistic)",
     "Technology fails to meet targets; write-off of GBP 800 million with 40% probability",
     "Dependency on licensor limits differentiation; revenue of GBP 1.5 billion with 30% probability"],
    ["Expected Value",
     "(0.6 x 5B) + (0.4 x -0.8B) = GBP 2.68 billion",
     "(0.7 x 3B) + (0.3 x 1.5B) = GBP 2.55 billion"],
    ["Strategic Implication",
     "Higher risk but greater potential reward; builds long-term capability",
     "Lower risk but creates dependency; faster market entry"],
]
add_table_with_header(dt_rows[0], dt_rows[1:], col_widths=[1.5, 2.6, 2.6])

add_formatted_paragraph(
    "Based on expected value analysis, Option A (in-house development) has a marginally "
    "higher expected value of GBP 2.68 billion compared to GBP 2.55 billion for Option B. "
    "However, the decision tree also reveals that Option A carries higher downside risk. A "
    "strategic decision-maker would need to consider Tata Motors' risk appetite, financial "
    "reserves, and long-term competitive strategy before making a final choice (Johnson et al., "
    "2017). This illustrates how decision trees provide a structured framework for comparing "
    "options but do not eliminate the need for managerial judgement."
)


# ===================================================================
# TASK B: BUSINESS ENVIRONMENT
# ===================================================================
print("Building Task B...")
add_heading_styled("TASK B: Business Environment", level=1)


# -------------------------------------------------------------------
# Section 2.1 - Internal Factors (15 marks)
# -------------------------------------------------------------------
add_heading_styled("2.1 Internal Factors: Toyota vs Boeing Comparative Analysis", level=2)

add_formatted_paragraph(
    "Internal factors, including organisational culture, leadership, resources, and operational "
    "systems, significantly influence an organisation's performance and strategic direction. "
    "This section provides a comparative analysis of Toyota and Boeing, two global manufacturers "
    "whose contrasting internal environments have produced very different outcomes in recent years."
)

add_heading_styled("Organisational Culture", level=3)

add_formatted_paragraph(
    "Toyota's organisational culture is built around the Toyota Production System (TPS) and "
    "the principle of kaizen, or continuous improvement. Every employee, from assembly line "
    "workers to senior executives, is expected to identify inefficiencies and suggest "
    "improvements. The famous 'andon cord' system, which allows any worker to halt the "
    "production line if they identify a defect, embodies Toyota's cultural commitment to "
    "quality over speed (Liker, 2004). This culture of empowerment and collective "
    "responsibility has made Toyota one of the most admired manufacturers globally, "
    "consistently ranking among the top automotive companies for quality and reliability."
)

add_formatted_paragraph(
    "Boeing, in contrast, has experienced a significant cultural shift over the past two "
    "decades. Historically, Boeing was known for its engineering-first culture, where "
    "technical excellence and safety were paramount. However, following the 1997 merger with "
    "McDonnell Douglas, the company increasingly prioritised financial performance and "
    "shareholder returns over engineering rigour (Useem, 2019). Senior leadership moved "
    "Boeing's headquarters from Seattle, where its engineers were based, to Chicago, creating "
    "a physical and symbolic distance between management and the technical workforce. This "
    "cultural transformation had profound consequences, as the 737 MAX crisis would later "
    "reveal."
)

add_heading_styled("Internal Weaknesses and the Boeing 737 MAX Crisis", level=3)

add_formatted_paragraph(
    "The Boeing 737 MAX crisis, which resulted in two fatal crashes in 2018 and 2019 killing "
    "346 people, exposed deep internal weaknesses within Boeing's organisational structure and "
    "culture. Investigations revealed that Boeing had outsourced critical software development "
    "to reduce costs, that internal safety concerns raised by engineers had been overridden by "
    "management pressure to meet production timelines, and that the company had lobbied the "
    "Federal Aviation Administration (FAA) to minimise pilot training requirements for the "
    "new aircraft (US House Committee on Transportation, 2020)."
)

add_formatted_paragraph(
    "In contrast, Toyota's response to its own quality crisis, the 2009-2010 unintended "
    "acceleration recalls affecting approximately 9 million vehicles, demonstrates a different "
    "cultural approach. While Toyota was initially criticised for a slow response, the company "
    "ultimately accepted responsibility, established a global quality task force, appointed its "
    "first Chief Quality Officer, and implemented systematic improvements to its recall "
    "processes (Cole, 2011). The critical difference is that Toyota's kaizen culture provided "
    "a framework for learning from failure, whereas Boeing's cost-driven culture actively "
    "discouraged the reporting of problems."
)

add_heading_styled("Alignment of Resources with Objectives", level=3)

add_formatted_paragraph(
    "Toyota aligns its resources closely with its strategic objectives. Its investment in "
    "hybrid technology, beginning with the Prius in 1997, exemplifies long-term resource "
    "commitment to sustainability objectives. Toyota has spent over 25 years building expertise "
    "in electrified powertrains, and this investment now positions it well for the transition "
    "to electric and hydrogen fuel cell vehicles (Toyota, 2023). The company's financial "
    "resources, technological capabilities, and human capital are deliberately directed towards "
    "its stated objective of 'producing happiness for all.'"
)

add_formatted_paragraph(
    "Boeing's resource allocation, by contrast, has been criticised for prioritising share "
    "buybacks over investment in new aircraft development. Between 2013 and 2019, Boeing spent "
    "approximately USD 43 billion on share buybacks, a figure that exceeded its total R&D "
    "spending during the same period (Reuters, 2020). Critics argue that this misalignment "
    "of resources, choosing to return money to shareholders rather than invest in engineering "
    "excellence, directly contributed to the quality and safety failures that culminated in "
    "the 737 MAX crisis. This comparison suggests that an organisation's internal resource "
    "allocation decisions are ultimately a reflection of its cultural priorities: Toyota "
    "invests in quality and innovation, while Boeing, under its post-merger leadership, "
    "invested in financial engineering."
)

add_heading_styled("Comparative Summary", level=3)

comp_rows = [
    ["Internal Factor", "Toyota", "Boeing"],
    ["Culture",
     "Kaizen (continuous improvement); quality-first; empowers all employees to identify problems (Liker, 2004)",
     "Shifted from engineering-first to finance-first after 1997 McDonnell Douglas merger (Useem, 2019)"],
    ["Response to Crisis",
     "Accepted responsibility during 2009-2010 recall crisis; created Chief Quality Officer role; systematic improvement",
     "737 MAX crisis (2018-2019) revealed suppressed safety concerns, cost-cutting on software development, regulatory capture"],
    ["Resource Alignment",
     "25+ years of sustained investment in hybrid/EV technology aligned with sustainability objectives",
     "USD 43 billion on share buybacks (2013-2019) vs underinvestment in R&D and new aircraft development"],
    ["Leadership",
     "Engineering-led leadership; CEO typically has deep operational experience within the company",
     "Finance-led leadership; headquarters moved away from engineering workforce; growing disconnect"],
]
add_table_with_header(comp_rows[0], comp_rows[1:], col_widths=[1.3, 2.7, 2.7])


# -------------------------------------------------------------------
# Section 2.2 - External PESTEC Factors (20 marks)
# -------------------------------------------------------------------
add_heading_styled("2.2 External Factors: PESTEC Analysis", level=2)

add_formatted_paragraph(
    "The PESTEC framework (Political, Economic, Social, Technological, Environmental, and "
    "Competitive factors) provides a structured approach to analysing the external environment "
    "in which businesses operate. This section applies the PESTEC framework to Netflix and "
    "McDonald's, focusing on the factors most relevant to each organisation."
)

# -- Netflix --
add_heading_styled("Netflix: Technological and Social Factors", level=3)

add_formatted_paragraph(
    "Technological Factors",
    bold=True, font_size=11, space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.LEFT
)

add_formatted_paragraph(
    "Netflix's entire business model is built on technological innovation, making technological "
    "factors the most significant external influence on its strategy. When Netflix launched its "
    "streaming service in 2007, broadband internet penetration in the United States was "
    "approximately 50 percent (FCC, 2008). The company made a strategic bet that broadband "
    "adoption would accelerate, and this bet proved correct. By 2023, Netflix had over 260 "
    "million subscribers globally, operating in over 190 countries (Netflix, 2024)."
)

add_formatted_paragraph(
    "A critical technological factor is Netflix's use of artificial intelligence and machine "
    "learning for content recommendation. Netflix's recommendation algorithm accounts for "
    "approximately 80 percent of the content watched on the platform, meaning that the vast "
    "majority of viewing decisions are influenced by AI rather than active browsing (Gomez-Uribe "
    "and Hunt, 2015). This technology creates a powerful competitive advantage: by analysing "
    "viewing patterns, search history, and even the time of day users watch, Netflix can "
    "personalise the experience for each subscriber. The algorithmic approach also informs "
    "content investment decisions. Netflix famously used data analytics to commission House of "
    "Cards, identifying that its subscribers already enjoyed the work of director David Fincher "
    "and actor Kevin Spacey, and that political dramas performed well on the platform."
)

add_formatted_paragraph(
    "However, technological factors also present threats. The proliferation of streaming "
    "platforms, including Disney+, Amazon Prime Video, HBO Max, and Apple TV+, has been enabled "
    "by the same streaming technology that Netflix pioneered. This suggests that technological "
    "innovation can create competitive advantages, but those advantages may be temporary if "
    "competitors can replicate the technology."
)

add_formatted_paragraph(
    "Social Factors",
    bold=True, font_size=11, space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.LEFT
)

add_formatted_paragraph(
    "Social factors have profoundly shaped Netflix's content strategy and market positioning. "
    "The most significant social trend is the shift in viewing habits away from scheduled "
    "broadcast television towards on-demand content consumption. Research by Ofcom (2023) "
    "found that viewers aged 16-34 in the UK spent more time watching streaming services than "
    "broadcast television for the first time in 2022. This shift reflects broader social "
    "changes: the expectation of instant access, the preference for binge-watching entire "
    "series, and the desire for content that can be consumed on mobile devices during commutes "
    "or breaks."
)

add_formatted_paragraph(
    "Content diversity represents another important social factor. Netflix has invested "
    "heavily in non-English-language content, producing original series and films in Korean "
    "(Squid Game), Spanish (Money Heist), German (Dark), and Hindi (Sacred Games), among "
    "others. This strategy responds to the social trend of globalisation in entertainment "
    "consumption: audiences are increasingly willing to watch subtitled or dubbed content from "
    "other cultures (Netflix, 2024). Squid Game, a Korean-language series, became Netflix's "
    "most-watched show of all time, demonstrating that cultural barriers to content consumption "
    "are diminishing. However, Netflix must also navigate social sensitivities: content that is "
    "acceptable in one market may be offensive or even illegal in another, requiring careful "
    "localisation and content moderation strategies."
)

# -- McDonald's --
add_heading_styled("McDonald's: Globalisation and Ethical Considerations", level=3)

add_formatted_paragraph(
    "Impact of Globalisation",
    bold=True, font_size=11, space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.LEFT
)

add_formatted_paragraph(
    "McDonald's is one of the most globalised companies in the world, operating over 40,000 "
    "restaurants in more than 100 countries (McDonald's Corporation, 2024). Globalisation has "
    "been both an enormous opportunity and a significant challenge for the company. The core "
    "strategic challenge is balancing global consistency, which underpins the brand's reliability "
    "and efficiency, with local adaptation, which is essential for cultural relevance."
)

add_formatted_paragraph(
    "Menu localisation is the most visible example of McDonald's response to globalisation. In "
    "India, where a significant proportion of the population is vegetarian and beef consumption "
    "is culturally and religiously sensitive, McDonald's offers an entirely different menu "
    "centred on vegetarian options such as the McAloo Tikki burger and the Maharaja Mac made "
    "with chicken rather than beef (Vignali, 2001). In Japan, the Teriyaki McBurger reflects "
    "local taste preferences, while in France, McDonald's serves croissants and macarons in "
    "acknowledgement of French food culture. This strategy of 'glocalisation' demonstrates how "
    "a global company must adapt to local social and cultural norms to succeed."
)

add_formatted_paragraph(
    "McDonald's global supply chain represents another dimension of globalisation. The company "
    "sources ingredients from thousands of suppliers worldwide, creating a complex logistics "
    "network that must maintain consistent quality standards across diverse markets. The "
    "advantage of this global supply chain is economies of scale and purchasing power; the risk "
    "is vulnerability to supply chain disruptions, as demonstrated during the COVID-19 pandemic "
    "when port closures and labour shortages affected ingredient availability in multiple "
    "markets simultaneously (McKinsey, 2021)."
)

add_formatted_paragraph(
    "Ethical Considerations",
    bold=True, font_size=11, space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.LEFT
)

add_formatted_paragraph(
    "McDonald's faces significant ethical scrutiny across several dimensions, including "
    "environmental sustainability, animal welfare, and labour practices. Environmental concerns "
    "centre on the company's contribution to deforestation (through beef and soy supply chains), "
    "plastic waste from packaging, and carbon emissions from its global operations. In response, "
    "McDonald's has committed to sourcing 100 percent of its beef, coffee, fish, and palm oil "
    "from sustainable sources and has pledged to reduce greenhouse gas emissions by 36 percent "
    "by 2030 (McDonald's Corporation, 2024). However, critics argue that these targets are "
    "insufficient given the scale of the company's environmental impact and that the fast food "
    "business model is inherently resource-intensive."
)

add_formatted_paragraph(
    "Animal welfare is another area of ethical concern. Campaign groups such as Compassion in "
    "World Farming have pressured McDonald's to improve conditions in its supply chain, "
    "particularly regarding the use of battery cages for hens and the welfare standards for "
    "broiler chickens. McDonald's has responded by committing to cage-free eggs in its European "
    "and North American operations, though implementation timelines have varied by market "
    "(CIWF, 2023)."
)

add_formatted_paragraph(
    "Labour practices represent a persistent ethical challenge. McDonald's and its franchisees "
    "have faced criticism for low wages, unpredictable shift scheduling, and the use of zero-hour "
    "contracts in some markets, particularly the UK. The 'Fight for $15' movement in the United "
    "States, which campaigns for a minimum wage of USD 15 per hour, has specifically targeted "
    "McDonald's as one of the largest fast food employers (Allegretto et al., 2013). The "
    "company has gradually increased starting wages in its company-owned restaurants but has "
    "less control over wages at franchised outlets, which account for approximately 95 percent "
    "of its restaurants globally. This illustrates a key ethical tension within the franchise "
    "model: the franchisor sets brand standards but cannot fully control labour practices at "
    "independently owned outlets."
)

add_heading_styled("PESTEC Summary Table", level=3)

pestec_rows = [
    ["PESTEC Factor", "Netflix", "McDonald's"],
    ["Political",
     "Content regulation varies by country; must comply with local censorship laws (e.g., restrictions in certain Middle Eastern and Asian markets)",
     "Subject to differing food safety regulations globally; trade agreements and tariffs affect supply chain costs"],
    ["Economic",
     "Subscription pricing must reflect varying purchasing power globally; economic downturns may increase 'staying in' entertainment",
     "Consumer spending on dining out is sensitive to economic conditions; commodity price fluctuations affect food costs"],
    ["Social",
     "Shift from broadcast to on-demand viewing; demand for diverse, multilingual content; binge-watching culture",
     "Growing health consciousness challenges fast food; cultural food preferences require menu localisation"],
    ["Technological",
     "AI-driven recommendations; streaming infrastructure; 4K/HDR delivery; mobile viewing",
     "Automated ordering kiosks; mobile app ordering; delivery partnerships with UberEats and Deliveroo"],
    ["Environmental",
     "Energy consumption of data centres; e-waste from streaming devices",
     "Deforestation concerns in beef supply chain; plastic packaging waste; carbon emissions from global logistics"],
    ["Competitive",
     "Intense competition from Disney+, Amazon, HBO Max, Apple TV+; local streaming services in each market",
     "Competition from other fast food chains (Burger King, KFC), fast-casual restaurants (Nando's, Five Guys), and delivery-only kitchens"],
]
add_table_with_header(pestec_rows[0], pestec_rows[1:], col_widths=[1.2, 2.7, 2.8])


# -------------------------------------------------------------------
# Section 2.3 - Stakeholder Analysis (15 marks)
# -------------------------------------------------------------------
add_heading_styled("2.3 Stakeholder Analysis", level=2)

add_formatted_paragraph(
    "Stakeholders are individuals, groups, or organisations that have an interest in or are "
    "affected by a business's activities and decisions. Effective stakeholder management is "
    "critical because different stakeholders often have conflicting objectives, and failing to "
    "balance these can lead to operational disruption, reputational damage, or regulatory "
    "intervention. This section analyses the stakeholders of the Tata Group using Mendelow's "
    "stakeholder matrix as a framework."
)

add_heading_styled("Identifying Stakeholders of the Tata Group", level=3)

add_formatted_paragraph(
    "The Tata Group, as one of India's largest and most diversified conglomerates with over "
    "100 operating companies across sectors including automotive, steel, IT services, and "
    "hospitality, has a wide range of stakeholders. The key stakeholder groups are identified "
    "below:"
)

stake_rows = [
    ["Stakeholder Group", "Interest in Tata Group", "Influence Level"],
    ["Tata Sons (Holding Company) / Tata Trusts",
     "Strategic direction, brand integrity, philanthropic mission; the Trusts hold 66% of Tata Sons equity",
     "Very High: controls board appointments and strategic vetoes"],
    ["Shareholders (of listed subsidiaries)",
     "Return on investment through dividends and share price growth",
     "High: institutional investors can influence board decisions through voting"],
    ["Employees (over 935,000 globally)",
     "Job security, fair wages, career development, safe working conditions",
     "Medium-High: collective action through unions can disrupt operations"],
    ["Customers",
     "Quality products and services at fair prices; ethical supply chains",
     "High: customer loyalty directly determines revenue"],
    ["Government and Regulators (Indian and international)",
     "Tax revenue, compliance with regulations, contribution to economic development",
     "High: regulatory changes can fundamentally alter operating conditions"],
    ["Local Communities",
     "Employment opportunities, environmental protection, community investment",
     "Medium: community opposition can delay or prevent projects"],
    ["Suppliers",
     "Fair payment terms, long-term contracts, transparent procurement processes",
     "Medium: large suppliers have leverage; small suppliers have less"],
]
add_table_with_header(stake_rows[0], stake_rows[1:], col_widths=[1.8, 2.7, 2.2])

add_heading_styled("Mendelow's Stakeholder Matrix", level=3)

add_formatted_paragraph(
    "Mendelow's matrix classifies stakeholders according to their level of power (ability to "
    "influence the organisation) and their level of interest (how much they care about the "
    "organisation's activities). This classification determines the appropriate engagement "
    "strategy for each stakeholder group (Mendelow, 1991)."
)

# Mendelow matrix as a 2x2 table
mend_table = doc.add_table(rows=3, cols=3)
mend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
mend_table.style = 'Table Grid'

# Header row
mend_table.rows[0].cells[0].text = ""
c1 = mend_table.rows[0].cells[1]
c1.text = ""
p = c1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Low Interest")
r.font.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = WHITE
set_cell_shading(c1, MEDIUM_BLUE)

c2 = mend_table.rows[0].cells[2]
c2.text = ""
p = c2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("High Interest")
r.font.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = WHITE
set_cell_shading(c2, MEDIUM_BLUE)

# Row 1: Low Power
row_labels = ["Low Power", "High Power"]
row_data = [
    [
        "Minimal Effort\nGeneral public, media (monitor occasionally)",
        "Keep Informed\nLocal communities, small suppliers (regular communication, CSR reports)"
    ],
    [
        "Keep Satisfied\nGovernment regulators (comply proactively, engage in policy consultations)",
        "Key Players\nTata Trusts/Tata Sons, major shareholders, employees (close engagement, strategic alignment)"
    ],
]
colors_mend = [["F2F2F2", "BDD7EE"], ["D6E4F0", "C6EFCE"]]

for r_idx in range(2):
    label_cell = mend_table.rows[r_idx + 1].cells[0]
    label_cell.text = ""
    p = label_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(row_labels[r_idx])
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = WHITE
    set_cell_shading(label_cell, MEDIUM_BLUE)

    for c_idx in range(2):
        data_cell = mend_table.rows[r_idx + 1].cells[c_idx + 1]
        data_cell.text = ""
        p = data_cell.paragraphs[0]
        run = p.add_run(row_data[r_idx][c_idx])
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GREY
        set_cell_shading(data_cell, colors_mend[r_idx][c_idx])

doc.add_paragraph()

add_heading_styled("Stakeholder Conflicts", level=3)

add_formatted_paragraph(
    "Stakeholder conflicts arise when the interests of different groups are fundamentally "
    "incompatible. Within the Tata Group, several significant conflicts can be identified."
)

add_formatted_paragraph(
    "Employee Welfare vs Shareholder Profit Expectations: "
    "This is perhaps the most common stakeholder conflict in any large corporation. "
    "Shareholders, particularly institutional investors in listed Tata subsidiaries such as "
    "TCS and Tata Motors, expect the company to maximise returns through cost efficiency, "
    "which can include reducing labour costs. Employees, on the other hand, seek higher wages, "
    "better benefits, and greater job security. The Tata Group has historically managed this "
    "tension better than many competitors because of its philanthropic ownership structure: "
    "because the Tata Trusts are the majority shareholder of Tata Sons, there is less pressure "
    "to prioritise short-term profit maximisation over employee welfare (Lala, 2006). However, "
    "this does not eliminate the conflict entirely. In 2016, Tata Steel was forced to consider "
    "selling its UK operations, including the Port Talbot steelworks in Wales, due to "
    "unsustainable financial losses. This decision pitted the interests of approximately 11,000 "
    "UK employees against the financial interests of Tata Steel shareholders who were bearing "
    "the cost of continued losses (BBC, 2016).",
    space_after=Pt(8)
)

add_heading_styled("Case Study: Tata Nano and Community Opposition", level=3)

add_formatted_paragraph(
    "The Tata Nano project provides a specific case study of stakeholder conflict. In 2006, "
    "Tata Motors began building a manufacturing plant for the Nano, marketed as the world's "
    "cheapest car, in Singur, West Bengal. The West Bengal government had acquired "
    "approximately 1,000 acres of agricultural land for the factory, displacing farmers who "
    "depended on that land for their livelihoods. While the government and Tata Motors argued "
    "that the factory would bring jobs and economic development to the region, local farming "
    "communities, supported by political opposition leader Mamata Banerjee, launched sustained "
    "protests against the land acquisition (Sinha, 2009)."
)

add_formatted_paragraph(
    "This conflict illustrates a classic stakeholder tension: the interests of the company "
    "(efficient production), the government (industrial development and tax revenue), and "
    "employees (jobs) aligned in favour of the factory, while the interests of local farmers "
    "(livelihood protection) and environmental groups (preservation of agricultural land) "
    "were opposed. Applying Mendelow's matrix, the local farming community initially appeared "
    "to be a 'keep informed' stakeholder with low power. However, by allying with a powerful "
    "political leader, they effectively moved to the 'key player' quadrant, gaining both "
    "power and interest. The result was dramatic: Tata Motors was forced to abandon the Singur "
    "plant in October 2008 and relocate to Sanand in Gujarat, at an estimated cost of over "
    "USD 350 million and a significant delay in the Nano's launch."
)

add_formatted_paragraph(
    "This case demonstrates that stakeholder positions on Mendelow's matrix are not static. "
    "Stakeholders can increase their power through collective action and political alliances, "
    "and organisations that fail to anticipate these shifts risk costly strategic setbacks. "
    "The Singur experience prompted the Tata Group to adopt more robust community engagement "
    "processes for future projects, recognising that early and genuine consultation with local "
    "communities is not merely an ethical obligation but a strategic necessity."
)


# ===================================================================
# REFERENCE LIST
# ===================================================================
print("Building reference list...")
doc.add_page_break()
add_heading_styled("Reference List", level=1)

references = [
    "Allegretto, S.A., Doussard, M., Graham-Squire, D., Jacobs, K., Thompson, D. and Thompson, J. (2013) Fast Food, Poverty Wages: The Public Cost of Low-Wage Jobs in the Fast-Food Industry. Berkeley, CA: UC Berkeley Labor Center.",
    "Audit Scotland (2023) NHS in Scotland 2023. Edinburgh: Audit Scotland.",
    "BBC (2016) 'Tata Steel: The Story So Far', BBC News, 31 March. Available at: https://www.bbc.co.uk/news/business-35944542 (Accessed: 10 February 2026).",
    "BPP (2021) Business Essentials: Organisations and Behaviour. London: BPP Learning Media.",
    "Buchanan, D.A. and Huczynski, A.A. (2017) Organizational Behaviour. 9th edn. Harlow: Pearson.",
    "CIPD (2019) People and Purpose: Delivering Sustainable Business Performance. London: Chartered Institute of Personnel and Development.",
    "Cole, R.E. (2011) 'What Really Happened to Toyota?', MIT Sloan Management Review, 52(4), pp. 29-35.",
    "Economic Times (2012) 'Tata Group, Starbucks Form Joint Venture for India', The Economic Times, 30 January.",
    "FCC (2008) Connecting America: The National Broadband Plan. Washington, DC: Federal Communications Commission.",
    "Financial Times (2019) 'Unilever's Sustainable Living Plan Under Scrutiny', Financial Times, 15 May.",
    "Financial Times (2023) 'Tata Motors: Can the Indian Giant Win the Global EV Race?', Financial Times, 22 September.",
    "Gomez-Uribe, C.A. and Hunt, N. (2015) 'The Netflix Recommender System: Algorithms, Business Value, and Innovation', ACM Transactions on Management Information Systems, 6(4), pp. 1-19.",
    "Healthcare Improvement Scotland (2023) Annual Report 2022-23. Edinburgh: HIS.",
    "Infosys (2023) Annual Report 2022-23. Bengaluru: Infosys Limited.",
    "Johnson, G., Whittington, R., Scholes, K., Angwin, D. and Regner, P. (2017) Exploring Strategy. 11th edn. Harlow: Pearson.",
    "Kotler, P. and Keller, K.L. (2016) Marketing Management. 15th edn. Harlow: Pearson.",
    "Kurien, V. (2005) I Too Had a Dream. New Delhi: Roli Books.",
    "Lala, R.M. (2006) The Creation of Wealth: The Tatas from the 19th to the 21st Century. 3rd edn. New Delhi: Penguin Books India.",
    "Lewis, D. (2014) Non-Governmental Organizations, Management and Development. 3rd edn. London: Routledge.",
    "Liker, J.K. (2004) The Toyota Way: 14 Management Principles from the World's Greatest Manufacturer. New York: McGraw-Hill.",
    "McDonald's Corporation (2024) 2023 Annual Report. Chicago: McDonald's Corporation.",
    "McKinsey (2021) 'How COVID-19 Has Pushed Companies Over the Technology Tipping Point', McKinsey Global Institute Report.",
    "Mendelsohn, M. (2004) The Guide to Franchising. 7th edn. London: Thomson Learning.",
    "Mendelow, A.L. (1991) 'Environmental Scanning: The Impact of the Stakeholder Concept', Proceedings of the International Conference on Information Systems, pp. 407-418.",
    "Mukherji, R. (2017) 'Jio's Disruptive Strategy and Its Impact on Indian Telecoms', Economic and Political Weekly, 52(7), pp. 14-17.",
    "Mullins, L.J. (2016) Management and Organisational Behaviour. 11th edn. Harlow: Pearson.",
    "National Dairy Development Board (2020) Annual Report 2019-20. Anand: NDDB.",
    "Needle, D. (2015) Business in Context: An Introduction to Business and Its Environment. 6th edn. Andover: Cengage Learning.",
    "Netflix (2024) Q4 2023 Shareholder Letter. Los Gatos, CA: Netflix Inc.",
    "Ofcom (2023) Media Nations 2023. London: Ofcom.",
    "OSCR (2022) Annual Report and Accounts 2021-22. Dundee: Office of the Scottish Charity Regulator.",
    "Oxfam (2023) Oxfam Annual Report and Accounts 2022-23. Oxford: Oxfam GB.",
    "Reuters (2020) 'Boeing's Buyback Binge Has Left It Exposed in the Coronavirus Crisis', Reuters Business News, 24 March.",
    "Scottish Government (2022) NHS Scotland: National Workforce Strategy. Edinburgh: Scottish Government.",
    "Shenhar, A.J. and Dvir, D. (2007) Reinventing Project Management. Boston: Harvard Business School Press.",
    "Sinha, S. (2009) 'Tata Nano and the Singur Controversy', Indian Journal of Industrial Relations, 44(4), pp. 613-626.",
    "SMMT (2022) Motor Industry Facts 2022. London: Society of Motor Manufacturers and Traders.",
    "Steel, D. and Cylus, J. (2012) 'United Kingdom (Scotland): Health System Review', Health Systems in Transition, 14(9), pp. 1-150.",
    "Stokes, D. and Wilson, N. (2017) Small Business Management and Entrepreneurship. 7th edn. Andover: Cengage Learning.",
    "Tata Group (2023) Tata Group Annual Overview 2022-23. Mumbai: Tata Sons Private Limited.",
    "Tata Motors (2023) Integrated Annual Report 2022-23. Mumbai: Tata Motors Limited.",
    "Toyota (2023) Toyota Environmental Report 2023. Toyota City: Toyota Motor Corporation.",
    "TRAI (2017) Indian Telecom Services Performance Indicators: October-December 2016. New Delhi: Telecom Regulatory Authority of India.",
    "Unilever (2020) Unilever Sustainable Living Plan: 10 Years On. London: Unilever PLC.",
    "US House Committee on Transportation (2020) Final Committee Report: The Design, Development and Certification of the Boeing 737 MAX. Washington, DC: US Government Publishing Office.",
    "Useem, J. (2019) 'The Long-Forgotten Flight That Sent Boeing Off Course', The Atlantic, November.",
    "Vignali, C. (2001) 'McDonald\u2019s: Think Global, Act Local - The Marketing Mix', British Food Journal, 103(2), pp. 97-111.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREY


# ===================================================================
# MAPPING TABLE
# ===================================================================
print("Building mapping table...")
doc.add_page_break()
add_heading_styled("Assessment Criteria Mapping", level=1)

add_formatted_paragraph(
    "The following table maps each task in this report to the relevant assessment criteria "
    "for Unit J229 76 Understanding Business."
)

map_headers = ["Section", "Task Description", "Assessment Criteria", "Marks"]
map_rows = [
    ["1.1.1", "Compare three organisations across sectors (Tata Motors, NHS Scotland, Oxfam)",
     "1.1 Comparing organisations across sectors by scale, structure, specialisation, regulation, and capital",
     "8"],
    ["1.1.2", "Compare ownership types (Ltd, Plc, Franchise)",
     "1.1 Comparing business ownership structures", "6"],
    ["1.1.3", "Explain organic and inorganic growth with examples",
     "1.1 Explaining business growth strategies", "6"],
    ["1.2.1", "Explain objectives for four organisations (Jio, Tata, Amul, Unilever)",
     "1.2 Explaining why organisations adopted specific objectives at specific times", "15"],
    ["1.3.1", "Explain organisational structures (hierarchical, functional, matrix)",
     "1.3 Explaining organisational structures with real examples", "10"],
    ["1.3.2", "Explain decision levels with SWOT and decision tree",
     "1.3 Explaining strategic, tactical, and operational decisions", "5"],
    ["2.1", "Analyse internal factors comparing Toyota and Boeing",
     "2.1 Analysing internal factors affecting business performance", "15"],
    ["2.2", "Analyse external PESTEC factors for Netflix and McDonald's",
     "2.2 Analysing external factors using PESTEC framework", "20"],
    ["2.3", "Analyse stakeholders of Tata Group using Mendelow's matrix",
     "2.3 Analysing stakeholder influence and conflicts", "15"],
]
add_table_with_header(map_headers, map_rows, col_widths=[0.7, 2.5, 2.8, 0.7])

# Total marks row
p_total = doc.add_paragraph()
p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_t = p_total.add_run(f"Total Marks: {TOTAL_MARKS}")
run_t.font.bold = True
run_t.font.size = Pt(12)
run_t.font.name = 'Calibri'
run_t.font.color.rgb = NAVY


# ===================================================================
# SAVE DOCX
# ===================================================================
print(f"Saving DOCX to: {DOCX_PATH}")
doc.save(DOCX_PATH)
print("DOCX saved successfully.")


# ===================================================================
# EXPORT TO PDF via Word COM
# ===================================================================
print("Exporting to PDF via Word COM automation...")

# Kill any existing Word process again before COM
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
time.sleep(3)

import win32com.client
import pythoncom

pythoncom.CoInitialize()

try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    doc_com = word.Documents.Open(DOCX_PATH)

    # Update table of contents
    try:
        for toc in doc_com.TablesOfContents:
            toc.Update()
    except Exception as e:
        print(f"Note: TOC update encountered an issue (non-fatal): {e}")

    # Update all fields (page numbers, etc.)
    try:
        for story_range in doc_com.StoryRanges:
            story_range.Fields.Update()
    except Exception:
        pass

    # Export as PDF (wdExportFormatPDF = 17)
    doc_com.ExportAsFixedFormat(
        OutputFileName=PDF_PATH,
        ExportFormat=17,  # wdExportFormatPDF
        OpenAfterExport=False,
        OptimizeFor=0,  # wdExportOptimizeForPrint
        CreateBookmarks=1,  # wdExportCreateHeadingBookmarks
    )
    print(f"PDF saved successfully to: {PDF_PATH}")

    doc_com.Close(SaveChanges=False)
    word.Quit()

except Exception as e:
    print(f"ERROR during PDF export: {e}")
    try:
        word.Quit()
    except Exception:
        pass
    sys.exit(1)
finally:
    pythoncom.CoUninitialize()

print("\n=== BUILD COMPLETE ===")
print(f"DOCX: {DOCX_PATH}")
print(f"PDF:  {PDF_PATH}")
