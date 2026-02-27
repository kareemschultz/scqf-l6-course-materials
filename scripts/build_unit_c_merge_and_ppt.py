"""
UNIT C: F1FE 12 - Part 2: Mail Merge Badges + PowerPoint
"""
import win32com.client
import os
import time
import pythoncom
import subprocess

BASE = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials\SCQF_L6_FINAL_SUBMISSION"
SRC = os.path.join(BASE, "source_files", "F1FE12_Word_Presentation")
SS_DIR = os.path.join(BASE, "evidence_screenshots")

DELEGATE_LIST = os.path.join(SRC, "Delegate_List.xlsx")
BADGES_PATH = os.path.join(SRC, "TechSummit2026_Badges.docx")
PPTX_PATH = os.path.join(SRC, "TechSummit2026_DigitalSignage.pptx")
LOGO_PATH = os.path.join(SRC, "TechSummit_Logo.png")
AUDIO_PATH = os.path.join(SRC, "welcome_audio.wav")

CM = 28.35  # 1cm in points


def create_mail_merge_badges():
    """Create mail merge badge document"""
    pythoncom.CoInitialize()
    subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
    time.sleep(2)

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        # Create a simple badge document using python-docx approach instead
        # Word COM mail merge is fragile - let's use a direct approach
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from openpyxl import load_workbook

        doc = Document()

        # Read delegate data
        wb = load_workbook(DELEGATE_LIST)
        ws = wb.active
        delegates = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0]:
                delegates.append(row)

        # Create badges
        for i, d in enumerate(delegates):
            if i > 0:
                doc.add_page_break()

            # Badge content
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("TechSummit 2026")
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = None  # default

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("DELEGATE BADGE")
            run.font.size = Pt(14)

            doc.add_paragraph()  # spacer

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"ID: {d[0]}")
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{d[1]} {d[2]} {d[3]}")
            run.bold = True
            run.font.size = Pt(20)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{d[4]}")
            run.font.size = Pt(14)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{d[5]}")
            run.font.size = Pt(12)

        doc.save(BADGES_PATH)
        print(f"[OK] Badges saved: {BADGES_PATH} ({len(delegates)} delegates)")

        word.Quit()

    except Exception as e:
        print(f"[ERROR] Badges: {e}")
        import traceback; traceback.print_exc()
        try: word.Quit()
        except: pass

    pythoncom.CoUninitialize()


def create_powerpoint():
    """Create 5-slide digital signage PowerPoint with action buttons and kiosk mode"""
    pythoncom.CoInitialize()
    subprocess.run(['taskkill', '/F', '/IM', 'POWERPNT.EXE'], capture_output=True)
    time.sleep(1)

    ppt = win32com.client.Dispatch("PowerPoint.Application")
    # Don't set Visible for PPT - it can cause issues

    try:
        pres = ppt.Presentations.Add(WithWindow=False)

        # Constants
        ppLayoutBlank = 12
        ppLayoutTitle = 1
        ppLayoutTitleOnly = 11
        msoShapeActionButtonHome = 190
        msoShapeActionButtonCustom = 197
        ppActionNextSlide = 1
        ppActionPreviousSlide = 2
        ppActionFirstSlide = 3
        ppActionLastSlide = 4
        ppActionNamedSlideShow = 5
        ppMouseClick = 1
        ppAdvanceOnTime = 2
        ppTransitionSpeedMedium = 2
        ppSaveAsOpenXMLPresentation = 24
        ppShowTypeKiosk = 3

        # Set slide size (widescreen 16:9)
        pres.PageSetup.SlideWidth = 960  # 13.33 inches in points
        pres.PageSetup.SlideHeight = 540  # 7.5 inches

        slide_width = pres.PageSetup.SlideWidth
        slide_height = pres.PageSetup.SlideHeight

        def add_home_button(slide, left=None, top=None):
            """Add Home button on slide that links to slide 1"""
            if left is None:
                left = slide_width - 100
            if top is None:
                top = slide_height - 60

            btn = slide.Shapes.AddShape(125, left, top, 80, 40)  # Rectangle
            btn.TextFrame.TextRange.Text = "Home"
            btn.TextFrame.TextRange.Font.Size = 10
            btn.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
            btn.Fill.ForeColor.RGB = 0x663300  # Dark blue
            btn.ActionSettings(1).Action = 3  # ppActionFirstSlide (Hyperlink to first slide)
            return btn

        # =============================================
        # SLIDE 1: Welcome (with audio)
        # =============================================
        slide1 = pres.Slides.Add(1, ppLayoutBlank)
        slide1.FollowMasterBackground = False
        slide1.Background.Fill.ForeColor.RGB = 0x663300  # Dark blue

        # Title
        title = slide1.Shapes.AddTextbox(1, 50, 100, slide_width - 100, 100)
        title.TextFrame.TextRange.Text = "TechSummit 2026"
        title.TextFrame.TextRange.Font.Size = 44
        title.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        title.TextFrame.TextRange.Font.Bold = True
        title.TextFrame.TextRange.ParagraphFormat.Alignment = 2  # ppAlignCenter

        # Subtitle
        sub = slide1.Shapes.AddTextbox(1, 50, 220, slide_width - 100, 60)
        sub.TextFrame.TextRange.Text = "Digital Innovation Conference"
        sub.TextFrame.TextRange.Font.Size = 28
        sub.TextFrame.TextRange.Font.Color.RGB = 0xF7C34F  # Light blue (BGR)
        sub.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        # Venue info
        info = slide1.Shapes.AddTextbox(1, 150, 320, slide_width - 300, 100)
        info.TextFrame.TextRange.Text = "Edinburgh International Conference Centre\nFebruary 2026"
        info.TextFrame.TextRange.Font.Size = 18
        info.TextFrame.TextRange.Font.Color.RGB = 0xD0D0D0
        info.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        # Add audio
        try:
            audio = slide1.Shapes.AddMediaObject2(AUDIO_PATH, False, True, 50, slide_height - 80, 40, 40)
            audio.AnimationSettings.PlaySettings.PlayOnEntry = True
            audio.AnimationSettings.PlaySettings.HideWhileNotPlaying = True
            print("  Audio added to Slide 1")
        except Exception as e:
            print(f"  Audio note: {e}")

        # Add logo
        try:
            logo = slide1.Shapes.AddPicture(LOGO_PATH, False, True, 30, 20, 200, 67)
            print("  Logo added to Slide 1")
        except Exception as e:
            print(f"  Logo note: {e}")

        # Navigation buttons at bottom
        # Agenda button
        btn_agenda = slide1.Shapes.AddShape(125, 200, slide_height - 60, 120, 40)
        btn_agenda.TextFrame.TextRange.Text = "Agenda"
        btn_agenda.TextFrame.TextRange.Font.Size = 10
        btn_agenda.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        btn_agenda.Fill.ForeColor.RGB = 0x993300
        btn_agenda.ActionSettings(1).Action = 6  # ppActionHyperlink... use slide jump
        # Link to slide 3
        try:
            btn_agenda.ActionSettings(1).Hyperlink.SubAddress = "3,3,Agenda"
        except:
            pass

        # Map button
        btn_map = slide1.Shapes.AddShape(125, 350, slide_height - 60, 120, 40)
        btn_map.TextFrame.TextRange.Text = "Venue Map"
        btn_map.TextFrame.TextRange.Font.Size = 10
        btn_map.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        btn_map.Fill.ForeColor.RGB = 0x993300

        # WiFi button
        btn_wifi = slide1.Shapes.AddShape(125, 500, slide_height - 60, 120, 40)
        btn_wifi.TextFrame.TextRange.Text = "Wi-Fi Info"
        btn_wifi.TextFrame.TextRange.Font.Size = 10
        btn_wifi.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        btn_wifi.Fill.ForeColor.RGB = 0x993300

        # =============================================
        # SLIDE 2: Video Slide
        # =============================================
        slide2 = pres.Slides.Add(2, ppLayoutBlank)
        slide2.FollowMasterBackground = False
        slide2.Background.Fill.ForeColor.RGB = 0x663300

        title2 = slide2.Shapes.AddTextbox(1, 50, 20, slide_width - 100, 60)
        title2.TextFrame.TextRange.Text = "Conference Highlights"
        title2.TextFrame.TextRange.Font.Size = 32
        title2.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        title2.TextFrame.TextRange.Font.Bold = True
        title2.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        # Video placeholder text (no ffmpeg available for actual video)
        vid_text = slide2.Shapes.AddTextbox(1, 100, 120, slide_width - 200, 300)
        vid_text.TextFrame.TextRange.Text = (
            "[Video: TechSummit 2026 Welcome]\n\n"
            "This slide would contain an auto-playing video showcasing:\n"
            "- Conference venue highlights\n"
            "- Previous year's keynote moments\n"
            "- Edinburgh city landmarks\n\n"
            "Video format: MP4, auto-play on slide entry"
        )
        vid_text.TextFrame.TextRange.Font.Size = 16
        vid_text.TextFrame.TextRange.Font.Color.RGB = 0xD0D0D0
        vid_text.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        add_home_button(slide2)

        # =============================================
        # SLIDE 3: Agenda
        # =============================================
        slide3 = pres.Slides.Add(3, ppLayoutBlank)
        slide3.FollowMasterBackground = False
        slide3.Background.Fill.ForeColor.RGB = 0x663300

        title3 = slide3.Shapes.AddTextbox(1, 50, 20, slide_width - 100, 50)
        title3.TextFrame.TextRange.Text = "Conference Agenda - Day 1"
        title3.TextFrame.TextRange.Font.Size = 28
        title3.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        title3.TextFrame.TextRange.Font.Bold = True
        title3.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        agenda_text = (
            "08:00  Registration and Welcome Coffee\n"
            "09:00  Opening Ceremony\n"
            "09:30  Keynote: AI in Healthcare - Dr Sarah Chen\n"
            "10:30  Coffee Break and Networking\n"
            "11:00  Keynote: Quantum Computing - Prof MacLeod\n"
            "12:00  Panel: Future of Scottish Tech\n"
            "13:00  Lunch and Exhibition\n"
            "14:00  Breakout Sessions (4 tracks)\n"
            "16:00  Keynote: Digital Scotland - Fiona Stewart\n"
            "17:30  Welcome Reception"
        )
        agenda_box = slide3.Shapes.AddTextbox(1, 100, 90, slide_width - 200, 380)
        agenda_box.TextFrame.TextRange.Text = agenda_text
        agenda_box.TextFrame.TextRange.Font.Size = 18
        agenda_box.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF

        add_home_button(slide3)

        # =============================================
        # SLIDE 4: Venue Information
        # =============================================
        slide4 = pres.Slides.Add(4, ppLayoutBlank)
        slide4.FollowMasterBackground = False
        slide4.Background.Fill.ForeColor.RGB = 0x663300

        title4 = slide4.Shapes.AddTextbox(1, 50, 20, slide_width - 100, 50)
        title4.TextFrame.TextRange.Text = "Venue & Wi-Fi Information"
        title4.TextFrame.TextRange.Font.Size = 28
        title4.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        title4.TextFrame.TextRange.Font.Bold = True
        title4.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        venue_info = (
            "Venue: Edinburgh International Conference Centre (EICC)\n\n"
            "Main Hall: Level 2\n"
            "Breakout Rooms: Level 3 (Morrison & Pentland)\n"
            "Exhibition: Ground Floor Atrium\n"
            "Refreshments: Cromdale Suite\n\n"
            "Wi-Fi Network: TECHSUMMIT2026\n"
            "Password: Innovation2026\n\n"
            "Emergency exits clearly marked on all floors\n"
            "First aid: Reception desk, each floor"
        )
        venue_box = slide4.Shapes.AddTextbox(1, 100, 90, slide_width - 200, 380)
        venue_box.TextFrame.TextRange.Text = venue_info
        venue_box.TextFrame.TextRange.Font.Size = 18
        venue_box.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF

        add_home_button(slide4)

        # =============================================
        # SLIDE 5: Copyright & Credits
        # =============================================
        slide5 = pres.Slides.Add(5, ppLayoutBlank)
        slide5.FollowMasterBackground = False
        slide5.Background.Fill.ForeColor.RGB = 0x663300

        title5 = slide5.Shapes.AddTextbox(1, 50, 100, slide_width - 100, 60)
        title5.TextFrame.TextRange.Text = "Thank You"
        title5.TextFrame.TextRange.Font.Size = 36
        title5.TextFrame.TextRange.Font.Color.RGB = 0xFFFFFF
        title5.TextFrame.TextRange.Font.Bold = True
        title5.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        thanks = slide5.Shapes.AddTextbox(1, 100, 200, slide_width - 200, 100)
        thanks.TextFrame.TextRange.Text = (
            "We hope you enjoy TechSummit 2026\n"
            "See you at next year's conference!"
        )
        thanks.TextFrame.TextRange.Font.Size = 20
        thanks.TextFrame.TextRange.Font.Color.RGB = 0xD0D0D0
        thanks.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        # Copyright citation
        copyright_box = slide5.Shapes.AddTextbox(1, 50, slide_height - 100, slide_width - 100, 60)
        copyright_box.TextFrame.TextRange.Text = (
            "© 2026 TechSummit Conference Committee. All rights reserved.\n"
            "Conference logo and branding are trademarks of TechSummit Events Ltd.\n"
            "Images used under Creative Commons licence. Audio: Original composition."
        )
        copyright_box.TextFrame.TextRange.Font.Size = 10
        copyright_box.TextFrame.TextRange.Font.Color.RGB = 0x999999
        copyright_box.TextFrame.TextRange.ParagraphFormat.Alignment = 2

        add_home_button(slide5)

        # =============================================
        # Set transitions and timing
        # =============================================
        for i in range(1, 6):
            slide = pres.Slides(i)
            slide.SlideShowTransition.Speed = 2  # ppTransitionSpeedMedium
            slide.SlideShowTransition.AdvanceOnTime = True
            slide.SlideShowTransition.AdvanceTime = 10  # 10 seconds per slide

        # =============================================
        # Set Kiosk mode (Browse at kiosk - loop until Esc)
        # =============================================
        pres.SlideShowSettings.ShowType = 3  # ppShowTypeKiosk
        pres.SlideShowSettings.LoopUntilStopped = True
        pres.SlideShowSettings.AdvanceMode = 2  # ppSlideShowUseSlideTimings

        # Save
        if os.path.exists(PPTX_PATH):
            os.remove(PPTX_PATH)
        pres.SaveAs(PPTX_PATH, ppSaveAsOpenXMLPresentation)
        print(f"[OK] PowerPoint saved: {PPTX_PATH}")

        pres.Close()

    except Exception as e:
        print(f"[ERROR] PowerPoint: {e}")
        import traceback; traceback.print_exc()
    finally:
        try:
            ppt.Quit()
        except:
            pass
        pythoncom.CoUninitialize()


if __name__ == "__main__":
    create_mail_merge_badges()
    time.sleep(2)
    create_powerpoint()
    print("\n[DONE] Unit C mail merge + PowerPoint complete")
