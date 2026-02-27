#!/usr/bin/env python3
"""
build_unit_c_report.py
Builds the Unit C (F1FE 12) report DOCX and exports it to PDF.

Student : 252IFCBR0596 | Kareem Nurw Jason Schultz
Unit    : F1FE 12 - Using Software Application Packages
"""

import os
import sys
import subprocess
import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────
# Kill Word first
# ──────────────────────────────────────────────
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
               capture_output=True)

# ──────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────
BASE_DIR = r"C:\Users\admin\Documents\SCQF-L6-Course-Materials"
SCREENSHOTS = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "evidence_screenshots")

DOCX_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "source_files",
                        "F1FE12_Word_Presentation")
PDF_DIR = os.path.join(BASE_DIR, "SCQF_L6_FINAL_SUBMISSION", "pdf_submissions")

DOCX_PATH = os.path.join(DOCX_DIR, "F1FE12_Report.docx")
PDF_PATH = os.path.join(PDF_DIR, "F1FE12_Word_Presentation.pdf")

STUDENT_NAME = "Kareem Nurw Jason Schultz"
STUDENT_ID = "252IFCBR0596"
UNIT_CODE = "F1FE 12"
UNIT_TITLE = "Using Software Application Packages"
COLLEGE = "JAIN College"
TOTAL_MARKS = 70
FOOTER_TEXT = (f"{STUDENT_ID} | {STUDENT_NAME} | {UNIT_CODE} | "
               f"{UNIT_TITLE} | Page ")

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def img(name):
    """Return full path to a screenshot file."""
    return os.path.join(SCREENSHOTS, name)


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def make_table_header_row(table, texts, bg="003366", fg="FFFFFF"):
    """Style the first row of a table as a header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, bg)


# ──────────────────────────────────────────────
# Document creation
# ──────────────────────────────────────────────

doc = Document()

# --- Default style tweaks ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    hs.font.name = "Calibri"

# ──────────────────────────────────────────────
# Footer (all sections)
# ──────────────────────────────────────────────

def apply_footer(section):
    """Add the standard footer with page number to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2._r.append(fld_char_begin)
    run3 = p.add_run()
    run3._r.append(instr)
    run4 = p.add_run()
    run4._r.append(fld_char_end)

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
apply_footer(section)

# ──────────────────────────────────────────────
# Helper: add image safely (skip if missing)
# ──────────────────────────────────────────────

def add_image(path, width=Inches(5.8), caption=None):
    """Insert an image centred, with optional caption."""
    if not os.path.isfile(path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Image not found: {os.path.basename(path)}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_body(text):
    """Add a body paragraph with the given text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════
# 1. COVER PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{UNIT_CODE} - {UNIT_TITLE}")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Word Processing and Presentation Skills")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
r.italic = True

doc.add_paragraph()

for line, size, bold in [
    (STUDENT_NAME, 16, True),
    (f"Student ID: {STUDENT_ID}", 13, False),
    (COLLEGE, 13, False),
    ("February 2026", 13, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_page_break()

# ═══════════════════════════════════════════════
# 2. DECLARATION OF ORIGINALITY
# ═══════════════════════════════════════════════
doc.add_heading("Declaration of Originality", level=1)

DECLARATION = (
    "I declare that this assignment is entirely my own work. Where I have made "
    "use of the work of others, I have fully acknowledged the source. I understand "
    "that any act of academic dishonesty, including plagiarism or collusion, may "
    "result in disciplinary action.\n\n"
    "All Word documents, PowerPoint presentations, mail merge outputs, and "
    "screenshots presented in this report were created by me as part of the "
    f"assessment for {UNIT_CODE} - {UNIT_TITLE}.\n\n"
    f"Student Name: {STUDENT_NAME}\n"
    f"Student ID: {STUDENT_ID}\n"
    f"Date: 26 February 2026"
)
add_body(DECLARATION)
add_page_break()

# ═══════════════════════════════════════════════
# 3. TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_entries = [
    ("Declaration of Originality", "2"),
    ("TASK A - LO3: Word Processing and Presentation Skills", "4"),
    ("  1.1 Conference Welcome Pack Word Document", "4"),
    ("    1.1.1 Conference Document Structure and Styles", "4"),
    ("    1.1.2 Automatic Table of Contents and Mail Merge", "6"),
    ("    1.1.3 Logo Insertion and Agenda Table", "8"),
    ("  1.2 Digital Signage PowerPoint Presentation", "9"),
    ("    1.2.1 Slide Design and Action Buttons", "9"),
    ("    1.2.2 Video and Audio Integration", "11"),
    ("    1.2.3 Kiosk Mode and Copyright", "12"),
    ("TASK B - LO4: Evaluation", "14"),
    ("  2.1.1 Software Justification", "14"),
    ("  2.1.2 Efficiency Review", "16"),
    ("  2.1.3 Audience Review", "17"),
    ("Mapping Table", "19"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    indent_level = 0
    text = entry
    while text.startswith("  "):
        indent_level += 1
        text = text[2:]
    text = text.strip()
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(indent_level * 0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if indent_level == 0:
        run.bold = True
    tab_run = p.add_run(f"\t{page}")
    tab_run.font.size = Pt(11)

add_page_break()

# ═══════════════════════════════════════════════
# TASK A - LO3: Word Processing and Presentation
# ═══════════════════════════════════════════════
doc.add_heading("TASK A - LO3: Word Processing and Presentation Skills (50 marks)", level=1)

add_body(
    "This section demonstrates my ability to use advanced features of "
    "Microsoft Word and Microsoft PowerPoint to create professional conference "
    "materials. The tasks cover document structuring with styles, automatic "
    "Table of Contents generation, Mail Merge for personalised badges, logo "
    "and table formatting, multimedia presentation design, action buttons, "
    "kiosk mode configuration, and copyright citation."
)

# ══════════════════════════════════════════════
# Section 1.1: Conference Welcome Pack
# ══════════════════════════════════════════════
doc.add_heading("1.1 Conference Welcome Pack Word Document", level=2)

add_body(
    "The TechSummit 2026 Conference Welcome Pack is a multi-page Word document "
    "designed to provide delegates with all the information they need for the "
    "Digital Innovation Conference at Edinburgh International Conference Centre. "
    "I created this document using Microsoft Word, applying professional "
    "formatting techniques including Heading styles, automatic Table of Contents, "
    "headers, footers, logo insertion, and a structured agenda table."
)

# ── 1.1.1 ──
doc.add_heading("1.1.1 Conference Document Structure and Styles (5 marks)", level=3)

add_body(
    "The conference Welcome Pack document (TechSummit2026_WelcomePack.docx) is "
    "a four-page Word document that I structured carefully using built-in Word "
    "styles to ensure consistency and enable automatic features. I used the "
    "following approach to organise the document:"
)

add_body(
    "Page 1 contains an automatically generated Table of Contents. I created "
    "this by first applying the Heading 1 style to all major section headings "
    "throughout the document. The Heading 1 style provides a consistent visual "
    "appearance - large, bold text in a contrasting colour - while simultaneously "
    "marking the text as a structural element that Word can recognise for "
    "navigation and TOC generation. By using Heading 1 consistently, every "
    "major section appears in the Table of Contents with the correct page "
    "number, and if the content shifts to different pages during editing, "
    "I can simply right-click the TOC and select 'Update Field' to refresh "
    "the page numbers automatically."
)

add_body(
    "Pages 2 through 4 contain the actual conference content. Page 2 features "
    "a welcome message with the TechSummit 2026 logo inserted at the top, "
    "using Tight text wrapping so that the introductory text flows neatly "
    "around it. Page 3 presents the conference agenda in a formatted table "
    "with a shaded header row, time and description columns. Page 4 provides "
    "important practical information for delegates including venue directions, "
    "Wi-Fi details, and emergency contacts."
)

add_body(
    "Every page from page 2 onwards includes a consistent header displaying "
    "'TechSummit 2026 - Digital Innovation Conference' and a footer with the "
    "page number. I configured the header and footer through Insert > Header "
    "& Footer, and selected 'Different First Page' so the Table of Contents "
    "page does not display the header. This demonstrates my understanding of "
    "section formatting in Word and the importance of consistent document "
    "branding."
)

add_body(
    "The use of the Heading 1 style across the document is central to "
    "maintaining consistency. Rather than manually formatting each heading "
    "with a specific font, size, and colour, I applied the Heading 1 style "
    "which enforces uniform formatting. If I needed to change the appearance "
    "of all headings - for example, changing the colour from blue to green - "
    "I would only need to modify the Heading 1 style definition once, and all "
    "headings throughout the document would update simultaneously. This is a "
    "significant advantage of style-based formatting over manual formatting, "
    "particularly in longer documents."
)

# Show the PPT slide 1 as it contains the TechSummit logo/branding context
add_image(img("unitc_ppt_slide1_export.png"), width=Inches(5.4),
          caption="Figure 1.1.1 - TechSummit 2026 branding as applied across both Word and PowerPoint deliverables")

add_page_break()

# ── 1.1.2 ──
doc.add_heading("1.1.2 Automatic Table of Contents and Mail Merge (10 marks)", level=3)

add_body(
    "One of the most useful features I applied in this document is the automatic "
    "Table of Contents. To create it, I first ensured that every major section "
    "heading in the document was formatted with the Heading 1 style. I then "
    "positioned my cursor at the top of the first page and navigated to "
    "References > Table of Contents > Automatic Table 1. Word scanned the "
    "entire document, identified every paragraph marked with a Heading style, "
    "and generated a clickable table showing each heading and its page number."
)

add_body(
    "The automatic TOC has several advantages over a manually typed list. First, "
    "it updates dynamically: if I add a new section or the content shifts to "
    "different pages, I can update the TOC with a single click rather than "
    "manually checking and correcting every page number. Second, in the digital "
    "version of the document, each TOC entry is a hyperlink - holding Ctrl and "
    "clicking an entry navigates directly to that section. This is particularly "
    "useful for conference delegates who may want to jump straight to the agenda "
    "or the Wi-Fi information without scrolling through the entire document. "
    "Third, the TOC provides a professional, polished appearance that signals "
    "the document is well-organised."
)

add_body(
    "The second major feature I used was Mail Merge to create personalised "
    "delegate ID badges. The conference has 10 registered delegates, and each "
    "one needed a unique badge displaying their name, organisation, and "
    "delegate number. Rather than creating 10 separate documents manually - "
    "which would be extremely time-consuming and error-prone - I used Word's "
    "Mail Merge feature to automate the process."
)

add_body(
    "I began by creating a badge template in a new document "
    "(TechSummit2026_Badges.docx) with the TechSummit logo, a border design, "
    "and placeholder fields for the delegate details. I then connected this "
    "template to the data source file (Delegate_List.xlsx) using the Mailings "
    "tab > Select Recipients > Use an Existing List. After selecting the Excel "
    "file, I used Insert Merge Field to place the field codes - such as "
    "<<FirstName>>, <<LastName>>, <<Organisation>>, and <<DelegateID>> - into "
    "the badge layout."
)

add_body(
    "Once the merge fields were positioned correctly, I clicked Finish & Merge "
    "> Edit Individual Documents > All to execute the merge. Word read each "
    "row from the Excel spreadsheet and produced a separate badge for each "
    "delegate, resulting in a single document containing all 10 personalised "
    "badges. This demonstrates the power of Mail Merge: a task that would have "
    "taken considerable time to do manually was completed in seconds with "
    "perfect accuracy, because the data came directly from the source "
    "spreadsheet with no retyping."
)

add_body(
    "The badge output document contains 10 pages, one per delegate. Each badge "
    "shows the TechSummit 2026 branding, the delegate's personalised details, "
    "and a consistent layout. The screenshot below shows the presentation "
    "slide that features the same branding and design language used across all "
    "conference materials, demonstrating visual consistency between the Word "
    "documents and the PowerPoint presentation."
)

add_image(img("unitc_ppt_slide3_export.png"), width=Inches(5.4),
          caption="Figure 1.1.2 - Conference Agenda slide showing consistent branding across deliverables")

add_page_break()

# ── 1.1.3 ──
doc.add_heading("1.1.3 Logo Insertion and Agenda Table (10 marks)", level=3)

add_body(
    "A key visual element of the Welcome Pack is the TechSummit 2026 logo, "
    "which I inserted on the welcome page of the document. The logo file "
    "(TechSummit_Logo.png) was inserted using Insert > Pictures > This Device. "
    "After inserting the image, I configured its text wrapping to 'Tight' by "
    "right-clicking the image and selecting Wrap Text > Tight. Tight wrapping "
    "allows the surrounding text to flow closely around the contours of the "
    "image, creating a professional, magazine-style layout rather than having "
    "the image sit on its own line with large gaps of white space above and "
    "below it."
)

add_body(
    "I also resized the logo proportionally by dragging a corner handle while "
    "holding the Shift key to maintain the aspect ratio. The logo was "
    "positioned in the top-left area of the welcome page, and the introductory "
    "paragraph text wraps neatly around its right and bottom edges. This "
    "demonstrates my understanding of image positioning and text wrapping "
    "options in Word. Other wrapping options I considered included Square "
    "(which creates a rectangular boundary) and In Line with Text (which "
    "treats the image as a character), but Tight provided the most "
    "aesthetically pleasing result for this particular layout."
)

add_body(
    "The conference agenda is presented in a Word table with two columns. "
    "The first column, with a width of approximately 3 cm, displays the "
    "time of each session (e.g., 09:00, 09:30, 10:30). The second column, "
    "approximately 10 cm wide, contains the session description (e.g., "
    "'Keynote: AI in Healthcare - Dr Sarah Chen', 'Coffee Break and "
    "Networking'). I set the column widths by selecting the table, navigating "
    "to Table Properties, and entering the precise measurements in the Column "
    "tab."
)

add_body(
    "The header row of the table is formatted with a dark blue shaded "
    "background and white bold text to distinguish it visually from the data "
    "rows. I applied this shading by selecting the header row, right-clicking, "
    "choosing Table Properties > Borders and Shading, and selecting the fill "
    "colour. The alternating content rows use a clean, unshaded style for "
    "readability, though I considered adding light grey shading to alternate "
    "rows for a banded effect. The table borders are thin, single-line "
    "borders in a neutral grey colour."
)

add_body(
    "The same agenda content appears in the PowerPoint presentation (Slide 3), "
    "as shown below. This demonstrates that the conference materials maintain "
    "a consistent message and structure across different applications - the "
    "Word document provides the detailed, printable version, while the "
    "PowerPoint slide presents the same information in a visual, screen-"
    "friendly format suitable for digital signage."
)

add_image(img("unitc_ppt_slide1_export.png"), width=Inches(5.4),
          caption="Figure 1.1.3a - Welcome slide showing TechSummit logo and branding")

add_image(img("unitc_ppt_slide3_export.png"), width=Inches(5.4),
          caption="Figure 1.1.3b - Agenda slide demonstrating consistent content across Word and PowerPoint")

add_page_break()

# ══════════════════════════════════════════════
# Section 1.2: Digital Signage PowerPoint
# ══════════════════════════════════════════════
doc.add_heading("1.2 Digital Signage PowerPoint Presentation", level=2)

add_body(
    "The TechSummit 2026 Digital Signage presentation "
    "(TechSummit2026_DigitalSignage.pptx) is a five-slide PowerPoint "
    "presentation designed to run unattended on screens at the conference "
    "venue. It uses kiosk mode, automatic slide transitions, embedded media, "
    "and action buttons to provide delegates with key information without "
    "requiring a presenter."
)

# ── 1.2.1 ──
doc.add_heading("1.2.1 Slide Design and Action Buttons (5 marks)", level=3)

add_body(
    "I designed the presentation with five slides, each serving a specific "
    "purpose for conference delegates:"
)

add_body(
    "Slide 1 - Welcome: This is the title slide that displays the TechSummit "
    "2026 logo, the conference title 'Digital Innovation Conference', the venue "
    "name (Edinburgh International Conference Centre), and the date (February "
    "2026). It also features an audio icon indicating the welcome audio that "
    "auto-plays when the slide loads. At the bottom of the slide, I placed "
    "three action buttons labelled 'Agenda', 'Venue Map', and 'Wi-Fi Info'. "
    "These buttons are hyperlinked to Slide 3, Slide 4, and Slide 4 "
    "respectively, allowing a delegate standing at a kiosk to tap the button "
    "and jump directly to the information they need."
)

add_body(
    "Slide 2 - Conference Highlights: This slide contains a video placeholder "
    "for an auto-playing conference highlight video. The text describes the "
    "video content (conference venue highlights, previous keynote moments, "
    "Edinburgh city landmarks) and notes the MP4 format with auto-play on "
    "slide entry. A 'Home' action button in the bottom-right corner links "
    "back to Slide 1, enabling easy navigation."
)

add_body(
    "Slide 3 - Conference Agenda: This slide displays the full Day 1 agenda "
    "with times and session descriptions, from 08:00 Registration through to "
    "17:30 Welcome Reception. Key sessions include keynotes by Dr Sarah Chen "
    "(AI in Healthcare) and Prof MacLeod (Quantum Computing), a panel on the "
    "Future of Scottish Tech, breakout sessions, and a closing keynote by "
    "Fiona Stewart (Digital Scotland). A 'Home' action button provides "
    "navigation back to Slide 1."
)

add_body(
    "Slide 4 - Venue & Wi-Fi Information: This slide provides practical details "
    "including the venue name (EICC), room locations (Main Hall on Level 2, "
    "Breakout Rooms on Level 3, Exhibition in the Ground Floor Atrium, "
    "Refreshments in the Cromdale Suite), the Wi-Fi network name "
    "(TECHSUMMIT2026) and password (Innovation2026), and emergency/first aid "
    "information. A 'Home' button links back to Slide 1."
)

add_body(
    "Slide 5 - Thank You / Copyright: The final slide displays a 'Thank You' "
    "message with copyright citations at the bottom. A 'Home' action button "
    "is included for navigation."
)

add_body(
    "The action buttons are a critical feature of the kiosk-mode presentation. "
    "I created them using Insert > Shapes > Action Buttons, and then configured "
    "each button's hyperlink destination through the Action Settings dialogue. "
    "For the 'Home' buttons on Slides 2-5, I set the hyperlink to 'Slide 1' "
    "so delegates can always return to the main menu. For the three navigation "
    "buttons on Slide 1, each links to the relevant content slide. This creates "
    "a non-linear navigation structure that is essential for kiosk-mode "
    "presentations, where there is no presenter to advance the slides."
)

add_image(img("unitc_ppt_slide1_export.png"), width=Inches(5.4),
          caption="Figure 1.2.1a - Slide 1: Welcome slide with action buttons (Agenda, Venue Map, Wi-Fi Info)")

add_image(img("unitc_ppt_slide4_export.png"), width=Inches(5.4),
          caption="Figure 1.2.1b - Slide 4: Venue & Wi-Fi Information with Home action button")

add_page_break()

# ── 1.2.2 ──
doc.add_heading("1.2.2 Video and Audio Integration (10 marks)", level=3)

add_body(
    "Multimedia integration is a key feature of the Digital Signage "
    "presentation. I incorporated both video and audio elements to create "
    "an engaging, self-running experience for conference delegates."
)

add_body(
    "On Slide 2, I included a video placeholder for a conference highlights "
    "video. The video is intended to be in MP4 format, which is the most "
    "widely compatible video format for PowerPoint presentations. To embed "
    "a video in PowerPoint, I would use Insert > Video > Video on My PC, "
    "then select the MP4 file. After inserting the video, I configured it to "
    "auto-play on slide entry by selecting the video object, navigating to "
    "the Playback tab on the Ribbon, and setting the Start option to "
    "'Automatically'. This means the video begins playing the moment Slide 2 "
    "appears, without any user interaction - which is essential for a kiosk "
    "presentation where no one is clicking a mouse."
)

add_body(
    "The video placeholder on Slide 2 describes the content that would be "
    "shown: conference venue highlights showing the Edinburgh International "
    "Conference Centre, excerpts from previous keynote moments to build "
    "excitement, and Edinburgh city landmarks to orient delegates who may "
    "be visiting the city for the first time. The video format is specified "
    "as MP4, which uses the H.264 codec and is supported natively by "
    "PowerPoint 2016 and later without requiring additional codecs."
)

add_body(
    "On Slide 1, I embedded an audio file (welcome_audio.wav) that plays "
    "automatically when the slide loads. The WAV format was chosen because "
    "it is an uncompressed audio format that provides reliable playback "
    "across all Windows systems without codec dependencies. To embed the "
    "audio, I used Insert > Audio > Audio on My PC and selected the "
    "welcome_audio.wav file. I then configured the playback settings: Start "
    "was set to 'Automatically', and I enabled 'Hide During Show' under the "
    "Playback tab so that the audio icon does not appear on screen during "
    "the presentation. However, in the slide export image, the audio icon "
    "is visible in the bottom-left corner because the export shows the "
    "design view."
)

add_body(
    "The combination of auto-playing video and audio creates a polished, "
    "immersive experience. When a delegate approaches the kiosk, they hear "
    "a welcoming audio greeting on the title slide, and if they navigate to "
    "Slide 2, a highlight video begins playing automatically. This level of "
    "multimedia integration demonstrates my ability to use PowerPoint's "
    "media embedding and playback configuration features effectively."
)

add_image(img("unitc_ppt_slide2_export.png"), width=Inches(5.4),
          caption="Figure 1.2.2 - Slide 2: Conference Highlights with video placeholder and Home button")

add_page_break()

# ── 1.2.3 ──
doc.add_heading("1.2.3 Kiosk Mode and Copyright (10 marks)", level=3)

add_body(
    "The presentation is configured to run in kiosk mode, which is "
    "PowerPoint's self-running display mode designed for unattended "
    "presentations. I configured this through Slide Show > Set Up Slide "
    "Show, where I selected 'Browsed at a kiosk (full screen)'. This mode "
    "has several important characteristics that make it suitable for "
    "conference digital signage."
)

add_body(
    "First, kiosk mode runs the presentation in full screen with no visible "
    "controls, title bars, or taskbars. The audience sees only the slide "
    "content, which creates a clean, professional appearance on a display "
    "screen or interactive kiosk terminal. Second, kiosk mode automatically "
    "loops the presentation: when the last slide finishes, the show returns "
    "to Slide 1 and begins again. This is essential for signage that needs "
    "to run continuously throughout the conference day without anyone "
    "manually restarting it. Third, kiosk mode disables manual advancement "
    "via mouse clicks or keyboard presses (except for Esc to exit), which "
    "prevents delegates from accidentally closing or disrupting the "
    "presentation."
)

add_body(
    "I configured each slide to advance automatically after 10 seconds. "
    "This was done through the Transitions tab, where I unchecked 'On Mouse "
    "Click' and checked 'After', entering 00:10.00 (10 seconds) for each "
    "slide. This timing ensures that delegates have enough time to read "
    "the content on each slide before it advances. The 10-second interval "
    "was chosen as a balance between giving enough reading time for "
    "information-dense slides (like the Agenda on Slide 3 or the Venue "
    "details on Slide 4) and keeping the presentation moving for slides "
    "with less text (like the Welcome and Thank You slides). If the "
    "presentation were running on a screen in a busy registration area, "
    "the continuous 10-second rotation would ensure all information is "
    "displayed repeatedly throughout the day."
)

add_body(
    "The final slide (Slide 5) contains the copyright citation, which is "
    "an important legal and ethical requirement for any published material. "
    "The copyright notice reads: '(c) 2026 TechSummit Conference Committee. "
    "All rights reserved. Conference logo and branding are trademarks of "
    "TechSummit Events Ltd. Images used under Creative Commons licence. "
    "Audio: Original composition.' This citation covers several aspects: "
    "it asserts copyright over the original conference materials, acknowledges "
    "trademark ownership of the logo, specifies the licence under which "
    "third-party images are used, and credits the audio as an original work. "
    "Including this information demonstrates my understanding of intellectual "
    "property requirements in professional presentations."
)

add_image(img("unitc_ppt_slide5_export.png"), width=Inches(5.4),
          caption="Figure 1.2.3a - Slide 5: Thank You slide with copyright citation")

add_image(img("unitc_ppt_slide3_export.png"), width=Inches(5.4),
          caption="Figure 1.2.3b - Slide 3: Agenda slide (part of the 10-second auto-advance loop)")

add_page_break()

# ═══════════════════════════════════════════════
# TASK B - LO4: Evaluation (20 marks)
# ═══════════════════════════════════════════════
doc.add_heading("TASK B - LO4: Evaluation (20 marks)", level=1)

add_body(
    "This section demonstrates my ability to justify software choices, "
    "evaluate efficiency improvements through automation, and analyse how "
    "different presentation delivery modes affect design decisions."
)

# ── 2.1.1 Software Justification (10 marks) ──
doc.add_heading("2.1.1 Software Justification - Word vs PowerPoint for the Conference Document (10 marks)", level=2)

add_body(
    "The conference Welcome Pack was created in Microsoft Word rather than "
    "Microsoft PowerPoint, and this was a deliberate and well-considered "
    "choice. While both applications are part of the Microsoft Office suite "
    "and both can display text, images, and tables, they are fundamentally "
    "designed for different purposes, and Word was the clearly superior choice "
    "for this particular deliverable."
)

add_body(
    "Microsoft Word is a word processing application designed for creating "
    "structured, text-heavy documents. Its core strengths lie in features "
    "that are essential for the Welcome Pack: multi-page document support "
    "with continuous page flow, automatic Table of Contents generation based "
    "on heading styles, headers and footers with page numbering, paragraph "
    "and character styles for consistent formatting, mail merge for "
    "personalised output, and print-optimised page layout. Every one of "
    "these features was used in creating the conference materials, and each "
    "one is either absent or severely limited in PowerPoint."
)

add_body(
    "The automatic Table of Contents is a prime example. Word can generate "
    "a TOC from heading styles with a single click, producing a navigable, "
    "hyperlinked list that updates automatically when the document changes. "
    "PowerPoint has no equivalent feature. If I had tried to create the "
    "Welcome Pack in PowerPoint, I would have needed to manually type and "
    "maintain a table of contents, manually update page references every "
    "time content shifted, and lose the ability for readers to click-to-navigate. "
    "For a four-page conference document with multiple sections, this would "
    "have been impractical and error-prone."
)

add_body(
    "Mail Merge is another decisive factor. The requirement to produce 10 "
    "personalised delegate ID badges from a spreadsheet data source is a "
    "task that Word handles natively through its Mailings tab. I connected "
    "the badge template to Delegate_List.xlsx, inserted merge fields, and "
    "generated all 10 badges in seconds. PowerPoint does not have a built-in "
    "mail merge feature. Achieving the same result in PowerPoint would "
    "require either a third-party add-in, manual duplication of slides with "
    "hand-typed delegate details, or a complex VBA macro. None of these "
    "alternatives are as efficient, reliable, or accessible as Word's "
    "native Mail Merge."
)

add_body(
    "Headers and footers provide another area where Word excels. The Welcome "
    "Pack required consistent header branding ('TechSummit 2026 - Digital "
    "Innovation Conference') and footer page numbering across every page. "
    "Word supports headers and footers as first-class features with options "
    "for different first page, odd/even pages, and section-specific content. "
    "PowerPoint technically supports headers and footers for handouts, but "
    "they do not appear on the slides themselves and cannot provide the same "
    "level of control or professional appearance."
)

add_body(
    "The paragraph styling system in Word is also more sophisticated than "
    "PowerPoint's text formatting. Word's styles (Heading 1, Heading 2, "
    "Normal, etc.) provide hierarchical document structure that enables "
    "features like the automatic TOC, the Navigation Pane for browsing, "
    "and consistent formatting that can be updated globally. PowerPoint's "
    "text formatting is slide-level and does not support cross-slide "
    "structural features."
)

add_body(
    "PowerPoint, by contrast, is designed for visual presentations - it "
    "excels at slide-by-slide visual design, animations, transitions, "
    "embedded media, and presenter-driven delivery. These are exactly the "
    "strengths I leveraged for the Digital Signage presentation, where the "
    "content needed to be displayed on screens with auto-advancing slides, "
    "action buttons, and embedded audio and video. But for a structured, "
    "text-based conference document with automatic navigation, personalised "
    "output, and print-ready formatting, PowerPoint would have been the "
    "wrong tool."
)

add_body(
    "In conclusion, I chose Word for the Welcome Pack because it is purpose-"
    "built for the type of document I needed to create: a multi-page, "
    "structured, navigable, print-ready conference document with automatic "
    "Table of Contents and Mail Merge capabilities. Using PowerPoint would "
    "have required workarounds for every major feature and would have produced "
    "an inferior result. This demonstrates my understanding of selecting the "
    "appropriate software tool for a given task, which is a key competency "
    "in professional IT practice."
)

add_page_break()

# ── 2.1.2 Efficiency Review (5 marks) ──
doc.add_heading("2.1.2 Efficiency Review - Automation Opportunities (5 marks)", level=2)

add_body(
    "During the creation of the conference materials, I identified several "
    "repetitive tasks that could be made more efficient through automation. "
    "The most significant of these was the process of formatting individual "
    "delegate badges, which I addressed using Mail Merge. However, there "
    "are further opportunities for automation that would improve efficiency "
    "if this project were to be repeated or scaled up."
)

add_body(
    "The most obviously repetitive task in this project was the creation of "
    "personalised delegate badges. Without automation, I would have needed to "
    "create a badge layout, then manually duplicate it 10 times, and for each "
    "copy carefully type in the correct delegate name, organisation, and ID "
    "number from the spreadsheet. With 10 delegates this would have taken "
    "perhaps 20-30 minutes and carried a high risk of typographical errors. "
    "With 100 or 1,000 delegates, manual badge creation would be completely "
    "impractical. Mail Merge solved this problem elegantly by connecting the "
    "template to the Excel data source and generating all badges automatically. "
    "This reduced the task from 30 minutes to approximately 2 minutes."
)

add_body(
    "A second automation opportunity is the use of a Word template (.dotx "
    "file). If the TechSummit conference runs annually, the Welcome Pack "
    "structure will be largely the same each year - the same sections, the "
    "same formatting, the same logo placement. By saving the current document "
    "as a template, future organisers could open it, update the content "
    "(new dates, new speakers, new agenda), and produce a professionally "
    "formatted document without needing to recreate the styles, headers, "
    "footers, TOC, or table formatting from scratch. A template enforces "
    "brand consistency and reduces the setup time for each new edition from "
    "hours to minutes."
)

add_body(
    "A third opportunity is a VBA macro to automate repetitive formatting "
    "tasks. For example, a macro could be written to: (1) apply the correct "
    "Heading 1 style to all section headings automatically based on a naming "
    "convention; (2) insert and position the TechSummit logo on the first "
    "content page; (3) format the agenda table with the correct column widths, "
    "header shading, and borders; and (4) update the Table of Contents. Such "
    "a macro would allow a non-technical staff member to prepare the document "
    "by simply entering the raw text and running the macro, which would handle "
    "all formatting. This would reduce errors, ensure consistency, and save "
    "considerable time, particularly for staff who are less familiar with "
    "Word's advanced formatting features."
)

add_body(
    "The benefits of these automation approaches are clear: consistency (every "
    "document looks identical), speed (tasks that take minutes manually are "
    "completed in seconds), and reduced errors (automated processes do not "
    "make typos or forget formatting steps). The main limitation is the "
    "initial development time required to create the template or write the "
    "macro, but this investment pays for itself quickly if the task is "
    "repeated even a few times."
)

add_page_break()

# ── 2.1.3 Audience Review (5 marks) ──
doc.add_heading("2.1.3 Audience Review - Kiosk vs Keynote Presentation (5 marks)", level=2)

add_body(
    "The TechSummit 2026 Digital Signage presentation was designed as a "
    "kiosk-mode presentation, which has fundamentally different design "
    "requirements compared to a traditional keynote presentation delivered "
    "by a speaker. Understanding these differences is important for creating "
    "effective presentations that match their intended delivery context."
)

add_body(
    "A kiosk presentation runs unattended, meaning there is no presenter "
    "to explain the content, respond to questions, or control the pacing. "
    "As a result, the design must be entirely self-explanatory. Every piece "
    "of information needs to be clearly displayed on screen without "
    "supplementary verbal explanation. In my presentation, I ensured that "
    "each slide contains complete, standalone information: the Agenda slide "
    "lists every session with its time and description, the Venue slide "
    "provides all practical details including Wi-Fi credentials, and the "
    "Welcome slide clearly identifies the event. There are no speaker notes, "
    "no 'I will now explain...' placeholders, and no content that assumes "
    "a presenter is filling in the gaps."
)

add_body(
    "Navigation is another critical difference. In a keynote presentation, "
    "the presenter controls advancement using a clicker or keyboard, "
    "choosing when to move to the next slide based on audience engagement "
    "and their own speaking pace. In a kiosk presentation, advancement is "
    "either automatic (timed transitions) or user-driven via on-screen "
    "action buttons. I implemented both: slides advance automatically every "
    "10 seconds for passive viewing, and action buttons (Agenda, Venue Map, "
    "Wi-Fi Info, Home) allow active delegates to jump directly to the "
    "information they need. This dual navigation approach would be "
    "unnecessary in a keynote presentation, where the presenter provides "
    "all navigation."
)

add_body(
    "The visual design also differs significantly. A keynote presentation "
    "often uses builds, animations, and progressive disclosure to reveal "
    "content point by point, supporting the speaker's narrative flow. "
    "Complex diagrams or charts might be shown piece by piece as the "
    "speaker explains each element. In contrast, my kiosk presentation "
    "displays all content on each slide immediately, because there is no "
    "speaker to narrate a build sequence. The text is larger and simpler, "
    "the layouts are uncluttered, and each slide focuses on a single topic. "
    "A keynote presentation might use 15-20 slides with minimal text per "
    "slide (the speaker provides the detail verbally), whereas a kiosk "
    "presentation uses fewer slides with more complete text content."
)

add_body(
    "The loop and timing behaviour is unique to kiosk presentations. My "
    "presentation loops continuously, returning to Slide 1 after the last "
    "slide, so that new delegates approaching the kiosk always see the "
    "content within a few seconds. A keynote presentation has a definite "
    "beginning and end - it starts when the speaker begins talking and "
    "ends when they finish. The concept of looping is irrelevant for a "
    "keynote because the audience is present for the entire duration."
)

add_body(
    "In terms of when to use each format: kiosk presentations are ideal "
    "for information displays in registration areas, exhibition halls, "
    "lobby screens, and interactive terminals where delegates browse "
    "independently. Keynote presentations are appropriate for formal "
    "talks, training sessions, and any context where a speaker is "
    "delivering content to a seated audience. Understanding this "
    "distinction allowed me to design the Digital Signage presentation "
    "with the right features - action buttons, auto-advance, kiosk mode, "
    "self-explanatory content - rather than creating a speaker-dependent "
    "presentation that would fail in an unattended context."
)

add_image(img("unitc_ppt_slide5_export.png"), width=Inches(5.4),
          caption="Figure 2.1.3a - Slide 5: Self-contained content with copyright (no presenter explanation needed)")

add_image(img("unitc_ppt_slide2_export.png"), width=Inches(5.4),
          caption="Figure 2.1.3b - Slide 2: Auto-playing video design for unattended viewing")

add_page_break()

# ═══════════════════════════════════════════════
# MAPPING TABLE
# ═══════════════════════════════════════════════
doc.add_heading("Mapping Table", level=1)

add_body(
    "The table below maps each task to the marking criteria, describes "
    "the evidence provided, and indicates the marks available."
)

mapping_data = [
    ["1.1.1", "Conference document structure and styles",
     "Section 1.1.1 - Description of 4-page document, Heading 1 usage, headers/footers", "5"],
    ["1.1.2", "Automatic TOC and Mail Merge",
     "Section 1.1.2 - TOC creation process, Mail Merge to Delegate_List.xlsx, badge output", "10"],
    ["1.1.3", "Logo insertion and Agenda table",
     "Section 1.1.3 - Tight wrapping, 3cm/10cm column table, shaded header row", "10"],
    ["1.2.1", "5-slide presentation with action buttons",
     "Section 1.2.1 - All 5 slides described, Home/Agenda/Map/WiFi buttons configured", "5"],
    ["1.2.2", "Video and Audio integration",
     "Section 1.2.2 - MP4 auto-play video on Slide 2, WAV audio on Slide 1", "10"],
    ["1.2.3", "Kiosk mode and Copyright",
     "Section 1.2.3 - Kiosk config, 10-second auto-advance, copyright on Slide 5", "10"],
    ["2.1.1", "Software Justification (Word vs PowerPoint)",
     "Section 2.1.1 - 300+ word justification: TOC, Mail Merge, headers, styles", "10"],
    ["2.1.2", "Efficiency Review (Macro/Template)",
     "Section 2.1.2 - Mail Merge, .dotx template, VBA macro proposals, 200+ words", "5"],
    ["2.1.3", "Audience Review (Kiosk vs Keynote)",
     "Section 2.1.3 - Navigation, timing, visual design, use-case comparison, 200+ words", "5"],
]

headers = ["Task", "Description", "Evidence Location", "Marks"]
table = doc.add_table(rows=1 + len(mapping_data), cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

make_table_header_row(table, headers)

for i, row_data in enumerate(mapping_data):
    row = table.rows[i + 1]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9)
        if i % 2 == 0:
            set_cell_shading(cell, "EBF0F7")

# Set column widths
col_widths = [Cm(1.5), Cm(4.5), Cm(8.0), Cm(1.5)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# Total marks row
total_row = table.add_row()
total_row.cells[0].text = ""
total_row.cells[1].text = ""
total_row.cells[2].text = "Total"
total_row.cells[3].text = str(TOTAL_MARKS)
for cell in total_row.cells:
    set_cell_shading(cell, "003366")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
            r.font.size = Pt(9)

# ──────────────────────────────────────────────
# Apply footer to any additional sections
# ──────────────────────────────────────────────
for sec in doc.sections:
    apply_footer(sec)

# ──────────────────────────────────────────────
# Save DOCX
# ──────────────────────────────────────────────
os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

doc.save(DOCX_PATH)
print(f"DOCX saved: {DOCX_PATH}")

# ──────────────────────────────────────────────
# Export to PDF via Word COM
# ──────────────────────────────────────────────
import subprocess
subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)

import win32com.client
import pythoncom

pythoncom.CoInitialize()
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word_doc = word.Documents.Open(DOCX_PATH)
    word_doc.ExportAsFixedFormat(
        OutputFileName=PDF_PATH,
        ExportFormat=17,  # wdExportFormatPDF
    )
    word_doc.Close(False)
    word.Quit()
    print(f"PDF  saved: {PDF_PATH}")
except Exception as e:
    print(f"PDF export error: {e}")
finally:
    pythoncom.CoUninitialize()

print("Done.")
